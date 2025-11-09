import datetime
import difflib
import json
import os
import random
import re
import sys
import tempfile
import traceback
from io import BytesIO
from pathlib import Path
import google.generativeai as genai
import mammoth
import pandas as pd
import PyPDF2
from docx import Document
from flask import (Blueprint, Flask, flash, jsonify, redirect,
                   render_template, request, send_file, send_from_directory, url_for)
from flask_login import (LoginManager, UserMixin, current_user,
                         login_required, login_user, logout_user)
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import func, text
from werkzeug.security import check_password_hash, generate_password_hash
from xhtml2pdf import pisa

# Define the root directory (parent of backend)
ROOT_DIR = Path(__file__).parent.parent

# Add the root directory to the Python path
sys.path.insert(0, str(ROOT_DIR))

# Flask app with proper template and static configuration
app = Flask(__name__, 
            template_folder=os.path.join(ROOT_DIR, 'frontend'),
            static_folder=os.path.join(ROOT_DIR, 'img'))

# Add secret key for session
app.secret_key = '70315ac5aa5fee60d775f2ad95d2a163b24a4ba65583df64620acc3a283ccbc8'  # Replace with a real secret key in production

# Initialize Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'  # Specify your login route

@login_manager.user_loader
def load_user(user_id):
    if isinstance(user_id, str):
        if user_id.startswith('guru_'):
            actual_id = int(user_id.split('_')[1])
            return db.session.get(Guru, actual_id)
        elif user_id.startswith('siswa_'):
            actual_id = int(user_id.split('_')[1])
            return db.session.get(Siswa, actual_id)
    return None

# Konfigurasi MySQL
app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+pymysql://root:@localhost/asesment_diagnostik_db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploads')
app.config['SQLALCHEMY_ECHO'] = True  # Tambahan untuk debug SQL queries
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

db = SQLAlchemy(app)

# =========================================
# KONFIGURASI GEMINI AI
# =========================================
# Gunakan environment variable untuk API key
my_api_key_gemini = 'AIzaSyCcx9iCB6oGXUPGoW2Ot3Mk7JjIreeCzRQ' 
genai.configure(api_key=my_api_key_gemini)
model = genai.GenerativeModel('gemini-2.5-flash-lite')

# =========================================
# MST ADAPTIVE SYSTEM CONSTANTS
# =========================================

# Technology Levels (5 Levels) - Updated for 5 Stage MST
TECHNOLOGY_LEVELS = {
    1: "Awareness",      # L1 - Mengenali/mengingat istilah, alat dasar (Medium only)
    2: "Literacy",       # L2 - Menjelaskan fungsi, cara kerja sederhana (Easy, Hard)
    3: "Capability",     # L3 - Menggunakan fitur dasar untuk tugas rutin (Medium, Hard)
    4: "Creativity",     # L4 - Merancang/membuat solusi/produk baru sederhana (Easy, Medium, Hard)
    5: "Criticism"       # L5 - Mengevaluasi dampak, efektivitas, memberi justifikasi (Easy, Medium, Hard)
}

# Difficulty Levels per Technology Level (original MST system)
DIFFICULTY_LEVELS = {
    1: ["Medium"],                    # L1 Awareness: Medium only
    2: ["Easy", "Hard"],             # L2 Literacy: Easy, Hard  
    3: ["Medium", "Hard"],           # L3 Capability: Medium, Hard
    4: ["Easy", "Medium", "Hard"],   # L4 Creativity: Easy, Medium, Hard
    5: ["Easy", "Medium", "Hard"]    # L5 Criticism: Easy, Medium, Hard
}

# Conversion functions between p-value and difficulty level
def difficulty_to_p_value(difficulty):
    """Convert difficulty level (Easy/Medium/Hard) to p-value (probability)"""
    conversion_map = {
        "Easy": 0.85,      # Higher probability = easier question
        "Medium": 0.60,    # Medium probability
        "Hard": 0.25       # Lower probability = harder question
    }
    return conversion_map.get(difficulty, 0.60)  # Default to Medium

def p_value_to_difficulty(p_value):
    """Convert p-value (probability) to difficulty level (Easy/Medium/Hard)"""
    if p_value >= 0.75:
        return "Easy"
    elif p_value >= 0.45:
        return "Medium"
    else:
        return "Hard"

# MST Routing Logic - Simple 5 Stage System with User's Taxonomy  
# Implements 4 Scenarios: High Ability (L1->L2H->L3H->L4H->L5H), Medium-High (L1->L2H->L3H->L4M->L5M->STOP),
# Medium-Low (L1->L2H->L3M->L4E->STOP), Low Ability (L1->L2E->STOP)
MST_ROUTING = {
    # Stage 1: L1 Awareness (Medium only) - Starting point for all students
    1: {
        "technology_level": 1, 
        "difficulty": "Medium",
        "pass": {"stage": 2, "technology_level": 2, "difficulty": "Hard"},    # L1M Pass -> L2H (Stage 2)
        "fail": {"stage": 2, "technology_level": 2, "difficulty": "Easy"}     # L1M Fail -> L2E (Stage 2)  
    },
    # Stage 2: L2 Literacy (Easy or Hard based on Stage 1 result)
    2: {
        "Hard": {  # L2H from L1M Pass
            "pass": {"stage": 3, "technology_level": 3, "difficulty": "Hard"},    # L2H Pass -> L3H (Stage 3)
            "fail": {"stage": 3, "technology_level": 3, "difficulty": "Medium"}   # L2H Fail -> L3M (Stage 3)
        },
        "Easy": {  # L2E from L1M Fail
            "pass": {"stage": 3, "technology_level": 3, "difficulty": "Medium"},  # L2E Pass -> L3M (Stage 3)
            "fail": {"stage": "STOP", "diagnosis": "< L2"}                        # L2E Fail -> STOP (Diagnosis: < L2)
        }
    },
    # Stage 3: L3 Capability (Medium or Hard based on Stage 2 result)
    3: {
        "Hard": {  # L3H from L2H Pass  
            "pass": {"stage": 4, "technology_level": 4, "difficulty": "Hard"},    # L3H Pass -> L4H (Stage 4)
            "fail": {"stage": 4, "technology_level": 4, "difficulty": "Medium"}   # L3H Fail -> L4M (Stage 4)
        },
        "Medium": {  # L3M from L2H Fail or L2E Pass
            "pass": {"stage": 4, "technology_level": 4, "difficulty": "Medium"},  # L3M Pass -> L4M (Stage 4)  
            "fail": {"stage": 4, "technology_level": 4, "difficulty": "Easy"}     # L3M Fail -> L4E (Stage 4)
        }
    },
    # Stage 4: L4 Creativity (Easy, Medium, or Hard based on Stage 3 result)
    4: {
        "Hard": {  # L4H from L3H Pass (High Ability Path)
            "pass": {"stage": 5, "technology_level": 5, "difficulty": "Hard"},    # L4H Pass -> L5H (Stage 5)
            "fail": {"stage": 5, "technology_level": 5, "difficulty": "Medium"}   # L4H Fail -> L5M (Stage 5)
        },
        "Medium": {  # L4M from L3H Fail or L3M Pass (Medium-High Path)
            "pass": {"stage": 5, "technology_level": 5, "difficulty": "Medium"},  # L4M Pass -> L5M (Stage 5)
            "fail": {"stage": 5, "technology_level": 5, "difficulty": "Easy"}     # L4M Fail -> L5E (Stage 5)
        },
        "Easy": {  # L4E from L3M Fail (Lower Path)
            "pass": {"stage": 5, "technology_level": 5, "difficulty": "Easy"},    # L4E Pass -> L5E (Stage 5)
            "fail": {"stage": "STOP", "diagnosis": "L3"}                          # L4E Fail -> STOP (Diagnosis: L3)
        }
    },
    # Stage 5: L5 Criticism (Easy, Medium, or Hard based on Stage 4 result) - Final stage
    5: {
        "Hard": {  # L5H from L4H Pass (Highest ability path)
            "pass": {"stage": "STOP", "diagnosis": "L5"},                         # L5H Pass -> STOP (Diagnosis: L5)
            "fail": {"stage": "STOP", "diagnosis": "L4"}                          # L5H Fail -> STOP (Diagnosis: L4)
        },
        "Medium": {  # L5M from L4H Fail or L4M Pass (Medium-high path)
            "pass": {"stage": "STOP", "diagnosis": "L5"},                         # L5M Pass -> STOP (Diagnosis: L5)
            "fail": {"stage": "STOP", "diagnosis": "L4"}                          # L5M Fail -> STOP (Diagnosis: L4)
        },
        "Easy": {  # L5E from L4M Fail or L4E Pass (Lower ability path)
            "pass": {"stage": "STOP", "diagnosis": "L5"},                         # L5E Pass -> STOP (Diagnosis: L5)
            "fail": {"stage": "STOP", "diagnosis": "L4"}                          # L5E Fail -> STOP (Diagnosis: L4)
        }
    }
}

# =========================================
# MST ADAPTIVE FUNCTIONS
# =========================================

def enforce_alignment_with_objectives(questions, module_components):
    """Ensure each question's modul_reference maps to one of the extracted objectives/outcomes.
    If missing or not matching, auto-map to the closest objective/outcome.
    """
    try:
        tujuan_list = module_components.get("tujuan_pembelajaran", []) or []
        capaian_list = module_components.get("capaian_pembelajaran", []) or []
        allowed = [t.strip() for t in (tujuan_list + capaian_list) if isinstance(t, str) and t.strip()]
        if not allowed:
            return questions

        normalized_allowed = {a: a for a in allowed}
        allowed_lower = {a.lower(): a for a in allowed}

        def choose_best_reference(text):
            if not text or not text.strip():
                return allowed[0]
            txt = text.strip()
            # Exact (case-insensitive) match
            if txt.lower() in allowed_lower:
                return allowed_lower[txt.lower()]
            # Fuzzy match
            candidates = difflib.get_close_matches(txt, allowed, n=1, cutoff=0.6)
            return candidates[0] if candidates else allowed[0]

        for q in questions:
            current_ref = q.get("modul_reference", "")
            best = choose_best_reference(current_ref)
            q["modul_reference"] = best
        return questions
    except Exception as e:
        print(f"WARNING: enforce_alignment_with_objectives error: {e}")
        return questions

def get_questions_for_module(collection_id, technology_level, difficulty, limit=3):
    """
    Get questions for a specific module (technology level + difficulty)
    """
    try:
        questions = db.session.query(Question).join(
            CollectionQuestion, CollectionQuestion.question_id == Question.id
        ).filter(
            CollectionQuestion.collection_id == collection_id,
            Question.technology_level == technology_level,
            Question.difficulty == difficulty,
            Question.is_validated == True
        ).limit(limit).all()
        
        return questions
    except Exception as e:
        print(f"Error getting questions for module: {str(e)}")
        return []

def calculate_module_score(student_id, questions):
    """
    Calculate score for a module based on answered questions
    Returns percentage (0-100)
    """
    if not questions:
        return 0
    
    question_ids = [q.id for q in questions]
    correct_answers = db.session.query(SiswaAnswer).filter(
        SiswaAnswer.siswa_id == student_id,
        SiswaAnswer.question_id.in_(question_ids),
        SiswaAnswer.is_correct == True
    ).count()
    
    total_questions = len(questions)
    percentage = (correct_answers / total_questions) * 100 if total_questions > 0 else 0
    return percentage

# =========================================
# MST ADAPTIVE TRACKING SYSTEM
# =========================================

class MSTState(db.Model):
    """Track MST adaptive test state per siswa per collection"""
    __tablename__ = 'mst_state'
    
    id = db.Column(db.Integer, primary_key=True)
    siswa_id = db.Column(db.Integer, db.ForeignKey('siswa.id'), nullable=False)
    collection_id = db.Column(db.Integer, db.ForeignKey('question_collections.id'), nullable=False)
    current_stage = db.Column(db.Integer, default=1)  # 1-5
    current_level = db.Column(db.String(10), nullable=False)  # Easy, Medium, Hard
    questions_answered = db.Column(db.Integer, default=0)
    questions_correct = db.Column(db.Integer, default=0)
    stage_completed = db.Column(db.Boolean, default=False)
    final_diagnosis = db.Column(db.String(100))  # "Level X (Category)" atau "Di Bawah Level Y"
    test_completed = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    updated_at = db.Column(db.DateTime, server_default=db.func.now(), onupdate=db.func.now())
    
    # Relationships
    siswa = db.relationship('Siswa', backref='mst_states')
    collection = db.relationship('QuestionCollection', backref='mst_states')

def get_mst_state(siswa_id, collection_id):
    """Get atau create MST state untuk siswa & collection"""
    state = MSTState.query.filter_by(siswa_id=siswa_id, collection_id=collection_id).first()
    if not state:
        # Validasi siswa exists
        siswa = Siswa.query.get(siswa_id)
        if not siswa:
            raise ValueError(f"Siswa dengan ID {siswa_id} tidak ditemukan dalam database")
        
        # Validasi collection exists
        collection = QuestionCollection.query.get(collection_id)
        if not collection:
            raise ValueError(f"Collection dengan ID {collection_id} tidak ditemukan dalam database")
        
        # Start dengan Stage 1: L1 Medium
        state = MSTState(
            siswa_id=siswa_id,
            collection_id=collection_id,
            current_stage=1,
            current_level="Medium"
        )
        db.session.add(state)
        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            raise Exception(f"Gagal membuat MST state: {str(e)}")
    return state

def determine_next_stage(state):
    """
    Implementasi alur MST adaptif 5 Stage sesuai taxonomy user - 1 SOAL per stage
    Returns: (next_stage, next_level, should_stop, diagnosis)
    
    4 Skenario Alur MST:
    1. High Ability: L1M(PASS)->L2H(PASS)->L3H(PASS)->L4H(PASS)->L5H(PASS/FAIL)->STOP
    2. Medium-High: L1M(PASS)->L2H(PASS)->L3H(FAIL)->L4M(PASS)->L5M(PASS/FAIL)->STOP
    3. Medium-Low: L1M(PASS)->L2H(FAIL)->L3M(PASS)->L4E(PASS/FAIL)->STOP
    4. Low Ability: L1M(FAIL)->L2E(FAIL)->STOP
    """
    stage = state.current_stage
    level = state.current_level
    correct_answers = state.questions_correct
    
    # 1 soal per stage, jadi passed = (correct_answers >= 1)
    passed = (correct_answers >= 1)
    
    print(f"[MST] Stage {stage} L{stage}{level[0]}: {correct_answers} correct = {'PASS' if passed else 'FAIL'}")
    
    # Get routing info from MST_ROUTING
    if stage in MST_ROUTING:
        stage_config = MST_ROUTING[stage]
        
        # Stage 1 is special - no difficulty branching
        if stage == 1:
            route = stage_config["pass"] if passed else stage_config["fail"]
        else:
            # Stages 2-5 have difficulty-based routing
            if level in stage_config:
                route = stage_config[level]["pass"] if passed else stage_config[level]["fail"]
            else:
                print(f"[MST ERROR] Stage {stage} level {level} not found in routing config")
                return (None, None, True, "Configuration Error")
        
        # Process route result
        if route["stage"] == "STOP":
            return (None, None, True, route["diagnosis"])
        else:
            return (route["stage"], route["difficulty"], False, None)
    
    # Fallback - should not happen with proper config
    print(f"[MST ERROR] Stage {stage} not found in MST_ROUTING")
    return (None, None, True, f"L{stage}")
    # Default fallback (seharusnya tidak tercapai dengan diagram baru)
    return (None, None, True, f"Level {max(1, stage-1)}")

# =========================================
# PROGRESS TRACKING SYSTEM (DATABASE-BASED untuk Production)
# =========================================

# Model untuk progress tracking yang compatible dengan multiple workers
class ProgressTracker(db.Model):
    __tablename__ = 'progress_tracker'
    
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.String(50), nullable=False, index=True)
    current_step = db.Column(db.Integer, default=1)
    step_1_status = db.Column(db.String(20), default='pending')
    step_1_message = db.Column(db.String(255), default='Mengunggah file modul ajar')
    step_2_status = db.Column(db.String(20), default='pending')
    step_2_message = db.Column(db.String(255), default='Menganalisis struktur dokumen')
    step_3_status = db.Column(db.String(20), default='pending')
    step_3_message = db.Column(db.String(255), default='Mengekstrak tujuan pembelajaran')
    step_4_status = db.Column(db.String(20), default='pending')
    step_4_message = db.Column(db.String(255), default='Memproses dengan AI Gemini')
    step_5_status = db.Column(db.String(20), default='pending')
    step_5_message = db.Column(db.String(255), default='Menyimpan soal ke database')
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    updated_at = db.Column(db.DateTime, server_default=db.func.now(), onupdate=db.func.now())

    def __repr__(self):
        return f'<ProgressTracker {self.user_id}: Step {self.current_step}>'

def update_progress(user_id, step, status="active", message=""):
    """Update progress untuk user tertentu (Database-based untuk production)"""
    try:
        # Cari atau buat progress record untuk user
        progress = ProgressTracker.query.filter_by(user_id=str(user_id)).first()
        
        if not progress:
            # Buat progress record baru
            progress = ProgressTracker(
                user_id=str(user_id),
                current_step=1
            )
            db.session.add(progress)
            db.session.flush()  # Flush to get the ID
        
        # Update step spesifik
        if step in [1, 2, 3, 4, 5]:
            setattr(progress, f'step_{step}_status', status)
            if message:
                setattr(progress, f'step_{step}_message', message)
            progress.current_step = step
            progress.updated_at = datetime.datetime.now()
        
        db.session.commit()
        print(f"Progress updated in DB for user {user_id}: Step {step} - {status}")
        
    except Exception as e:
        db.session.rollback()
        print(f"Error updating progress for user {user_id}: {str(e)}")
        # Fallback ke memory-based tracking jika DB error
        global progress_tracker_fallback
        if 'progress_tracker_fallback' not in globals():
            progress_tracker_fallback = {}
        
        if user_id not in progress_tracker_fallback:
            progress_tracker_fallback[user_id] = {
                "current_step": 1,
                "steps": {
                    1: {"status": "pending", "message": "Mengunggah file modul ajar"},
                    2: {"status": "pending", "message": "Menganalisis struktur dokumen"},
                    3: {"status": "pending", "message": "Mengekstrak tujuan pembelajaran"},
                    4: {"status": "pending", "message": "Memproses dengan AI Gemini"},
                    5: {"status": "pending", "message": "Menyimpan soal ke database"}
                },
                "timestamp": datetime.datetime.now()
            }
        
        if step in progress_tracker_fallback[user_id]["steps"]:
            progress_tracker_fallback[user_id]["steps"][step]["status"] = status
            if message:
                progress_tracker_fallback[user_id]["steps"][step]["message"] = message
            progress_tracker_fallback[user_id]["current_step"] = step
            progress_tracker_fallback[user_id]["timestamp"] = datetime.datetime.now()

def get_progress(user_id):
    """Ambil progress untuk user tertentu (Database-based untuk production)"""
    try:
        progress = ProgressTracker.query.filter_by(user_id=str(user_id)).first()
        
        if progress:
            # Format data sesuai dengan format lama
            steps_data = {}
            for i in range(1, 6):
                steps_data[i] = {
                    "status": getattr(progress, f'step_{i}_status'),
                    "message": getattr(progress, f'step_{i}_message')
                }
            
            result = {
                "current_step": progress.current_step,
                "steps": steps_data,
                "timestamp": progress.updated_at
            }
            
            # Commit untuk menutup transaksi read-only
            db.session.commit()
            return result
        else:
            # Commit untuk menutup transaksi read-only bahkan jika tidak ada data
            db.session.commit()
            return None
            
    except Exception as e:
        print(f"Error getting progress for user {user_id}: {str(e)}")
        # Rollback pada error
        db.session.rollback()
        # Fallback ke memory-based tracking jika DB error
        global progress_tracker_fallback
        if 'progress_tracker_fallback' in globals():
            return progress_tracker_fallback.get(user_id, None)
        return None

def clear_progress(user_id):
    """Bersihkan progress setelah selesai (Database-based untuk production)"""
    try:
        progress = ProgressTracker.query.filter_by(user_id=str(user_id)).first()
        if progress:
            db.session.delete(progress)
            db.session.commit()
            print(f"Progress cleared from DB for user {user_id}")
    except Exception as e:
        db.session.rollback()
        print(f"Error clearing progress for user {user_id}: {str(e)}")
        # Fallback ke memory-based tracking jika DB error
        global progress_tracker_fallback
        if 'progress_tracker_fallback' in globals() and user_id in progress_tracker_fallback:
            del progress_tracker_fallback[user_id]

# =========================================
# OPTIMIZED MODULE EXTRACTION FUNCTIONS
# =========================================

def validate_kurikulum_merdeka_modul_ajar(content_text):
    """
    Validasi KHUSUS untuk Modul Ajar Kurikulum Merdeka - Hanya menerima yang memiliki komponen wajib
    """
    if not content_text or len(content_text.strip()) < 400:
        return False, "Dokumen terlalu pendek untuk modul ajar Kurikulum Merdeka (minimal 400 karakter)"
    
    content_lower = content_text.lower()
    missing_components = []
    validation_score = 0
    
    # 1. WAJIB: Header "MODUL AJAR" - Indikator utama Kurikulum Merdeka (lebih fleksibel)
    modul_ajar_patterns = [
        r'modul\s+ajar',
        r'modul\s+pembelajaran',
        r'teaching\s+module',
        r'learning\s+module',
        r'rpp\s+kurikulum\s+merdeka',
        r'rencana\s+pelaksanaan\s+pembelajaran',
        r'lesson\s+plan',
        r'rpp\s+merdeka',
        r'kurikulum\s+merdeka',
        r'merdeka\s+belajar',
        r'perangkat\s+ajar',
        r'bahan\s+ajar',
        r'materi\s+pembelajaran',
        r'desain\s+pembelajaran'
    ]
    
    has_modul_header = any(re.search(pattern, content_lower) for pattern in modul_ajar_patterns)
    if has_modul_header:
        validation_score += 25
        print("[OK] Header MODUL AJAR ditemukan")
    else:
        missing_components.append("Header 'MODUL AJAR'")
        print("[X] Header MODUL AJAR tidak ditemukan")
    
    # 2. WAJIB: Identitas Modul (Mata Pelajaran, Kelas, Fase) - lebih fleksibel
    identitas_patterns = [
        r'mata\s+pelajaran\s*[:]\s*',
        r'subject\s*[:]\s*',
        r'fase\s*/?\s*kelas\s*[:]\s*',
        r'phase\s*/?\s*class\s*[:]\s*',
        r'fase\s*[:]\s*[a-g]',
        r'kelas\s*[:]\s*[ivx\d]+',
        r'grade\s*[:]\s*\d+',
        r'mapel\s*[:]\s*',
        r'pelajaran\s*[:]\s*',
        r'kelas\s+\d+',
        r'semester\s*[:]\s*',
        r'tingkat\s*[:]\s*',
        r'jenjang\s*[:]\s*',
        r'sekolah\s+dasar',
        r'sekolah\s+menengah',
        r'smp|sma|smk|sd',
        r'elementary|junior|senior|high\s+school'
    ]
    
    has_identitas = any(re.search(pattern, content_lower) for pattern in identitas_patterns)
    if has_identitas:
        validation_score += 15
        print("[OK] Identitas Modul (Mata Pelajaran/Kelas/Fase) ditemukan")
    else:
        missing_components.append("Identitas Modul (Mata Pelajaran, Kelas, Fase)")
        print("[X] Identitas Modul tidak lengkap")
    
    # 3. WAJIB: Komponen Inti - Tujuan Pembelajaran (lebih fleksibel)
    tujuan_patterns = [
        r'tujuan\s+pembelajaran',
        r'learning\s+objectives?',
        r'capaian\s+pembelajaran',
        r'learning\s+outcomes?',
        r'i\.\s*tujuan\s+pembelajaran',
        r'1\.\s*tujuan\s+pembelajaran',
        r'objektif\s+pembelajaran',
        r'target\s+pembelajaran',
        r'sasaran\s+pembelajaran',
        r'kompetensi\s+dasar',
        r'indikator\s+pencapaian',
        r'setelah\s+mengikuti\s+pembelajaran',
        r'siswa\s+mampu',
        r'peserta\s+didik\s+dapat',
        r'akan\s+dapat\s+memahami',
        r'diharapkan\s+siswa'
    ]
    
    has_tujuan = any(re.search(pattern, content_lower) for pattern in tujuan_patterns)
    if has_tujuan:
        validation_score += 20
        print("[OK] Tujuan Pembelajaran ditemukan")
    else:
        missing_components.append("Tujuan Pembelajaran")
        print("[X] Tujuan Pembelajaran tidak ditemukan")
    
    # 4. WAJIB: Kompetensi Awal - Komponen khas Kurikulum Merdeka (lebih fleksibel)
    kompetensi_patterns = [
        r'kompetensi\s+awal',
        r'prerequisite\s+competenc',
        r'kemampuan\s+prasyarat',
        r'ii\.\s*kompetensi\s+awal',
        r'2\.\s*kompetensi\s+awal',
        r'kemampuan\s+dasar',
        r'pengetahuan\s+awal',
        r'prasyarat\s+pembelajaran',
        r'kemampuan\s+prerequisit',
        r'bekal\s+awal',
        r'pengetahuan\s+sebelumnya',
        r'sudah\s+dipelajari',
        r'telah\s+menguasai',
        r'kemampuan\s+yang\s+dimiliki'
    ]
    
    has_kompetensi_awal = any(re.search(pattern, content_lower) for pattern in kompetensi_patterns)
    if has_kompetensi_awal:
        validation_score += 20
        print("[OK] Kompetensi Awal ditemukan")
    else:
        missing_components.append("Kompetensi Awal")
        print("[X] Kompetensi Awal tidak ditemukan")
    
    # 5. WAJIB: Pemahaman Bermakna - Komponen khas Kurikulum Merdeka
    bermakna_patterns = [
        r'pemahaman\s+bermakna',
        r'meaningful\s+understanding',
        r'essential\s+understanding',
        r'iii\.\s*pemahaman\s+bermakna',
        r'3\.\s*pemahaman\s+bermakna'
    ]
    
    has_pemahaman_bermakna = any(re.search(pattern, content_lower) for pattern in bermakna_patterns)
    if has_pemahaman_bermakna:
        validation_score += 15
        print("[OK] Pemahaman Bermakna ditemukan")
    else:
        missing_components.append("Pemahaman Bermakna")
        print("[X] Pemahaman Bermakna tidak ditemukan")
    
    # 6. OPSIONAL BONUS: Profil Pelajar Pancasila (ciri khas Kurikulum Merdeka)
    profil_patterns = [
        r'profil\s+pelajar\s+pancasila',
        r'pancasila\s+student\s+profile',
        r'p5\s*[:]\s*',
        r'dimensi\s+profil\s+pelajar',
        r'karakter\s+pelajar\s+pancasila'
    ]
    
    has_profil_pancasila = any(re.search(pattern, content_lower) for pattern in profil_patterns)
    if has_profil_pancasila:
        validation_score += 10
        print("[OK] BONUS: Profil Pelajar Pancasila ditemukan")
    
    # 7. OPSIONAL BONUS: Pertanyaan Pemantik
    pemantik_patterns = [
        r'pertanyaan\s+pemantik',
        r'essential\s+questions?',
        r'driving\s+questions?',
        r'guiding\s+questions?',
        r'iv\.\s*pertanyaan\s+pemantik'
    ]
    
    has_pertanyaan_pemantik = any(re.search(pattern, content_lower) for pattern in pemantik_patterns)
    if has_pertanyaan_pemantik:
        validation_score += 5
        print("[OK] BONUS: Pertanyaan Pemantik ditemukan")
    
    # 8. VALIDASI KONTEN PENDIDIKAN
    edu_indicators = [
        'peserta didik', 'siswa', 'murid', 'pelajar',
        'guru', 'pengajar', 'fasilitator', 'educator',
        'pembelajaran', 'belajar', 'mengajar', 'learning', 'teaching'
    ]
    
    edu_count = sum(1 for indicator in edu_indicators if indicator in content_lower)
    if edu_count >= 5:
        validation_score += 5
        print(f"[OK] Konteks pendidikan memadai ({edu_count} indikator ditemukan)")
    else:
        missing_components.append("Konteks pendidikan yang memadai")
        print(f"[X] Konteks pendidikan kurang ({edu_count} indikator)")
    
    # VALIDASI AKHIR
    print(f"\n=== HASIL VALIDASI MODUL AJAR KURIKULUM MERDEKA ===")
    print(f"Skor Validasi: {validation_score}/100")
    
    # Threshold yang lebih fleksibel: minimal 55 dari 100 untuk komponen wajib (turunkan agar mengurangi false negative)
    if validation_score >= 55:
        print("[DITERIMA] File memenuhi standar Modul Ajar Kurikulum Merdeka")
        return True, f"Modul Ajar Kurikulum Merdeka valid (skor: {validation_score}/100)"
    else:
        print("[DITOLAK] File tidak memenuhi standar Modul Ajar Kurikulum Merdeka")
        print(f"Komponen yang hilang atau kurang: {', '.join(missing_components)}")
        return False, f"Bukan Modul Ajar Kurikulum Merdeka yang valid. Komponen hilang: {', '.join(missing_components[:3])}"

def validate_file_format_and_content(file_stream, file_ext, filename):
    """
    Validasi format file dan konten secara fleksibel
    """
    try:
        # 1. Validasi ukuran file (lebih toleran)
        # Guard: file_stream bisa None jika client mengirim tanpa payload atau terjadi error saat baca
        if file_stream is None:
            return False, "File kosong atau tidak terbaca dari permintaan (stream None). Pastikan Anda memilih file yang benar."

        file_size = len(file_stream)
        if file_size < 512:  # Dikurangi dari 1KB ke 512 bytes
            return False, "File terlalu kecil. File pembelajaran minimal berukuran 512 bytes"
        
        if file_size > 15 * 1024 * 1024:  # Dinaikkan dari 10MB ke 15MB
            return False, "File terlalu besar. Maksimal ukuran file adalah 15MB"
        
        # 2. Validasi nama file (lebih fleksibel) - OPSIONAL
        filename_keywords = ['modul', 'pembelajaran', 'ajar', 'lesson', 'rpp', 'materi', 'bahan', 'silabus']
        has_relevant_filename = any(keyword in filename.lower() for keyword in filename_keywords)
        # Tidak reject jika nama file tidak sesuai, hanya beri peringatan
        
        # 3. Ekstrak dan validasi konten
        content_text = ""
        if file_ext == 'pdf':
            content_text = extract_text_from_pdf_bytes(file_stream)
        elif file_ext in ['doc', 'docx']:
            content_text = extract_text_from_docx_bytes(file_stream, file_ext)
        
        if not content_text:
            return False, "Tidak dapat mengekstrak teks dari file. Pastikan file tidak rusak atau terproteksi"
        
        # 4. Validasi konten khusus Modul Ajar Kurikulum Merdeka
        is_valid, message = validate_kurikulum_merdeka_modul_ajar(content_text)
        if not is_valid:
            return False, message
        
        # Tambahkan warning jika nama file tidak sesuai tapi tetap lanjutkan
        if not has_relevant_filename:
            print(f"WARNING: Nama file '{filename}' tidak mengindikasikan materi pembelajaran, tapi konten valid")
            
        return True, content_text
        
    except Exception as e:
        # Tangani error generik dengan pesan yang lebih ramah
        return False, f"Error validasi file: {str(e)}"

def check_content_quality_score(module_components):
    """
    Periksa kualitas ekstraksi untuk memastikan layak diproses (versi lebih toleran)
    """
    quality_score = 0
    issues = []
    
    # 1. Mata Pelajaran (20 poin) - lebih toleran
    mata_pelajaran = module_components.get("mata_pelajaran", "")
    if mata_pelajaran and len(mata_pelajaran.strip()) > 2:  # Dikurangi dari 5 ke 2
        quality_score += 20
    elif mata_pelajaran:  # Ada tapi pendek, beri setengah poin
        quality_score += 10
        issues.append("Mata pelajaran terdeteksi tapi kurang jelas")
    else:
        issues.append("Mata pelajaran tidak teridentifikasi")
    
    # 2. Tujuan Pembelajaran (40 poin - paling penting tapi lebih fleksibel)
    tujuan_list = module_components.get("tujuan_pembelajaran", [])
    if tujuan_list and len(tujuan_list) >= 1:  # Dikurangi dari 2 ke 1
        # Periksa kualitas setiap tujuan
        valid_tujuan = [t for t in tujuan_list if len(t.strip()) > 10]  # Dikurangi dari 15 ke 10
        if len(valid_tujuan) >= 2:
            quality_score += 40
        elif len(valid_tujuan) >= 1:
            quality_score += 30  # Dinaikkan dari 25 ke 30
            issues.append("Tujuan pembelajaran ditemukan tapi bisa lebih detail")
        else:
            quality_score += 15
            issues.append("Tujuan pembelajaran ditemukan tapi sangat singkat")
    else:
        issues.append("Tujuan pembelajaran tidak ditemukan")
    
    # 3. Kompetensi Awal (15 poin) - lebih toleran
    kompetensi_awal = module_components.get("kompetensi_awal", "")
    if kompetensi_awal and len(kompetensi_awal.strip()) > 20:  # Dikurangi dari 30 ke 20
        quality_score += 15
    elif kompetensi_awal and len(kompetensi_awal.strip()) > 5:
        quality_score += 8  # Beri poin parsial
        issues.append("Kompetensi awal terdeteksi tapi kurang detail")
    else:
        quality_score += 5  # Beri poin minimal bahkan jika tidak ada
        issues.append("Kompetensi awal tidak teridentifikasi")
    
    # 4. Pemahaman Bermakna (15 poin) - lebih toleran
    pemahaman_list = module_components.get("pemahaman_bermakna", [])
    if pemahaman_list and len(pemahaman_list) >= 1:
        quality_score += 15
    else:
        quality_score += 8  # Beri poin parsial meski tidak ada
        issues.append("Pemahaman bermakna tidak ditemukan")
    
    # 5. Target Peserta Didik (10 poin) - lebih toleran
    target_peserta = module_components.get("target_peserta_didik", "")
    if target_peserta and len(target_peserta.strip()) > 10:  # Dikurangi dari 20 ke 10
        quality_score += 10
    else:
        quality_score += 5  # Beri poin parsial
        issues.append("Target peserta didik tidak teridentifikasi")
    
    return quality_score, issues

def extract_kurikulum_merdeka_components(content_text):
    """
    Ekstrak komponen khusus Modul Ajar Kurikulum Merdeka dengan pattern yang tepat
    """
    results = {
        "mata_pelajaran": "",
        "topik_utama": "",
        "kelas": "",
        "fase": "",
        "kompetensi_awal": "",
        "tujuan_pembelajaran": [],
        "capaian_pembelajaran": [],
        "pemahaman_bermakna": [],
        "target_peserta_didik": "",
        "profil_pelajar_pancasila": [],
        "pertanyaan_pemantik": []
    }
    
    content = content_text
    content_lower = content_text.lower()
    
    print("\n=== EKSTRAKSI KOMPONEN KURIKULUM MERDEKA ===")
    
    # 1. EKSTRAK IDENTITAS MODUL - Pattern khusus Kurikulum Merdeka
    # Mata Pelajaran
    mata_pelajaran_patterns = [
        r'mata\s+pelajaran\s*[:]\s*([^\n\r]+)',
        r'subject\s*[:]\s*([^\n\r]+)',
        r'modul\s+ajar\s*\n\s*([^\n\r]+)',
        r'nama\s+modul\s*[:]\s*([^\n\r]+)'
    ]
    
    for pattern in mata_pelajaran_patterns:
        match = re.search(pattern, content, re.IGNORECASE)
        if match:
            results["mata_pelajaran"] = match.group(1).strip()
            results["topik_utama"] = match.group(1).strip()
            print(f"[OK] Mata Pelajaran: {results['mata_pelajaran']}")
            break
    
    # Fase dan Kelas
    fase_kelas_patterns = [
        r'fase\s*/?\s*kelas\s*[:]\s*([^\n\r]+)',
        r'fase\s*[:]\s*([a-g])\s*[\(/]*\s*kelas\s*([ivx\d\-]+)',
        r'kelas\s*[:]\s*([ivx\d\-\s]+)',
        r'grade\s*[:]\s*(\d+)',
        r'phase\s*[:]\s*([a-g])'
    ]
    
    for pattern in fase_kelas_patterns:
        match = re.search(pattern, content, re.IGNORECASE)
        if match:
            if len(match.groups()) > 1:
                results["fase"] = match.group(1).strip()
                results["kelas"] = match.group(2).strip()
            else:
                results["kelas"] = match.group(1).strip()
            print(f"[OK] Kelas/Fase: {results['kelas']} (Fase: {results.get('fase', 'N/A')})")
            break
    
    # 2. EKSTRAK KOMPETENSI AWAL - Komponen khas Kurikulum Merdeka
    kompetensi_patterns = [
        r'ii\.\s*kompetensi\s+awal\s*[:\n](.*?)(?=\n\s*(?:iii\.|tujuan\s+pembelajaran|pemahaman\s+bermakna|$))',
        r'kompetensi\s+awal\s*[:\n](.*?)(?=\n\s*(?:[ivx]+\.|tujuan\s+pembelajaran|pemahaman\s+bermakna|profil\s+pelajar|$))',
        r'prerequisite\s+competenc[^\n]*[:\n](.*?)(?=\n\s*(?:learning\s+objective|meaningful\s+understanding|$))',
        r'kemampuan\s+prasyarat\s*[:\n](.*?)(?=\n\s*(?:tujuan|pembelajaran|$))'
    ]
    
    for i, pattern in enumerate(kompetensi_patterns, 1):
        kompetensi_match = re.search(pattern, content, re.IGNORECASE | re.DOTALL)
        if kompetensi_match:
            kompetensi_text = kompetensi_match.group(1).strip()
            
            # Clean dan format text
            lines = [line.strip() for line in kompetensi_text.split('\n') if line.strip()]
            clean_text = ' '.join(lines)
            
            # Batasi panjang
            if len(clean_text) > 500:
                sentences = re.split(r'[.!?]+', clean_text)
                key_sentences = [s.strip() for s in sentences[:3] if len(s.strip()) > 10]
                results["kompetensi_awal"] = '. '.join(key_sentences) + '.'
            else:
                results["kompetensi_awal"] = clean_text
                
            print(f"[OK] Kompetensi Awal ditemukan (pattern {i}): {len(results['kompetensi_awal'])} karakter")
            break
    
    # 3. EKSTRAK TUJUAN PEMBELAJARAN - Komponen utama Kurikulum Merdeka  
    tujuan_patterns = [
        # Pattern 1: Dalam struktur Komponen Inti
        r'komponen\s+inti.*?i\.\s*tujuan\s+pembelajaran\s*[:\n](.*?)(?=\n\s*(?:ii\.|kompetensi\s+awal|pemahaman\s+bermakna|$))',
        # Pattern 2: Roman numeral langsung
        r'i\.\s*tujuan\s+pembelajaran\s*[:\n](.*?)(?=\n\s*(?:ii\.|kompetensi\s+awal|pemahaman\s+bermakna|$))',
        # Pattern 3: Header tujuan pembelajaran langsung
        r'tujuan\s+pembelajaran\s*[:\n](.*?)(?=\n\s*(?:kompetensi\s+awal|pemahaman\s+bermakna|pertanyaan\s+pemantik|profil\s+pelajar|ii\.|2\.|$))',
        # Pattern 4: Dengan bullets/numbering
        r'tujuan\s+pembelajaran[^\n]*\n((?:\s*[•\-\*\d\.]\s*[^\n]+\n?)+)',
        # Pattern 5: Setelah pembelajaran siswa dapat
        r'(?:setelah\s+pembelajaran.*?|tujuan\s+pembelajaran.*?)(?:siswa\s+(?:dapat|mampu)|peserta\s+didik\s+(?:dapat|mampu))(.*?)(?=\n\s*(?:kompetensi\s+awal|pemahaman\s+bermakna|$))'
    ]
    
    tujuan_found = False
    for i, pattern in enumerate(tujuan_patterns, 1):
        if tujuan_found:
            break
            
        tujuan_match = re.search(pattern, content, re.IGNORECASE | re.DOTALL)
        if tujuan_match:
            tujuan_text = tujuan_match.group(1).strip()
            print(f"[OK] Tujuan pembelajaran ditemukan (pattern {i})")
            
            # Strategy 1: Ekstrak bullet points atau numbering
            bullets = re.findall(r'(?:^|\n)\s*[•\-\*\d\.]\s*([^\n•\-\*]+)', tujuan_text, re.MULTILINE)
            bullets = [b.strip() for b in bullets if len(b.strip()) > 20]
            
            if len(bullets) >= 2:
                results["tujuan_pembelajaran"] = bullets[:6]
                print(f"   ->  {len(bullets)} tujuan dalam format bullets")
                tujuan_found = True
                break
            
            # Strategy 2: Cari kalimat dengan action verbs pembelajaran
            learning_verbs = ['dapat', 'mampu', 'menjelaskan', 'menganalisis', 'menerapkan', 
                            'memahami', 'mengidentifikasi', 'mendemonstrasikan', 'membuat', 
                            'merancang', 'mengevaluasi', 'menciptakan']
            
            learning_sentences = []
            for verb in learning_verbs:
                sentences = re.findall(rf'[^\n.!?]*(?:siswa\s+{verb}|peserta\s+didik\s+{verb})[^\n.!?]*[.!?\n]?', 
                                     tujuan_text, re.IGNORECASE)
                learning_sentences.extend([s.strip().rstrip('.!?\n') for s in sentences if len(s.strip()) > 25])
            
            # Remove duplicates and limit
            unique_sentences = list(dict.fromkeys(learning_sentences))[:6]
            
            if len(unique_sentences) >= 2:
                results["tujuan_pembelajaran"] = unique_sentences
                print(f"   ->  {len(unique_sentences)} tujuan dengan action verbs")
                tujuan_found = True
                break
            
            # Strategy 3: Split by lines and find meaningful objectives
            lines = [line.strip() for line in tujuan_text.split('\n') if line.strip()]
            meaningful_lines = []
            
            for line in lines:
                if (len(line) > 25 and len(line) < 300 and
                    any(verb in line.lower() for verb in learning_verbs[:8])):
                    meaningful_lines.append(line)
            
            if len(meaningful_lines) >= 2:
                results["tujuan_pembelajaran"] = meaningful_lines[:5]
                print(f"   ->  {len(meaningful_lines)} tujuan dari lines")
                tujuan_found = True
                break
    
    # Fallback untuk tujuan pembelajaran jika tidak ditemukan
    if not tujuan_found:
        print("⚠ Mencari tujuan pembelajaran dengan fallback...")
        all_learning_sentences = re.findall(
            r'[^\n.!?]*(?:siswa\s+(?:dapat|mampu)|peserta\s+didik\s+(?:dapat|mampu))[^\n.!?]*[.!?\n]?', 
            content, re.IGNORECASE)
        
        filtered_sentences = []
        for sentence in all_learning_sentences:
            clean_sentence = sentence.strip().rstrip('.!?\n')
            if 30 < len(clean_sentence) < 200:
                filtered_sentences.append(clean_sentence)
        
        if len(filtered_sentences) >= 2:
            results["tujuan_pembelajaran"] = filtered_sentences[:4]
            print(f"   ->  FALLBACK: {len(filtered_sentences)} tujuan ditemukan")

    # 3b. EKSTRAK CAPAIAN PEMBELAJARAN (jika tersedia terpisah dari tujuan)
    capaian_patterns = [
        r'(?:capaian\s+pembelajaran|cp)(?:\s*[:\n]|\s+-\s*)(.*?)(?=\n\s*(?:pemahaman\s+bermakna|kompetensi\s+awal|indikator|tujuan\s+pembelajaran|profil\s+pelajar|pertanyaan\s+pemantik|[ivx]+\.|\d+\.|$))',
        r'capaian\s+pembelajaran[^\n]*\n((?:\s*[•\-\*\d\.]+\s*[^\n]+\n?)+)'
    ]
    capaian_found = False
    for i, pattern in enumerate(capaian_patterns, 1):
        if capaian_found:
            break
        cp_match = re.search(pattern, content, re.IGNORECASE | re.DOTALL)
        if cp_match:
            cp_text = cp_match.group(1).strip()
            # Extract bullets or meaningful lines
            bullets = re.findall(r'(?:^|\n)\s*[•\-\*\d\.]\s*([^\n•\-\*]+)', cp_text, re.MULTILINE)
            bullets = [b.strip() for b in bullets if len(b.strip()) > 15]
            if bullets:
                results["capaian_pembelajaran"] = bullets[:6]
                print(f"[OK] Capaian Pembelajaran: {len(bullets)} item (pattern {i})")
                capaian_found = True
                break
            # Otherwise split lines
            lines = [line.strip() for line in cp_text.split('\n') if len(line.strip()) > 20]
            if lines:
                results["capaian_pembelajaran"] = lines[:6]
                print(f"[OK] Capaian Pembelajaran (lines): {len(lines)} item (pattern {i})")
                capaian_found = True
                break
    
    # 4. EKSTRAK PEMAHAMAN BERMAKNA - Komponen khas Kurikulum Merdeka
    bermakna_patterns = [
        r'(?:komponen\s+inti.*?)?(?:ii\.|2\.)\s*pemahaman\s+bermakna\s*[:\n](.*?)(?=\n\s*(?:iii\.|3\.|pertanyaan\s+pemantik|profil\s+pelajar|target\s+peserta|$))',
        r'pemahaman\s+bermakna\s*[:\n](.*?)(?=\n\s*(?:pertanyaan\s+pemantik|profil\s+pelajar|target\s+peserta|iii\.|3\.|kegiatan\s+pembelajaran|$))',
        r'meaningful\s+understanding\s*[:\n](.*?)(?=\n\s*(?:essential\s+question|student\s+profile|$))',
        r'essential\s+understanding\s*[:\n](.*?)(?=\n\s*(?:question|profile|$))'
    ]
    
    for i, pattern in enumerate(bermakna_patterns, 1):
        bermakna_match = re.search(pattern, content, re.IGNORECASE | re.DOTALL)
        if bermakna_match:
            bermakna_text = bermakna_match.group(1).strip()
            print(f"[OK] Pemahaman bermakna ditemukan (pattern {i})")
            
            # Strategy 1: Extract bullets/numbering
            bullets = re.findall(r'(?:^|\n)\s*[•\-\*\d\.]\s*([^\n•\-\*]+)', bermakna_text, re.MULTILINE)
            bullets = [b.strip() for b in bullets if len(b.strip()) > 15]
            
            if len(bullets) >= 2:
                results["pemahaman_bermakna"] = bullets[:4]
                print(f"   ->  {len(bullets)} pemahaman bermakna dalam format bullets")
                break
            
            # Strategy 2: Split by lines
            lines = [line.strip() for line in bermakna_text.split('\n') 
                    if line.strip() and len(line.strip()) > 20]
            
            if len(lines) >= 2:
                results["pemahaman_bermakna"] = lines[:3]
                print(f"   ->  {len(lines)} pemahaman bermakna dari lines")
                break
            
            # Strategy 3: Use entire text if meaningful
            if 50 < len(bermakna_text) < 400:
                results["pemahaman_bermakna"] = [bermakna_text]
                print("   ->  1 pemahaman bermakna (text utuh)")
                break
    
    # 5. EKSTRAK PROFIL PELAJAR PANCASILA - Ciri khas Kurikulum Merdeka
    profil_patterns = [
        r'profil\s+pelajar\s+pancasila\s*[:\n](.*?)(?=\n\s*(?:pertanyaan\s+pemantik|kegiatan\s+pembelajaran|$))',
        r'pancasila\s+student\s+profile\s*[:\n](.*?)(?=\n\s*(?:essential\s+question|learning\s+activities|$))',
        r'dimensi\s+profil\s+pelajar\s*[:\n](.*?)(?=\n\s*(?:pemantik|kegiatan|$))',
        r'p5\s*[:]\s*(.*?)(?=\n\s*(?:pemantik|kegiatan|$))'
    ]
    
    for i, pattern in enumerate(profil_patterns, 1):
        profil_match = re.search(pattern, content, re.IGNORECASE | re.DOTALL)
        if profil_match:
            profil_text = profil_match.group(1).strip()
            
            # Extract dimensi atau karakteristik
            dimensi_list = re.findall(r'(?:[•\-\*\d\.]\s*)?([^•\-\*\n\d\.][^:\n]{15,})', profil_text)
            if dimensi_list:
                results["profil_pelajar_pancasila"] = [d.strip() for d in dimensi_list[:6]]
                print(f"[OK] Profil Pelajar Pancasila: {len(dimensi_list)} dimensi")
            break
    
    # 6. EKSTRAK PERTANYAAN PEMANTIK
    pemantik_patterns = [
        r'pertanyaan\s+pemantik\s*[:\n](.*?)(?=\n\s*(?:kegiatan\s+pembelajaran|model\s+pembelajaran|$))',
        r'essential\s+questions?\s*[:\n](.*?)(?=\n\s*(?:learning\s+activities|teaching\s+model|$))',
        r'driving\s+questions?\s*[:\n](.*?)(?=\n\s*(?:activities|model|$))'
    ]
    
    for pattern in pemantik_patterns:
        pemantik_match = re.search(pattern, content, re.IGNORECASE | re.DOTALL)
        if pemantik_match:
            pemantik_text = pemantik_match.group(1).strip()
            
            # Extract questions
            questions = re.findall(r'[•\-\*\d\.]*\s*([^•\-\*\n\d\.][^?\n]*\?)', pemantik_text)
            if questions:
                results["pertanyaan_pemantik"] = [q.strip() for q in questions[:5]]
                print(f"[OK] Pertanyaan Pemantik: {len(questions)} pertanyaan")
            break
    
    # 7. TARGET PESERTA DIDIK
    target_patterns = [
        r'target\s+peserta\s+didik\s*[:\n](.*?)(?=\n\s*(?:model\s+pembelajaran|kegiatan\s+pembelajaran|$))',
        r'sasaran\s+peserta\s+didik\s*[:\n](.*?)(?=\n\s*(?:model|kegiatan|$))',
        r'v\.\s*target\s+peserta\s+didik\s*[:\n](.*?)(?=\n\s*(?:vi\.|6\.|$))'
    ]
    
    for pattern in target_patterns:
        target_match = re.search(pattern, content, re.IGNORECASE | re.DOTALL)
        if target_match:
            target_text = target_match.group(1).strip()
            
            # Clean dan format text
            lines = [line.strip() for line in target_text.split('\n') if line.strip()]
            clean_text = ' '.join(lines)[:400]  # Batasi 400 karakter
            
            if len(clean_text) > 20:
                results["target_peserta_didik"] = clean_text
                print(f"[OK] Target Peserta Didik: {len(clean_text)} karakter")
            break
    
    # SUMMARY HASIL EKSTRAKSI
    print(f"\n=== RINGKASAN EKSTRAKSI KURIKULUM MERDEKA ===")
    print(f"Mata Pelajaran: {results['mata_pelajaran'] or 'Tidak terdeteksi'}")
    print(f"Kelas/Fase: {results['kelas'] or 'Tidak terdeteksi'}")
    print(f"Tujuan Pembelajaran: {len(results['tujuan_pembelajaran'])} item")
    print(f"Kompetensi Awal: {'OK' if results['kompetensi_awal'] else 'X'}")
    print(f"Pemahaman Bermakna: {len(results['pemahaman_bermakna'])} item")
    print(f"Profil Pelajar Pancasila: {len(results['profil_pelajar_pancasila'])} dimensi")
    print(f"Pertanyaan Pemantik: {len(results['pertanyaan_pemantik'])} pertanyaan")
    print(f"Target Peserta Didik: {'OK' if results['target_peserta_didik'] else 'X'}")
    print("=" * 50)
    
    return results

# Fungsi backward compatibility  
def extract_specific_module_components(content_text):
    """
    DEPRECATED: Gunakan extract_kurikulum_merdeka_components() untuk ekstraksi yang lebih akurat
    """
    return extract_kurikulum_merdeka_components(content_text)

def validate_kurikulum_merdeka_components(module_components):
    """
    Validasi khusus komponen Modul Ajar Kurikulum Merdeka yang telah diekstraksi
    """
    validation_report = {
        "detected": [],
        "missing": [],
        "quality_score": 0
    }
    
    print("\n=== VALIDASI KOMPONEN KURIKULUM MERDEKA ===")
    
    # 1. VALIDASI MATA PELAJARAN & IDENTITAS MODUL (20 poin)
    mata_pelajaran = module_components.get("mata_pelajaran", "")
    kelas = module_components.get("kelas", "")
    
    if mata_pelajaran and len(mata_pelajaran.strip()) > 3:
        validation_report["detected"].append("Mata Pelajaran")
        validation_report["quality_score"] += 15
        print(f"[OK] Mata Pelajaran: {mata_pelajaran}")
        
        if kelas and len(kelas.strip()) > 1:
            validation_report["quality_score"] += 5
            print(f"[OK] Kelas/Fase: {kelas}")
    else:
        validation_report["missing"].append("Mata Pelajaran")
        print("[X] Mata Pelajaran tidak teridentifikasi")
    
    # 2. VALIDASI TUJUAN PEMBELAJARAN (30 poin - paling penting)
    tujuan_list = module_components.get("tujuan_pembelajaran", [])
    if tujuan_list and len(tujuan_list) >= 1:
        meaningful_tujuan = [t for t in tujuan_list if len(t.strip()) > 20]
        
        if len(meaningful_tujuan) >= 3:
            validation_report["detected"].append(f"Tujuan Pembelajaran ({len(tujuan_list)} item)")
            validation_report["quality_score"] += 30
            print(f"[OK] Tujuan Pembelajaran: {len(tujuan_list)} item berkualitas")
        elif len(meaningful_tujuan) >= 2:
            validation_report["detected"].append(f"Tujuan Pembelajaran ({len(tujuan_list)} item)")
            validation_report["quality_score"] += 25
            print(f"[OK] Tujuan Pembelajaran: {len(tujuan_list)} item (cukup)")
        elif len(meaningful_tujuan) >= 1:
            validation_report["detected"].append("Tujuan Pembelajaran (minimal)")
            validation_report["quality_score"] += 15
            print(f"⚠ Tujuan Pembelajaran: hanya {len(meaningful_tujuan)} item berkualitas")
        else:
            validation_report["missing"].append("Tujuan Pembelajaran")
            print("[X] Tujuan Pembelajaran tidak memadai")
    else:
        validation_report["missing"].append("Tujuan Pembelajaran")
        print("[X] Tujuan Pembelajaran tidak ditemukan")
    
    # 3. VALIDASI KOMPETENSI AWAL (25 poin - khas Kurikulum Merdeka)
    kompetensi_awal = module_components.get("kompetensi_awal", "")
    if kompetensi_awal and len(kompetensi_awal.strip()) > 50:
        validation_report["detected"].append("Kompetensi Awal")
        validation_report["quality_score"] += 25
        print(f"[OK] Kompetensi Awal: {len(kompetensi_awal)} karakter")
    elif kompetensi_awal and len(kompetensi_awal.strip()) > 20:
        validation_report["detected"].append("Kompetensi Awal (singkat)")
        validation_report["quality_score"] += 15
        print(f"⚠ Kompetensi Awal: terlalu singkat ({len(kompetensi_awal)} karakter)")
    else:
        validation_report["missing"].append("Kompetensi Awal")
        print("[X] Kompetensi Awal tidak memadai")
    
    # 4. VALIDASI PEMAHAMAN BERMAKNA (20 poin - khas Kurikulum Merdeka)
    bermakna_list = module_components.get("pemahaman_bermakna", [])
    if bermakna_list and len(bermakna_list) > 0:
        meaningful_bermakna = [item for item in bermakna_list if len(item.strip()) > 15]
        
        if len(meaningful_bermakna) >= 2:
            validation_report["detected"].append(f"Pemahaman Bermakna ({len(bermakna_list)} item)")
            validation_report["quality_score"] += 20
            print(f"[OK] Pemahaman Bermakna: {len(bermakna_list)} item")
        elif len(meaningful_bermakna) >= 1:
            validation_report["detected"].append("Pemahaman Bermakna (minimal)")
            validation_report["quality_score"] += 10
            print(f"⚠ Pemahaman Bermakna: hanya {len(meaningful_bermakna)} item berkualitas")
        else:
            validation_report["missing"].append("Pemahaman Bermakna")
            print("[X] Pemahaman Bermakna tidak berkualitas")
    else:
        validation_report["missing"].append("Pemahaman Bermakna")
        print("[X] Pemahaman Bermakna tidak ditemukan")
    
    # 5. BONUS: Target Peserta Didik (5 poin)
    target_peserta = module_components.get("target_peserta_didik", "")
    if target_peserta and len(target_peserta.strip()) > 30:
        validation_report["quality_score"] += 5
        print(f"[OK] BONUS: Target Peserta Didik ditemukan")
    
    # Laporan akhir
    print(f"\n--- HASIL VALIDASI KOMPONEN ---")
    print(f"Terdeteksi: {', '.join(validation_report['detected'])}")
    if validation_report["missing"]:
        print(f"Hilang: {', '.join(validation_report['missing'])}")
    print(f"Skor Kualitas: {validation_report['quality_score']}/100")
    
    return validation_report

# Fungsi backward compatibility
def validate_detected_keywords(module_components):
    """
    DEPRECATED: Gunakan validate_kurikulum_merdeka_components() untuk validasi yang lebih ketat
    """
    return validate_kurikulum_merdeka_components(module_components)
    
    # Bonus points for quality content
    bonus = 0
    if len(tujuan_list) >= 3:  # Bonus jika tujuan pembelajaran banyak
        bonus += 5
    if len(bermakna_list) >= 2:  # Bonus jika pemahaman bermakna banyak
        bonus += 5
    if module_components.get("target_peserta_didik"):  # Bonus jika ada target
        bonus += 5
        
    validation_report["quality_score"] = min(100, base_score + bonus)
    
    return validation_report

def log_extraction_results(module_components):
    """
    Log hasil ekstraksi untuk debugging dan monitoring
    """
    print("=== HASIL EKSTRAKSI MODUL AJAR ===")
    print(f"Mata Pelajaran: {module_components.get('mata_pelajaran', 'TIDAK TERDETEKSI')}")
    print(f"Topik Utama: {module_components.get('topik_utama', 'TIDAK TERDETEKSI')}")
    print(f"Kelas: {module_components.get('kelas', 'TIDAK TERDETEKSI')}")
    
    # Detail Tujuan Pembelajaran
    tujuan_items = module_components.get('tujuan_pembelajaran', [])
    print(f"Tujuan Pembelajaran: {len(tujuan_items)} item")
    if tujuan_items:
        for i, tujuan in enumerate(tujuan_items[:3], 1):
            print(f"  {i}. {tujuan[:80]}{'...' if len(tujuan) > 80 else ''}")
        if len(tujuan_items) > 3:
            print(f"  ... dan {len(tujuan_items) - 3} tujuan lainnya")
    
    # Detail Pemahaman Bermakna
    bermakna_items = module_components.get('pemahaman_bermakna', [])
    print(f"Pemahaman Bermakna: {len(bermakna_items)} item")
    if bermakna_items:
        for i, bermakna in enumerate(bermakna_items[:3], 1):
            print(f"  {i}. {bermakna[:80]}{'...' if len(bermakna) > 80 else ''}")
    
    print(f"Kompetensi Awal: {'OK' if module_components.get('kompetensi_awal') else 'X'}")
    print(f"Target Peserta Didik: {'OK' if module_components.get('target_peserta_didik') else 'X'}")
    
    # Validasi
    validation = validate_detected_keywords(module_components)
    print(f"Quality Score: {validation['quality_score']:.1f}%")
    if validation["missing"]:
        print(f"Komponen yang tidak terdeteksi: {', '.join(validation['missing'])}")
    print("=====================================")
    
    return validation

def create_optimized_prompt_with_good_structure(module_components):
    """
    Buat prompt singkat dan fokus berdasarkan komponen modul ajar untuk MST 5 Stage
    """
    
    # Format tujuan pembelajaran
    tujuan_list = []
    for tujuan in module_components.get("tujuan_pembelajaran", [])[:4]:
        tujuan_list.append(f"• {tujuan.strip()}")
    
    # Format pemahaman bermakna
    pemahaman_list = []
    for pemahaman in module_components.get("pemahaman_bermakna", [])[:3]:
        pemahaman_list.append(f"• {pemahaman.strip()}")

    # Prompt baru (disediakan oleh pengguna) — fokus pada kejelasan, anti-ambigu, dan format JSON valid
    prompt = f"""
Anda adalah seorang ahli pembuat soal asesmen diagnostik yang bertugas membantu guru untuk membuat soal yang akan diberikan pada siswa SMA. Soal dibuat berdasarkan Tujuan Pembelajaran dan komponen modul ajar yang telah diekstrak.

ANALISIS MODUL AJAR:
========================================
MODUL/ELEMEN AJAR:
{module_components.get("mata_pelajaran", "Informatika")} - {module_components.get("topik_utama", "Pembelajaran Teknologi")}
Kelas: {module_components.get("kelas", "X (Sepuluh)")}

KOMPETENSI AWAL:
{module_components.get("kompetensi_awal", "Siswa memiliki kemampuan dasar dalam berpikir logis dan pemecahan masalah")}

TUJUAN PEMBELAJARAN:
{chr(10).join(tujuan_list) if tujuan_list else "• Memahami konsep dasar materi pembelajaran"}

PEMAHAMAN BERMAKNA:
{chr(10).join(pemahaman_list) if pemahaman_list else "• Mengembangkan pemahaman konseptual yang mendalam"}

CAPAIAN PEMBELAJARAN (bila tersedia):
{chr(10).join([f"• {cp.strip()}" for cp in module_components.get("capaian_pembelajaran", [])[:4]]) if module_components.get("capaian_pembelajaran") else "• (Tidak tersedia di dokumen; gunakan Tujuan Pembelajaran sebagai acuan utama)"}

TARGET PESERTA DIDIK:
{module_components.get("target_peserta_didik", "Peserta didik reguler/tipikal: umum, tidak ada kesulitan dalam mencerna dan memahami materi ajar")}
========================================

Dalam membuat soal untuk sesi multi-tahap, gunakan taksonomi kompetensi teknologi berikut sebagai panduan untuk menentukan tipe dan tingkat kesulitan soal pada setiap level:

SPESIFIKASI SOAL PER LEVEL:
========================================
Level 1 → Kesadaran Teknologi
Tingkat kesulitan: Medium
Fokus soal: definisi, istilah, identifikasi teknologi dasar
Jenis pengetahuan: knowledge that
Contoh soal: "Alat yang digunakan untuk menyimpan data berbasis internet adalah..."

Level 2 → Literasi Teknologi
Tingkat kesulitan: Easy atau Hard
Fokus soal: klasifikasi, hubungan antar teknologi, penjelasan fungsi
Jenis pengetahuan: knowledge that
Contoh soal: "Manakah teknologi yang termasuk komunikasi sinkron?"

Level 3 → Kemampuan Teknologi
Tingkat kesulitan: Medium atau Hard
Fokus soal: aplikasi praktis sederhana, langkah-langkah dasar teknologi
Jenis pengetahuan: knowledge that + how (level SMA)
Contoh soal: "Urutkan langkah membuat email: 1. login, 2. klik compose, 3. isi pesan, 4. kirim"

Level 4 → Kreativitas Teknologi (Dasar)
Tingkat kesulitan: Easy, Medium, atau Hard
Fokus soal: identifikasi masalah sederhana, pemecahan masalah dasar
Jenis pengetahuan: knowledge that + how (level SMA)
Contoh soal: "Jika komputer tidak bisa terhubung internet, langkah pertama yang dilakukan adalah..."

Level 5 → Kritik Teknologi
Tingkat kesulitan: Easy, Medium, atau Hard
Fokus soal: membandingkan teknologi sederhana, memilih solusi yang tepat
Jenis pengetahuan: knowledge that + how + why (level SMA)
Contoh soal: "Untuk menyimpan foto keluarga, mana yang lebih aman: flash disk atau cloud storage?"

ATURAN PENTING - KEJELASAN DAN KETEPATAN:
========================================
✓ Setiap soal merujuk langsung pada konten modul yang dianalisis
✓ Gunakan terminologi spesifik dari modul
✓ 4 pilihan (A, B, C, D), maksimal 15 kata per opsi
✓ Posisi jawaban benar ACAK dan bervariasi
✓ Pengecoh masuk akal, cerminkan miskonsepsi umum
✓ Bahasa lugas, sesuai jenjang siswa SMA kelas X
✓ Kalimat soal maksimal 3 baris
✓ Fokus pada konsep dan aplikasi, bukan coding kompleks
✓ Gunakan contoh teknologi yang familiar untuk siswa SMA

ATURAN KHUSUS ANTI-AMBIGUITAS:
========================================
✓ WAJIB: Gunakan kata-kata PASTI dan DEFINITIF dalam soal
✓ WAJIB: Jawaban benar harus 100% AKURAT tanpa keraguan
✓ WAJIB: Setiap opsi harus JELAS BENAR atau JELAS SALAH
✓ WAJIB: Hindari interpretasi ganda dalam pertanyaan
✓ WAJIB: Gunakan istilah teknis yang PRESISI dan STANDAR
✓ WAJIB: Jawaban harus dapat DIVERIFIKASI dari sumber terpercaya

✗ LARANGAN KERAS - KATA & FRASA YANG DILARANG:
========================================
✗ JANGAN PERNAH gunakan kata: "mungkin", "kemungkinan", "biasanya", "umumnya", "sebaiknya"
✗ JANGAN gunakan frasa: "dapat berupa", "salah satunya adalah", "antara lain", "misalnya"
✗ JANGAN gunakan: "seharusnya", "lebih baik jika", "direkomendasikan", "disarankan"
✗ JANGAN gunakan kata tanya ambigu: "manakah yang lebih baik", "yang paling cocok"
✗ JANGAN gunakan perbandingan subjektif: "lebih mudah", "paling efektif", "terbaik"
✗ JANGAN gunakan kuantifikasi tidak pasti: "beberapa", "banyak", "sedikit"

✗ LARANGAN UMUM:
========================================
✗ JANGAN gunakan: "Modul Ajar", "Kompetensi Awal", "Tujuan Pembelajaran", "Siswa", "Peserta didik" dalam teks soal
✗ JANGAN buat soal generik yang lepas dari modul
✗ JANGAN buat jawaban benar selalu yang terpanjang
✗ JANGAN buat soal coding/programming yang terlalu teknis
✗ JANGAN gunakan istilah teknis tingkat universitas
✗ JANGAN buat soal yang memerlukan pengetahuan di luar kurikulum SMA
✗ JANGAN buat opsi jawaban yang memerlukan asumsi tambahan
✗ JANGAN buat soal dengan jawaban benar lebih dari satu

KAIDAH BAHASA INDONESIA (KBBI/EYD):
========================================
- Ikuti kaidah KBBI/EYD untuk tanda baca (koma, titik, titik dua, titik koma, tanda kurung, tanda kutip, tanda tanya, tanda seru).
- Gunakan huruf kapital secara tepat (awalan kalimat, nama diri, akronim baku) dan hindari kapitalisasi berlebihan.
- Penempatan spasi benar: spasi setelah koma/titik/titik dua; tidak ada spasi sebelum tanda baca penutup.
- Teks konsisten: satu konsep per soal, kalimat ringkas (≤2 klausa), suara aktif, hindari negasi ganda.
- Hindari campuran bahasa yang tidak perlu; gunakan padanan bahasa Indonesia baku jika tersedia.

ATURAN ANGKA & MATEMATIKA:
========================================
- Pemisah ribuan menggunakan titik: 1.000; 25.600; 1.250.000.
- Pecahan desimal menggunakan koma: 3,5; 12,75.
- Simbol satuan mengikuti standar (MB, GB, Hz) tanpa titik akhir dan tanpa spasi sebelum satuan (contoh: 5 GB, 2,4 GHz).
- Soal matematika tidak terlalu rumit: gunakan angka kecil/sederhana, operasi dasar, dan konteks realistis.
- Jika menyebut angka ribuan di soal/opsi/penjelasan, wajib gunakan titik sebagai pemisah ribuan.

PEDOMAN PENULISAN SOAL DEFINITIF:
========================================
✓ Gunakan kata kerja pasti: "adalah", "berfungsi", "digunakan untuk"
✓ Gunakan fakta teknis objektif yang dapat diverifikasi
✓ Sebutkan spesifikasi atau karakteristik yang konkret  
✓ Gunakan istilah standar industri/akademik yang baku
✓ Pastikan hanya ada SATU jawaban yang benar-benar tepat
✓ Buat pengecoh yang jelas salah untuk yang memahami materi

CONTOH SOAL YANG DIHINDARI:
========================================
✗ Coding dengan sintaks pemrograman
✗ Konfigurasi server atau database kompleks  
✗ Analisis algoritma tingkat lanjut
✗ Konsep networking di level enterprise
✗ Pertanyaan opini atau preferensi subjektif
✗ Soal dengan jawaban bergantung interpretasi
✗ Pertanyaan dengan informasi yang kurang lengkap

FORMAT OUTPUT - JSON ARRAY:
========================================
PENTING: Berikan response dalam format JSON array yang valid. Jangan tambahkan teks apapun sebelum atau sesudah JSON.

CONTOH FORMAT YANG BENAR (DEFINITIF & JELAS):
[{ {
    "level": 1,
    "difficulty": "Medium",
    "question_type": "multiple_choice", 
    "soal": "Fungsi utama CPU dalam sistem komputer adalah",
    "options": [
        "Menyimpan data secara permanen",
        "Memproses instruksi dan data", 
        "Menampilkan output ke layar",
        "Menyediakan daya listrik"
    ],
    "jawaban_benar": "Memproses instruksi dan data",
    "taxonomy_indicator": "Mengingat dan mengenali istilah/perangkat dasar TIK",
    "explanation": "CPU (Central Processing Unit) adalah unit pemrosesan pusat yang bertugas mengeksekusi instruksi program dan memproses data dalam komputer",
    "modul_reference": "• Menjelaskan komponen dasar komputer"
} }]

KRITERIA KUALITAS SOAL:
========================================
✓ SOAL harus menggunakan FAKTA TEKNIS yang dapat diverifikasi
✓ OPSI JAWABAN harus spesifik dan tidak ambigu
✓ JAWABAN BENAR harus 100% akurat menurut standar industri/akademik
✓ PENGECOH harus jelas salah bagi yang memahami konsep
✓ PENJELASAN harus memberikan alasan teknis yang solid
✓ TIDAK ADA opsi yang memerlukan interpretasi subjektif

VALIDASI SETIAP SOAL:
========================================
1. Apakah pertanyaan dapat dijawab dengan PASTI dari materi?
2. Apakah jawaban benar dapat DIBUKTIKAN secara objektif?
3. Apakah pengecoh jelas salah untuk yang memahami topik?
4. Apakah tidak ada kata ambiguitas ("mungkin", "biasanya", dll)?
5. Apakah istilah teknis yang digunakan STANDAR dan BAKU?

PASTIKAN FORMAT JSON:
========================================
- Mulai langsung dengan '[' dan akhiri dengan ']'
- Tidak ada teks penjelasan sebelum atau sesudah JSON
- Gunakan double quotes untuk semua string
- Tidak ada trailing comma setelah elemen terakhir
- Jawaban benar harus tepat salah satu dari opsi yang ada
- Semua string harus bebas dari karakter khusus yang merusak JSON

PANDUAN KHUSUS UNTUK KURIKULUM SMA KELAS X:
========================================
TOPIK YANG SESUAI:
- Pengenalan komputer dan perangkat digital
- Sistem operasi dasar (Windows, Android, iOS)
- Aplikasi perkantoran (Word, Excel, PowerPoint)
- Internet dan browsing aman
- Media sosial dan etika digital
- Keamanan digital dasar
- Multimedia sederhana (foto, video, audio)

HINDARI TOPIK KOMPLEKS:
- Pemrograman dengan sintaks spesifik
- Database management tingkat lanjut  
- Network administration
- Server configuration
- Advanced cybersecurity
- Machine learning atau AI development
 - Machine learning atau AI development

KETERKAITAN TAKSONOMI & TUJUAN (WAJIB):
========================================
Setiap soal HARUS terhubung jelas ke taksonomi dan tujuan/capaian pembelajaran:
- Gunakan field "level" (1–5) sesuai spesifikasi level di atas.
- Gunakan field "difficulty" dengan nilai: "Easy", "Medium", atau "Hard" sesuai mapping level.
- Tambahkan field "taxonomy_indicator": satu kalimat ringkas yang mewakili indikator taksonomi pada level tersebut.
- Tambahkan field "modul_reference": tulis persis salah satu butir Capaian Pembelajaran (CP) atau Tujuan Pembelajaran (TP) yang paling relevan (copy apa adanya).

DAFTAR INDIKATOR TAKSONOMI TEKNOLOGI (RINGKAS):
- Level 1 (Kesadaran): Mengingat dan mengenali istilah/perangkat dasar TIK.
- Level 2 (Literasi): Menjelaskan fungsi dan hubungan sederhana antarteknologi.
- Level 3 (Kemampuan): Menerapkan langkah/fitur dasar untuk tugas praktis.
- Level 4 (Kreativitas): Merancang atau memilih solusi sederhana untuk masalah nyata.
- Level 5 (Kritik): Mengevaluasi pilihan/implikasi teknologi disertai alasan yang jelas.

CHECKLIST KUALITAS SEBELUM MENGIRIM JAWABAN:
========================================
Pastikan SETIAP soal memenuhi kriteria berikut:

 PERTANYAAN:
    - Tidak menggunakan kata: mungkin, biasanya, umumnya, sebaiknya, kemungkinan
    - Menggunakan fakta teknis yang dapat diverifikasi dari sumber standar
    - Bahasa jelas, langsung, dan definitif
    - Dapat dijawab dengan pasti berdasarkan pengetahuan teknis

 OPSI JAWABAN:
    - Tidak ada opsi yang memerlukan interpretasi subjektif  
    - Jawaban benar 100% akurat menurut standar industri
    - Pengecoh jelas salah bagi yang memahami konsep
    - Tidak ada ambiguitas dalam formulasi opsi

 TERMINOLOGY:
    - Menggunakan istilah teknis baku dan standar
    - Tidak ada jargon yang dapat diinterpretasi berbeda
    - Konsisten dengan terminologi yang diajarkan di SMA
    - Sesuai dengan kurikulum dan standar pendidikan

 VERIFIKASI AKHIR:
    - Soal dapat dikerjakan oleh siswa SMA kelas X
    - Berkaitan langsung dengan materi modul ajar
    - Tidak memerlukan pengetahuan di luar scope SMA
    - Jawaban dapat dipertanggungjawabkan secara akademis

TARGET: 55 soal total dalam format JSON array yang valid.

Distribusi wajib per level & kesulitan (HARUS PERSIS, 5 per kombinasi yang diizinkan):
- Level 1 (Kesadaran): Medium 5 soal
- Level 2 (Literasi): Easy 5 soal, Hard 5 soal (total 10)
- Level 3 (Kemampuan): Medium 5 soal, Hard 5 soal (total 10)
- Level 4 (Kreativitas): Easy 5 soal, Medium 5 soal, Hard 5 soal (total 15)
- Level 5 (Kritik): Easy 5 soal, Medium 5 soal, Hard 5 soal (total 15)

Catatan: Pastikan setiap soal memakai field "level" (1–5) dan "difficulty" ("Easy" | "Medium" | "Hard") sesuai distribusi di atas. Jumlah per kombinasi level-kesulitan harus tepat 5.
"""
    return prompt

def extract_hybrid_module_components(content_text):
    """
    Fungsi hybrid yang mencoba ekstraksi spesifik terlebih dahulu,
    jika gagal akan menggunakan fallback ke metode yang lebih umum
    """
    # Coba ekstraksi spesifik Kurikulum Merdeka terlebih dahulu
    specific_components = extract_kurikulum_merdeka_components(content_text)
    validation = validate_kurikulum_merdeka_components(specific_components)
    
    # Jika quality score rendah, gunakan metode fallback
    if validation["quality_score"] < 30:  # Threshold rendah untuk fallback
        print(f"Ekstraksi spesifik kurang optimal (score: {validation['quality_score']:.1f}%), menggunakan metode fallback...")
        
        # Fallback: gunakan metode ekstraksi yang lebih umum
        fallback_components = extract_educational_components(content_text, return_dict=True)
        
        # Convert ke format yang sama dengan specific_components dan ambil yang terbaik
        hybrid_components = {
            "mata_pelajaran": specific_components.get("mata_pelajaran") or "Informatika", 
            "topik_utama": specific_components.get("topik_utama") or fallback_components.get("Modul/Elemen Ajar", "")[:50],
            "kelas": specific_components.get("kelas") or "X (Sepuluh)",
            "kompetensi_awal": specific_components.get("kompetensi_awal") or fallback_components.get("Kompetensi Awal", "Siswa memiliki kemampuan dasar dalam berpikir logis")[:400],
            "tujuan_pembelajaran": specific_components.get("tujuan_pembelajaran") or [fallback_components.get("Tujuan Pembelajaran", "Memahami konsep dasar materi pembelajaran")[:200]],
            "pemahaman_bermakna": specific_components.get("pemahaman_bermakna") or [fallback_components.get("Pemahaman Bermakna", "Mengembangkan pemahaman konseptual yang mendalam")[:300]],
            "target_peserta_didik": specific_components.get("target_peserta_didik") or fallback_components.get("Target Peserta Didik", "Peserta didik reguler/tipikal")[:300],
            "model_pembelajaran": "Pembelajaran berbasis masalah dan proyek"
        }
        
        print("Menggunakan metode hybrid - menggabungkan ekstraksi spesifik dengan fallback")
        return hybrid_components, {"quality_score": 60, "method": "hybrid"}
    else:
        print(f"Ekstraksi spesifik berhasil (score: {validation['quality_score']:.1f}%)")
        return specific_components, validation

# =========================================
# RESPONSE HELPERS
# =========================================
def extract_response_text(response):
    """Safely extract text from Gemini response across possible shapes.
    Returns an empty string if unavailable.
    """
    try:
        if response is None:
            return ""
        text = getattr(response, "text", None)
        if isinstance(text, str) and text.strip():
            return text.strip()
        # Try candidates -> content.parts[*].text
        candidates = getattr(response, "candidates", None)
        if candidates:
            parts_texts = []
            for c in candidates:
                content = getattr(c, "content", None)
                parts = getattr(content, "parts", None) if content else None
                if parts:
                    for p in parts:
                        t = getattr(p, "text", None)
                        if isinstance(t, str) and t.strip():
                            parts_texts.append(t.strip())
            if parts_texts:
                return "\n".join(parts_texts)
        return ""
    except Exception:
        return ""

# =========================================
# DATABASE MODELS
# =========================================
# Model untuk database guru
class Guru(db.Model, UserMixin):
    __tablename__ = 'guru'
    
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), unique=True, nullable=False)
    email = db.Column(db.String(100), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    nama = db.Column(db.String(100), nullable=False)
    has_seen_welcome_guide = db.Column(db.Boolean, default=False, nullable=False)
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    
    # Property untuk identifikasi tipe pengguna
    @property
    def user_type(self):
        return 'guru'
    
    # Override get_id method untuk Flask-Login
    def get_id(self):
        return f"guru_{self.id}"
    
    def set_password(self, password):
        self.password_hash = generate_password_hash(password)
    
    def check_password(self, password):
        return check_password_hash(self.password_hash, password)
    
    # Metode untuk mengakses hasil siswa
    def get_hasil_siswa(self):
        hasil_siswa = db.session.query(SiswaResult).join(
            Siswa, SiswaResult.siswa_id == Siswa.id
        ).join(
            SiswaAnswer, SiswaAnswer.siswa_id == Siswa.id
        ).join(
            Question, SiswaAnswer.question_id == Question.id
        ).filter(
            Question.guru_id == self.id
        ).distinct(
            SiswaResult.id
        ).all()
        
        return hasil_siswa

# Model untuk pertanyaan yang dibuat oleh guru
class Question(db.Model):
    __tablename__ = 'questions'
    
    id = db.Column(db.Integer, primary_key=True)
    guru_id = db.Column(db.Integer, db.ForeignKey('guru.id'), nullable=False)
    collection_id = db.Column(db.Integer, db.ForeignKey('question_collections.id'), nullable=True)
    level = db.Column(db.Integer, nullable=False)  # Deprecated, use technology_level instead
    technology_level = db.Column(db.Integer, nullable=False, default=1)  # L1-L5 (Awareness, Literacy, Capability, Creativity, Criticism)
    difficulty = db.Column(db.String(10), nullable=False, default='Medium')  # Easy, Medium, Hard
    soal = db.Column(db.Text, nullable=False)
    jawaban_benar = db.Column(db.Text, nullable=True)
    options = db.Column(db.Text, nullable=True)
    question_type = db.Column(db.String(20), default='multiple_choice')
    p = db.Column(db.Float, nullable=False)
    bobot = db.Column(db.Integer, nullable=False)
    explanation = db.Column(db.Text)
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    version = db.Column(db.Integer, default=1)
    is_current = db.Column(db.Boolean, default=True)
    is_validated = db.Column(db.Boolean, default=False) # NEW: Kolom is_validated
    
    
    # Relasi dengan guru
    guru = db.relationship('Guru', backref='questions')

# Model untuk database siswa
class Siswa(db.Model, UserMixin):
    __tablename__ = 'siswa'
    
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), unique=True, nullable=False)
    email = db.Column(db.String(100), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    nama = db.Column(db.String(100), nullable=False)
    kelas = db.Column(db.String(100), nullable=True)  # Changed from foreign key to string field
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    
    # Relasi dengan jawaban siswa
    jawaban = db.relationship('SiswaAnswer', backref='siswa', lazy=True)
    
    # Relasi dengan hasil siswa
    hasil = db.relationship('SiswaResult', backref='siswa', lazy=True)
    
    # Property untuk identifikasi tipe pengguna
    @property
    def user_type(self):
        return 'siswa'
    
    # Override get_id method untuk Flask-Login
    def get_id(self):
        return f"siswa_{self.id}"
    
    def set_password(self, password):
        self.password_hash = generate_password_hash(password)
    
    def check_password(self, password):
        return check_password_hash(self.password_hash, password)
    
    # Metode untuk mendapatkan hasil sendiri
    def get_my_result(self, kelas_id):
        return SiswaResult.query.filter_by(
            siswa_id=self.id,
            kelas_id=kelas_id
        ).first()

# Model untuk hasil siswa (pengganti SiswaResult)
class SiswaResult(db.Model):
    __tablename__ = 'siswa_results'
    
    id = db.Column(db.Integer, primary_key=True)
    siswa_id = db.Column(db.Integer, db.ForeignKey('siswa.id'), nullable=False)
    collection_id = db.Column(db.Integer, db.ForeignKey('question_collections.id'), nullable=True)  # <-- TAMBAHKAN INI
    correct = db.Column(db.Integer, default=0)
    incorrect = db.Column(db.Integer, default=0)
    current_level = db.Column(db.Integer, default=1)
    # MST Adaptive fields
    current_stage = db.Column(db.Integer, default=1)  # Stage 1-5
    current_path = db.Column(db.String(20), default='START')  # START, UPPER, MIDDLE, LOWER
    technology_level_achieved = db.Column(db.Integer, default=1)  # Final diagnosis L1-L5
    is_adaptive_test_complete = db.Column(db.Boolean, default=False)
    adaptive_test_data = db.Column(db.Text)  # JSON data for tracking MST path
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    updated_at = db.Column(db.DateTime, server_default=db.func.now(), onupdate=db.func.now())

# Model untuk jawaban siswa (pengganti SiswaAnswer)
class SiswaAnswer(db.Model):
    __tablename__ = 'siswa_answers'
    
    id = db.Column(db.Integer, primary_key=True)
    siswa_id = db.Column(db.Integer, db.ForeignKey('siswa.id'), nullable=False)
    question_id = db.Column(db.Integer, db.ForeignKey('questions.id'), nullable=False)
    collection_id = db.Column(db.Integer, db.ForeignKey('question_collections.id'), nullable=True)  # For MST tracking
    siswa_answer = db.Column(db.Text, nullable=False)
    selected_answer = db.Column(db.Text, nullable=True)  # Alias for siswa_answer
    is_correct = db.Column(db.Boolean, default=False)
    level = db.Column(db.String(20), nullable=False, default='Awareness')  # MST level name
    difficulty = db.Column(db.String(10), nullable=False, default='Medium')  # Easy, Medium, Hard
    stage = db.Column(db.Integer, nullable=False, default=1)  # MST Stage 1-5
    answered_at = db.Column(db.DateTime, server_default=db.func.now())
    
    # Relasi dengan Question
    question = db.relationship('Question', backref="siswa_answers")

results_bp = Blueprint('results', __name__)


# =========================================
# TABEL RELASI MANY-TO-MANY (harus didefinisikan SEBELUM model yang menggunakannya)
# =========================================
collection_students = db.Table('collection_students',
    db.Column('id', db.Integer, primary_key=True),
    db.Column('collection_id', db.Integer, db.ForeignKey('question_collections.id', ondelete='CASCADE'), nullable=False),
    db.Column('siswa_id', db.Integer, db.ForeignKey('siswa.id', ondelete='CASCADE'), nullable=False),
    db.Column('added_at', db.DateTime, server_default=db.func.now()),
    db.UniqueConstraint('collection_id', 'siswa_id', name='unique_collection_student')
)

# Model untuk koleksi soal
class QuestionCollection(db.Model):
    __tablename__ = 'question_collections'
    
    id = db.Column(db.Integer, primary_key=True)
    guru_id = db.Column(db.Integer, db.ForeignKey('guru.id'), nullable=False)
    name = db.Column(db.String(100), nullable=False)
    description = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    
    # Relasi dengan guru
    guru = db.relationship('Guru', backref='collections')
    
    # Tambahkan relasi dengan siswa
    students = db.relationship('Siswa', secondary=collection_students, 
                              backref=db.backref('collections', lazy='dynamic'))

# Model untuk relasi many-to-many antara Question dan QuestionCollection
class CollectionQuestion(db.Model):
    __tablename__ = 'collection_questions'
    
    id = db.Column(db.Integer, primary_key=True)
    collection_id = db.Column(db.Integer, db.ForeignKey('question_collections.id'), nullable=False)
    question_id = db.Column(db.Integer, db.ForeignKey('questions.id'), nullable=False)
    
    # Relasi
    collection = db.relationship('QuestionCollection', backref='collection_questions')
    question = db.relationship('Question', backref='collection_questions')


class QuestionVersion(db.Model):
    __tablename__ = 'question_versions'
    
    id = db.Column(db.Integer, primary_key=True)
    question_id = db.Column(db.Integer, db.ForeignKey('questions.id'), nullable=False)
    soal = db.Column(db.Text, nullable=False)
    options = db.Column(db.Text)
    jawaban_benar = db.Column(db.Text)
    explanation = db.Column(db.Text)
    p = db.Column(db.Float)
    version = db.Column(db.Integer)
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    
    # Relasi ke pertanyaan utama
    question = db.relationship('Question', backref='versions')

# Fungsi pemeriksaan apakah pengguna adalah guru
def is_guru():
    return hasattr(current_user, 'kelas')

def get_hasil_siswa(self):
    # Since there's no Kelas relationship, just get all student results
    hasil_siswa = SiswaResult.query.all()
    return hasil_siswa

@app.route('/guru/welcome')
@login_required
def guru_welcome_guide():
    # Pastikan hanya guru yang bisa mengakses halaman ini
    if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
        flash("Akses ditolak. Halaman ini hanya untuk guru.", "error")
        return redirect(url_for('login'))
    return render_template('guru/guru_welcome.html') # File HTML baru yang akan kita buat

@app.route('/api/guru/mark_welcome_seen', methods=['POST'])
@login_required
def mark_welcome_seen():
    if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
        return jsonify({'success': False, 'message': 'Akses ditolak'}), 403

    try:
        # Cari guru yang sedang login di database
        guru = Guru.query.get(current_user.id)
        if guru:
            # Update statusnya menjadi sudah melihat panduan
            guru.has_seen_welcome_guide = True
            db.session.commit()
            return jsonify({'success': True, 'message': 'Status panduan berhasil diperbarui.'})
        return jsonify({'success': False, 'message': 'Guru tidak ditemukan.'}), 404
    except Exception as e:
        db.session.rollback()
        print(f"Error saat menandai panduan selesai: {str(e)}")
        return jsonify({'success': False, 'message': 'Terjadi kesalahan server.'}), 500

@app.route('/guru/panduan')
@login_required
def guru_panduan_lengkap():
    if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
        flash("Akses ditolak.", "error")
        return redirect(url_for('login'))
    return render_template('guru/guru_panduan.html')

# Rute untuk guru mengakses hasil siswa
@results_bp.route('/guru/hasil-siswa', methods=['GET'])
@login_required
def guru_access_hasil_siswa():
    # Pastikan pengguna adalah guru
    if not is_guru():
        return jsonify({'message': 'Akses ditolak. Hanya guru yang dapat mengakses hasil siswa.'}), 403
    
    # Ambil hasil siswa dari kelas yang dimiliki guru
    hasil_siswa = current_user.get_hasil_siswa()
    
    # Format hasil untuk response
    hasil_formatted = []
    for hasil in hasil_siswa:
        # Dapatkan data siswa
        siswa = Siswa.query.get(hasil.siswa_id)
        
        hasil_formatted.append({
            'id': hasil.id,
            'siswa': {
                'id': siswa.id,
                'nama': siswa.nama,
                'username': siswa.username,
                'kelas': siswa.kelas
            },
            'correct': hasil.correct,
            'incorrect': hasil.incorrect,
            'current_level': hasil.current_level,
            'created_at': hasil.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            'updated_at': hasil.updated_at.strftime('%Y-%m-%d %H:%M:%S')
        })
    
    return jsonify({
        'message': 'Berhasil mengambil hasil siswa',
        'hasil': hasil_formatted
    }), 200

# Rute untuk guru mengakses detail hasil siswa tertentu
@results_bp.route('/guru/hasil-siswa/<int:siswa_id>', methods=['GET'])
@login_required
def guru_access_detail_hasil_siswa(siswa_id):
    if not is_guru():
        return jsonify({'message': 'Akses ditolak. Hanya guru yang dapat mengakses hasil siswa.'}), 403
    
    # Check if student exists
    siswa = Siswa.query.get(siswa_id)
    if not siswa:
        return jsonify({'message': 'Siswa tidak ditemukan.'}), 404
    
    # Get student results
    hasil = SiswaResult.query.filter_by(siswa_id=siswa_id).all()
    
    # Get student answers
    jawaban = SiswaAnswer.query.filter_by(siswa_id=siswa_id).all()
    
    # Format results
    hasil_formatted = []
    for h in hasil:
        hasil_formatted.append({
            'id': h.id,
            'correct': h.correct,
            'incorrect': h.incorrect,
            'current_level': h.current_level,
            'created_at': h.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            'updated_at': h.updated_at.strftime('%Y-%m-%d %H:%M:%S')
        })
    
    # Format answers
    jawaban_formatted = []
    for j in jawaban:
        jawaban_formatted.append({
            'id': j.id,
            'question_id': j.question_id,
            'siswa_answer': j.siswa_answer,
            'is_correct': j.is_correct,
            'level': j.level,
            'answered_at': j.answered_at.strftime('%Y-%m-%d %H:%M:%S')
        })
    
    return jsonify({
        'message': 'Berhasil mengambil detail hasil siswa',
        'siswa': {
            'id': siswa.id,
            'nama': siswa.nama,
            'username': siswa.username,
            'kelas': siswa.kelas
        },
        'hasil': hasil_formatted,
        'jawaban': jawaban_formatted
    }), 200

# Rute untuk siswa melihat hasil sendiri
@results_bp.route('/siswa/hasil-saya', methods=['GET'])
@login_required
def siswa_access_hasil_sendiri():
    if is_guru():
        return jsonify({'message': 'Akses ditolak. Rute ini hanya untuk siswa.'}), 403
    
    # Get student results
    hasil = SiswaResult.query.filter_by(siswa_id=current_user.id).first()
    
    if not hasil:
        return jsonify({'message': 'Hasil tidak ditemukan.'}), 404
    
    # Get student answers
    jawaban = SiswaAnswer.query.filter_by(siswa_id=current_user.id).all()
    
    # Format answers
    jawaban_formatted = []
    for j in jawaban:
        jawaban_formatted.append({
            'question_id': j.question_id,
            'siswa_answer': j.siswa_answer,
            'is_correct': j.is_correct,
            'level': j.level,
            'answered_at': j.answered_at.strftime('%Y-%m-%d %H:%M:%S')
        })
    
    return jsonify({
        'message': 'Berhasil mengambil hasil',
        'hasil': {
            'correct': hasil.correct,
            'incorrect': hasil.incorrect,
            'current_level': hasil.current_level,
            'kelas': current_user.kelas  # Use the student's kelas field
        },
        'jawaban': jawaban_formatted
    }), 200

# Fungsi untuk mendaftarkan blueprint ke aplikasi
# =========================================
# AUTH ROUTES


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')

        user = Guru.query.filter_by(email=email).first()
        
        if not user:
            user = Siswa.query.filter_by(email=email).first()

        if user and user.check_password(password):
            login_user(user)
            flash(f'Selamat datang, {user.nama}!', 'success')

            # --- PERUBAHAN 2: Modifikasi logika redirect untuk guru ---
            if hasattr(user, 'user_type') and user.user_type == 'guru':
                # Cek apakah guru sudah pernah melihat panduan
                if not user.has_seen_welcome_guide:
                    # Jika belum, arahkan ke halaman panduan
                    return redirect(url_for('guru_welcome_guide'))
                else:
                    # Jika sudah, arahkan ke dashboard guru seperti biasa
                    return redirect(url_for('guru'))
            elif hasattr(user, 'user_type') and user.user_type == 'siswa':
                return redirect(url_for('DashboardSiswa'))
            else:
                flash("Login berhasil, tetapi peran tidak dikenal.", "info")
                return redirect(url_for('index'))
        else:
            flash("Email atau password salah", "error")
            return redirect(url_for('login'))

    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('Anda telah berhasil logout', 'success')
    return redirect(url_for('login'))

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        nama = request.form.get('nama')
        email = request.form.get('email')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')
        role = request.form.get('role')
        
        # Ambil kelas dari form (akan ada di request jika role adalah siswa)
        kelas = request.form.get('kelas') 

        # Log untuk debugging
        print(f"Register attempt: Nama={nama}, Email={email}, Role={role}, Kelas={kelas}")

        if password != confirm_password:
            flash("Password dan konfirmasi password tidak sama!", "error")
            return render_template('register.html', 
                                   nama=nama, email=email, role=role, kelas=kelas)
        
        # Cek duplikasi berdasarkan peran
        if role == 'guru':
            existing_user = Guru.query.filter_by(email=email).first()
            if existing_user:
                flash("Email sudah terdaftar sebagai Guru.", "error")
                return render_template('register.html', 
                                       nama=nama, email=email, role=role, kelas=kelas)
            
            new_user = Guru(
                username=email,  # Gunakan email sebagai username untuk kompatibilitas
                email=email,
                nama=nama
            )
            new_user.set_password(password)
            db.session.add(new_user)
            db.session.commit()
            flash('Registrasi guru berhasil! Silakan login.', 'success')
            return redirect(url_for('login'))
        
        elif role == 'siswa': # role == 'siswa'
            existing_user = Siswa.query.filter_by(email=email).first()
            if existing_user:
                flash("Email sudah terdaftar sebagai Siswa.", "error")
                return render_template('register.html', 
                                       nama=nama, email=email, role=role, kelas=kelas)
            
            # Pastikan kelas dipilih untuk siswa
            if not kelas:
                flash("Kelas harus dipilih untuk siswa!", "error")
                return render_template('register.html', 
                                       nama=nama, email=email, role=role, kelas=kelas)

            new_user = Siswa(
                username=email,  # Gunakan email sebagai username untuk kompatibilitas
                email=email,
                nama=nama,
                kelas=kelas # Simpan kelas siswa
            )
            new_user.set_password(password)
            db.session.add(new_user)
            db.session.commit()
            flash('Registrasi siswa berhasil! Silakan login.', 'success')
            return redirect(url_for('login'))
        else: # Handle case where role is not 'guru' or 'siswa' (shouldn't happen with proper frontend)
            flash("Peran registrasi tidak valid.", "error")
            return render_template('register.html', 
                                   nama=nama, email=email, role=role, kelas=kelas)
            
    # Untuk permintaan GET, render template register
    # Peran default di sini bisa 'siswa' agar form siswa langsung muncul
    return render_template('register.html', role='siswa')

# =========================================
# TRANSISI level DENGAN PROBABILITAS
# =========================================
# Fungsi next_level yang baru berdasarkan diagram
def next_level(current_level, is_correct):
    """Menentukan level selanjutnya berdasarkan MST 5 Stage Technology Assessment."""
    if current_level == 1:
        return 2  # L1 selalu ke L2
    elif current_level == 2:
        return 3 if is_correct else 3  # L2 selalu ke L3 (sudah difilter easy/hard di routing)
    elif current_level == 3:
        return 4 if is_correct else 4  # L3 selalu ke L4 (sudah difilter medium/hard di routing)
    elif current_level == 4:
        return 5 if is_correct else 5  # L4 selalu ke L5 (sudah difilter oleh routing)
    elif current_level == 5:
        return None  # L5 adalah stage akhir
    else:
        return None  # Untuk keamanan

# Level akhir hanya L5 dalam sistem MST 5-stage
FINAL_LEVELS = [5]

# =========================================
# FUNGSI MENENTUKAN BOBOT (ASSIGN_WEIGHT)
# =========================================
def assign_weight(p: float) -> int:
    if p > 0.90:
        return 1
    elif 0.90 >= p > 0.62:
        return 2
    elif abs(p - 0.62) < 1e-6:
        return 3
    elif 0.62 > p > 0.20:
        return 4
    else:
        return 5

# =========================================
# FUNGSI KETERANGAN LEVEL
# =========================================
def get_level_description(level):
    """Mengembalikan deskripsi teks untuk level numerik berdasarkan MST 5 Stage Technology Assessment."""
    descriptions = {
        1: "L1 - Kesadaran Teknologi: Pengenalan teknologi dasar, identifikasi perangkat dan aplikasi",
        2: "L2 - Literasi Teknologi: Pemahaman cara kerja teknologi, literasi digital dasar", 
        3: "L3 - Kemampuan Teknologi: Penerapan teknologi untuk memecahkan masalah praktis",
        4: "L4 - Kreativitas Teknologi: Penggunaan teknologi secara kreatif dan inovatif",
        5: "L5 - Evaluasi Kritis Teknologi: Evaluasi dan kritik teknologi secara mendalam"
    }
    return descriptions.get(level, "Level Tidak Dikenal")

def get_level_description_short(level):
    """Mengembalikan deskripsi singkat untuk level numerik berdasarkan MST 5 Stage Technology Assessment."""
    descriptions = {
        1: "L1 - Kesadaran Teknologi",
        2: "L2 - Literasi Teknologi", 
        3: "L3 - Kemampuan Teknologi",
        4: "L4 - Kreativitas Teknologi",
        5: "L5 - Evaluasi Kritis Teknologi"
    }
    return descriptions.get(level, "Level Tidak Dikenal")

def get_technology_taxonomy(level):
    """Mengembalikan taksonomi teknologi untuk level numerik berdasarkan MST 5 Stage (format lengkap)."""
    technology_mapping = {
        1: "Knowledge That - Awareness level: Definisi, istilah, identifikasi teknologi dasar",
        2: "Knowledge That - Literacy level: Pemahaman cara kerja, literasi digital dasar", 
        3: "Knowledge That + How - Capability level: Penerapan teknologi untuk problem solving praktis",
        4: "Knowledge That + How + Create - Creativity level: Penggunaan teknologi secara kreatif dan inovatif",
        5: "Knowledge That + How + Why + Evaluate - Critical evaluation level: Evaluasi dan kritik teknologi mendalam"
    }
    return technology_mapping.get(level, "Tidak Dikenal")

def get_technology_taxonomy_short(level):
    """Mengembalikan taksonomi teknologi untuk level numerik berdasarkan MST 5 Stage (format singkat)."""
    technology_mapping = {
        1: "Know That",
        2: "Know That", 
        3: "Know That + How",
        4: "Know That + How + Create",
        5: "Know That + How + Why + Evaluate"
    }
    return technology_mapping.get(level, "Tidak Dikenal")

# Legacy functions for backward compatibility (deprecated)

# ===============================
# NORMALIZATION HELPERS
# ===============================
def normalize_level(value, min_level=1, max_level=5):
    """Normalize various level formats to an integer within [min_level, max_level].
    Accepts integers, strings like 'L5', 'Level 5', 'L5 - ...', '< L2'.
    Any value >= max_level maps to max_level; values below min map to min.
    """
    try:
        if value is None:
            return min_level
        if isinstance(value, int):
            if value >= max_level:
                return max_level
            if value <= min_level:
                return min_level
            return value
        s = str(value)
        m = re.search(r"([1-9])", s)
        if m:
            n = int(m.group(1))
            if n >= max_level:
                return max_level
            if n <= min_level:
                return min_level
            return n
    except Exception:
        pass
    return min_level


# =========================================
# FUNGSI EVALUASI JAWABAN
# =========================================
def is_answer_correct(user_answer, expected_answer):
    """
    Evaluasi jawaban pilihan ganda dengan membandingkan jawaban pengguna dengan jawaban yang benar
    """
    # Membersihkan dan menormalisasi jawaban
    user_answer = user_answer.lower().strip()
    expected_answer = expected_answer.lower().strip()
    
    # Untuk pilihan ganda, jawaban harus sama persis (tidak case-sensitive)
    return user_answer == expected_answer

# =========================================
# FUNGSI UNTUK CARI DAN TAMBAH KOLOM
# =========================================
def add_column_if_missing(table_name, column_name, column_definition):
    """Helper function to add column if it doesn't exist"""
    try:
        exists = db.session.execute(text(
            f"SELECT COUNT(*) FROM information_schema.COLUMNS WHERE TABLE_SCHEMA = DATABASE() AND TABLE_NAME='{table_name}' AND COLUMN_NAME='{column_name}'"
        )).scalar()
        
        if not exists:
            db.session.execute(text(f"ALTER TABLE {table_name} ADD COLUMN {column_definition}"))
            print(f"✅ Added column {column_name} to {table_name}")
            return True
        else:
            print(f"✓ Column {column_name} already exists in {table_name}")
            return False
    except Exception as e:
        print(f"❌ Error adding column {column_name} to {table_name}: {e}")
        raise e

def check_and_add_columns():
    """
    Cek dan tambahkan kolom yang diperlukan untuk sistem MST
    Migrasi lengkap dari: migrate_mst_database.py, add_missing_columns.py
    """
    try:
        print("🔧 Memulai migrasi database untuk sistem MST...")
        
        # 1. MIGRASI TABEL QUESTIONS
        print("\n📋 Migrasi tabel questions...")
        
        # Kolom dasar untuk questions (dari add_missing_columns.py)
        add_column_if_missing('questions', 'jawaban_benar', 'jawaban_benar TEXT NULL')
        add_column_if_missing('questions', 'options', 'options TEXT NULL')
        add_column_if_missing('questions', 'question_type', "question_type VARCHAR(20) DEFAULT 'multiple_choice'")
        add_column_if_missing('questions', 'is_validated', 'is_validated BOOLEAN DEFAULT FALSE')
        
        # Kolom MST untuk questions (dari migrate_mst_database.py)
        add_column_if_missing('questions', 'technology_level', "technology_level INT DEFAULT 1 COMMENT 'L1-L5: Awareness, Literacy, Capability, Creativity, Criticism'")
        add_column_if_missing('questions', 'difficulty', "difficulty VARCHAR(10) DEFAULT 'Medium' COMMENT 'Easy, Medium, Hard'")
        
        # Update existing questions dengan default values
        db.session.execute(text("""
            UPDATE questions
            SET technology_level = COALESCE(technology_level,
                CASE
                    WHEN level <= 2 THEN 1
                    WHEN level <= 4 THEN 2
                    WHEN level <= 6 THEN 3
                    ELSE 4
                END),
                difficulty = COALESCE(difficulty, 'Medium')
            WHERE technology_level IS NULL OR difficulty IS NULL
        """))
        
        # 2. MIGRASI TABEL SISWA_RESULTS
        print("\n👥 Migrasi tabel siswa_results...")
        
        add_column_if_missing('siswa_results', 'collection_id', 'collection_id INT NULL')
        add_column_if_missing('siswa_results', 'current_stage', "current_stage INT DEFAULT 1 COMMENT 'Stage 1-5 dalam MST'")
        add_column_if_missing('siswa_results', 'current_path', "current_path VARCHAR(20) DEFAULT 'START' COMMENT 'START, UPPER, MIDDLE, LOWER'")
        add_column_if_missing('siswa_results', 'technology_level_achieved', "technology_level_achieved INT DEFAULT 1 COMMENT 'Final diagnosis L1-L5'")
        add_column_if_missing('siswa_results', 'is_adaptive_test_complete', "is_adaptive_test_complete BOOLEAN DEFAULT FALSE COMMENT 'Apakah tes adaptif sudah selesai'")
        add_column_if_missing('siswa_results', 'adaptive_test_data', "adaptive_test_data TEXT COMMENT 'JSON data untuk tracking MST path'")
        
        # 3. MIGRASI TABEL SISWA_ANSWERS
        print("\n📝 Migrasi tabel siswa_answers...")
        
        add_column_if_missing('siswa_answers', 'collection_id', 'collection_id INT NULL')
        add_column_if_missing('siswa_answers', 'selected_answer', 'selected_answer TEXT NULL')
        add_column_if_missing('siswa_answers', 'technology_level', "technology_level INT DEFAULT 1 COMMENT 'L1-L5 saat jawaban dibuat'")
        add_column_if_missing('siswa_answers', 'difficulty', "difficulty VARCHAR(10) DEFAULT 'Medium' COMMENT 'Easy, Medium, Hard'")
        add_column_if_missing('siswa_answers', 'stage', "stage INT DEFAULT 1 COMMENT 'Stage MST saat menjawab'")
        
        # Ubah tipe data kolom level dari INT ke VARCHAR (dari add_missing_columns.py)
        try:
            level_col_info = db.session.execute(text("SHOW COLUMNS FROM siswa_answers WHERE Field = 'level'")).fetchone()
            if level_col_info and 'int' in level_col_info[1].lower():
                print("🔄 Mengubah tipe data kolom level dari INT ke VARCHAR(20)...")
                db.session.execute(text("ALTER TABLE siswa_answers MODIFY COLUMN level VARCHAR(20) NOT NULL DEFAULT 'Awareness'"))
                print("✅ Tipe data kolom level berhasil diubah")
        except Exception as e:
            print(f"⚠️ Warning: Tidak dapat mengubah tipe kolom level: {e}")
        
        # Update existing answers dengan default values
        db.session.execute(text("""
            UPDATE siswa_answers
            SET technology_level = COALESCE(technology_level,
                CASE
                    WHEN level LIKE '%1%' OR level LIKE '%Awareness%' THEN 1
                    WHEN level LIKE '%2%' OR level LIKE '%Literacy%' THEN 2
                    WHEN level LIKE '%3%' OR level LIKE '%Capability%' THEN 3
                    WHEN level LIKE '%4%' OR level LIKE '%Creativity%' THEN 4
                    WHEN level LIKE '%5%' OR level LIKE '%Criticism%' THEN 5
                    ELSE 1
                END),
                difficulty = COALESCE(difficulty, 'Medium'),
                stage = COALESCE(stage, 1)
            WHERE technology_level IS NULL OR difficulty IS NULL OR stage IS NULL
        """))
        
        # 4. TAMBAHKAN FOREIGN KEY CONSTRAINTS (dari add_missing_columns.py)
        print("\n🔗 Menambahkan foreign key constraints...")
        
        try:
            # Foreign key untuk siswa_answers.collection_id
            fk_exists = db.session.execute(text("""
                SELECT COUNT(*) FROM information_schema.KEY_COLUMN_USAGE 
                WHERE TABLE_SCHEMA = DATABASE() 
                AND TABLE_NAME = 'siswa_answers' 
                AND COLUMN_NAME = 'collection_id' 
                AND REFERENCED_TABLE_NAME = 'question_collections'
            """)).scalar()
            
            if not fk_exists:
                db.session.execute(text("""
                    ALTER TABLE siswa_answers 
                    ADD CONSTRAINT fk_siswa_answers_collection 
                    FOREIGN KEY (collection_id) 
                    REFERENCES question_collections(id)
                """))
                print("✅ Foreign key constraint untuk siswa_answers.collection_id ditambahkan")
        except Exception as e:
            print(f"⚠️ Warning: Tidak dapat menambahkan foreign key constraint: {e}")
        
        db.session.commit()
        print("\n✅ Semua migrasi database berhasil diselesaikan!")
        
        # SUMMARY
        print("\n📊 SUMMARY PERUBAHAN DATABASE:")
        print("=" * 50)
        print("TABEL QUESTIONS:")
        print("  + jawaban_benar (TEXT) - Jawaban yang benar")
        print("  + options (TEXT) - Pilihan jawaban")
        print("  + question_type (VARCHAR) - Tipe soal")
        print("  + is_validated (BOOLEAN) - Status validasi soal")
        print("  + technology_level (INT) - Level teknologi L1-L5")
        print("  + difficulty (VARCHAR) - Easy, Medium, Hard")
        print("\nTABEL SISWA_RESULTS:")
        print("  + collection_id (INT) - ID koleksi soal")
        print("  + current_stage (INT) - Stage MST saat ini")
        print("  + current_path (VARCHAR) - Path MST (START/UPPER/MIDDLE/LOWER)")
        print("  + technology_level_achieved (INT) - Diagnosis akhir L1-L5")
        print("  + is_adaptive_test_complete (BOOLEAN) - Status tes selesai")
        print("  + adaptive_test_data (TEXT) - Data tracking MST")
        print("\nTABEL SISWA_ANSWERS:")
        print("  + collection_id (INT) - ID koleksi soal")
        print("  + selected_answer (TEXT) - Jawaban yang dipilih")
        print("  + technology_level (INT) - Level saat menjawab")
        print("  + difficulty (VARCHAR) - Kesulitan soal")
        print("  + stage (INT) - Stage MST saat menjawab")
        print("  ~ level (VARCHAR) - Diubah dari INT ke VARCHAR")
        print("=" * 50)
        
        return True
        
    except Exception as e:
        db.session.rollback()
        print(f"❌ Error dalam migrasi database: {str(e)}")
        return False

# =========================================
# FUNGSI UNTUK KONEKSI DATABASE
# =========================================
def test_db_connection():
    try:
        # Cek apakah koneksi ke database bisa dibuat
        db.session.execute(text("SELECT 1"))
        print("Database connection successful")
        return True
    except Exception as e:
        print(f"Database connection error: {str(e)}")
        return False

# Inisialisasi database dengan app context
def init_db():
    print("Initializing database...")
    db.create_all()
    # Cek dan tambahkan kolom yang diperlukan
    check_and_add_columns()
    print("Database initialized successfully")

# =========================================
# ROUTES
# =========================================

# Favicon route - menggunakan Education.ico sebagai favicon
@app.route('/favicon.ico')
def favicon():
    try:
        return send_file(
            os.path.join(ROOT_DIR, 'img', 'Education.ico'), 
            mimetype='image/vnd.microsoft.icon'
        )
    except FileNotFoundError:
        # Fallback jika file tidak ditemukan
        print("Warning: Education.ico not found in img folder")
        return '', 404

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/guru')
@login_required # Tambahkan decorator login_required jika belum ada untuk memastikan user terotentikasi
def guru():
    # Pastikan yang mengakses halaman guru adalah user dengan user_type 'guru'
    if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
        flash("Akses ditolak. Hanya guru yang dapat mengakses halaman ini.", "error")
        return redirect(url_for('login'))
    return render_template('guru/Dashboard_guru.html')

@app.route('/DashboardSiswa')
@login_required # Tambahkan decorator login_required jika belum ada
def DashboardSiswa():
    # Pastikan yang mengakses halaman DashboardSiswa adalah user dengan user_type 'siswa'
    if not hasattr(current_user, 'user_type') or current_user.user_type != 'siswa':
        flash("Akses ditolak. Hanya siswa yang dapat mengakses halaman ini.", "error")
        return redirect(url_for('login'))
    return render_template('siswa/Dashboard_siswa.html')

@app.route('/mst-adaptive')
@login_required
def mst_adaptive_page():
    # Pastikan yang mengakses halaman MST adalah siswa
    if not hasattr(current_user, 'user_type') or current_user.user_type != 'siswa':
        flash("Akses ditolak. Hanya siswa yang dapat mengakses MST Adaptive Test.", "error")
        return redirect(url_for('login'))
    return render_template('siswa/qna_siswa.html')

@app.route('/qna')
def qna():
    return render_template('siswa/qna_siswa.html')

@app.route('/mst-guide')
def mst_guide():
    return render_template('siswa/mst_guide.html')


# =========================================
# RESULTS PAGE
# =========================================

@app.route('/api/siswa/test_status', methods=['GET'])
@login_required
def get_test_status():
    user_id = request.args.get("user_id")
    collection_id = request.args.get("collection_id")
    
    # Cek keamanan: pastikan user_id sesuai dengan pengguna saat ini
    if str(current_user.id) != str(user_id):
        return jsonify({"success": False, "message": "Akses tidak diizinkan"}), 403
    
    try:
        # Ambil hasil siswa untuk koleksi ini
        result = SiswaResult.query.filter_by(
            siswa_id=user_id,
            collection_id=collection_id
        ).first()
        
        if not result:
            return jsonify({
                "success": True,
                "is_completed": False,
                "current_level": 1
            })
        
        # Cek apakah current_level adalah level final (L5 dalam MST 5-stage)
        is_completed = result.current_level == 5
        
        return jsonify({
            "success": True,
            "is_completed": is_completed,
            "current_level": result.current_level,
            "correct": result.correct,
            "incorrect": result.incorrect
        })
        
    except Exception as e:
        print(f"Error mendapatkan status tes: {str(e)}")
        return jsonify({
            "success": False,
            "message": f"Error: {str(e)}"
        }), 500

@app.route('/api/siswa/results/<int:collection_id>', methods=['GET'])
@login_required
def get_siswa_test_results(collection_id):
    try:
        user_id = request.args.get("user_id")
        
        # For teacher access (viewing any student's results)
        if hasattr(current_user, 'user_type') and current_user.user_type == 'guru':
            # Check if collection belongs to this teacher
            collection = QuestionCollection.query.filter_by(
                id=collection_id, 
                guru_id=current_user.id
            ).first()
            
            if not collection:
                return jsonify({"success": False, "message": "Access denied - Collection not found or not owned by you"}), 403
                
            # If user_id is provided, get that specific student's results
            # Otherwise, return all results for this collection
            if user_id:
                # Verify this student is assigned to this collection
                student_assigned = db.session.query(collection_students).filter_by(
                    collection_id=collection_id,
                    siswa_id=user_id
                ).first() is not None
                
                if not student_assigned:
                    return jsonify({"success": False, "message": "Student not assigned to this collection"}), 404
            
        # For student access (viewing own results only)
        elif hasattr(current_user, 'user_type') and current_user.user_type == 'siswa':
            # Students can only view their own results
            if str(current_user.id) != str(user_id):
                return jsonify({"success": False, "message": "You can only view your own results"}), 403
                
            # Verify student has access to this collection
            student_assigned = db.session.query(collection_students).filter_by(
                collection_id=collection_id,
                siswa_id=current_user.id
            ).first() is not None
            
            if not student_assigned:
                return jsonify({"success": False, "message": "You don't have access to this collection"}), 403
        else:
            return jsonify({"success": False, "message": "Invalid user type"}), 403
        
        # Get student results for this collection
        result_query = SiswaResult.query.filter_by(
            collection_id=collection_id
        )
        
        if user_id:
            result_query = result_query.filter_by(siswa_id=user_id)
            
        results = result_query.all()
        
        if not results:
            # Return empty result instead of error if no results found
            return jsonify({
                "success": True,
                "message": "No results found",
                "correct": 0,
                "incorrect": 0,
                "current_level": 1,
                "answers": []
            })
            
        # Format results
        formatted_results = []
        for result in results:
            # Get student
            student = Siswa.query.get(result.siswa_id)
            
            # Get answer history
            answer_history = get_student_answer_history(result.siswa_id, collection_id)
            
            formatted_results.append({
                "student_id": result.siswa_id,
                "student_name": student.nama if student else "Unknown",
                "student_class": student.kelas if student else "Unknown",
                "correct": result.correct,
                "incorrect": result.incorrect,
                "current_level": result.current_level,
                "answers": answer_history
            })
        
        # If single student requested, return first result
        if user_id:
            if formatted_results:
                return jsonify({
                    "success": True,
                    "correct": formatted_results[0]["correct"],
                    "incorrect": formatted_results[0]["incorrect"],
                    "current_level": formatted_results[0]["current_level"],
                    "answers": formatted_results[0]["answers"]
                })
            else:
                return jsonify({
                    "success": True,
                    "message": "No results found",
                    "correct": 0,
                    "incorrect": 0,
                    "current_level": 1,
                    "answers": []
                })
        
        # Return all results for teacher
        return jsonify({
            "success": True,
            "results": formatted_results
        })
        
    except Exception as e:
        print(f"Error getting student test results: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"Server error: {str(e)}"
        }), 500

# =========================================
# SISTEM REKOMENDASI BELAJAR KONTEKSTUAL
# =========================================

def analyze_student_weaknesses(student_id, collection_id):
    """
    Menganalisis kelemahan siswa berdasarkan jawaban yang salah untuk memberikan rekomendasi belajar yang tepat
    """
    try:
        # Ambil semua jawaban salah siswa untuk collection ini
        wrong_answers = db.session.query(SiswaAnswer, Question)\
            .join(Question, SiswaAnswer.question_id == Question.id)\
            .join(CollectionQuestion, CollectionQuestion.question_id == Question.id)\
            .filter(
                SiswaAnswer.siswa_id == student_id,
                CollectionQuestion.collection_id == collection_id,
                SiswaAnswer.is_correct == False
            ).all()
        
        if not wrong_answers:
            return {"message": "Tidak ada jawaban salah ditemukan", "recommendations": []}
        
        # Kategorisasi berdasarkan level dan topik
        weakness_analysis = {
            "concept_gaps": [],      # Kesalahan konsep dasar
            "application_gaps": [],  # Kesalahan aplikasi
            "analysis_gaps": [],     # Kesalahan analisis
            "level_statistics": {},  # Statistik per level
            "topic_statistics": {}   # Statistik per topik
        }
        
        for answer, question in wrong_answers:
            level = question.level
            topic = extract_topic_from_question(question.soal)
            
            # Update statistik level
            if level not in weakness_analysis["level_statistics"]:
                weakness_analysis["level_statistics"][level] = {"count": 0, "questions": []}
            weakness_analysis["level_statistics"][level]["count"] += 1
            weakness_analysis["level_statistics"][level]["questions"].append({
                "question": question.soal,
                "correct_answer": question.jawaban_benar,
                "student_answer": answer.siswa_answer
            })
            
            # Update statistik topik
            if topic not in weakness_analysis["topic_statistics"]:
                weakness_analysis["topic_statistics"][topic] = {"count": 0, "questions": []}
            weakness_analysis["topic_statistics"][topic]["count"] += 1
            weakness_analysis["topic_statistics"][topic]["questions"].append({
                "question": question.soal,
                "level": level,
                "correct_answer": question.jawaban_benar
            })
            
            # Kategorisasi berdasarkan level
            if level in [1, 2]:  # Level dasar - konsep
                weakness_analysis["concept_gaps"].append({
                    "question": question.soal,
                    "topic": topic,
                    "level": level
                })
            elif level in [3, 4]:  # Level menengah - aplikasi
                weakness_analysis["application_gaps"].append({
                    "question": question.soal,
                    "topic": topic,
                    "level": level
                })
            else:  # Level tinggi - analisis
                weakness_analysis["analysis_gaps"].append({
                    "question": question.soal,
                    "topic": topic,
                    "level": level
                })
        
        return weakness_analysis
        
    except Exception as e:
        print(f"Error analyzing student weaknesses: {str(e)}")
        return {"error": str(e), "recommendations": []}

def extract_topic_from_question(question_text):
    """
    Mengekstrak topik dari teks soal untuk kategorisasi
    """
    question_lower = question_text.lower()
    
    # Mapping topik berdasarkan kata kunci
    topic_keywords = {
        "Jaringan Komputer": [
            "jaringan", "network", "internet", "router", "switch", "hub", 
            "protokol", "tcp", "ip", "http", "https", "dns", "dhcp",
            "topologi", "lan", "wan", "wifi", "ethernet"
        ],
        "Keamanan Siber": [
            "keamanan", "security", "password", "enkripsi", "virus", "malware",
            "firewall", "antivirus", "phishing", "hacker", "cybersecurity"
        ],
        "Sistem Operasi": [
            "sistem operasi", "windows", "linux", "android", "ios", "file system",
            "command", "shell", "directory", "folder", "registry"
        ],
        "Aplikasi Perkantoran": [
            "microsoft office", "word", "excel", "powerpoint", "spreadsheet",
            "dokumen", "presentasi", "formula", "chart", "table"
        ],
        "Database": [
            "database", "basis data", "sql", "query", "tabel", "field",
            "record", "primary key", "foreign key", "join", "select"
        ],
        "Programming": [
            "program", "coding", "algoritma", "variabel", "function", "loop",
            "if", "else", "array", "string", "integer", "boolean"
        ],
        "Hardware": [
            "hardware", "cpu", "ram", "storage", "hard disk", "ssd",
            "processor", "motherboard", "vga", "monitor", "keyboard", "mouse"
        ],
        "Internet & Web": [
            "browser", "website", "html", "css", "javascript", "url",
            "domain", "hosting", "email", "social media", "cloud"
        ]
    }
    
    # Cari topik yang paling sesuai
    topic_scores = {}
    for topic, keywords in topic_keywords.items():
        score = sum(1 for keyword in keywords if keyword in question_lower)
        if score > 0:
            topic_scores[topic] = score
    
    # Return topik dengan score tertinggi, atau "Umum" jika tidak ada yang cocok
    if topic_scores:
        return max(topic_scores, key=topic_scores.get)
    else:
        return "Teknologi Informasi Umum"

def generate_contextual_recommendations(weakness_analysis):
    """
    Membuat rekomendasi belajar berdasarkan analisis kelemahan siswa
    """
    recommendations = []
    
    # Rekomendasi berdasarkan statistik topik
    if weakness_analysis.get("topic_statistics"):
        # Urutkan topik berdasarkan jumlah kesalahan terbanyak
        sorted_topics = sorted(
            weakness_analysis["topic_statistics"].items(),
            key=lambda x: x[1]["count"],
            reverse=True
        )
        
        for topic, data in sorted_topics[:3]:  # Ambil 3 topik teratas
            recommendation = create_topic_recommendation(topic, data["count"], data["questions"])
            if recommendation:
                recommendations.append(recommendation)
    
    # Rekomendasi berdasarkan level kesulitan
    if weakness_analysis.get("level_statistics"):
        level_recommendations = create_level_recommendations(weakness_analysis["level_statistics"])
        recommendations.extend(level_recommendations)
    
    # Rekomendasi berdasarkan pola kesalahan
    pattern_recommendations = create_pattern_recommendations(weakness_analysis)
    recommendations.extend(pattern_recommendations)
    
    return recommendations

def create_topic_recommendation(topic, error_count, questions):
    """
    Membuat rekomendasi spesifik untuk topik tertentu
    """
    # Template rekomendasi per topik
    topic_recommendations = {
        "Jaringan Komputer": {
            "materials": [
                "📖 Bab: Konsep Dasar Jaringan Komputer",
                "🎥 Video: Pengenalan Topologi Jaringan",
                "📝 Latihan: Mengidentifikasi Perangkat Jaringan",
                "🔧 Praktik: Konfigurasi Jaringan Sederhana"
            ],
            "focus_areas": [
                "Memahami definisi dan fungsi jaringan komputer",
                "Mengenal jenis-jenis topologi jaringan",
                "Memahami protokol komunikasi data",
                "Mengidentifikasi perangkat jaringan dan fungsinya"
            ]
        },
        "Keamanan Siber": {
            "materials": [
                "📖 Bab: Keamanan Digital dan Cyber Security",
                "🎥 Video: Jenis-jenis Ancaman Siber",
                "📝 Latihan: Membuat Password yang Kuat",
                "🛡️ Praktik: Penggunaan Antivirus dan Firewall"
            ],
            "focus_areas": [
                "Memahami jenis-jenis ancaman keamanan digital",
                "Menerapkan praktik keamanan password",
                "Mengenal malware dan cara pencegahannya",
                "Memahami konsep enkripsi dasar"
            ]
        },
        "Sistem Operasi": {
            "materials": [
                "📖 Bab: Sistem Operasi dan Manajemen File",
                "🎥 Video: Fungsi dan Komponen Sistem Operasi",
                "📝 Latihan: Navigasi File dan Folder",
                "⚙️ Praktik: Pengaturan Sistem Operasi"
            ],
            "focus_areas": [
                "Memahami fungsi utama sistem operasi",
                "Menguasai manajemen file dan folder",
                "Mengenal interface dan command line",
                "Memahami proses dan multitasking"
            ]
        },
        "Aplikasi Perkantoran": {
            "materials": [
                "📖 Bab: Aplikasi Produktivitas Perkantoran",
                "🎥 Video: Tutorial Microsoft Office/LibreOffice",
                "📝 Latihan: Membuat Dokumen dan Presentasi",
                "📊 Praktik: Pengolahan Data dengan Spreadsheet"
            ],
            "focus_areas": [
                "Menguasai pengolah kata (Word Processing)",
                "Memahami spreadsheet dan formula",
                "Membuat presentasi yang efektif",
                "Integrasi antar aplikasi perkantoran"
            ]
        },
        "Database": {
            "materials": [
                "📖 Bab: Konsep Basis Data",
                "🎥 Video: Pengenalan Database dan SQL",
                "📝 Latihan: Merancang Tabel Database",
                "🗄️ Praktik: Query Sederhana"
            ],
            "focus_areas": [
                "Memahami konsep database dan DBMS",
                "Mengenal struktur tabel dan relasi",
                "Memahami primary key dan foreign key",
                "Dasar-dasar query SQL"
            ]
        },
        "Hardware": {
            "materials": [
                "📖 Bab: Perangkat Keras Komputer",
                "🎥 Video: Komponen-komponen Komputer",
                "📝 Latihan: Identifikasi Hardware",
                "🔧 Praktik: Perawatan Perangkat Keras"
            ],
            "focus_areas": [
                "Mengenal komponen utama komputer",
                "Memahami fungsi CPU, RAM, dan Storage",
                "Mengenal perangkat input dan output",
                "Dasar-dasar troubleshooting hardware"
            ]
        }
    }
    
    # Ambil rekomendasi untuk topik, atau buat generic jika tidak ada
    if topic in topic_recommendations:
        rec = topic_recommendations[topic]
    else:
        rec = {
            "materials": [
                f"📖 Pelajari kembali materi: {topic}",
                f"🎥 Cari video tutorial tentang: {topic}",
                f"📝 Latihan soal-soal: {topic}",
                f"🔍 Riset lebih dalam tentang: {topic}"
            ],
            "focus_areas": [
                f"Pahami konsep dasar {topic}",
                f"Pelajari aplikasi praktis {topic}",
                f"Latihan problem solving {topic}"
            ]
        }
    
    return {
        "type": "topic_focus",
        "topic": topic,
        "priority": "High" if error_count >= 3 else "Medium",
        "error_count": error_count,
        "title": f"📚 Fokus Belajar: {topic}",
        "description": f"Anda memiliki {error_count} kesalahan pada topik ini. Perlu penguatan pemahaman konsep.",
        "materials": rec["materials"],
        "focus_areas": rec["focus_areas"],
        "estimated_time": "2-3 jam belajar",
        "difficulty": "Sesuai dengan level kesalahan Anda"
    }

def create_level_recommendations(level_statistics):
    """
    Membuat rekomendasi berdasarkan level kesulitan
    """
    recommendations = []
    
    for level, data in level_statistics.items():
        if data["count"] >= 2:  # Jika ada 2+ kesalahan di level ini
            level_info = get_level_description_with_study_tips(level)
            
            recommendation = {
                "type": "level_focus",
                "level": level,
                "priority": "High" if data["count"] >= 3 else "Medium",
                "error_count": data["count"],
                "title": f"🎯 {level_info['title']}",
                "description": f"Perkuat kemampuan di level ini ({data['count']} kesalahan ditemukan)",
                "study_approach": level_info["study_tips"],
                "practice_suggestions": level_info["practice_suggestions"],
                "estimated_time": level_info["estimated_time"]
            }
            recommendations.append(recommendation)
    
    return recommendations

def get_level_description_with_study_tips(level):
    """
    Memberikan deskripsi level beserta tips belajar
    """
    level_guides = {
        1: {
            "title": "Level 1 - Kesadaran Teknologi (Dasar)",
            "study_tips": [
                "🔍 Fokus pada definisi dan terminologi dasar",
                "📖 Baca dan hafalkan konsep-konsep kunci",
                "🎯 Gunakan flashcard untuk mengingat istilah",
                "📝 Buat catatan ringkas dengan kata kunci"
            ],
            "practice_suggestions": [
                "Latihan soal pilihan ganda tentang definisi",
                "Quiz terminologi teknologi",
                "Identifikasi komponen-komponen teknologi"
            ],
            "estimated_time": "1-2 jam"
        },
        2: {
            "title": "Level 2 - Literasi Teknologi (Pemahaman)",
            "study_tips": [
                "🧠 Pahami hubungan antar konsep",
                "📊 Buat mind map untuk menghubungkan ide",
                "💭 Jelaskan konsep dengan kata-kata sendiri",
                "🔄 Bandingkan dan kontraskan konsep serupa"
            ],
            "practice_suggestions": [
                "Latihan menjelaskan fungsi teknologi",
                "Klasifikasi jenis-jenis teknologi",
                "Analisis hubungan cause-effect"
            ],
            "estimated_time": "2-3 jam"
        },
        3: {
            "title": "Level 3 - Kemampuan Teknologi (Aplikasi)",
            "study_tips": [
                "🛠️ Praktik langsung dengan tools",
                "📋 Ikuti step-by-step tutorial",
                "🎮 Gunakan simulator atau virtual lab",
                "✍️ Tulis prosedur dengan kata-kata sendiri"
            ],
            "practice_suggestions": [
                "Tutorial hands-on dengan software",
                "Simulasi konfigurasi sistem",
                "Latihan troubleshooting sederhana"
            ],
            "estimated_time": "3-4 jam"
        },
        4: {
            "title": "Level 4 - Kreativitas Teknologi (Problem Solving)",
            "study_tips": [
                "🧩 Latihan pemecahan masalah bertahap",
                "🎯 Identifikasi pola dalam masalah",
                "💡 Brainstorm multiple solutions",
                "🔧 Praktik debugging dan troubleshooting"
            ],
            "practice_suggestions": [
                "Case study problem solving",
                "Debugging exercise",
                "Design thinking untuk solusi teknologi"
            ],
            "estimated_time": "4-5 jam"
        },
        5: {
            "title": "Level 5 - Kemampuan Lanjut (Analisis Mendalam)",
            "study_tips": [
                "🔬 Analisis sistem secara holistik",
                "📈 Evaluasi pros dan cons solusi",
                "🎓 Pelajari best practices industri",
                "💼 Pahami konteks bisnis/organisasi"
            ],
            "practice_suggestions": [
                "Studi kasus industri",
                "Analisis perbandingan teknologi",
                "Evaluasi efektivitas sistem"
            ],
            "estimated_time": "5-6 jam"
        }
    }
    
    return level_guides.get(level, {
        "title": f"Level {level} - Kemampuan Khusus",
        "study_tips": ["Pelajari materi sesuai level kesulitan"],
        "practice_suggestions": ["Latihan soal sesuai level"],
        "estimated_time": "3-4 jam"
    })

def create_pattern_recommendations(weakness_analysis):
    """
    Membuat rekomendasi berdasarkan pola kesalahan
    """
    recommendations = []
    
    # Analisis pola berdasarkan jenis kesalahan
    concept_gaps = len(weakness_analysis.get("concept_gaps", []))
    application_gaps = len(weakness_analysis.get("application_gaps", []))
    analysis_gaps = len(weakness_analysis.get("analysis_gaps", []))
    
    # Rekomendasi jika banyak kesalahan konsep dasar
    if concept_gaps >= 3:
        recommendations.append({
            "type": "learning_pattern",
            "pattern": "conceptual_weakness",
            "priority": "Critical",
            "title": "🏗️ Perkuat Fondasi Konsep Dasar",
            "description": f"Ditemukan {concept_gaps} kesalahan pada konsep dasar. Penting untuk memperkuat pemahaman fundamental.",
            "action_plan": [
                "📖 Mulai dari definisi dan terminologi dasar",
                "🎯 Gunakan metode pembelajaran visual (diagram, infografis)",
                "📝 Buat rangkuman konsep dengan bahasa sendiri",
                "👥 Diskusi dengan teman atau guru untuk klarifikasi",
                "🔄 Review berkala konsep yang sudah dipelajari"
            ],
            "warning": "⚠️ Konsep dasar yang lemah akan mempengaruhi pembelajaran level yang lebih tinggi",
            "estimated_time": "5-7 hari belajar intensif"
        })
    
    # Rekomendasi jika banyak kesalahan aplikasi
    if application_gaps >= 3:
        recommendations.append({
            "type": "learning_pattern", 
            "pattern": "application_weakness",
            "priority": "High",
            "title": "⚙️ Tingkatkan Kemampuan Aplikasi",
            "description": f"Ditemukan {application_gaps} kesalahan dalam mengaplikasikan konsep. Perlu lebih banyak praktik.",
            "action_plan": [
                "🛠️ Perbanyak latihan hands-on dan praktik langsung",
                "📋 Ikuti tutorial step-by-step",
                "🎮 Gunakan simulator atau tools pembelajaran interaktif",
                "🔧 Latihan troubleshooting kasus sederhana",
                "📊 Buat project kecil untuk menerapkan konsep"
            ],
            "estimated_time": "3-5 hari dengan praktik rutin"
        })
    
    # Rekomendasi jika banyak kesalahan analisis
    if analysis_gaps >= 2:
        recommendations.append({
            "type": "learning_pattern",
            "pattern": "analysis_weakness", 
            "priority": "Medium",
            "title": "🧠 Kembangkan Kemampuan Analisis",
            "description": f"Ditemukan {analysis_gaps} kesalahan dalam analisis tingkat tinggi. Perlu pengembangan critical thinking.",
            "action_plan": [
                "🔍 Latihan case study dan problem solving",
                "🧩 Break down masalah kompleks menjadi bagian kecil",
                "💭 Praktik berpikir sistematis dan logis",
                "📊 Analisis cause-effect dalam teknologi",
                "🎯 Evaluasi multiple solutions untuk satu masalah"
            ],
            "estimated_time": "1-2 minggu dengan latihan konsisten"
        })
    
    return recommendations

@app.route('/api/siswa/learning-recommendations/<int:collection_id>', methods=['GET'])
@login_required
def get_learning_recommendations(collection_id):
    """
    Endpoint untuk mendapatkan rekomendasi belajar berdasarkan hasil jawaban siswa
    """
    try:
        user_id = request.args.get("user_id")
        
        # Validasi akses
        if hasattr(current_user, 'user_type') and current_user.user_type == 'guru':
            # Teacher dapat mengakses rekomendasi semua siswa di collection mereka
            collection = QuestionCollection.query.filter_by(
                id=collection_id, 
                guru_id=current_user.id
            ).first()
            
            if not collection:
                return jsonify({"success": False, "message": "Collection not found"}), 404
                
            if not user_id:
                return jsonify({"success": False, "message": "user_id required for teacher access"}), 400
                
        elif hasattr(current_user, 'user_type') and current_user.user_type == 'siswa':
            # Siswa hanya dapat mengakses rekomendasi sendiri
            user_id = current_user.id
            
            # Verify student has access to this collection
            student_assigned = db.session.query(collection_students).filter_by(
                collection_id=collection_id,
                siswa_id=current_user.id
            ).first() is not None
            
            if not student_assigned:
                return jsonify({"success": False, "message": "Access denied"}), 403
        else:
            return jsonify({"success": False, "message": "Invalid user type"}), 403
        
        # Analisis kelemahan siswa
        weakness_analysis = analyze_student_weaknesses(user_id, collection_id)
        
        if "error" in weakness_analysis:
            return jsonify({
                "success": False,
                "message": f"Error analyzing weaknesses: {weakness_analysis['error']}"
            }), 500
        
        # Generate rekomendasi kontekstual
        recommendations = generate_contextual_recommendations(weakness_analysis)
        
        # Get collection info untuk konteks
        collection = QuestionCollection.query.get(collection_id)
        student = Siswa.query.get(user_id)
        
        # Summary statistik
        total_errors = sum(data["count"] for data in weakness_analysis.get("topic_statistics", {}).values())
        weak_topics = list(weakness_analysis.get("topic_statistics", {}).keys())[:3]
        weak_levels = [level for level, data in weakness_analysis.get("level_statistics", {}).items() if data["count"] >= 2]
        
        return jsonify({
            "success": True,
            "student_name": student.nama if student else "Unknown",
            "collection_name": collection.name if collection else "Unknown",
            "analysis_summary": {
                "total_errors": total_errors,
                "weak_topics": weak_topics,
                "weak_levels": weak_levels,
                "concept_gaps": len(weakness_analysis.get("concept_gaps", [])),
                "application_gaps": len(weakness_analysis.get("application_gaps", [])),
                "analysis_gaps": len(weakness_analysis.get("analysis_gaps", []))
            },
            "recommendations": recommendations,
            "detailed_analysis": weakness_analysis
        })
        
    except Exception as e:
        print(f"Error generating learning recommendations: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"Server error: {str(e)}"
        }), 500

@app.route('/api/guru/learning-recommendations/<int:collection_id>/<int:student_id>', methods=['GET'])
@login_required
def get_student_learning_recommendations_for_teacher(collection_id, student_id):
    """
    Endpoint khusus untuk guru mendapatkan rekomendasi belajar siswa
    """
    try:
        # Validasi akses guru
        if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
            return jsonify({"success": False, "message": "Access denied"}), 403
            
        # Validasi koleksi milik guru
        collection = QuestionCollection.query.filter_by(
            id=collection_id, 
            guru_id=current_user.id
        ).first()
        
        if not collection:
            return jsonify({"success": False, "message": "Collection not found"}), 404
            
        # Validasi siswa exists dan assigned ke collection ini
        student = Siswa.query.get(student_id)
        if not student:
            return jsonify({"success": False, "message": "Student not found"}), 404
            
        student_assigned = db.session.query(collection_students).filter_by(
            collection_id=collection_id,
            siswa_id=student_id
        ).first() is not None
        
        if not student_assigned:
            return jsonify({"success": False, "message": "Student not assigned to this collection"}), 403
        
        # Analisis kelemahan siswa
        weakness_analysis = analyze_student_weaknesses(student_id, collection_id)
        
        if "error" in weakness_analysis:
            return jsonify({
                "success": False,
                "message": f"Error analyzing weaknesses: {weakness_analysis['error']}"
            }), 500
        
        # Generate rekomendasi kontekstual
        recommendations = generate_contextual_recommendations(weakness_analysis)
        
        # Get additional context
        result = SiswaResult.query.filter_by(
            siswa_id=student_id,
            collection_id=collection_id
        ).first()
        
        # Summary statistik
        total_errors = sum(data["count"] for data in weakness_analysis.get("topic_statistics", {}).values())
        weak_topics = list(weakness_analysis.get("topic_statistics", {}).keys())[:3]
        weak_levels = [level for level, data in weakness_analysis.get("level_statistics", {}).items() if data["count"] >= 2]
        
        return jsonify({
            "success": True,
            "student_name": student.nama,
            "collection_name": collection.name,
            "student_result": {
                "current_level": result.current_level if result else 0,
                "correct": result.correct if result else 0,
                "incorrect": result.incorrect if result else 0,
                "accuracy": round((result.correct / (result.correct + result.incorrect) * 100), 2) if result and (result.correct + result.incorrect) > 0 else 0
            },
            "analysis_summary": {
                "total_errors": total_errors,
                "weak_topics": weak_topics,
                "weak_levels": weak_levels,
                "concept_gaps": len(weakness_analysis.get("concept_gaps", [])),
                "application_gaps": len(weakness_analysis.get("application_gaps", [])),
                "analysis_gaps": len(weakness_analysis.get("analysis_gaps", []))
            },
            "recommendations": recommendations,
            "detailed_analysis": weakness_analysis
        })
        
    except Exception as e:
        print(f"Error generating teacher learning recommendations: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"Server error: {str(e)}"
        }), 500

@app.route('/api/guru/collection-recommendations-summary/<int:collection_id>', methods=['GET'])
@login_required
def get_collection_recommendations_summary(collection_id):
    """
    Endpoint untuk guru mendapatkan ringkasan rekomendasi semua siswa dalam koleksi
    """
    try:
        # Validasi akses guru
        if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
            return jsonify({"success": False, "message": "Access denied"}), 403
            
        # Validasi koleksi milik guru
        collection = QuestionCollection.query.filter_by(
            id=collection_id, 
            guru_id=current_user.id
        ).first()
        
        if not collection:
            return jsonify({"success": False, "message": "Collection not found"}), 404
        
        # Ambil semua siswa yang assigned ke koleksi ini
        # --- PERBAIKAN 1: Menggunakan 'added_at' ---
        assigned_students = db.session.query(
            Siswa, collection_students.c.added_at  # Diganti dari assigned_at
        ).join(
            collection_students, 
            Siswa.id == collection_students.c.siswa_id
        ).filter(
            collection_students.c.collection_id == collection_id
        ).all()
        # --- AKHIR PERBAIKAN 1 ---
        
        students_summary = []
        collection_topics_count = {}
        collection_levels_count = {}
        total_students_analyzed = 0
        
        for student, added_at in assigned_students: # Diganti dari assigned_at
        
            # Cek apakah siswa sudah mengerjakan tes
            result = SiswaResult.query.filter_by(
                siswa_id=student.id,
                collection_id=collection_id
            ).first()
            
            # --- PERBAIKAN 2: Logika Status Tes ---
            status_tes = "belum_mengerjakan"
            current_level_result = 1
            if result:
                current_level_result = result.current_level or 1
                total_jawaban = (result.correct or 0) + (result.incorrect or 0)
                # Gunakan flag is_adaptive_test_complete dari database
                is_completed = result.is_adaptive_test_complete 
                
                if is_completed:
                    # Jika tes selesai (mencapai STOP node di level manapun L2, L4, atau L5)
                    status_tes = "sudah_mengerjakan"
                elif total_jawaban > 0:
                    # Jika belum selesai TAPI sudah menjawab (masih di L1, L2, L3, L4)
                    status_tes = "sedang_mengerjakan"
                # else: status_tes tetap "belum_mengerjakan"
            # --- AKHIR PERBAIKAN 2 ---

            # Hanya analisis jika tes sudah selesai
            if status_tes == "sudah_mengerjakan":
                weakness_analysis = analyze_student_weaknesses(student.id, collection_id)
                
                if "error" not in weakness_analysis:
                    total_students_analyzed += 1
                    
                    recommendations = generate_contextual_recommendations(weakness_analysis)
                    
                    priority_count = {"Critical": 0, "High": 0, "Medium": 0}
                    for rec in recommendations:
                        priority = rec.get("priority", "Medium")
                        if priority in priority_count:
                            priority_count[priority] += 1
                    
                    for topic, data in weakness_analysis.get("topic_statistics", {}).items():
                        if topic not in collection_topics_count:
                            collection_topics_count[topic] = 0
                        collection_topics_count[topic] += data["count"]
                    
                    for level, data in weakness_analysis.get("level_statistics", {}).items():
                        if level not in collection_levels_count:
                            collection_levels_count[level] = 0
                        collection_levels_count[level] += data["count"]
                    
                    students_summary.append({
                        "student_id": student.id,
                        "student_name": student.nama,
                        "student_class": student.kelas,
                        "status": status_tes,
                        "assigned_date": added_at.strftime('%Y-%m-%d %H:%M:%S') if added_at else None,
                        "result": {
                            "current_level": result.current_level,
                            "correct": result.correct,
                            "incorrect": result.incorrect,
                            "accuracy": round((result.correct / (result.correct + result.incorrect) * 100), 2) if (result.correct + result.incorrect) > 0 else 0
                        },
                        "weakness_summary": {
                            "total_errors": sum(data["count"] for data in weakness_analysis.get("topic_statistics", {}).values()),
                            "weak_topics": list(weakness_analysis.get("topic_statistics", {}).keys())[:3],
                            "weak_levels": [level for level, data in weakness_analysis.get("level_statistics", {}).items() if data["count"] >= 2]
                        },
                        "recommendations_count": len(recommendations),
                        "priority_summary": priority_count
                    })
                else:
                    # Error saat analisis
                    students_summary.append({
                        "student_id": student.id,
                        "student_name": student.nama,
                        "student_class": student.kelas,
                        "status": "error_analysis",
                        "assigned_date": added_at.strftime('%Y-%m-%d %H:%M:%S') if added_at else None,
                        "result": { "current_level": current_level_result, "correct": result.correct, "incorrect": result.incorrect, "accuracy": 0 },
                        "recommendations_count": 0,
                        "priority_summary": {"Critical": 0, "High": 0, "Medium": 0},
                        "error": weakness_analysis.get("error", "Unknown error")
                    })
            else:
                # Jika status 'belum_mengerjakan' atau 'sedang_mengerjakan'
                students_summary.append({
                    "student_id": student.id,
                    "student_name": student.nama,
                    "student_class": student.kelas,
                    "status": status_tes,
                    "assigned_date": added_at.strftime('%Y-%m-%d %H:%M:%S') if added_at else None,
                    "result": {
                        "current_level": current_level_result,
                        "correct": result.correct if result else 0,
                        "incorrect": result.incorrect if result else 0,
                        "accuracy": round((result.correct / (result.correct + result.incorrect) * 100), 2) if result and (result.correct + result.incorrect) > 0 else 0
                    },
                    "recommendations_count": 0,
                    "priority_summary": {"Critical": 0, "High": 0, "Medium": 0}
                })
        
        most_common_topics = sorted(collection_topics_count.items(), key=lambda x: x[1], reverse=True)[:5]
        most_problematic_levels = sorted(collection_levels_count.items(), key=lambda x: x[1], reverse=True)[:3]
        
        return jsonify({
            "success": True,
            "collection_name": collection.name,
            "collection_description": collection.description,
            "total_students": len(assigned_students),
            "students_completed": len([s for s in students_summary if s["status"] == "sudah_mengerjakan"]),
            "students_pending": len([s for s in students_summary if s["status"] == "sedang_mengerjakan"]), # Hanya sedang mengerjakan
            "students_summary": students_summary,
            "collection_analysis": {
                "total_students_analyzed": total_students_analyzed,
                "most_common_weak_topics": [{"topic": topic, "error_count": count} for topic, count in most_common_topics],
                "most_problematic_levels": [{"level": level, "error_count": count} for level, count in most_problematic_levels],
                "recommendations_needed": {
                    "critical": sum(1 for s in students_summary if s.get("priority_summary", {}).get("Critical", 0) > 0),
                    "high": sum(1 for s in students_summary if s.get("priority_summary", {}).get("High", 0) > 0),
                    "medium": sum(1 for s in students_summary if s.get("priority_summary", {}).get("Medium", 0) > 0)
                }
            }
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"Error generating collection recommendations summary: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"Server error: {str(e)}"
        }), 500

# Helper function to get student answer history
def generate_personalized_tip_for_answer(question_text, user_answer, correct_answer, explanation, level):
    """Generate personalized tip for a specific wrong answer using AI"""
    try:
        prompt = f"""
        Kamu adalah GURU yang memberikan tips belajar singkat untuk siswa yang salah menjawab soal.

        SOAL: {question_text}
        
        JAWABAN SISWA: {user_answer}
        JAWABAN BENAR: {correct_answer}
        PENJELASAN: {explanation}
        LEVEL SOAL: {level}

        TUGASMU:
        Buat 1 tips belajar yang SANGAT SINGKAT (maksimal 2 kalimat) yang spesifik untuk soal ini.
        Tips harus praktis dan langsung bisa diterapkan siswa.

        GAYA BAHASA:
        [OK] Gunakan "kamu" atau "Bapak/Ibu sarankan"
        [OK] Bahasa guru yang peduli
        [OK] Konkret dan spesifik untuk materi soal ini
        [OK] Positif dan memotivasi

        HINDARI:
        [DITOLAK] Tips umum yang bisa dipakai untuk semua soal
        [DITOLAK] Kalimat panjang
        [DITOLAK] Bullet points atau numbering
        [DITOLAK] Pengulangan penjelasan yang sudah ada

        CONTOH YANG BAGUS:
        "Untuk soal seperti ini, coba buat tabel perbandingan antara konsep A dan B. Ini akan bantu kamu lihat perbedaannya dengan jelas."
        "Bapak/Ibu sarankan kamu latih lagi rumus ini sambil tulis langkah-langkahnya. Jangan lupa cek satuan ya!"

        Berikan 1 tips yang tepat untuk soal ini:
        """
        
        response = model.generate_content(prompt)
        tip = response.text.strip()
        
        # Bersihkan dari bullet points atau numbering jika ada
        tip = re.sub(r'^[\d\.\-\*\•]+\s*', '', tip)
        tip = tip.replace('\n', ' ').strip()
        
        return tip
    except Exception as e:
        print(f"Error generating personalized tip: {str(e)}")
        # Fallback ke tips berdasarkan level jika AI gagal
        if level <= 2:
            return "Untuk menguasai konsep dasar ini, buat catatan ringkas dengan kata-katamu sendiri. Ini akan bantu kamu ingat lebih lama."
        elif level <= 4:
            return "Pelajari kembali materi ini dengan mencari tahu mengapa jawaban yang benar itu benar. Coba jelaskan ke diri sendiri atau teman."
        else:
            return "Untuk soal analisis seperti ini, buat tabel perbandingan atau diagram. Visual akan bantu kamu lihat hubungan antar konsepnya."

def get_student_answer_history(student_id, collection_id):
    try:
        # Get collection questions
        collection_question_ids = db.session.query(CollectionQuestion.question_id)\
            .filter(CollectionQuestion.collection_id == collection_id).all()
        collection_question_ids = [id[0] for id in collection_question_ids]
        
        if not collection_question_ids:
            return []
            
        # Get answers for this student and these questions, ordered by level and then by question id
        answers = db.session.query(SiswaAnswer, Question)\
            .join(Question, SiswaAnswer.question_id == Question.id)\
            .filter(
                SiswaAnswer.siswa_id == student_id,
                Question.id.in_(collection_question_ids)
            )\
            .order_by(Question.level.asc(), Question.id.asc())\
            .all()
            
        # Format answers
        formatted_answers = []
        for answer, question in answers:
            answer_dict = {
                "question_id": question.id,
                "question": question.soal,
                "user_answer": answer.siswa_answer,
                "correct_answer": question.jawaban_benar,
                "is_correct": answer.is_correct,
                "explanation": question.explanation,
                "level": answer.level,
                "answered_at": answer.answered_at.isoformat() if answer.answered_at else None,
                "stage": answer.stage if hasattr(answer, 'stage') else None, # Ambil stage dari SiswaAnswer
                "difficulty": answer.difficulty if hasattr(answer, 'difficulty') else None # Ambil difficulty dari SiswaAnswer
            }
            
            # Generate personalized tip untuk jawaban yang salah
            if not answer.is_correct:
                personalized_tip = generate_personalized_tip_for_answer(
                    question.soal,
                    answer.siswa_answer,
                    question.jawaban_benar,
                    question.explanation,
                    answer.level
                )
                answer_dict["personalized_tip"] = personalized_tip
            
            formatted_answers.append(answer_dict)
            
        return formatted_answers
    except Exception as e:
        print(f"Error getting student answer history: {str(e)}")
        return []

@app.route('/api/collection-analytics/answers', methods=['GET'])
@login_required
def get_filtered_collection_answers():
    try:
        # Required parameter
        collection_id = request.args.get('collection_id')
        
        if not collection_id:
            return jsonify({"success": False, "message": "Collection ID is required"}), 400
            
        # Check if collection belongs to current teacher
        collection = QuestionCollection.query.filter_by(
            id=collection_id, 
            guru_id=current_user.id
        ).first()
        
        if not collection:
            return jsonify({"success": False, "message": "Collection not found or not owned by you"}), 403
            
        # Get collection questions
        collection_question_ids = db.session.query(CollectionQuestion.question_id)\
            .filter(CollectionQuestion.collection_id == collection_id).all()
        collection_question_ids = [id[0] for id in collection_question_ids]
        
        if not collection_question_ids:
            return jsonify({"success": True, "answers": [], "message": "No questions in this collection"})
            
        # Build query for student answers
        query = db.session.query(
            SiswaAnswer, Question, Siswa
        ).join(
            Question, SiswaAnswer.question_id == Question.id
        ).join(
            Siswa, SiswaAnswer.siswa_id == Siswa.id
        ).filter(
            Question.id.in_(collection_question_ids)
        )
        
        # Apply filters
        # 1. Filter by student ID
        student_id = request.args.get('student_id')
        if student_id:
            query = query.filter(SiswaAnswer.siswa_id == student_id)
            
        # 2. Filter by class
        class_filter = request.args.get('class')
        if class_filter and class_filter != 'all':
            query = query.filter(Siswa.kelas == class_filter)
            
        # 3. Filter by level
        level_filter = request.args.get('level')
        if level_filter and level_filter != 'all':
            query = query.filter(SiswaAnswer.level == level_filter)
            
        # 4. Filter by correctness
        is_correct_filter = request.args.get('is_correct')
        if is_correct_filter:
            if is_correct_filter.lower() == 'true':
                query = query.filter(SiswaAnswer.is_correct == True)
            elif is_correct_filter.lower() == 'false':
                query = query.filter(SiswaAnswer.is_correct == False)
                
        # 5. Filter by date range
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        if start_date:
            try:
                start_datetime = datetime.datetime.fromisoformat(start_date)
                query = query.filter(SiswaAnswer.answered_at >= start_datetime)
            except ValueError:
                # Ignore invalid date format
                pass
                
        if end_date:
            try:
                end_datetime = datetime.datetime.fromisoformat(end_date)
                query = query.filter(SiswaAnswer.answered_at <= end_datetime)
            except ValueError:
                # Ignore invalid date format
                pass
                
        # 6. Apply time period filter (shorthand for date range)
        period_filter = request.args.get('period')
        if period_filter and not (start_date or end_date):
            now = datetime.datetime.now()
            
            if period_filter == 'today':
                today_start = now.replace(hour=0, minute=0, second=0, microsecond=0)
                query = query.filter(SiswaAnswer.answered_at >= today_start)
            elif period_filter == 'yesterday':
                yesterday_start = (now - datetime.timedelta(days=1)).replace(hour=0, minute=0, second=0, microsecond=0)
                yesterday_end = now.replace(hour=0, minute=0, second=0, microsecond=0)
                query = query.filter(SiswaAnswer.answered_at >= yesterday_start, SiswaAnswer.answered_at < yesterday_end)
            elif period_filter == 'week':
                week_start = (now - datetime.timedelta(days=7))
                query = query.filter(SiswaAnswer.answered_at >= week_start)
            elif period_filter == 'month':
                month_start = (now - datetime.timedelta(days=30))
                query = query.filter(SiswaAnswer.answered_at >= month_start)
                
        # Sort by most recent first
        query = query.order_by(SiswaAnswer.answered_at.desc())
        
        # Get paginated results
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 50))
        
        # Execute query with pagination
        paginated_results = query.paginate(page=page, per_page=per_page, error_out=False)
        
        # Get total count for stats
        total_count = paginated_results.total
        total_pages = paginated_results.pages
        
        # Format results
        formatted_answers = []
        for answer, question, student in paginated_results.items:
            formatted_answers.append({
                "id": answer.id,
                "student_id": student.id,
                "student_name": student.nama,
                "student_class": student.kelas,
                "question_id": question.id,
                "question": question.soal,
                "user_answer": answer.siswa_answer,
                "correct_answer": question.jawaban_benar,
                "is_correct": answer.is_correct,
                "explanation": question.explanation,
                "level": answer.level,
                "answered_at": answer.answered_at.isoformat() if answer.answered_at else None
            })
            
        # Get unique classes for class filter
        available_classes = db.session.query(Siswa.kelas).distinct().all()
        available_classes = [cls[0] for cls in available_classes if cls[0]]
        
        return jsonify({
            "success": True,
            "answers": formatted_answers,
            "pagination": {
                "total": total_count,
                "page": page,
                "per_page": per_page,
                "total_pages": total_pages
            },
            "filters": {
                "available_classes": available_classes
            }
        })
            
    except Exception as e:
        print(f"Error getting filtered answers: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"Server error: {str(e)}"
        }), 500
    
@app.route('/api/collection-analytics/level-analysis', methods=['GET'])
@login_required
def get_level_analysis():
    try:
        # Required parameter
        collection_id = request.args.get('collection_id')
        
        if not collection_id:
            return jsonify({"success": False, "message": "Collection ID is required"}), 400
            
        # Check if collection belongs to current teacher
        collection = QuestionCollection.query.filter_by(
            id=collection_id, 
            guru_id=current_user.id
        ).first()
        
        if not collection:
            return jsonify({"success": False, "message": "Collection not found or not owned by you"}), 403
            
        # Get students assigned to this collection
        student_ids = db.session.query(collection_students.c.siswa_id)\
            .filter(collection_students.c.collection_id == collection_id)\
            .all()
        student_ids = [id[0] for id in student_ids]
        
        if not student_ids:
            return jsonify({
                "success": True, 
                "message": "No students assigned to this collection",
                "level_distribution": {},
                "level_transitions": {}
            })
            
        # Get current level distribution
        level_distribution = db.session.query(
            SiswaResult.current_level, 
            db.func.count(SiswaResult.siswa_id)
        ).filter(
            SiswaResult.collection_id == collection_id,
            SiswaResult.siswa_id.in_(student_ids)
        ).group_by(
            SiswaResult.current_level
        ).all()
        
        # Format level distribution
        formatted_distribution = {str(level): count for level, count in level_distribution}
        
        # For levels not present in results, set count to 0
        for level in range(1, 8):  # levels 1-7
            if str(level) not in formatted_distribution:
                formatted_distribution[str(level)] = 0
                
        # Analyze level transitions
        # For this, we need to look at the sequence of answers for each student
        level_transitions = {
            "1-2": 0,
            "2-3": 0,
            "2-4": 0,
            "3-5": 0,
            "3-6": 0,
            "4-6": 0,
            "4-7": 0
        }
        
        # Get collection questions
        collection_question_ids = db.session.query(CollectionQuestion.question_id)\
            .filter(CollectionQuestion.collection_id == collection_id).all()
        collection_question_ids = [id[0] for id in collection_question_ids]
        
        if collection_question_ids:
            # For each student, get answers in chronological order
            for student_id in student_ids:
                answers = db.session.query(SiswaAnswer)\
                    .filter(
                        SiswaAnswer.siswa_id == student_id,
                        SiswaAnswer.question_id.in_(collection_question_ids)
                    )\
                    .order_by(SiswaAnswer.answered_at)\
                    .all()
                    
                if answers:
                    # Track level transitions
                    prev_level = None
                    for answer in answers:
                        curr_level = answer.level
                        
                        if prev_level is not None and prev_level != curr_level:
                            transition_key = f"{prev_level}-{curr_level}"
                            if transition_key in level_transitions:
                                level_transitions[transition_key] += 1
                                
                        prev_level = curr_level
        
        # Additional statistics for MST 5-stage system
        # Count students in basic (L1-L2), intermediate (L3), and advanced (L4-L5) levels
        basic_count = formatted_distribution.get("1", 0) + formatted_distribution.get("2", 0)
        intermediate_count = formatted_distribution.get("3", 0)
        advanced_count = formatted_distribution.get("4", 0) + formatted_distribution.get("5", 0)
        
        # Count completed students (reached final level L5)
        completed_count = formatted_distribution.get("5", 0)
        
        # Calculate percentages
        total_students = len(student_ids)
        basic_percent = (basic_count / total_students * 100) if total_students > 0 else 0
        intermediate_percent = (intermediate_count / total_students * 100) if total_students > 0 else 0
        advanced_percent = (advanced_count / total_students * 100) if total_students > 0 else 0
        completed_percent = (completed_count / total_students * 100) if total_students > 0 else 0
        
        return jsonify({
            "success": True,
            "level_distribution": formatted_distribution,
            "level_transitions": level_transitions,
            "statistics": {
                "total_students": total_students,
                "basic_levels": {
                    "count": basic_count,
                    "percentage": round(basic_percent, 2)
                },
                "intermediate_levels": {
                    "count": intermediate_count,
                    "percentage": round(intermediate_percent, 2)
                },
                "advanced_levels": {
                    "count": advanced_count,
                    "percentage": round(advanced_percent, 2)
                },
                "completed": {
                    "count": completed_count,
                    "percentage": round(completed_percent, 2)
                }
            }
        })
        
    except Exception as e:
        print(f"Error getting level analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"Server error: {str(e)}"
        }), 500
    
@app.route('/api/collection-analytics/class-analysis', methods=['GET'])
@login_required
def get_class_analysis():
    try:
        # Required parameter
        collection_id = request.args.get('collection_id')
        
        if not collection_id:
            return jsonify({"success": False, "message": "Collection ID is required"}), 400
            
        # Check if collection belongs to current teacher
        collection = QuestionCollection.query.filter_by(
            id=collection_id, 
            guru_id=current_user.id
        ).first()
        
        if not collection:
            return jsonify({"success": False, "message": "Collection not found or not owned by you"}), 403
            
        # Get students assigned to this collection
        students = db.session.query(Siswa)\
            .join(
                collection_students,
                collection_students.c.siswa_id == Siswa.id
            )\
            .filter(
                collection_students.c.collection_id == collection_id
            )\
            .all()
            
        if not students:
            return jsonify({
                "success": True, 
                "message": "No students assigned to this collection",
                "class_stats": {}
            })
            
        # Group students by class
        class_groups = {}
        for student in students:
            if student.kelas not in class_groups:
                class_groups[student.kelas] = []
            class_groups[student.kelas].append(student.id)
            
        # For each class, get performance statistics
        class_stats = {}
        for kelas, student_ids in class_groups.items():
            # Get results for these students
            results = db.session.query(SiswaResult)\
                .filter(
                    SiswaResult.collection_id == collection_id,
                    SiswaResult.siswa_id.in_(student_ids)
                )\
                .all()
                
            # Calculate stats
            total_correct = sum(r.correct or 0 for r in results)
            total_incorrect = sum(r.incorrect or 0 for r in results)
            total_answers = total_correct + total_incorrect
            accuracy = (total_correct / total_answers * 100) if total_answers > 0 else 0
            
            # Count students by level
            level_counts = {str(n): 0 for n in range(1, 8)}
            for result in results:
                level = str(result.current_level or 1)
                if level in level_counts:
                    level_counts[level] += 1
                    
            # Count completed students (in levels 5-7)
            completed_count = sum(
                1 for r in results if r.current_level in [5, 6, 7]
            )
            
            class_stats[kelas] = {
                "student_count": len(student_ids),
                "total_correct": total_correct,
                "total_incorrect": total_incorrect,
                "accuracy": round(accuracy, 2),
                "level_distribution": level_counts,
                "completed_count": completed_count,
                "completion_rate": round((completed_count / len(student_ids) * 100), 2) if student_ids else 0
            }
            
        return jsonify({
            "success": True,
            "class_stats": class_stats
        })
        
    except Exception as e:
        print(f"Error getting class analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"Server error: {str(e)}"
        }), 500
    
@app.route('/api/collection-analytics/question-analysis', methods=['GET'])
@login_required
def get_question_analysis():
    try:
        # Required parameter
        collection_id = request.args.get('collection_id')
        
        if not collection_id:
            return jsonify({"success": False, "message": "Collection ID is required"}), 400
            
        # Check if collection belongs to current teacher
        collection = QuestionCollection.query.filter_by(
            id=collection_id, 
            guru_id=current_user.id
        ).first()
        
        if not collection:
            return jsonify({"success": False, "message": "Collection not found or not owned by you"}), 403
            
        # Get questions in this collection
        questions = db.session.query(Question)\
            .join(
                CollectionQuestion,
                CollectionQuestion.question_id == Question.id
            )\
            .filter(
                CollectionQuestion.collection_id == collection_id
            )\
            .all()
            
        if not questions:
            return jsonify({
                "success": True, 
                "message": "No questions in this collection",
                "questions": []
            })
            
        # For each question, get answer statistics
        question_stats = []
        for question in questions:
            # Get all answers for this question
            answers = db.session.query(SiswaAnswer)\
                .filter(
                    SiswaAnswer.question_id == question.id
                )\
                .all()
                
            # Calculate statistics
            total_answers = len(answers)
            correct_answers = sum(1 for a in answers if a.is_correct)
            incorrect_answers = total_answers - correct_answers
            accuracy = (correct_answers / total_answers * 100) if total_answers > 0 else 0
            
            # Get class breakdown
            class_breakdown = {}
            for answer in answers:
                student = Siswa.query.get(answer.siswa_id)
                if student and student.kelas:
                    if student.kelas not in class_breakdown:
                        class_breakdown[student.kelas] = {
                            "correct": 0,
                            "incorrect": 0,
                            "total": 0
                        }
                    class_breakdown[student.kelas]["total"] += 1
                    if answer.is_correct:
                        class_breakdown[student.kelas]["correct"] += 1
                    else:
                        class_breakdown[student.kelas]["incorrect"] += 1
            
            # Calculate accuracy per class
            for kelas, stats in class_breakdown.items():
                stats["accuracy"] = round((stats["correct"] / stats["total"] * 100), 2) if stats["total"] > 0 else 0
                
            question_stats.append({
                "level": question.level,
                "soal": question.soal,
                "options": json.loads(question.options) if question.options else [],
                "jawaban_benar": question.jawaban_benar,
                "bobot": question.bobot,
                "total_answers": total_answers,
                "correct_answers": correct_answers,
                "incorrect_answers": incorrect_answers,
                "accuracy": round(accuracy, 2),
                "class_breakdown": class_breakdown
            })
            
        # Sort questions by level
        question_stats.sort(key=lambda q: q["level"])
        
        # Filter by level if requested
        level_filter = request.args.get('level')
        if level_filter and level_filter != 'all':
            question_stats = [q for q in question_stats if q["level"] == int(level_filter)]
            
        return jsonify({
            "success": True,
            "questions": question_stats
        })
        
    except Exception as e:
        print(f"Error getting question analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"Server error: {str(e)}"
        }), 500
    

@app.route('/api/siswa/refresh_answers/<int:collection_id>', methods=['GET'])
@login_required
def refresh_siswa_answers(collection_id):
    user_id = request.args.get("user_id")
    force = request.args.get("force") == "true"
    
    # Cek keamanan: pastikan user_id sesuai dengan pengguna saat ini
    if str(current_user.id) != str(user_id):
        return jsonify({"success": False, "message": "Akses ditolak"}), 403
    
    try:
        # Ambil hasil siswa untuk koleksi ini
        result = SiswaResult.query.filter_by(
            siswa_id=user_id,
            collection_id=collection_id
        ).first()
        
        if not result:
            return jsonify({
                "success": False,
                "message": "Tidak ada hasil ditemukan"
            }), 404
        
        # Dapatkan daftar question_id yang terkait dengan collection ini
        collection_question_ids = db.session.query(CollectionQuestion.question_id)\
            .filter(CollectionQuestion.collection_id == collection_id).all()
        collection_question_ids = [id[0] for id in collection_question_ids]
        
        # Hitung jawaban yang ada di database
        answer_count = db.session.query(func.count(SiswaAnswer.id)).filter(
            SiswaAnswer.siswa_id == user_id,
            SiswaAnswer.question_id.in_(collection_question_ids) if collection_question_ids else False
        ).scalar() or 0
        
        # Hitung total jawaban benar dari actual records
        correct_count = db.session.query(func.count(SiswaAnswer.id)).filter(
            SiswaAnswer.siswa_id == user_id,
            SiswaAnswer.is_correct == True,
            SiswaAnswer.question_id.in_(collection_question_ids) if collection_question_ids else False
        ).scalar() or 0
        
        # Hitung total jawaban salah dari actual records
        incorrect_count = db.session.query(func.count(SiswaAnswer.id)).filter(
            SiswaAnswer.siswa_id == user_id,
            SiswaAnswer.is_correct == False,
            SiswaAnswer.question_id.in_(collection_question_ids) if collection_question_ids else False
        ).scalar() or 0
        
        # Cek inkonsistensi
        has_inconsistency = (result.correct != correct_count or result.incorrect != incorrect_count)
        
        if has_inconsistency or force:
            # Jika ada inkonsistensi atau force=true, perbarui data
            result.correct = correct_count
            result.incorrect = incorrect_count
            db.session.commit()
            
            return jsonify({
                "success": True,
                "message": "Data jawaban telah diperbarui",
                "changes": {
                    "old_correct": result.correct if has_inconsistency else correct_count,
                    "new_correct": correct_count,
                    "old_incorrect": result.incorrect if has_inconsistency else incorrect_count,
                    "new_incorrect": incorrect_count
                }
            })
        
        return jsonify({
            "success": True,
            "message": "Data jawaban sudah konsisten",
            "stats": {
                "correct": correct_count,
                "incorrect": incorrect_count,
                "total": answer_count
            }
        })
        
    except Exception as e:
        print(f"Error menyegarkan jawaban: {str(e)}")
        db.session.rollback()
        return jsonify({
            "success": False,
            "message": f"Error: {str(e)}"
        }), 500

@app.route('/api/questions/<int:question_id>/set_validation', methods=['POST'])
@login_required
def set_question_validation_status(question_id):
    try:
        # Pastikan pengguna adalah guru
        if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
            return jsonify({'success': False, 'message': 'Akses ditolak. Hanya guru yang dapat mengubah status validasi.'}), 403

        question = Question.query.get(question_id)
        if not question:
            return jsonify({'success': False, 'message': 'Soal tidak ditemukan'}), 404

        # Pastikan guru yang login adalah pemilik soal
        if question.guru_id != current_user.id:
            return jsonify({'success': False, 'message': 'Anda tidak memiliki izin untuk mengubah status validasi soal ini'}), 403

        data = request.get_json()
        # Menggunakan 'is_validated' dari body request, default ke nilai saat ini jika tidak ada
        new_status = data.get('is_validated', not question.is_validated) 

        question.is_validated = bool(new_status)
        db.session.commit()

        return jsonify({'success': True, 'message': f'Status validasi soal ID {question_id} berhasil diperbarui.', 'is_validated': question.is_validated}), 200

    except Exception as e:
        db.session.rollback()
        print(f"Error updating question validation status: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f"Terjadi kesalahan server: {str(e)}"}), 500

    
@app.route('/result', methods=['GET'])
@login_required  # Ensure user is logged in
def results():
    # Get collection_id from query parameters if available
    collection_id = request.args.get("collection_id")
    
    # Use the currently logged-in user instead of query param
    if not current_user.is_authenticated:
        return redirect(url_for('login'))
    
    # For siswa (student) users
    if hasattr(current_user, 'user_type') and current_user.user_type == 'siswa':
        # Get the student's results for the specified collection
        user_result = SiswaResult.query.filter_by(
            siswa_id=current_user.id,
            collection_id=collection_id if collection_id else None
        ).first()
        
        if not user_result:
            # Create a new result record if none exists
            return render_template('siswa/result.html', 
                                correct=0, 
                                incorrect=0, 
                                current_level=1,
                                accuracy=0,
                                answers=[],
                                collection_name="N/A",
                                collection_id=collection_id,
                                avg_time="N/A")
        
        # Calculate accuracy
        total = user_result.correct + user_result.incorrect
        accuracy = round((user_result.correct / total * 100) if total > 0 else 0, 2)
        
        # Get user answer history with questions for this collection
        query = db.session.query(
            SiswaAnswer, Question
        ).join(
            Question, SiswaAnswer.question_id == Question.id
        ).filter(
            SiswaAnswer.siswa_id == current_user.id
        )
        
        # Add collection filter if specified
        if collection_id:
            query = query.filter(Question.collection_id == collection_id)
            
        # Get collection name if available
        collection_name = "N/A"
        if collection_id:
            collection = QuestionCollection.query.get(collection_id)
            if collection:
                collection_name = collection.name
        
        # Execute query and limit to most recent answers
        answers = query.order_by(SiswaAnswer.answered_at.desc()).limit(10).all()
        
        # Format answers for template
        answer_history = []
        for ua, q in answers:
            answer_history.append({
                'question': q.soal,
                'user_answer': ua.siswa_answer,
                'is_correct': ua.is_correct,
                'level': ua.level,
                'explanation': q.explanation
            })
        
        # Calculate average answer time (if needed)
        avg_time = "N/A"  # Implement logic here if you want to track this
        
        return render_template('siswa/result.html', 
                              correct=user_result.correct, 
                              incorrect=user_result.incorrect,
                              current_level=user_result.current_level,
                              accuracy=accuracy,
                              answers=answer_history,
                              collection_name=collection_name,
                              collection_id=collection_id,
                              avg_time=avg_time)
    else:
        # For non-student users or if user_type not defined
        flash("Only students can view their results", "warning")
        return redirect(url_for('index'))
# =========================================
# ENDPOINT RESET DATABASE 
# =========================================
@app.route('/reset_database', methods=['POST'])
def reset_database_endpoint():
    try:
        # BAGIAN YANG DIMODIFIKASI - MULAI
        # Dapatkan daftar ID soal yang ada dalam koleksi (yang harus dipertahankan)
        preserved_question_ids = db.session.query(CollectionQuestion.question_id).distinct().all()
        preserved_question_ids = [id[0] for id in preserved_question_ids]
        
        if preserved_question_ids:
            print(f"Mempertahankan {len(preserved_question_ids)} soal dalam koleksi")
        
        # Pertama, hapus semua jawaban pengguna yang tidak terkait dengan soal dalam koleksi
        print("Menghapus jawaban pengguna...")
        if preserved_question_ids:
            # Hapus jawaban siswa hanya untuk soal yang tidak dalam koleksi
            SiswaAnswer.query.filter(~SiswaAnswer.question_id.in_(preserved_question_ids)).delete(synchronize_session=False)
        else:
            # Jika tidak ada soal yang dipertahankan, hapus semua jawaban
            SiswaAnswer.query.delete()
        db.session.commit()
        
        # Kedua, reset current_level semua user ke level 1
        print("Mengatur ulang level pengguna...")
        SiswaResult.query.update({
            SiswaResult.current_level: 1,
            SiswaResult.correct: 0,
            SiswaResult.incorrect: 0
        })
        db.session.commit()
        
        # Hapus relasi collection_questions untuk soal yang akan dihapus
        # tapi jangan hapus koleksi itu sendiri
        # Relasi untuk soal yang ada dalam koleksi tidak perlu dihapus
        if preserved_question_ids:
            print("Mempertahankan relasi collection_questions untuk soal yang dipertahankan...")
        else:
            print("Menghapus semua relasi collection_questions...")
            CollectionQuestion.query.delete()
            db.session.commit()
        
        # Terakhir, hapus semua soal yang tidak dalam koleksi
        print("Menghapus soal yang tidak dalam koleksi...")
        if preserved_question_ids:
            # Hapus soal yang tidak ada dalam koleksi
            Question.query.filter(~Question.id.in_(preserved_question_ids)).delete(synchronize_session=False)
        else:
            # Jika tidak ada soal yang dipertahankan, hapus semua soal
            Question.query.delete()
        db.session.commit()
        
        message = "Database berhasil direset. Silahkan upload file kembali untuk regenerasi soal."
        if preserved_question_ids:
            message += f" {len(preserved_question_ids)} soal dalam koleksi dipertahankan."
        
        print(message)
        return jsonify({"success": True, "message": message}), 200
        # BAGIAN YANG DIMODIFIKASI - SELESAI
    except Exception as e:
        print(f"Error menghapus data lama: {str(e)}")
        db.session.rollback()
        return jsonify({"success": False, "message": f"Error menghapus data lama: {str(e)}"}), 500

# Add this to the Flask app code
@app.route('/collection-analytics/<int:collection_id>')
@login_required
def collection_analytics(collection_id):
    # Verify collection existence and ownership
    collection = QuestionCollection.query.get_or_404(collection_id)
    
    # Only allow the teacher who owns the collection to view analytics
    if collection.guru_id != current_user.id:
        flash('Anda tidak memiliki akses ke analisis koleksi ini', 'error')
        return redirect(url_for('guru'))
    
    return render_template('guru/collection-analytics.html', collection_id=collection_id)

# Add these endpoints to your Flask app

# Get detailed student answer data for a collection
@app.route('/api/siswa_answers', methods=['GET'])
@login_required
def get_siswa_answers():
    try:
        user_id = request.args.get('user_id')
        collection_id = request.args.get('collection_id')
        
        if not user_id or not collection_id:
            return jsonify({'success': False, 'message': 'Missing required parameters'}), 400
        
        # Get question IDs in this collection
        collection_question_ids = db.session.query(CollectionQuestion.question_id)\
            .filter(CollectionQuestion.collection_id == collection_id).all()
        collection_question_ids = [id[0] for id in collection_question_ids]
        
        # Get answers for these questions
        siswa_answers = db.session.query(SiswaAnswer).join(
            Question, SiswaAnswer.question_id == Question.id
        ).filter(
            SiswaAnswer.siswa_id == user_id,
            Question.id.in_(collection_question_ids)
        ).all()
        
        # Format answers for response
        formatted_answers = []
        for answer in siswa_answers:
            formatted_answers.append({
                'id': answer.id,
                'siswa_id': answer.siswa_id,
                'question_id': answer.question_id,
                'siswa_answer': answer.siswa_answer,
                'is_correct': answer.is_correct,
                'level': answer.level,
                'answered_at': answer.answered_at.isoformat() if answer.answered_at else None
            })
        
        return jsonify({
            'success': True,
            'siswa_answers': formatted_answers
        })
        
    except Exception as e:
        print(f"Error getting student answers: {str(e)}")
        return jsonify({
            'success': False,
            'message': f"Error: {str(e)}"
        }), 500

# Get collection by ID
@app.route('/api/collections/<int:collection_id>', methods=['GET'])
@login_required
def get_collection(collection_id):
    try:
        collection = QuestionCollection.query.get(collection_id)
        if not collection:
            return jsonify({'success': False, 'message': 'Collection not found'}), 404
        
        # Check if user has access to this collection
        if collection.guru_id != current_user.id and current_user.user_type != 'guru':
            return jsonify({'success': False, 'message': 'Access denied'}), 403
        
        return jsonify({
            'success': True,
            'collection': {
                'id': collection.id,
                'name': collection.name,
                'description': collection.description,
                'created_at': collection.created_at.isoformat() if collection.created_at else None
            }
        })
        
    except Exception as e:
        print(f"Error getting collection: {str(e)}")
        return jsonify({
            'success': False,
            'message': f"Error: {str(e)}"
        }), 500

def enforce_mst_distribution(parsed_questions):
    """Return a new list of questions enforced to match the MST distribution table.

    Distribution (by level):
    L1: 10  -> E:5  M:5
    L2: 20  -> E:12 H:8
    L3: 20  -> E:5  M:8  H:7
    L4: 22  -> E:4  M:11 H:7
    L5: 22  -> E:11 M:7  H:4
    Total: 94

    Strategy:
    - Group parsed questions by (level, difficulty)
    - For each required slot, take existing unique questions if available
    - If insufficient, duplicate existing questions from same (level,difficulty) with a deterministic suffix
    - If still impossible (no question at that (level,diff)), borrow from same level different difficulty preferring Medium->Easy->Hard as fallback
    """
    # target map
    target = {
        1: {'total': 10, 'Easy': 5, 'Medium': 5, 'Hard': 0},
        2: {'total': 20, 'Easy': 12, 'Medium': 0, 'Hard': 8},
        3: {'total': 20, 'Easy': 5, 'Medium': 8, 'Hard': 7},
        4: {'total': 22, 'Easy': 4, 'Medium': 11, 'Hard': 7},
        5: {'total': 22, 'Easy': 11, 'Medium': 7, 'Hard': 4}
    }

    # group existing
    groups = {}
    for q in parsed_questions:
        lvl = int(q.get('level', 0)) if q.get('level') is not None else 0
        diff = q.get('difficulty', 'Medium')
        groups.setdefault((lvl, diff), [])
        groups[(lvl, diff)].append(q)

    final = []
    uid = 1
    for lvl in range(1, 6):
        for diff in ['Easy', 'Medium', 'Hard']:
            need = target[lvl].get(diff, 0)
            available = list(groups.get((lvl, diff), []))

            # If not enough available, try borrow from other difficulties in same level
            if len(available) < need:
                # borrow order: Medium -> Easy -> Hard for balance (prefer Medium)
                borrow_order = ['Medium', 'Easy', 'Hard']
                for borrow_diff in borrow_order:
                    if borrow_diff == diff:
                        continue
                    borrow_avail = groups.get((lvl, borrow_diff), [])
                    for b in borrow_avail:
                        if len(available) >= need:
                            break
                        available.append(b)
                    if len(available) >= need:
                        break

            # Now fill using unique items, then duplicate with suffix if necessary
            idx = 0
            while len(available) < need:
                # Try to duplicate an existing item from available if any
                if available:
                    src = available[idx % len(available)]
                    dup = src.copy()
                    dup['soal'] = dup.get('soal', '') + f" [variasi {uid}]"
                    dup['id'] = f"synthetic-{lvl}-{diff}-{uid}"
                    available.append(dup)
                    uid += 1
                else:
                    # As last resort, create a placeholder question
                    placeholder = {
                        'level': lvl,
                        'question_type': 'multiple_choice',
                        'soal': f'Placeholder soal level {lvl} ({diff}) - tambahkan ulang dari materi',
                        'options': ['A', 'B', 'C', 'D'],
                        'jawaban_benar': 'A',
                        'difficulty': diff,
                        'id': f'placeholder-{lvl}-{diff}-{uid}'
                    }
                    available.append(placeholder)
                    uid += 1

            # Take exactly `need` items and append to final
            for i in range(need):
                item = available[i]
                # ensure correctness of fields
                item['level'] = lvl
                item['difficulty'] = diff
                final.append(item)

    return final

 

def enforce_exactly_5_per_level(parsed_questions):
    """Enforce exactly 5 questions per level (1..5), regardless of difficulty.

    Strategy:
    - Group incoming questions by level, retain original difficulty if any
    - Selection preference within a level: Medium -> Easy -> Hard to keep balance
    - If less than 5 available, borrow from any remaining within level
    - If still short, duplicate with small variation suffix
    - As last resort, create placeholder questions

    Returns a list of 25 questions (5 per level).
    """
    parsed = parsed_questions or []
    # Group by level and difficulty for selection preference
    by_level = {lvl: {"Medium": [], "Easy": [], "Hard": [], "other": []} for lvl in range(1, 6)}
    for q in parsed:
        try:
            lvl = int(q.get('level', 0)) if q.get('level') is not None else 0
        except Exception:
            lvl = 0
        if 1 <= lvl <= 5:
            diff = (q.get('difficulty') or 'Medium').capitalize()
            if diff not in ("Easy", "Medium", "Hard"):
                by_level[lvl]["other"].append(q)
            else:
                by_level[lvl][diff].append(q)

    final = []
    uid = 1
    pref_order = ["Medium", "Easy", "Hard"]

    for lvl in range(1, 6):
        selected = []
        # take by preference
        for diff in pref_order:
            for q in by_level[lvl][diff]:
                if len(selected) >= 5:
                    break
                selected.append(q)
            if len(selected) >= 5:
                break
        # then take from 'other' bucket if still needed
        if len(selected) < 5:
            for q in by_level[lvl]["other"]:
                if len(selected) >= 5:
                    break
                selected.append(q)

        # duplicate if still less than 5
        idx = 0
        while len(selected) < 5:
            if selected:
                src = selected[idx % len(selected)]
                dup = src.copy()
                dup['soal'] = (dup.get('soal', '') or '') + f" [variasi {uid}]"
                dup['id'] = f"synthetic-{lvl}-{uid}"
                # keep difficulty if present, default to Medium
                if not dup.get('difficulty'):
                    dup['difficulty'] = 'Medium'
                selected.append(dup)
                uid += 1
                idx += 1
            else:
                # create placeholder
                placeholder = {
                    'level': lvl,
                    'question_type': 'multiple_choice',
                    'soal': f'Soal placeholder level {lvl} - tambahkan materi terkait',
                    'options': ['A', 'B', 'C', 'D'],
                    'jawaban_benar': 'A',
                    'difficulty': 'Medium',
                    'id': f'placeholder-{lvl}-{uid}'
                }
                selected.append(placeholder)
                uid += 1

        # normalize fields and add to final
        for item in selected[:5]:
            item['level'] = lvl
            if not item.get('difficulty'):
                item['difficulty'] = 'Medium'
            final.append(item)

    return final

def enforce_mst_distribution_5_each(parsed_questions):
    """Enforce exactly 5 questions for each allowed (level, difficulty) combo.

    Allowed difficulties per level are defined in DIFFICULTY_LEVELS.
    Target per combo: 5 questions.
    Total target = sum(5 * len(DIFFICULTY_LEVELS[level]) for level in 1..5) = 55.

    Strategy:
    - Group incoming questions by (level, difficulty)
    - For each allowed (level, diff), take up to 5 existing
    - If less than 5, borrow from the same level other difficulties (relabel to target diff)
    - If still less, duplicate with a small variation suffix
    - As last resort, create a placeholder question
    """
    # Group incoming
    groups = {}
    for q in (parsed_questions or []):
        try:
            lvl = int(q.get('level', 0)) if q.get('level') is not None else 0
        except Exception:
            lvl = 0
        diff = (q.get('difficulty') or 'Medium').capitalize()
        if diff not in ("Easy", "Medium", "Hard"):
            diff = 'Medium'
        groups.setdefault((lvl, diff), []).append(q)

    final = []
    uid = 1
    # Borrow order preference when filling gaps
    borrow_order = ['Medium', 'Easy', 'Hard']

    for lvl in range(1, 6):
        allowed_diffs = DIFFICULTY_LEVELS.get(lvl, [])
        for diff in allowed_diffs:
            need = 5
            selected = []
            # Take existing of same (lvl, diff)
            for item in list(groups.get((lvl, diff), [])):
                if len(selected) >= need:
                    break
                selected.append(item)

            # Borrow from same level other diffs (relabel to target diff)
            if len(selected) < need:
                for borrow_diff in borrow_order:
                    if borrow_diff == diff:
                        continue
                    for b in list(groups.get((lvl, borrow_diff), [])):
                        if len(selected) >= need:
                            break
                        clone = b.copy()
                        clone['difficulty'] = diff
                        selected.append(clone)
                    if len(selected) >= need:
                        break

            # Duplicate with small variation if still short
            idx = 0
            while len(selected) < need:
                if selected:
                    src = selected[idx % len(selected)]
                    dup = src.copy()
                    dup['soal'] = (dup.get('soal', '') or '') + f" [variasi {uid}]"
                    dup['id'] = f"synthetic-{lvl}-{diff}-{uid}"
                    dup['difficulty'] = diff
                    selected.append(dup)
                    uid += 1
                    idx += 1
                else:
                    # create placeholder
                    placeholder = {
                        'level': lvl,
                        'question_type': 'multiple_choice',
                        'soal': f'Soal placeholder level {lvl} ({diff}) - tambahkan materi terkait',
                        'options': ['A', 'B', 'C', 'D'],
                        'jawaban_benar': 'A',
                        'difficulty': diff,
                        'id': f'placeholder-{lvl}-{diff}-{uid}'
                    }
                    selected.append(placeholder)
                    uid += 1

            # Normalize and keep exactly 5
            for item in selected[:need]:
                item['level'] = lvl
                item['difficulty'] = diff
                final.append(item)

    return final

def emergency_json_parser(raw_text):
    """
    Emergency parser untuk mencoba mengekstrak soal dari response AI yang tidak valid JSON.
    Mengembalikan list soal yang sudah dipaksa mengikuti distribusi 5 per (level, kesulitan yang diizinkan) total 55.
    """
    try:
        questions = []

        # Cari pattern soal dengan berbagai format
        patterns = [
            r'"level"\s*:\s*(\d+).*?"soal"\s*:\s*"([^"]+)".*?"options"\s*:\s*\[(.*?)\].*?"jawaban_benar"\s*:\s*"([^"]+)' ,
            r'level.*?(\d+).*?soal.*?"([^"]+)".*?options.*?\[(.*?)\].*?jawaban_benar.*?"([^"]+)"',
        ]

        for pattern in patterns:
            matches = re.findall(pattern, raw_text, re.DOTALL | re.IGNORECASE)
            for match in matches:
                try:
                    level = int(match[0])
                    soal = match[1].strip()
                    options_text = match[2]
                    jawaban_benar = match[3].strip()

                    # Parse options
                    options = []
                    option_matches = re.findall(r'"([^"]+)"', options_text)
                    if len(option_matches) >= 4:
                        options = option_matches[:4]

                    if soal and options and jawaban_benar and 1 <= level <= 7:
                        # Default difficulty by level (approx)
                        if level <= 2:
                            default_difficulty = "Easy"
                        elif level <= 4:
                            default_difficulty = "Medium"
                        else:
                            default_difficulty = "Hard"

                        questions.append({
                            "level": level,
                            "question_type": "multiple_choice",
                            "soal": soal,
                            "options": options,
                            "jawaban_benar": jawaban_benar,
                            "difficulty": default_difficulty,
                            "explanation": f"Soal level {level} dari modul ajar"
                        })
                except (ValueError, IndexError):
                    continue

            if questions:
                break

        # Terapkan distribusi 5 per (level, difficulty) setelah emergency parse
        final_questions = enforce_mst_distribution_5_each(questions)
        return final_questions if final_questions else None

    except Exception as e:
        print(f"Emergency parser error: {str(e)}")
        return None

# =========================================
# PROGRESS TRACKING ENDPOINTS
# =========================================
@app.route('/api/progress/<int:user_id>', methods=['GET'])
@login_required
def get_upload_progress(user_id):
    """Endpoint untuk mengecek progress upload - Compatible dengan VPS multiple workers"""
    try:
        # Validasi user exists (gunakan Session.get untuk menghindari warning SQLAlchemy 2.0)
        if hasattr(current_user, 'user_type') and current_user.user_type == 'guru':
            user = db.session.get(Guru, user_id)
        else:
            user = db.session.get(Siswa, user_id)
            
        if not user:
            return jsonify({
                "success": False, 
                "message": "User tidak ditemukan",
                "progress": None
            }), 200  # Return 200 to stop endless polling
            
        # Pastikan user bisa akses progress sendiri atau guru bisa akses semua
        if not (current_user.id == user_id or (hasattr(current_user, 'user_type') and current_user.user_type == 'guru')):
            return jsonify({
                "success": False,
                "message": "Akses ditolak",
                "progress": None
            }), 200  # Return 200 to stop endless polling
        
        progress_data = get_progress(user_id)
        if not progress_data:
            # Return 200 dengan progress null untuk menghentikan polling yang tidak perlu
            return jsonify({
                "success": False,
                "message": "Tidak ada progress aktif",
                "progress": None,
                "status": "no_active_progress"
            }), 200
        
        return jsonify({
            "success": True,
            "progress": progress_data,
            "status": "active"
        })
        
    except Exception as e:
        print(f"Error getting progress for user {user_id}: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Server error",
            "progress": None,
            "status": "error"
        }), 200  # Return 200 to avoid endless error polling

@app.route('/api/progress/stop/<int:user_id>', methods=['POST'])
@login_required
def stop_progress_tracking(user_id):
    """Endpoint untuk menghentikan progress tracking"""
    try:
        # Validasi user permission
        if not (current_user.id == user_id or (hasattr(current_user, 'user_type') and current_user.user_type == 'guru')):
            return jsonify({
                "success": False,
                "message": "Akses ditolak"
            }), 403
        
        clear_progress(user_id)
        
        return jsonify({
            "success": True,
            "message": "Progress tracking dihentikan"
        })
        
    except Exception as e:
        print(f"Error stopping progress for user {user_id}: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Server error"
        }), 500

@app.route('/api/progress/cleanup', methods=['POST'])
@login_required
def cleanup_old_progress():
    """Endpoint untuk membersihkan progress tracking lama (hanya guru)"""
    try:
        if not (hasattr(current_user, 'user_type') and current_user.user_type == 'guru'):
            return jsonify({
                "success": False,
                "message": "Akses ditolak. Hanya guru yang dapat membersihkan progress."
            }), 403
        
        # Hapus progress yang lebih dari 1 jam
        from datetime import timedelta
        cutoff_time = datetime.datetime.now() - timedelta(hours=1)
        
        old_progress = ProgressTracker.query.filter(
            ProgressTracker.updated_at < cutoff_time
        ).all()
        
        count = 0
        for progress in old_progress:
            db.session.delete(progress)
            count += 1
        
        db.session.commit()
        
        return jsonify({
            "success": True,
            "message": f"Berhasil membersihkan {count} progress lama"
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"Error cleaning up old progress: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Server error"
        }), 500

# =========================================
# ENDPOINT UPLOAD & GENERATE (jumlah diatur oleh post-processing)
# =========================================
@app.route('/upload', methods=['POST'])
@login_required
def upload_file():
    if not current_user.is_authenticated or not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
        return jsonify({"message": "Hanya guru yang dapat menggenerate soal"}), 403

    user_id = current_user.id
    print(f"[{datetime.datetime.now()}] Upload file endpoint called by teacher: {current_user.username}")
    
    # Initialize progress tracking - STEP 1 START
    print(f"[PROGRESS] Initializing progress for user {user_id}")
    update_progress(user_id, 1, "pending", "Memulai proses upload...")
    
    if 'file' not in request.files:
        clear_progress(user_id)
        return jsonify({"message": "Tidak ada file yang diunggah"}), 400

    file = request.files['file']
    if file.filename == '':
        clear_progress(user_id)
        return jsonify({"message": "Tidak ada file yang dipilih"}), 400

    file_ext = file.filename.rsplit('.', 1)[-1].lower()
    if file_ext not in ['doc', 'docx', 'pdf']:
        clear_progress(user_id)
        return jsonify({"message": "Format file tidak didukung. Hanya file .doc, .docx, dan .pdf yang diizinkan"}), 400

    try:
        # STEP 1: Upload dan validasi file
        print(f"[PROGRESS] Step 1 - Reading file...")
        update_progress(user_id, 1, "active", "Membaca dan memvalidasi file...")
        file_stream = file.read()
        update_progress(user_id, 1, "completed", "File berhasil diunggah dan divalidasi")
        
        # STEP 2: Analisis struktur dokumen  
        print(f"[PROGRESS] Step 2 - Analyzing document structure...")
        update_progress(user_id, 2, "active", "Menganalisis struktur dan format dokumen...")
        print(f"[{datetime.datetime.now()}] Memulai validasi file...")
        
        # Validasi format dan konten
        is_valid_file, validation_result = validate_file_format_and_content(file_stream, file_ext, file.filename)
        
        if not is_valid_file:
            clear_progress(user_id)
            print(f"Validasi gagal: {validation_result}")
            
            # Buat pesan error yang lebih informatif berdasarkan jenis masalah
            if "minimal 400 karakter" in validation_result:
                error_message = "📄 File terlalu pendek! Modul Ajar Kurikulum Merdeka membutuhkan konten yang lebih lengkap (minimal 400 karakter). Pastikan dokumen berisi komponen-komponen yang diperlukan."
            elif "Header 'MODUL AJAR'" in validation_result:
                error_message = "📋 Header tidak ditemukan! File harus memiliki judul 'MODUL AJAR' di bagian atas dokumen sesuai format Kurikulum Merdeka."
            elif "Tujuan Pembelajaran" in validation_result:
                error_message = "🎯 Komponen Tujuan Pembelajaran tidak ditemukan! Modul Ajar Kurikulum Merdeka wajib memiliki bagian 'Tujuan Pembelajaran' yang jelas."
            elif "Kompetensi Awal" in validation_result:
                error_message = "📚 Komponen Kompetensi Awal tidak ditemukan! Ini adalah bagian penting dalam Kurikulum Merdeka yang menjelaskan kemampuan prasyarat siswa."
            elif "Pemahaman Bermakna" in validation_result:
                error_message = "💡 Komponen Pemahaman Bermakna tidak ditemukan! Bagian ini wajib ada untuk menjelaskan konsep inti yang ingin dipahami siswa."
            elif "Identitas Modul" in validation_result:
                error_message = "📝 Identitas Modul tidak lengkap! Pastikan ada informasi Mata Pelajaran, Kelas/Fase sesuai format Kurikulum Merdeka."
            elif "Bukan Modul Ajar Kurikulum Merdeka" in validation_result:
                error_message = "⚠️ File bukan Modul Ajar Kurikulum Merdeka yang valid! Sistem hanya menerima dokumen dengan format dan komponen sesuai standar Kurikulum Merdeka."
            else:
                error_message = f"[DITOLAK] File tidak memenuhi standar Modul Ajar Kurikulum Merdeka: {validation_result}"
                
            return jsonify({
                "success": False,
                "message": error_message,
                "error_type": "invalid_module_format",
                "suggestion": "💡 Pastikan file Anda adalah Modul Ajar Kurikulum Merdeka yang lengkap dengan komponen: Header MODUL AJAR, Identitas Modul, Tujuan Pembelajaran, Kompetensi Awal, dan Pemahaman Bermakna."
            }), 400
        
        content_text = validation_result
        print(f"[{datetime.datetime.now()}] File berhasil divalidasi. Panjang konten: {len(content_text)} karakter.")
        update_progress(user_id, 2, "completed", "Struktur dokumen berhasil dianalisis")
        
        # STEP 3: Ekstraksi tujuan pembelajaran
        print(f"[PROGRESS] Step 3 - Extracting learning objectives...")
        update_progress(user_id, 3, "active", "Mengekstrak tujuan pembelajaran dan komponen modul...")
        
        # VALIDASI PENDIDIKAN: Sudah dilakukan dalam validate_file_format_and_content
        # Tidak perlu validasi ganda lagi, langsung lanjut ke ekstraksi
        
        # EKSTRAKSI KOMPONEN: Gunakan ekstraksi hybrid dengan fallback
        print("Mengekstrak komponen modul ajar...")
        module_components, validation_result = extract_hybrid_module_components(content_text)

        # VALIDASI KUALITAS: Periksa kualitas ekstraksi
        quality_score, quality_issues = check_content_quality_score(module_components)
        
        print(f"Quality score: {quality_score}/100")
        print(f"Issues: {quality_issues}")
        
        # THRESHOLD: Minimal 60 poin untuk diproses (dinaikkan untuk memastikan kualitas)
        if quality_score < 60:
            clear_progress(user_id)
            
            # Buat pesan error yang lebih spesifik berdasarkan masalah yang ditemukan
            main_issues = []
            if "Tujuan pembelajaran tidak ditemukan" in quality_issues:
                main_issues.append("🎯 Tujuan Pembelajaran tidak teridentifikasi dengan jelas")
            if "Mata pelajaran tidak teridentifikasi" in quality_issues:
                main_issues.append("📚 Mata Pelajaran tidak terdeteksi")
            if "Kompetensi awal tidak teridentifikasi" in quality_issues:
                main_issues.append("📝 Kompetensi Awal tidak ditemukan")
            if "Pemahaman bermakna tidak ditemukan" in quality_issues:
                main_issues.append("💡 Pemahaman Bermakna tidak teridentifikasi")
                
            if main_issues:
                error_detail = "; ".join(main_issues)
                error_message = f"📋 Komponen Modul Ajar tidak lengkap (skor kualitas: {quality_score}/100)!\n\nMasalah yang ditemukan:\n{error_detail}\n\n💡 Pastikan dokumen Anda memiliki struktur Kurikulum Merdeka yang lengkap."
            else:
                error_message = f"📄 Kualitas ekstraksi komponen rendah (skor: {quality_score}/100). Struktur dokumen tidak sesuai format Kurikulum Merdeka atau teks tidak dapat dibaca dengan baik."
                
            return jsonify({
                "success": False,
                "message": error_message,
                "error_type": "low_quality_extraction",
                "quality_score": quality_score,
                "issues": quality_issues,
                "suggestion": "💡 Periksa format dokumen Anda. Pastikan menggunakan struktur Kurikulum Merdeka dengan komponen: I. Tujuan Pembelajaran, II. Kompetensi Awal, III. Pemahaman Bermakna"
            }), 400

        # Tambahkan logging dan validasi detail
        detailed_validation = log_extraction_results(module_components)

        # VALIDASI FINAL: Periksa komponen kunci (lebih fleksibel)
        tujuan_list = module_components.get("tujuan_pembelajaran", [])
        
        # Jika tidak ada tujuan sama sekali, coba ekstrak ulang atau buat default
        if not tujuan_list:
            print("WARNING: Tujuan pembelajaran tidak ditemukan, membuat tujuan default...")
            # Buat tujuan pembelajaran default berdasarkan mata pelajaran
            mata_pelajaran = module_components.get("mata_pelajaran", "")
            if mata_pelajaran:
                default_tujuan = f"Memahami konsep dasar {mata_pelajaran}"
                module_components["tujuan_pembelajaran"] = [default_tujuan]
                print(f"Tujuan default dibuat: {default_tujuan}")
            else:
                module_components["tujuan_pembelajaran"] = ["Memahami materi pembelajaran yang diberikan"]
                print("Tujuan sangat umum dibuat karena mata pelajaran tidak terdeteksi")
        
        # Set default mata pelajaran jika kosong
        if not module_components.get("mata_pelajaran"):
            module_components["mata_pelajaran"] = "Pembelajaran Umum"
            print("WARNING: Mata pelajaran tidak terdeteksi, menggunakan default 'Pembelajaran Umum'")

        print(f"[{datetime.datetime.now()}] Validasi berhasil. Melanjutkan ke pemrosesan AI...")
        update_progress(user_id, 3, "completed", "Tujuan pembelajaran berhasil diekstrak")
        
        # STEP 4: Proses dengan AI Gemini
        print(f"[PROGRESS] Step 4 - Processing with AI Gemini...")
        update_progress(user_id, 4, "active", "Memproses dengan AI Gemini untuk membuat soal...")

        # Buat DataFrame ringkas (mengganti convert_text_to_dataframe_improved)
        df = pd.DataFrame({
            'Komponen': ['Mata Pelajaran', 'Topik', 'Kelas', 'Jumlah Tujuan', 'Pemahaman Bermakna'],
            'Detail': [
                module_components.get("mata_pelajaran", "")[:30],
                module_components.get("topik_utama", "")[:30], 
                module_components.get("kelas", ""),
                str(len(module_components.get("tujuan_pembelajaran") or [])),
                str(len(module_components.get("pemahaman_bermakna") or []))
            ]
        })

        # PERUBAHAN: Gunakan prompt yang dioptimalkan tapi tetap struktur lama
        prompt = create_optimized_prompt_with_good_structure(module_components)
        if not isinstance(prompt, str) or not prompt.strip():
            clear_progress(user_id)
            return jsonify({
                "success": False,
                "message": "Gagal menyusun prompt dari dokumen. Pastikan komponen utama (Tujuan Pembelajaran, Kompetensi Awal, Pemahaman Bermakna) terdeteksi.",
                "error_type": "prompt_generation_failed"
            }), 500

        print(f"[{datetime.datetime.now()}] Mengirim prompt teroptimalkan ({len(prompt)} chars) ke Gemini AI...")
        print(f"Efisiensi: Ukuran prompt ~{len(prompt)} karakter (vs ~8000+ sebelumnya) - penghematan signifikan!")
        response = model.generate_content(prompt)
        # Pastikan raw_text selalu string agar aman diproses di langkah berikutnya
        raw_text = extract_response_text(response) or ""
        print(f"[{datetime.datetime.now()}] Respons dari Gemini AI diterima.")
        update_progress(user_id, 4, "completed", "AI berhasil menghasilkan soal")
        
        # STEP 5: Menyimpan ke database
        print(f"[PROGRESS] Step 5 - Saving to database...")
        update_progress(user_id, 5, "active", "Memproses dan menyimpan soal ke database...")
        
        print("Raw AI response:", (raw_text or "")[:500])
        
        # Enhanced JSON parsing with multiple strategies
        print(f"Response length: {len(raw_text or '')} chars")
        print(f"Response starts with: {(raw_text or '')[:100]}")
        print(f"Response ends with: {(raw_text or '')[-100:]}")
        
        # Check for JSON indicators
        json_indicators = ['[', '{', '"level"', '"soal"', '"options"']
        found_indicators = [indicator for indicator in json_indicators if indicator in (raw_text or '')]
        print(f"JSON indicators found: {found_indicators}")
        
        gen_questions = []
        try:
            # Method 1: Direct JSON parsing
            gen_questions = json.loads(raw_text)
            print("[DITERIMA] Berhasil parsing JSON langsung")
        except json.JSONDecodeError:
            print("[DITOLAK] Direct JSON parsing gagal, mencoba pembersihan...")
            
            try:
                # Method 2: Clean markdown and common issues
                cleaned_text = (raw_text or '').strip()
                # Remove markdown code blocks
                cleaned_text = re.sub(r'```(?:json)?\s*\n?', '', cleaned_text)
                cleaned_text = re.sub(r'```\s*$', '', cleaned_text)
                # Remove trailing commas
                cleaned_text = re.sub(r',(\s*[}\]])', r'\1', cleaned_text)
                # Remove comments
                cleaned_text = re.sub(r'//.*?\n', '\n', cleaned_text)
                
                gen_questions = json.loads(cleaned_text)
                print("[DITERIMA] Berhasil parsing JSON setelah pembersihan")
            except json.JSONDecodeError:
                print("[DITOLAK] Pembersihan gagal, mencoba ekstraksi regex...")
                
                try:
                    # Method 3: Extract JSON array using regex
                    json_pattern = r'\[\s*\{.*?\}\s*\]'
                    match = re.search(json_pattern, (raw_text or ''), re.DOTALL)
                    if match:
                        json_text = match.group()
                        # Clean the extracted JSON
                        json_text = re.sub(r',(\s*[}\]])', r'\1', json_text)
                        gen_questions = json.loads(json_text)
                        print("[DITERIMA] Berhasil parsing JSON dari ekstraksi regex")
                    else:
                        raise json.JSONDecodeError("No JSON array found", (raw_text or ''), 0)
                except json.JSONDecodeError:
                    print("[DITOLAK] Regex ekstraksi gagal, mencoba bracket matching...")
                    
                    try:
                        # Method 4: Find JSON by bracket matching
                        _rt = (raw_text or '')
                        start_bracket = _rt.find('[')
                        if start_bracket == -1:
                            raise json.JSONDecodeError("No opening bracket found", _rt, 0)
                        
                        # Count brackets to find matching closing bracket
                        bracket_count = 0
                        end_bracket = -1
                        
                        for i in range(start_bracket, len(_rt)):
                            if _rt[i] == '[':
                                bracket_count += 1
                            elif _rt[i] == ']':
                                bracket_count -= 1
                                if bracket_count == 0:
                                    end_bracket = i
                                    break
                        
                        if end_bracket == -1:
                            raise json.JSONDecodeError("No matching closing bracket found", _rt, 0)
                        
                        json_text = _rt[start_bracket:end_bracket+1]
                        # Clean the extracted JSON
                        json_text = re.sub(r',(\s*[}\]])', r'\1', json_text)
                        json_text = re.sub(r'//.*?\n', '\n', json_text)
                        
                        gen_questions = json.loads(json_text)
                        print("[DITERIMA] Berhasil parsing JSON dari bracket matching")
                    except json.JSONDecodeError as final_error:
                        print(f"[DITOLAK] Semua metode parsing gagal. Error terakhir: {str(final_error)}")
                        
                        # ENHANCED: Log raw response untuk debugging
                        print("=== RAW AI RESPONSE (first 1000 chars) ===")
                        print((raw_text or '')[:1000])
                        print("=== END RAW RESPONSE ===")
                        
                        # Method 5: Emergency fallback - manual parsing
                        try:
                            print("🚨 Mencoba emergency parsing...")
                            gen_questions = emergency_json_parser((raw_text or ''))
                            if gen_questions:
                                print(f"[DITERIMA] Emergency parsing berhasil: {len(gen_questions)} soal")
                            else:
                                raise Exception("Emergency parsing juga gagal")
                        except Exception as emergency_error:
                            print(f"[DITOLAK] Emergency parsing gagal: {str(emergency_error)}")
                            return jsonify({
                                "message": f"Gagal memproses response AI. Raw response length: {len(raw_text or '')}. Silakan coba lagi atau hubungi administrator."
                            }), 500

        if not isinstance(gen_questions, list):
            print(f"Diharapkan list, tapi dapat {type(gen_questions)}")
            return jsonify({"message": "AI tidak mengembalikan array soal yang valid"}), 500
            
        print(f"Berhasil mengekstrak {len(gen_questions)} soal")

        # Align modul_reference to extracted objectives/outcomes to enforce grounding
        try:
            before_align = len(gen_questions)
            gen_questions = enforce_alignment_with_objectives(gen_questions, module_components)
            print(f"Alignment modul_reference diterapkan pada {before_align} soal")
        except Exception as e:
            print(f"WARNING: Gagal menerapkan alignment modul_reference: {e}")
        
        # Enforce 5 questions per allowed (level, difficulty) combo (target total: 55)
        try:
            before_cnt = len(gen_questions)
            gen_questions = enforce_mst_distribution_5_each(gen_questions)
            after_cnt = len(gen_questions)
            print(f"Distribusi 5 per (level, difficulty) diterapkan: {before_cnt} -> {after_cnt} soal (target 55)")
        except Exception as e:
            print(f"WARNING: Gagal menerapkan distribusi 5-per-combo: {e}")
        
        if len(gen_questions) < 7:
            print(f"Jumlah soal tidak mencukupi: {len(gen_questions)}")
            return jsonify({"message": f"AI hanya menghasilkan {len(gen_questions)} soal, minimal diperlukan 7 soal"}), 500

        # Delete old questions not in collections
        try:
            preserved_question_ids_in_collections = db.session.query(CollectionQuestion.question_id).join(
                Question, Question.id == CollectionQuestion.question_id
            ).filter(
                Question.guru_id == current_user.id
            ).distinct().all()
            preserved_question_ids_in_collections = [id[0] for id in preserved_question_ids_in_collections]
            
            # Delete student answers for questions to be deleted
            siswa_answers_to_delete_q_ids = db.session.query(Question.id).filter(
                Question.guru_id == current_user.id,
                ~Question.id.in_(preserved_question_ids_in_collections)
            ).subquery()

            SiswaAnswer.query.filter(
                SiswaAnswer.question_id.in_(siswa_answers_to_delete_q_ids)
            ).delete(synchronize_session=False)
            db.session.commit()

            # Delete questions not in collections
            Question.query.filter(
                Question.guru_id == current_user.id,
                ~Question.id.in_(preserved_question_ids_in_collections)
            ).delete(synchronize_session=False)
            db.session.commit()
            
        except Exception as e:
            db.session.rollback()
            print(f"Error menghapus data lama: {str(e)}")
            import traceback
            traceback.print_exc()
            return jsonify({"message": f"Error menghapus data lama: {str(e)}"}), 500
            
        # STEP 5: Save to database
        update_progress(user_id, 5, "active", "Menyimpan soal ke database...")
        
        # Save new questions to database
        print("Menyimpan soal baru ke database untuk guru:", current_user.username)
        new_questions_from_db = [] 
        for q in gen_questions:
            try:
                # MST 5-Stage: level should be 1-5 (not 1-7)
                level_num = int(q.get("level"))
                if not (1 <= level_num <= 5):
                    print(f"⚠️  WARNING: Soal dengan level {level_num} dilewati (harus 1-5)")
                    continue
                
                # Handle both old format (p value) and new format (difficulty)
                if 'difficulty' in q:
                    difficulty = q.get('difficulty', 'Medium')
                    p_val = difficulty_to_p_value(difficulty)
                else:
                    # Backward compatibility: convert old p value to difficulty
                    p_val = float(q.get("p", 0.75))
                    difficulty = p_value_to_difficulty(p_val)
                
                bobot = assign_weight(p_val)
                
                options_str = None
                if 'options' in q and isinstance(q['options'], list):
                    options_str = json.dumps(q['options'])
                
                question_type = q.get('question_type', 'multiple_choice')
                
                # MST 5-Stage: level and technology_level are the same (1-5)
                technology_level = level_num  # Direct mapping for MST 5-stage
                
                new_question = Question(
                    guru_id=current_user.id,
                    level=level_num,                        # Keep for backward compatibility
                    technology_level=technology_level,      # PRIMARY field for MST 5-stage (L1-L5)
                    soal=q.get("soal"),
                    jawaban_benar=q.get("jawaban_benar", ""),
                    options=options_str,
                    question_type=question_type,
                    p=p_val,
                    bobot=bobot,
                    explanation=q.get("explanation", ""),
                    difficulty=difficulty
                )
                db.session.add(new_question)
                new_questions_from_db.append(new_question)
            except Exception as e:
                db.session.rollback()
                print(f"Error memproses soal: {str(e)}")
                import traceback
                traceback.print_exc()
                return jsonify({"message": f"Error memproses soal: {str(e)}"}), 500
            
        db.session.commit()
        print("Database commit berhasil: soal baru tersimpan.")

        # Prepare data for frontend
        questions_for_frontend = []
        for q_db in new_questions_from_db:
            options = []
            if q_db.options:
                try:
                    options = json.loads(q_db.options)
                except json.JSONDecodeError:
                    pass
            questions_for_frontend.append({
                "id": q_db.id,
                "guru_id": q_db.guru_id,
                "collection_id": q_db.collection_id,
                "level": q_db.level,
                "technology_level": q_db.technology_level,  # Include technology_level in response
                "soal": q_db.soal,
                "jawaban_benar": q_db.jawaban_benar,
                "options": options,
                "question_type": q_db.question_type,
                "p": float(q_db.p),
                "difficulty": q_db.difficulty,
                "bobot": q_db.bobot,
                "explanation": q_db.explanation,
                "created_at": q_db.created_at.isoformat() if q_db.created_at else None,
                "version": q_db.version,
                "is_current": q_db.is_current,
                "is_validated": q_db.is_validated
            })

        # Count questions by level
        level_counts = {}
        for q in questions_for_frontend:
            level = int(q.get("level"))
            level_counts[level] = level_counts.get(level, 0) + 1
                
        level_summary = ", ".join([f"level {level}: {count} soal" for level, count in sorted(level_counts.items())])

        # Complete step 5 and clear progress
        update_progress(user_id, 5, "completed", "Semua soal berhasil disimpan!")
        
        # Create success message without quality score
        message = f"[DITERIMA] Modul ajar berhasil divalidasi dan diproses! Total {len(questions_for_frontend)} soal pilihan ganda berhasil digenerate ({level_summary})"

        # Clear progress after successful completion
        clear_progress(user_id)

        return jsonify({
            "success": True,
            "message": message,  
            "data": questions_for_frontend,
            "validation_info": {
                "quality_score": quality_score,
                "components_found": len([k for k, v in module_components.items() if v]),
                "total_objectives": len(module_components.get("tujuan_pembelajaran", []))
            }
        }), 200

    except Exception as e:
        clear_progress(user_id)
        print(f"Exception in upload_file: {str(e)}")
        db.session.rollback()
        import traceback
        traceback.print_exc()
        
        # Buat pesan error yang lebih informatif berdasarkan jenis error
        error_str = str(e).lower()
        
        if "progressinterval is not defined" in error_str:
            error_message = "🔄 Terjadi kesalahan dalam sistem pelacakan progress. Silakan refresh halaman dan coba lagi."
        elif "timeout" in error_str or "time" in error_str:
            error_message = "⏱️ Pemrosesan memakan waktu terlalu lama. File terlalu besar atau kompleks. Silakan coba dengan file yang lebih sederhana."
        elif "json" in error_str or "decode" in error_str:
            error_message = "📄 Format respons tidak valid. Sistem AI mengalami gangguan sementara. Silakan coba lagi dalam beberapa menit."
        elif "connection" in error_str or "network" in error_str:
            error_message = "🌐 Gangguan koneksi ke sistem AI. Periksa koneksi internet Anda dan coba lagi."
        elif "memory" in error_str or "overflow" in error_str:
            error_message = "💾 File terlalu besar untuk diproses. Silakan gunakan file dengan ukuran lebih kecil (maksimal 15MB)."
        elif "permission" in error_str or "access" in error_str:
            error_message = "🔐 Tidak memiliki akses untuk memproses file. Pastikan Anda login sebagai guru dan file tidak terproteksi."
        elif "database" in error_str or "sql" in error_str:
            error_message = "💾 Gangguan pada database sistem. Tim teknis sedang memperbaiki. Silakan coba lagi dalam beberapa menit."
        elif "extract" in error_str or "parsing" in error_str:
            error_message = "📋 Gagal menganalisis isi dokumen. Pastikan file tidak rusak dan format sesuai (.pdf, .doc, .docx)."
        else:
            error_message = f"[DITOLAK] Terjadi kesalahan sistem yang tidak terduga. Silakan coba lagi atau hubungi administrator jika masalah berlanjut.\n\nDetail teknis: {str(e)}"
            
        return jsonify({
            "success": False, 
            "message": error_message,
            "error_type": "system_error",
            "suggestion": "💡 Coba refresh halaman dan upload ulang file Anda. Pastikan file adalah Modul Ajar Kurikulum Merdeka yang valid dan berukuran tidak terlalu besar."
        }), 500

# Function to validate if a document contains educational content
def check_educational_content(content_text):
    """
    Checks if the document content appears to be educational material.
    
    Returns:
        bool: True if the document is likely educational content, False otherwise
    """
    if not content_text or len(content_text.strip()) < 100:
        print("Document content too short for validation")
        return False
    
    # Convert to lowercase for case-insensitive matching
    content = content_text.lower()
    
    # Define educational content patterns and keywords
    edu_patterns = [
        # Headers or sections typically found in educational materials
        r'\b(modul|bab|materi|pembelajaran|pendidikan|pengajaran|pelatihan|belajar)\b',
        r'\b(tujuan|capaian|sasaran|kompetensi|kemampuan)\s+(pembelajaran|belajar|pendidikan)\b',
        r'\b(indikator|pencapaian|keberhasilan|penilaian)\s+(kompetensi|pembelajaran|belajar)\b',
        r'\b(peserta\s+didik|siswa|murid|pelajar|mahasiswa)\b',
        
        # Structural elements commonly found in educational materials
        r'\b(pendahuluan|isi|penutup|daftar\s+pustaka|referensi|latihan|evaluasi|penilaian)\b',
        r'\b(kelas|mata\s+pelajaran|mapel|bidang\s+studi|pelajaran|kurikulum|silabus)\b',
        
        # Educational activities
        r'\b(diskusi|latihan|tugas|proyek|aktivitas|kegiatan|praktikum|eksperimen)\b',
        
        # Educational roles
        r'\b(guru|pendidik|pengajar|fasilitator|instruktur|dosen|tutor)\b'
    ]
    
    # Educational keyword categories with weights
    edu_keyword_categories = {
        'structural': {
            'weight': 2, 
            'keywords': ['modul', 'bab', 'materi', 'pelajaran', 'kurikulum', 'silabus', 'pendahuluan', 
                        'isi', 'penutup', 'daftar pustaka', 'referensi']
        },
        'objectives': {
            'weight': 3, 
            'keywords': ['tujuan pembelajaran', 'capaian pembelajaran', 'kompetensi dasar', 
                        'indikator pencapaian', 'sasaran pembelajaran']
        },
        'learning': {
            'weight': 1, 
            'keywords': ['belajar', 'pembelajaran', 'pendidikan', 'pengajaran', 'pelatihan', 
                        'memahami', 'menguasai', 'menganalisis']
        },
        'audience': {
            'weight': 1, 
            'keywords': ['peserta didik', 'siswa', 'murid', 'pelajar', 'mahasiswa']
        },
        'activities': {
            'weight': 1, 
            'keywords': ['diskusi', 'latihan', 'tugas', 'proyek', 'aktivitas', 'kegiatan', 
                        'praktikum', 'eksperimen', 'evaluasi']
        },
        'roles': {
            'weight': 1, 
            'keywords': ['guru', 'pendidik', 'pengajar', 'fasilitator', 'instruktur', 'dosen', 'tutor']
        }
    }
    
    # Check pattern matches
    pattern_matches = 0
    for pattern in edu_patterns:
        if re.search(pattern, content):
            pattern_matches += 1
    
    # Check keyword matches with weights
    weighted_score = 0
    keywords_found = []
    
    for category, data in edu_keyword_categories.items():
        category_score = 0
        for keyword in data['keywords']:
            if keyword in content:
                category_score += 1
                keywords_found.append(keyword)
        
        # Apply category weight to the score
        weighted_score += (category_score * data['weight'])
    
    # Log the results
    print(f"Pattern matches: {pattern_matches}/8")
    print(f"Weighted keyword score: {weighted_score}")
    print(f"Educational keywords found: {', '.join(keywords_found[:10])}{' and more' if len(keywords_found) > 10 else ''}")
    
    # Validation criteria - document passes if it:
    # 1. Has at least 3 pattern matches, OR
    # 2. Has a weighted score of at least 5
    is_educational = (pattern_matches >= 3) or (weighted_score >= 5)
    
    return is_educational

# New function to extract educational components from document content
def extract_educational_components(content_text, return_dict=False):
    """
    Extracts educational components from content text with improved flexibility.
    
    Args:
        content_text: The text content to extract educational components from
        return_dict: If True, returns a dictionary; if False, returns a tuple of components
        
    Returns:
        Either a 5-tuple of (module_elements, initial_competencies, learning_objectives, 
        pemahaman_bermakna, target_peserta_didik) or a dictionary of components
    """
    # Initialize results dictionary with "Tidak tersedia" as default values
    results = {
        "Modul/Elemen Ajar": "Tidak tersedia",
        "Kompetensi Awal": "Tidak tersedia",
        "Tujuan Pembelajaran": "Tidak tersedia",
        "Pemahaman Bermakna": "Tidak tersedia",
        "Target Peserta Didik": "Tidak tersedia"
    }
    
    # Normalize text for easier processing
    content = content_text.lower()
    original_content = content_text  # Keep original case for extraction
    
    # Use entire document as Module/Elemen Ajar if we can't find specific sections
    # Extract the first 500 characters as a fallback
    if len(original_content) > 0:
        document_preview = original_content[:min(500, len(original_content))].strip()
        results["Modul/Elemen Ajar"] = document_preview
        
    # Auto-generate learning objectives if we can't find them
    # Look for keywords that might indicate learning content
    learning_keywords = ["belajar", "memahami", "menguasai", "pembelajaran", "algoritma", 
                        "program", "komputer", "pemrograman", "coding", "kode", "struktur",
                        "fungsi", "variabel", "class", "objek", "database", "data", "web"]
                        
    # Count learning-related keywords
    keyword_count = 0
    for keyword in learning_keywords:
        if keyword in content:
            keyword_count += 1
            
    # If document has learning content, generate a generic learning objective
    if keyword_count >= 3:
        results["Tujuan Pembelajaran"] = "Memahami konsep dasar dan mampu mengaplikasikan materi yang disajikan dalam modul pembelajaran ini."
    
    # ===== Module/Element Section =====
    module_patterns = [
        r'(?:modul|elemen|materi)(?:\s+ajar|\s+pembelajaran|\s+pokok)?(.*?)(?:kompetensi|tujuan|capaian|daftar pustaka|referensi)',
        r'(?:bab|materi|konten|isi)(?:\s+pembelajaran|\s+utama)?(.*?)(?:kompetensi|tujuan|capaian|daftar pustaka|referensi)',
        r'(?:pokok|inti)(?:\s+bahasan|\s+materi|\s+pembelajaran)?(.*?)(?:kompetensi|tujuan|capaian|daftar pustaka|referensi)',
        r'(?:i\.|1\.)(?:\s+identitas\s+modul)(.*?)(?:ii\.|2\.)'  # Captures sections starting with Roman numerals
    ]
    
    for pattern in module_patterns:
        match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
        if match:
            module_text = match.group(1).strip()
            if module_text:
                results["Modul/Elemen Ajar"] = clean_extracted_text(module_text)
                break
    
    # ===== Initial Competencies Section =====
    competency_patterns = [
        r'(?:kompetensi|kemampuan)(?:\s+awal|\s+dasar|\s+prasyarat)(?:\s+[\w\s]+)?[:\r\n](.*?)(?:tujuan|capaian|indikator|materi|bab|daftar pustaka|referensi)',
        r'(?:prerequisite|prasyarat)(?:\s+[\w\s]+)?[:\r\n](.*?)(?:tujuan|capaian|indikator|materi|bab|daftar pustaka|referensi)',
        r'(?:pengetahuan|keterampilan)(?:\s+awal|\s+dasar)(?:\s+[\w\s]+)?[:\r\n](.*?)(?:tujuan|capaian|indikator|materi|bab|daftar pustaka|referensi)',
        r'(?:ii\.|2\.)(?:\s+kompetensi\s+awal)(.*?)(?:iii\.|3\.)'  # Captures sections with Roman numerals
    ]
    
    for pattern in competency_patterns:
        match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
        if match:
            comp_text = match.group(1).strip()
            if comp_text:
                results["Kompetensi Awal"] = clean_extracted_text(comp_text)
                break
    
    # ===== Learning Objectives Section =====
    objective_patterns = [
        r'(?:tujuan|capaian)(?:\s+pembelajaran|\s+belajar|\s+pendidikan|\s+instruksional|\s+pelatihan)(?:\s+[\w\s]+)?[:\r\n](.*?)(?:kompetensi|materi|bab|daftar pustaka|referensi|metodologi|pemahaman|target)',
        r'(?:learning|instructional)(?:\s+objectives|\s+goals|\s+outcomes)(?:\s+[\w\s]+)?[:\r\n](.*?)(?:kompetensi|materi|bab|daftar pustaka|referensi|metodologi|pemahaman|target)',
        r'(?:indikator|pencapaian)(?:\s+keberhasilan|\s+pembelajaran|\s+belajar)(?:\s+[\w\s]+)?[:\r\n](.*?)(?:kompetensi|materi|bab|daftar pustaka|referensi|metodologi|pemahaman|target)',
        r'(?:iii\.|3\.)(?:\s+tujuan\s+pembelajaran)(.*?)(?:iv\.|4\.)'  # Captures sections with Roman numerals
    ]
    
    for pattern in objective_patterns:
        match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
        if match:
            obj_text = match.group(1).strip()
            if obj_text:
                results["Tujuan Pembelajaran"] = clean_extracted_text(obj_text)
                break
    
    # ===== NEW: Meaningful Understanding Section =====
    understanding_patterns = [
        r'(?:pemahaman\s+bermakna|bermakna\s+pemahaman|pemahaman\s+yang\s+bermakna)[:\r\n](.*?)(?:kompetensi|tujuan|capaian|materi|bab|daftar pustaka|referensi|target)',
        r'(?:meaningful\s+understanding|understanding\s+meaningful)[:\r\n](.*?)(?:kompetensi|tujuan|capaian|materi|bab|daftar pustaka|referensi|target)',
        r'(?:big\s+ideas|gagasan\s+utama|ide\s+pokok)[:\r\n](.*?)(?:kompetensi|tujuan|capaian|materi|bab|daftar pustaka|referensi|target)',
        r'(?:iv\.|4\.)(?:\s+pemahaman\s+bermakna)(.*?)(?:v\.|5\.)'  # Captures sections with Roman numerals
    ]
    
    for pattern in understanding_patterns:
        match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
        if match:
            understanding_text = match.group(1).strip()
            if understanding_text:
                results["Pemahaman Bermakna"] = clean_extracted_text(understanding_text)
                break
    
    # ===== NEW: Target Students Section =====
    target_patterns = [
        r'(?:target\s+peserta\s+didik|sasaran\s+peserta|peserta\s+didik\s+target)[:\r\n](.*?)(?:kompetensi|tujuan|capaian|materi|bab|daftar pustaka|referensi)',
        r'(?:target\s+audience|target\s+learners|intended\s+learners)[:\r\n](.*?)(?:kompetensi|tujuan|capaian|materi|bab|daftar pustaka|referensi)',
        r'(?:siswa\s+yang\s+dituju|karakteristik\s+peserta\s+didik)[:\r\n](.*?)(?:kompetensi|tujuan|capaian|materi|bab|daftar pustaka|referensi)',
        r'(?:v\.|5\.)(?:\s+target\s+peserta\s+didik)(.*?)(?:vi\.|6\.)'  # Captures sections with Roman numerals
    ]
    
    for pattern in target_patterns:
        match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
        if match:
            target_text = match.group(1).strip()
            if target_text:
                results["Target Peserta Didik"] = clean_extracted_text(target_text)
                break
    
    # ===== Keyword-based approach for sections that weren't found =====
    lines = original_content.split('\n')
    
    # Dictionary to map section keys to their search keywords
    section_keywords = {
        "Modul/Elemen Ajar": ['modul', 'elemen', 'materi', 'bab', 'pokok bahasan', 'identitas modul'],
        "Kompetensi Awal": ['kompetensi awal', 'prasyarat', 'kemampuan dasar', 'kompetensi dasar'],
        "Tujuan Pembelajaran": ['tujuan pembelajaran', 'capaian pembelajaran', 'learning objectives'],
        "Pemahaman Bermakna": ['pemahaman bermakna', 'bermakna', 'big ideas', 'gagasan utama', 'ide pokok'],
        "Target Peserta Didik": ['target peserta', 'sasaran peserta', 'karakteristik peserta', 'target siswa', 'audience']
    }
    
    # For sections still not found, try the keyword approach
    for section, keywords in section_keywords.items():
        if results[section] == "Tidak tersedia":
            paragraph_start = -1
            
            # First look for exact section headers
            for i, line in enumerate(lines):
                line_lower = line.lower().strip()
                # Check if line matches any of our section patterns
                if any(keyword in line_lower for keyword in keywords):
                    paragraph_start = i
                    break
            
            if paragraph_start >= 0:
                # Determine where the section might end - either at next section or after several lines
                paragraph_end = paragraph_start + 1
                
                # Look for the next potential section header
                for i in range(paragraph_start + 1, min(paragraph_start + 15, len(lines))):
                    line_lower = lines[i].lower().strip()
                    
                    # Check if this line might be the start of the next section
                    if (re.match(r'^[ivxlcdm]+\.|^\d+\.', line_lower) or  # Roman numerals or numbers followed by period
                        any(keyword in line_lower for section_kw in section_keywords.values() 
                            for keyword in section_kw if section_kw != keywords)):
                        paragraph_end = i
                        break
                
                # Extract and clean the section text
                section_text = '\n'.join(lines[paragraph_start:paragraph_end]).strip()
                if section_text:
                    results[section] = clean_extracted_text(section_text)
    
    # Second pass to try to find sections by scanning the whole document
    for section, value in results.items():
        if value == "Tidak tersedia":
            # Look for sentences that might contain relevant information
            for i, line in enumerate(lines):
                line_lower = line.lower()
                
                # Get keywords for this section
                keywords = section_keywords.get(section, [])
                
                # Check each keyword
                for keyword in keywords:
                    if keyword in line_lower:
                        # Extract several lines from this point
                        extract = '\n'.join(lines[i:min(i+8, len(lines))])
                        results[section] = clean_extracted_text(extract)
                        break
                
                if results[section] != "Tidak tersedia":
                    break
    
    # Return the appropriate format based on the return_dict parameter
    if return_dict:
        return results
    else:
        return (
            results["Modul/Elemen Ajar"],
            results["Kompetensi Awal"],
            results["Tujuan Pembelajaran"],
            results["Pemahaman Bermakna"],
            results["Target Peserta Didik"]
        )

def clean_extracted_text(text):
    """Clean up extracted text by removing extra whitespace, bullet points, etc."""
    if not text or text == "Tidak tersedia":
        return text
        
    # Remove excess whitespace but preserve paragraph breaks
    text = re.sub(r'[ \t]+', ' ', text)  # Collapse multiple spaces/tabs to single space
    text = re.sub(r'\n{3,}', '\n\n', text)  # Limit consecutive newlines to max 2
    
    # Handle common formatting patterns
    text = text.replace(' .', '.')  # Fix common spacing issues
    text = text.replace(' ,', ',')
    text = text.replace(' :', ':')
    
    # Preserve bullet points but normalize their format
    text = re.sub(r'^\s*[•○●♦◘■]\s*', '• ', text, flags=re.MULTILINE)
    
    # Convert numbered lists to consistent format but preserve them
    text = re.sub(r'^\s*(\d+)[\.\)]\s*', r'\1. ', text, flags=re.MULTILINE)
    
    # Add line breaks for readability in the prompt after sentence endings
    text = re.sub(r'([.!?])\s+([A-Z])', r'\1\n\2', text)
    
    # Handle brackets or parenthetical content 
    text = re.sub(r'\(\s*(.*?)\s*\)', r'(\1)', text)  # Clean up spaces in parentheses
    
    # Fix common OCR issues
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Add space between lowercase and uppercase
    
    return text.strip()
# Function to extract text from a PDF file
def extract_text_from_pdf(file_path):
    """Extract text from a PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")
        raise

def extract_text_from_pdf_bytes(file_bytes: bytes):
    """Extract text from PDF given raw bytes (no disk write)."""
    try:
        from io import BytesIO
        bio = BytesIO(file_bytes)
        pdf_reader = PyPDF2.PdfReader(bio)
        text = ""
        
        # Try to get PDF metadata
        metadata = pdf_reader.metadata
        if metadata:
            info_text = []
            for key, value in metadata.items():
                if key.startswith('/') and value and isinstance(value, str):
                    # Clean up the key (remove leading slash)
                    clean_key = key[1:] if key.startswith('/') else key
                    info_text.append(f"{clean_key}: {value}")
            
            if info_text:
                text += "\n".join(info_text) + "\n\n"
                
        # Extract text from all pages
        for page_num, page in enumerate(pdf_reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                else:
                    # If we can't extract text normally, try to get any text using a more generic approach
                    text += f"[Content from page {page_num} - text extraction limited]\n"
            except Exception as page_error:
                print(f"Warning: Error extracting text from page {page_num}: {str(page_error)}")
                text += f"[Content from page {page_num} - extraction error]\n"
                
        # If we got very little text, try an alternative extraction method
        if len(text.strip()) < 100:
            print("Warning: Limited text extracted from PDF, trying alternative extraction method")
            text += "[This PDF may contain scanned images or protected content. Limited text extraction.]\n"
            
        return text
    except Exception as e:
        print(f"Error extracting text from PDF bytes: {str(e)}")
        # Return a placeholder instead of raising, so processing can continue
        return "[PDF content could not be extracted. The file may be damaged or protected.]"


# Function to extract text from a Word document (.docx)
def extract_text_from_docx(file_path):
    """Extract text from a Word document (.docx)"""
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                if row_text:
                    text += "\n" + " | ".join(row_text)

        return text
    except Exception as e:
        print(f"Error extracting text from DOCX: {str(e)}")
        raise

def extract_text_from_docx_bytes(file_bytes: bytes, file_ext: str):
    """Extract text from a Word document (.doc or .docx) using raw bytes."""
    from io import BytesIO
    try:
        bio = BytesIO(file_bytes)
        text = ""
        
        # For DOCX format
        if file_ext == 'docx':
            doc = Document(bio)
            
            # Get document properties if available
            try:
                core_properties = doc.core_properties
                if core_properties:
                    properties_text = []
                    for prop_name in ['title', 'subject', 'author', 'keywords', 'category']:
                        prop_value = getattr(core_properties, prop_name, None)
                        if prop_value:
                            properties_text.append(f"{prop_name.title()}: {prop_value}")
                    
                    if properties_text:
                        text += "\n".join(properties_text) + "\n\n"
            except Exception as prop_error:
                print(f"Warning: Could not extract document properties: {str(prop_error)}")
            
            # Extract text from paragraphs with style information
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Check if paragraph has a style that might indicate a heading or important text
                    style_name = paragraph.style.name if paragraph.style else "Normal"
                    if "heading" in style_name.lower() or style_name.lower().startswith(("h1", "h2", "h3")):
                        text += f"[{style_name}] {paragraph.text}\n"
                    else:
                        text += paragraph.text + "\n"
            
            # Extract tables with better formatting
            for table in doc.tables:
                text += "\n[Table Content]\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_content = cell.text.strip().replace('\n', ' ')
                        row_text.append(cell_content)
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                text += "[End Table]\n"
            
            # Also try to extract header and footer content which might contain important info
            try:
                header_text = []
                for section in doc.sections:
                    for header in section.header.paragraphs:
                        if header.text.strip():
                            header_text.append(header.text)
                
                if header_text:
                    text += "\n[Header Content]\n" + "\n".join(header_text) + "\n"
            except Exception as header_error:
                print(f"Warning: Could not extract headers: {str(header_error)}")
                
        else:  # .doc fallback using mammoth
            try:
                result = mammoth.extract_raw_text(bio)
                text = result.value
            except Exception as mammoth_error:
                print(f"Warning: Error extracting .doc with mammoth: {str(mammoth_error)}")
                text = "[This .doc file could not be fully extracted. Limited content available.]"
        
        # If we got very little text, add a note
        if len(text.strip()) < 100:
            print("Warning: Limited text extracted from DOC/DOCX")
            text += "[This document may contain mostly images or protected content. Limited text extraction.]\n"
            
        return text
    except Exception as e:
        print(f"Error extracting text from DOC/DOCX bytes: {str(e)}")
        # Return a placeholder instead of raising, so processing can continue
        return "[DOC/DOCX content could not be extracted. The file may be damaged or protected.]"


# =========================================
# EDUCATIONAL COMPONENT EXTRACTION HELPERS
# =========================================
# NOTE: This is a legacy implementation that's been replaced by the more advanced version above.
# This version is kept for backward compatibility with older code that might call it.



def extract_module_elements(sections):
    """Extract module/subject matter elements, excluding administrative info."""
    
    # Target sections that contain actual learning content
    target_sections = [
        'komponen inti', 'tujuan pembelajaran', 'pemahaman bermakna',
        'kegiatan pembelajaran', 'materi pembelajaran', 'isi pembelajaran'
    ]
    
    # Avoid these administrative sections
    avoid_sections = [
        'informasi umum', 'identitas modul', 'sarana dan prasarana',
        'model pembelajaran', 'refleksi guru', 'lampiran', 'daftar pustaka'
    ]
    
    content_parts = []
    
    for section_name, section_content in sections.items():
        section_lower = section_name.lower()
        
        # Skip administrative sections
        if any(avoid in section_lower for avoid in avoid_sections):
            continue
            
        # Include target sections or sections with learning keywords
        if (any(target in section_lower for target in target_sections) or
            has_learning_content_keywords(section_content)):
            
            # Extract key learning topics/concepts
            topics = extract_learning_topics(section_content)
            if topics:
                content_parts.extend(topics)
    
    if content_parts:
        return '; '.join(content_parts[:5])  # Limit to 5 main topics
    
    return "Tidak tersedia"


def extract_initial_competencies(sections):
    """Extract initial competencies with high precision."""
    
    for section_name, section_content in sections.items():
        section_lower = section_name.lower()
        
        # Look specifically for competency sections
        if ('kompetensi awal' in section_lower or 
            'kemampuan awal' in section_lower or
            'prasyarat' in section_lower):
            
            return clean_competency_text(section_content)
    
    # If no dedicated section found, look for competency mentions in content
    for section_name, section_content in sections.items():
        if 'kompetensi' in section_content.lower():
            competency_sentences = extract_competency_sentences(section_content)
            if competency_sentences:
                return '; '.join(competency_sentences[:3])
    
    return "Tidak tersedia"


def extract_learning_objectives(sections):
    """Extract learning objectives with precision."""
    
    for section_name, section_content in sections.items():
        section_lower = section_name.lower()
        
        # Target objective sections specifically
        if ('tujuan pembelajaran' in section_lower or
            'capaian pembelajaran' in section_lower or
            'learning objective' in section_lower):
            
            objectives = extract_bulleted_items(section_content)
            if objectives:
                return '; '.join(objectives[:6])  # Limit to 6 main objectives
            else:
                return clean_extracted_text(section_content)
    
    return "Tidak tersedia"


def extract_meaningful_understanding(sections):
    """Extract meaningful understanding components."""
    
    for section_name, section_content in sections.items():
        section_lower = section_name.lower()
        
        if ('pemahaman bermakna' in section_lower or
            'meaningful understanding' in section_lower):
            
            understanding_items = extract_bulleted_items(section_content)
            if understanding_items:
                return '; '.join(understanding_items[:4])
            else:
                return clean_extracted_text(section_content)
    
    return "Tidak tersedia"


def extract_target_students(sections):
    """Extract target student characteristics."""
    
    for section_name, section_content in sections.items():
        section_lower = section_name.lower()
        
        if ('target peserta didik' in section_lower or
            'sasaran peserta' in section_lower or
            'karakteristik peserta' in section_lower):
            
            return clean_extracted_text(section_content)
    
    return "Tidak tersedia"


def has_learning_content_keywords(text):
    """Check if text contains learning content keywords."""
    learning_keywords = [
        'algoritma', 'pemrograman', 'variabel', 'fungsi', 'struktur kontrol',
        'diagram alir', 'pseudokode', 'bahasa c', 'ekspresi', 'perulangan'
    ]
    
    text_lower = text.lower()
    return any(keyword in text_lower for keyword in learning_keywords)


def extract_learning_topics(text):
    """Extract main learning topics from text."""
    topics = []
    
    # Look for common topic indicators
    topic_patterns = [
        r'mempelajari\s+([^.]+)',
        r'memahami\s+([^.]+)', 
        r'menguasai\s+([^.]+)',
        r'konsep\s+([^.]+)',
        r'materi\s+([^.]+)'
    ]
    
    for pattern in topic_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        topics.extend([match.strip() for match in matches])
    
    # Remove duplicates and clean
    unique_topics = []
    for topic in topics:
        cleaned_topic = clean_topic_text(topic)
        if cleaned_topic and cleaned_topic not in unique_topics:
            unique_topics.append(cleaned_topic)
    
    return unique_topics[:5]  # Limit to 5 topics


def extract_bulleted_items(text):
    """Extract bulleted or numbered items from text."""
    items = []
    
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        
        # Check for bullet points or numbers
        if (re.match(r'^[-•*]\s+', line) or 
            re.match(r'^\d+[\.)]\s+', line)):
            
            # Clean the item
            item = re.sub(r'^[-•*\d\.)]\s+', '', line)
            if item and len(item) > 10:  # Ensure meaningful content
                items.append(item.strip())
    
    return items


def extract_competency_sentences(text):
    """Extract sentences that describe competencies."""
    sentences = re.split(r'[.!?]+', text)
    competency_sentences = []
    
    competency_keywords = [
        'mampu', 'dapat', 'menguasai', 'memahami', 'menerapkan',
        'menganalisis', 'mengevaluasi', 'menciptakan'
    ]
    
    for sentence in sentences:
        sentence = sentence.strip()
        if (any(keyword in sentence.lower() for keyword in competency_keywords) and
            len(sentence) > 20):  # Ensure meaningful length
            competency_sentences.append(sentence)
    
    return competency_sentences


def clean_competency_text(text):
    """Clean competency text by removing administrative noise."""
    # Remove common administrative phrases
    admin_phrases = [
        r'materi pada unit.*berkaitan dengan',
        r'dalam unit.*siswa diajarkan',
        r'lewat unit ini.*diimplementasikan',
        r'dengan demikian.*untuk menghasilkan'
    ]
    
    cleaned = text
    for phrase in admin_phrases:
        cleaned = re.sub(phrase, '', cleaned, flags=re.IGNORECASE | re.DOTALL)
    
    return clean_extracted_text(cleaned)


def clean_topic_text(topic):
    """Clean extracted topic text."""
    # Remove common noise words and phrases
    topic = re.sub(r'\b(dengan|untuk|dalam|pada|yang|dari|dan|atau)\b', ' ', topic, flags=re.IGNORECASE)
    topic = re.sub(r'\s+', ' ', topic)
    topic = topic.strip()
    
    # Only return if meaningful length
    return topic if len(topic) > 5 else None




def extract_text_from_docx(file_path):
    """Extract text from a Word document (.doc or .docx)"""
    try:
        if file_path.endswith('.docx'):
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text += "\t".join(row_text) + "\n"
            return text
        else:  # For .doc files
            with open(file_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                return result.value
    except Exception as e:
        print(f"Error extracting text from Word document: {str(e)}")
        raise

def convert_text_to_dataframe_improved(text):
    """Improved function to convert extracted text to a dataframe"""
    try:
        # Split text into lines for processing
        lines = text.split('\n')
        
        # Filter out empty lines
        lines = [line.strip() for line in lines if line.strip()]
        
        # Look for table-like patterns in the text
        table_lines = []
        
        # Check each line for tab separators or multiple spaces indicating tabular data
        for line in lines:
            if '\t' in line or re.search(r'\s{2,}', line):
                table_lines.append(line)
        
        if table_lines:
            # Create a temporary file to process the data
            with tempfile.NamedTemporaryFile(delete=False, mode='w', suffix='.tsv', encoding='utf-8') as temp_file:
                for line in table_lines:
                    # Normalize spaces and tabs
                    normalized_line = re.sub(r'\s{2,}', '\t', line)
                    temp_file.write(normalized_line + '\n')
            
            try:
                # Try reading the file as a tab-separated values file
                df = pd.read_csv(temp_file.name, sep='\t', on_bad_lines='skip', encoding='utf-8')  # Updated from error_bad_lines
                os.unlink(temp_file.name)
                return df
            except Exception as e:
                # If that fails, try a different approach with more robust error handling
                print(f"First attempt at parsing failed: {str(e)}")
                try:
                    os.unlink(temp_file.name)
                except:
                    pass
                
                # Try an alternative approach using pandas read_fwf (fixed width format)
                with tempfile.NamedTemporaryFile(delete=False, mode='w', suffix='.txt', encoding='utf-8') as temp_file:
                    for line in table_lines:
                        temp_file.write(line + '\n')
                
                try:
                    # Try to infer the fixed width format
                    df = pd.read_fwf(temp_file.name, encoding='utf-8')
                    os.unlink(temp_file.name)
                    return df
                except Exception as e2:
                    print(f"Fixed width parsing failed: {str(e2)}")
                    os.unlink(temp_file.name)
                    
                    # Final fallback: create a simple dataframe with lines as rows
                    data = {"line": table_lines}
                    return pd.DataFrame(data)
        
        # If no tabular data found, return a dataframe with the content
        return pd.DataFrame({"content": [text]})
            
    except Exception as e:
        print(f"Error converting text to dataframe: {str(e)}")
        # Return an empty dataframe as fallback
        return pd.DataFrame()
    
# =========================================
# VALIDATION MATRIX HELPERS (requirements per stage/difficulty)
# =========================================

def get_required_validation_pairs():
    """Return required pairs of (technology_level -> list[difficulty]) for MST routing.
    Based on system design:
      - L1: Medium
      - L2: Hard, Easy
      - L3: Hard, Medium
      - L4: Hard, Medium, Easy
      - L5: Hard, Medium, Easy
    """
    return {
        1: ['Medium'],
        2: ['Hard', 'Easy'],
        3: ['Hard', 'Medium'],
        4: ['Hard', 'Medium', 'Easy'],
        5: ['Hard', 'Medium', 'Easy'],
    }

def check_validation_matrix(collection_id):
    """Check that the collection has at least one validated question for each required (level, difficulty) pair.
    Returns: (ready: bool, missing: list[str]) where items are like 'L3-Hard'.
    """
    required = get_required_validation_pairs()
    missing = []
    try:
        for lvl, diffs in required.items():
            for diff in diffs:
                count = db.session.query(Question).join(
                    CollectionQuestion, CollectionQuestion.question_id == Question.id
                ).filter(
                    CollectionQuestion.collection_id == int(collection_id),
                    Question.technology_level == int(lvl),
                    Question.difficulty == diff,
                    Question.is_validated == True
                ).count()
                if count == 0:
                    missing.append(f"L{lvl}-{diff}")
    except Exception as e:
        print(f"[ValidationMatrix] Error checking matrix: {e}")
        # Conservative: treat as missing all to block until fixed
        return False, ['internal-error']
    return (len(missing) == 0), missing

# =========================================
# GET QUESTION (per level)
# =========================================
@app.route('/get_question', methods=['GET'])
def get_question():
    """MST Adaptive Question Selection"""
    user_id = request.args.get("user_id", "default")
    collection_id = request.args.get("collection_id")
    print(f"[MST] Get question untuk user_id: {user_id}, collection_id: {collection_id}")
    
    try:
        # Validasi input
        if not collection_id:
            return jsonify({
                "status": "error",
                "message": "Collection ID tidak ditemukan dalam URL!"
            }), 400
            
        # Konversi user_id menjadi integer
        siswa_id = int(user_id) if user_id != "default" else None
        if siswa_id is None:
            return jsonify({
                "status": "error",
                "message": "Invalid user ID. Please login properly."
            }), 400
            
        # Verifikasi akses siswa ke collection
        has_access = db.session.query(collection_students).filter_by(
            collection_id=collection_id,  
            siswa_id=siswa_id
        ).first() is not None
        
        if not has_access:
            return jsonify({
                "status": "error",
                "message": "Anda tidak memiliki akses ke koleksi soal ini."
            }), 403
            
        # Get atau create MST state
        mst_state = get_mst_state(siswa_id, collection_id)
        
        # Cek pra-syarat: minimal 1 soal tervalidasi untuk setiap kombinasi level & kesulitan
        ready, missing = check_validation_matrix(collection_id)
        if not ready:
            missing_str = ", ".join(missing) if isinstance(missing, list) else "-"
            return jsonify({
                "status": "blocked",
                "message": "Tes belum dapat dimulai. Guru perlu memvalidasi minimal 1 soal untuk setiap kombinasi Level & Kesulitan (contoh: L1-Medium, L2-Easy, L2-Hard, ...). Yang belum tersedia: " + missing_str,
                "missing": missing
            }), 200
        
        # Cek apakah test sudah selesai
        if mst_state.test_completed:
            return jsonify({
                "status": "game over",
                "message": f"Tes selesai! Diagnosis: {mst_state.final_diagnosis}"
            }), 200
        
        print(f"[MST] Current state: Stage {mst_state.current_stage}, Level {mst_state.current_level}")
        
        # Dapatkan soal berdasarkan stage dan level saat ini
        target_level_map = {
            1: "Awareness",
            2: "Literacy", 
            3: "Capability",
            4: "Creativity",
            5: "Criticism"
        }
        
        target_level = target_level_map[mst_state.current_stage]
        target_difficulty = mst_state.current_level
        
        # Cari soal yang sesuai dengan stage dan difficulty
        questions = db.session.query(Question).join(
            CollectionQuestion,  
            CollectionQuestion.question_id == Question.id
        ).filter(
            Question.technology_level == mst_state.current_stage,  # Gunakan technology_level (1-5)
            Question.difficulty == target_difficulty,  # Gunakan difficulty (Easy/Medium/Hard)
            CollectionQuestion.collection_id == int(collection_id),
            Question.is_validated == True
        ).all()
        
        print(f"[MST] Mencari soal: Technology Level {mst_state.current_stage}, Difficulty {target_difficulty}")
        print(f"[MST] Ditemukan {len(questions)} soal")
        
        # Jika tidak ada soal, coba fallback mechanism
        if not questions:
            print(f"[MST] Tidak ada soal untuk Stage {mst_state.current_stage} {target_difficulty}, mencoba fallback...")
            
            # Coba cari soal dengan difficulty lain di stage yang sama
            fallback_difficulties = ['Medium', 'Easy', 'Hard']
            if target_difficulty in fallback_difficulties:
                fallback_difficulties.remove(target_difficulty)
            
            for fallback_diff in fallback_difficulties:
                questions = db.session.query(Question).join(
                    CollectionQuestion,  
                    CollectionQuestion.question_id == Question.id
                ).filter(
                    Question.technology_level == mst_state.current_stage,
                    Question.difficulty == fallback_diff,
                    CollectionQuestion.collection_id == int(collection_id),
                    Question.is_validated == True
                ).all()
                
                if questions:
                    print(f"[MST] Fallback berhasil: menggunakan {fallback_diff} untuk Stage {mst_state.current_stage}")
                    target_difficulty = fallback_diff  # Update difficulty yang digunakan
                    break
            
            # Jika masih tidak ada, tandai tes selesai dengan diagnosis berdasarkan progress
            if not questions:
                # Hitung diagnosis berdasarkan stage terakhir yang berhasil diselesaikan
                if mst_state.current_stage == 2:
                    diagnosis = "< L2 (Siswa belum mencapai Level Literacy)"
                elif mst_state.current_stage == 3:
                    diagnosis = "L2 (Literacy)"
                elif mst_state.current_stage == 4:
                    diagnosis = "L3 (Capability)"
                elif mst_state.current_stage == 5:
                    diagnosis = "L4 (Creativity)"
                else:
                    diagnosis = "L1 (Awareness)"
                
                # Update state dan tandai selesai
                mst_state.test_completed = True
                mst_state.final_diagnosis = diagnosis
                db.session.commit()
                
                return jsonify({
                    "status": "win",  
                    "message": f"Tes selesai! Diagnosis: {diagnosis}. (Catatan: Koleksi ini tidak memiliki cukup soal untuk semua stage MST)"
                }), 200

        # Pilih soal acak
        question = random.choice(questions)
        print(f"[MST] Soal dipilih ID: {question.id}")
        
        # Parse opsi dengan lebih robust
        options = []
        if hasattr(question, 'options') and question.options:
            try:
                if isinstance(question.options, str):
                    options = json.loads(question.options)
                elif isinstance(question.options, list):
                    options = question.options
                else:
                    options = []
            except json.JSONDecodeError as e:
                print(f"[MST] Error parsing options: {e}")
                options = []
        
        # Pastikan options adalah list dan tidak kosong
        if not options or not isinstance(options, list):
            options = ["Option A", "Option B", "Option C", "Option D"]  # Default options
            print(f"[MST] Menggunakan default options untuk question ID: {question.id}")
        
        response_data = {
            "status": "continue",  # Status untuk frontend
            "id": question.id,
            "level": target_level,  # Nama level (Awareness, dll)
            "technology_level": question.technology_level,  # Angka level (1-5)
            "difficulty": question.difficulty,
            "stage": mst_state.current_stage,
            "question": question.soal,  # Frontend mengharapkan field 'question'
            "options": options,
            "question_type": question.question_type or "multiple_choice",
            "bobot": question.bobot,
            "explanation": question.explanation,
            "jawaban_benar": question.jawaban_benar if question.question_type == 'multiple_choice' else None
        }
        
        return jsonify(response_data), 200
        
    except ValueError:
        return jsonify({
            "status": "error",
            "message": "Format ID tidak valid"
        }), 400
        
    except Exception as e:
        print(f"[MST] Error di get_question: {str(e)}")
        return jsonify({
            "status": "error",
            "message": f"Kesalahan server: {str(e)}"
        }), 500
# =========================================
# SUBMIT ANSWER
# =========================================
@app.route('/submit_answer', methods=['POST'])
def submit_answer():
    """MST Adaptive Answer Processing"""
    data = request.get_json(silent=True) or {}
    user_id = data.get("user_id", "default")
    raw_answer = data.get("answer", "")
    answer = raw_answer.strip() if isinstance(raw_answer, str) else str(raw_answer)
    question_id = data.get("question_id")
    collection_id = data.get("collection_id")

    print(f"[MST] ========== SUBMIT ANSWER REQUEST ==========")
    print(f"[MST] user_id: {user_id} (type: {type(user_id)})")
    print(f"[MST] jawaban: {answer[:80]}...")
    print(f"[MST] question_id: {question_id} (type: {type(question_id)})")
    print(f"[MST] collection_id: {collection_id} (type: {type(collection_id)})")
    print(f"[MST] ============================================")

    try:
        # Validasi dan konversi ID user
        try:
            siswa_id = int(user_id)
        except (TypeError, ValueError) as e:
            return jsonify({
                "status": "error",
                "message": f"User ID tidak valid: '{user_id}'. Pastikan Anda sudah login."
            }), 400

        # Validasi siswa exists
        siswa = Siswa.query.get(siswa_id)
        if not siswa:
            return jsonify({
                "status": "error",
                "message": f"Siswa dengan ID {siswa_id} tidak ditemukan. Silakan login ulang."
            }), 404

        print(f"[MST] OK Siswa validated: {siswa.nama} (ID: {siswa.id})")

        # Validasi input
        if question_id is None or collection_id is None:
            return jsonify({
                "status": "error",
                "message": "Data tidak lengkap: question_id atau collection_id hilang"
            }), 400

        # Konversi ID dengan safety check
        try:
            question_id_int = int(question_id)
            collection_id_int = int(collection_id)
        except (TypeError, ValueError):
            return jsonify({
                "status": "error",
                "message": "ID soal atau koleksi tidak valid"
            }), 400

        # Get MST state
        mst_state = get_mst_state(siswa_id, collection_id_int)
        print(f"[MST] MST State SEBELUM proses: Stage {mst_state.current_stage}, Level {mst_state.current_level}, Selesai? {mst_state.test_completed}")

        if mst_state.test_completed:
            return jsonify({
                "status": "game over",
                "message": f"Tes sudah selesai! Diagnosis: {mst_state.final_diagnosis}"
            }), 200

        # Get question
        question = Question.query.get(question_id_int)
        if not question:
            return jsonify({"status": "error", "message": "Soal tidak ditemukan"}), 404

        # Check if answer is correct - DENGAN NULL SAFETY
        correct_answer = (question.jawaban_benar or "").strip()
        user_answer_norm = answer.strip().lower()
        correct_answer_norm = correct_answer.lower()
        is_correct = (user_answer_norm == correct_answer_norm) if correct_answer else False

        print(f"[MST] Jawaban: '{answer}' vs '{correct_answer}' => {'BENAR' if is_correct else 'SALAH'}")

        # --- PERBAIKAN LOGIKA: Simpan SEMUA jawaban (benar dan salah) untuk riwayat ---
        try:
            level_name = TECHNOLOGY_LEVELS.get(mst_state.current_stage, str(mst_state.current_stage))
            user_answer_record = SiswaAnswer(
                siswa_id=siswa_id,
                question_id=question_id_int,
                collection_id=collection_id_int,
                siswa_answer=answer,
                selected_answer=answer,
                is_correct=is_correct, # Simpan status benar/salah
                stage=mst_state.current_stage,
                level=level_name,
                difficulty=question.difficulty if hasattr(question, 'difficulty') and question.difficulty else 'Medium'
            )
            db.session.add(user_answer_record)
            # Jangan commit dulu, tunggu update state MST
            print(f"[MST] Answer record DIBUAT (is_correct={is_correct}) (belum commit)")
        except Exception as e_ans:
            print(f"[MST] Error membuat record jawaban: {str(e_ans)}")
            db.session.rollback()
            return jsonify({"status": "error", "message": f"Error menyimpan jawaban: {str(e_ans)}"}), 500
        # --- PERBAIKAN SELESAI ---

        # Update MST state counters (setelah jawaban dibuat, sebelum menentukan stage berikutnya)
        # Note: counter ini HANYA untuk stage saat ini
        mst_state.questions_answered = (mst_state.questions_answered or 0) + 1
        if is_correct:
            mst_state.questions_correct = (mst_state.questions_correct or 0) + 1

        # Check if stage module is completed - MST ADAPTIF: 1 SOAL PER STAGE
        stage_question_limit = 1  # Sesuai skenario: 1 soal per stage

        if (mst_state.questions_answered or 0) >= stage_question_limit:
            # Evaluate stage completion
            next_stage, next_level_difficulty, should_stop, diagnosis = determine_next_stage(mst_state)
            print(f"[MST] Hasil determine_next_stage: next_stage={next_stage}, next_level_difficulty={next_level_difficulty}, should_stop={should_stop}, diagnosis={diagnosis}")

            # --- PERBAIKAN UTAMA DIMULAI: Update state jika tes selesai ---
            if should_stop:
                # Test selesai
                mst_state.test_completed = True
                mst_state.final_diagnosis = diagnosis
                print(f"[MST] TES SELESAI! Diagnosis: {diagnosis} akan disimpan.")
                try:
                    db.session.commit() # Commit semua perubahan (jawaban + state MST yang sudah final)
                    print("[MST] Commit Berhasil - Status tes selesai disimpan")
                except Exception as e_commit_stop:
                     print(f"[MST] Error commit saat tes selesai: {str(e_commit_stop)}")
                     db.session.rollback()
                     return jsonify({"status": "error", "message": f"Error finalisasi tes: {str(e_commit_stop)}"}), 500

                return jsonify({
                    "status": "game over", # Frontend akan redirect ke halaman hasil
                    "message": f"Tes selesai! Diagnosis akhir Anda adalah: {diagnosis}",
                    "explanation": question.explanation,
                    "is_correct": is_correct,
                    "correct_answer": correct_answer
                }), 200
            # --- PERBAIKAN UTAMA SELESAI ---
            else:
                # Lanjut ke stage/level berikutnya
                print(f"[MST] Lanjut ke Stage {next_stage} Difficulty {next_level_difficulty}")
                mst_state.current_stage = next_stage
                mst_state.current_level = next_level_difficulty # Simpan difficulty untuk stage berikutnya
                mst_state.questions_answered = 0  # Reset counter untuk stage baru
                mst_state.questions_correct = 0   # Reset counter untuk stage baru

                try:
                    # === SOLUSI TAMBAHAN: UPDATE SiswaResult SEKARANG ===
                    student_result = SiswaResult.query.filter_by(
                        siswa_id=siswa_id,
                        collection_id=collection_id_int
                    ).first()

                    if student_result:
                        # Update level siswa saat ini ke stage yang baru
                        student_result.current_level = next_stage 
                        print(f"[MST] SiswaResult diupdate: current_level={next_stage}")
                    # === SOLUSI TAMBAHAN SELESAI ===

                    db.session.commit() # Commit semua perubahan (jawaban + state MST + state SiswaResult)
                    print("[MST] Commit Berhasil - Lanjut ke stage berikutnya")
                except Exception as e_commit_next:
                 print(f"[MST] Error commit saat lanjut stage: {str(e_commit_next)}")
                 db.session.rollback()
                 return jsonify({"status": "error", "message": f"Error lanjut stage: {str(e_commit_next)}"}), 500

                return jsonify({
                    "status": "continue",
                    "message": f"{'Benar!' if is_correct else 'Salah.'} Lanjut ke Stage {next_stage} - {next_level_difficulty}",
                    "explanation": question.explanation,
                    "next_stage": next_stage,
                    "next_level": next_level_difficulty, # Kirim difficulty berikutnya
                    "is_correct": is_correct,
                    "correct_answer": correct_answer
                }), 200
        else:
            # Kasus ini seharusnya tidak terjadi jika stage_question_limit = 1
            # Jika limit > 1, commit jawaban saja
            try:
                db.session.commit()
                print("[MST] Commit Berhasil - Masih dalam stage yang sama (limit > 1)")
            except Exception as e_commit_same:
                print(f"[MST] Error commit saat masih di stage sama: {str(e_commit_same)}")
                db.session.rollback()
                return jsonify({"status": "error", "message": f"Error simpan jawaban: {str(e_commit_same)}"}), 500

            progress = f"{mst_state.questions_answered}/{stage_question_limit}"
            return jsonify({
                "status": "continue",
                "message": f"{'Benar!' if is_correct else 'Salah.'} Progress Stage {mst_state.current_stage}: {progress}",
                "explanation": question.explanation,
                "stage_progress": progress,
                "is_correct": is_correct,
                "correct_answer": correct_answer
            }), 200

    except ValueError as e:
        print(f"[MST] ValueError di submit_answer: {str(e)}")
        # Tidak perlu rollback karena belum ada perubahan DB signifikan
        return jsonify({
            "status": "error",
            "message": f"Format data tidak valid: {str(e)}"
        }), 400
    except Exception as e:
        print(f"[MST] Exception di submit_answer: {str(e)}")
        print(f"[MST] Error type: {type(e).__name__}")
        import traceback
        print(f"[MST] Traceback: {traceback.format_exc()}")
        db.session.rollback() # Rollback jika ada error tak terduga
        return jsonify({
            "status": "error",
            "message": f"Kesalahan server: {str(e)}"
        }), 500

# =========================================
# GET SUMMARY
# =========================================
@app.route('/get_summary', methods=['GET'])
def get_summary():
    user_id = request.args.get("user_id", "default")
    
    # Konversi user_id menjadi integer jika bukan "default"
    try:
        user_id = int(user_id) if user_id != "default" else None
    except ValueError:
        user_id = None
    
    if user_id is None:
        return jsonify({
            "status": "error", 
            "message": "ID pengguna tidak valid",
        }), 400
    
    try:
        # Ambil data hasil siswa
        user_result = SiswaResult.query.filter_by(siswa_id=user_id).first()
        if not user_result:
            # Return initial data instead of error
            return jsonify({
                "status": "success",
                "correct": 0,
                "incorrect": 0,
                "total_tests": 0,
                "completed_tests": 0
            }), 200
        
        # Hitung jumlah tes yang tersedia untuk siswa ini
        total_tests = db.session.query(func.count(collection_students.c.collection_id)).filter(
            collection_students.c.siswa_id == user_id
        ).scalar() or 0
        
        # Hitung jumlah tes yang sudah selesai
        # Tes selesai jika level saat ini adalah level final (5, 6, 7)
        completed_tests = SiswaResult.query.filter(
            SiswaResult.siswa_id == user_id,
            SiswaResult.current_level.in_([5, 6, 7])
        ).count()
        
        # Hitung jumlah jawaban benar dan salah
        total_correct = user_result.correct or 0
        total_incorrect = user_result.incorrect or 0
        
        return jsonify({
            "status": "success",
            "correct": total_correct,
            "incorrect": total_incorrect,
            "total_tests": total_tests,
            "completed_tests": completed_tests
        }), 200
    
    except Exception as e:
        print(f"Error getting summary: {str(e)}")
        return jsonify({
            "status": "error",
            "message": f"Error: {str(e)}"
        }), 500
    
@app.route("/get_questions", methods=["GET"])
@login_required # Pastikan hanya user yang login yang bisa mengakses
def get_questions():
    try:
        # Check if user is authenticated
        if not current_user.is_authenticated:
            return jsonify({"success": False, "message": "User not authenticated"}), 401
            
        # Check user type
        if not hasattr(current_user, 'user_type'):
            return jsonify({"success": False, "message": "User type not found"}), 403
            
        print(f"[get_questions] User ID: {current_user.id}, Type: {current_user.user_type}")
        print("Fetching questions from database...")
        
        # Ambil semua soal yang dimiliki oleh guru yang sedang login
        # Gunakan SQLAlchemy ORM daripada raw SQL untuk filter ini
        questions_raw = (
            Question.query
            .filter_by(guru_id=current_user.id, is_current=True)
            .order_by(Question.level.asc(), Question.id.asc())
            .all()
        )
        
        print(f"Fetched {len(questions_raw)} questions for guru_id {current_user.id}")
        
        questions_data = []
        for q in questions_raw:
            options = []
            if q.options:
                try:
                    options = json.loads(q.options)
                except json.JSONDecodeError:
                    pass # Abaikan jika tidak valid JSON
            
            questions_data.append({
                "id": q.id,
                "guru_id": q.guru_id,
                "collection_id": q.collection_id,
                "level": q.level,
                "soal": q.soal,
                "jawaban_benar": q.jawaban_benar,
                "options": options,
                "question_type": q.question_type,
                "p": float(q.p),
                "bobot": q.bobot,
                "explanation": q.explanation,
                "created_at": q.created_at.isoformat() if q.created_at else None,
                "version": q.version,
                "is_current": q.is_current,
                "is_validated": q.is_validated 
            })
            
        print(f"Successfully formatted {len(questions_data)} questions for response")
        return jsonify({"success": True, "data": questions_data})
        
    except Exception as e:
        print(f"Error getting questions: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "message": f"Error: {str(e)}"}), 500
    
@app.route("/api/check_questions_count", methods=["GET"])
@login_required
def check_questions_count():
    # Hanya guru yang bisa melihat ini
    if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
        return jsonify({"success": False, "message": "Akses ditolak."}), 403

    try:
        # Hitung soal milik guru yang sedang login
        total_count = db.session.query(db.func.count(Question.id)).filter_by(guru_id=current_user.id).scalar()
        
        # Hitung berdasarkan level milik guru yang sedang login
        level_counts = db.session.query(
            Question.level,  
            db.func.count(Question.id)
        ).filter_by(guru_id=current_user.id).group_by(Question.level).all()
        
        level_distribution = {str(level): count for level, count in level_counts}
        
        return jsonify({
            "success": True,
            "totalQuestions": total_count,
            "levelDistribution": level_distribution,
            "queryTime": 0
        })
    except Exception as e:
        print(f"Error checking questions count: {str(e)}")
        return jsonify({
            "success": False,
            "message": f"Error checking database: {str(e)}"
        }), 500
    
@app.route("/api/db_check", methods=["GET"])
def db_check():
    try:
        from sqlalchemy import text
        
        # Hitung total soal
        result = db.session.execute(text("SELECT COUNT(*) FROM questions"))
        total = result.scalar()
        
        # Hitung soal per level
        result = db.session.execute(text("SELECT level, COUNT(*) FROM questions GROUP BY level"))
        level_counts = {row[0]: row[1] for row in result}
        
        # Ambil sampel soal
        result = db.session.execute(text("SELECT id, level, soal FROM questions LIMIT 5"))
        samples = [{"id": row[0], "level": row[1], "soal": row[2]} for row in result]
        
        return jsonify({
            "success": True,
            "total_questions": total,
            "per_level": level_counts,
            "samples": samples
        })
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})
    
@app.route("/api/table_structure", methods=["GET"])
def check_table_structure():
    try:
        from sqlalchemy import text
        
        # Cek struktur tabel questions
        result = db.session.execute(text("DESCRIBE questions"))
        columns = [dict(row._mapping) for row in result]
        
        # Cek foreign key constraints
        result = db.session.execute(text("SELECT * FROM INFORMATION_SCHEMA.KEY_COLUMN_USAGE WHERE TABLE_NAME = 'questions'"))
        constraints = [dict(row._mapping) for row in result]
        
        return jsonify({
            "success": True,
            "table_structure": columns,
            "constraints": constraints
        })
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})
    
@app.route('/api/collections', methods=['GET'])
@login_required
def get_collections():
    # Ambil semua koleksi milik guru
    collections = QuestionCollection.query.filter_by(guru_id=current_user.id).all()
    
    result = []
    for collection in collections:
        # Hitung jumlah soal di tiap koleksi
        question_count = db.session.query(CollectionQuestion).filter_by(collection_id=collection.id).count()
        
        result.append({
            'id': collection.id,
            'name': collection.name,
            'description': collection.description,
            'question_count': question_count,
            'created_at': collection.created_at.isoformat() if collection.created_at else None
        })
    
    return jsonify({'success': True, 'collections': result})

@app.route('/api/collections', methods=['POST'])
@login_required
def create_collection():
    data = request.json
    
    # Validasi data
    if not data.get('name'):
        return jsonify({'success': False, 'message': 'Nama koleksi harus diisi'}), 400
    
    # Cek apakah nama koleksi sudah ada untuk guru ini
    existing_collection = QuestionCollection.query.filter_by(
        guru_id=current_user.id,
        name=data.get('name')
    ).first()
    
    if existing_collection:
        app.logger.warning(f"Attempt to create duplicate collection '{data.get('name')}' for guru {current_user.id}")
        return jsonify({
            'success': False, 
            'message': f'Koleksi dengan nama "{data.get("name")}" sudah ada. Silakan gunakan nama yang berbeda.'
        }), 400
    
    # Buat koleksi baru
    collection = QuestionCollection(
        guru_id=current_user.id,
        name=data.get('name'),
        description=data.get('description', '')
    )
    
    db.session.add(collection)
    db.session.commit()
    
    # Tambahkan soal ke koleksi jika ada
    if data.get('question_ids') and isinstance(data.get('question_ids'), list):
        for question_id in data.get('question_ids'):
            # Validasi bahwa soal ada
            question = Question.query.get(question_id)
            if question:
                collection_question = CollectionQuestion(
                    collection_id=collection.id,
                    question_id=question_id
                )
                db.session.add(collection_question)
        
        db.session.commit()
    
    return jsonify({
        'success': True, 
        'message': 'Koleksi berhasil dibuat', 
        'collection_id': collection.id
    })

@app.route('/api/collections/check-name', methods=['POST'])
@login_required
def check_collection_name():
    """Endpoint untuk memeriksa apakah nama koleksi sudah ada"""
    data = request.json
    
    if not data.get('name'):
        return jsonify({'success': False, 'message': 'Nama koleksi harus diisi'}), 400
    
    # Cek apakah nama koleksi sudah ada untuk guru ini
    existing_collection = QuestionCollection.query.filter_by(
        guru_id=current_user.id,
        name=data.get('name')
    ).first()
    
    # Log untuk debugging
    app.logger.info(f"Checking collection name '{data.get('name')}' for guru {current_user.id}: {'exists' if existing_collection else 'available'}")
    
    return jsonify({
        'success': True,
        'exists': existing_collection is not None,
        'message': 'Nama sudah digunakan' if existing_collection else 'Nama tersedia'
    })

@app.route('/api/collections/<int:collection_id>/details', methods=['GET'])
@login_required
def get_collection_details_for_student(collection_id):
    try:
        # Cek apakah siswa memiliki akses ke koleksi ini
        has_access = db.session.query(collection_students).filter_by(
            collection_id=collection_id, 
            siswa_id=current_user.id
        ).first() is not None
        
        if not has_access:
            return jsonify({'success': False, 'message': 'Anda tidak memiliki akses ke koleksi ini'}), 403
        
        # Ambil detail koleksi
        collection = QuestionCollection.query.get(collection_id)
        if not collection:
            return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
        
        # Ambil info tentang guru
        guru = Guru.query.get(collection.guru_id)
        
        # Ambil progress siswa untuk koleksi ini
        progress = SiswaResult.query.filter_by(
            siswa_id=current_user.id,
            collection_id=collection_id
        ).first()
        
        result = {
            'success': True,
            'collection': {
                'id': collection.id,
                'name': collection.name,
                'description': collection.description,
                'created_at': collection.created_at.isoformat() if collection.created_at else None
            },
            'guru': {
                'id': guru.id,
                'nama': guru.nama,
                'email': guru.email
            },
            'progress': {
                'current_level': progress.current_level if progress else 1,
                'correct': progress.correct if progress else 0,
                'incorrect': progress.incorrect if progress else 0,
                'is_completed': progress.current_level in [5, 6, 7] if progress else False
            }
        }
        
        return jsonify(result)
        
    except Exception as e:
        print(f"Error getting collection details: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500
    
@app.route('/api/collections/<int:collection_id>/questions', methods=['GET'])
@login_required
def get_collection_questions(collection_id):
    # Verifikasi kepemilikan koleksi
    collection = QuestionCollection.query.filter_by(id=collection_id, guru_id=current_user.id).first()
    if not collection:
        return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
    
    # Ambil soal dari koleksi dengan join
    questions = db.session.query(Question).join(
        CollectionQuestion, CollectionQuestion.question_id == Question.id
    ).filter(
        CollectionQuestion.collection_id == collection_id
    ).all()
    
    # Format soal untuk response
    question_data = []
    for q in questions:
        options = []
        if hasattr(q, 'options') and q.options:
            try:
                options = json.loads(q.options)
            except:
                options = []
        
        # Determine stage based on technology_level (MST 5-stage system)
        stage = None
        
        # MST 5-Stage: Use technology_level directly (1-5)
        if hasattr(q, 'technology_level') and q.technology_level is not None and q.technology_level >= 1:
            stage = f"L{q.technology_level}"
        else:
            # Fallback: Use old level field if technology_level not set
            # This is for backward compatibility with old questions
            if q.level:
                # Old level system (1-7) to new stage (L1-L5) mapping
                level_to_stage_map = {1: 'L1', 2: 'L2', 3: 'L3', 4: 'L3', 5: 'L4', 6: 'L4', 7: 'L5'}
                stage = level_to_stage_map.get(q.level, 'L1')
            else:
                stage = 'L1'  # Default fallback
        
        question_data.append({
            'id': q.id,
            'level': q.level,
            'stage': stage,  # NEW: Add stage field for proper MST 5-stage display
            'technology_level': getattr(q, 'technology_level', None),  # NEW: Include technology_level
            'soal': q.soal,
            'jawaban_benar': q.jawaban_benar,
            'options': options,
            'question_type': q.question_type if hasattr(q, 'question_type') else 'multiple_choice',
            'p': float(q.p) if q.p else 0.5,
            'difficulty': q.difficulty if hasattr(q, 'difficulty') and q.difficulty else p_value_to_difficulty(float(q.p) if q.p else 0.5),
            'bobot': q.bobot or 1,
            'explanation': q.explanation or '',
            'is_validated': q.is_validated # NEW: Sertakan status validasi
        })
    
    return jsonify({
        'success': True, 
        'collection': {
            'id': collection.id,
            'name': collection.name,
            'description': collection.description
        },
        'questions': question_data
    })


@app.route('/api/collections/<int:collection_id>/questions', methods=['POST'])
@login_required
def add_questions_to_collection(collection_id):
    # Verifikasi kepemilikan koleksi
    collection = QuestionCollection.query.filter_by(id=collection_id, guru_id=current_user.id).first()
    if not collection:
        return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
    
    data = request.json
    question_ids = data.get('question_ids', [])
    
    # Validasi data
    if not question_ids or not isinstance(question_ids, list):
        return jsonify({'success': False, 'message': 'ID soal tidak valid'}), 400
    
    # Tambahkan soal ke koleksi
    added_count = 0
    for question_id in question_ids:
        # Cek apakah soal sudah ada di koleksi
        existing = CollectionQuestion.query.filter_by(
            collection_id=collection_id, 
            question_id=question_id
        ).first()
        
        if not existing:
            # Cek apakah soal ada
            question = Question.query.get(question_id)
            if question:
                collection_question = CollectionQuestion(
                    collection_id=collection_id,
                    question_id=question_id
                )
                db.session.add(collection_question)
                added_count += 1
    
    db.session.commit()
    
    return jsonify({
        'success': True, 
        'message': f'{added_count} soal berhasil ditambahkan ke koleksi'
    })

@app.route('/api/collections/<int:collection_id>/description', methods=['PUT'])
@login_required
def update_collection_description(collection_id):
    if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
        return jsonify({"success": False, "message": "Unauthorized"}), 403
    
    try:
        # Validate input
        data = request.json
        if not data or 'description' not in data:
            return jsonify({"success": False, "message": "Description is required"}), 400
            
        # Get collection and ensure it belongs to the current teacher
        collection = QuestionCollection.query.filter_by(id=collection_id, guru_id=current_user.id).first()
        
        if not collection:
            return jsonify({"success": False, "message": "Koleksi tidak ditemukan"}), 404
        
        # Update description
        collection.description = data['description']
        db.session.commit()
        
        return jsonify({
            "success": True, 
            "message": "Deskripsi koleksi berhasil diperbarui",
            "collection": {
                "id": collection.id,
                "name": collection.name,
                "description": collection.description
            }
        })
    
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "message": str(e)}), 500

@app.route('/api/collections/<int:collection_id>', methods=['DELETE'])
@login_required
def delete_collection(collection_id):
    try:
        # Verifikasi kepemilikan koleksi
        collection = QuestionCollection.query.filter_by(id=collection_id, guru_id=current_user.id).first()
        if not collection:
            return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
        
        # Hapus semua data terkait secara berurutan
        # 1. Hapus siswa answers yang terkait collection ini
        db.session.execute(
            text("DELETE FROM siswa_answers WHERE collection_id = :cid"),
            {"cid": collection_id}
        )
        
        # 2. Hapus MST state yang terkait collection ini
        db.session.execute(
            text("DELETE FROM mst_state WHERE collection_id = :cid"),
            {"cid": collection_id}
        )
        
        # 3. Hapus relasi collection_students
        db.session.execute(
            text("DELETE FROM collection_students WHERE collection_id = :cid"),
            {"cid": collection_id}
        )
        
        # 4. Hapus relasi collection_questions
        db.session.execute(
            text("DELETE FROM collection_questions WHERE collection_id = :cid"),
            {"cid": collection_id}
        )
        
        # 5. Hapus koleksi itu sendiri
        db.session.delete(collection)
        db.session.commit()
        
        return jsonify({
            'success': True, 
            'message': 'Koleksi berhasil dihapus'
        })
    except Exception as e:
        db.session.rollback()
        print(f"Error deleting collection: {str(e)}")
        return jsonify({
            'success': False, 
            'message': f'Gagal menghapus koleksi: {str(e)}'
        }), 500

# Endpoint untuk menghapus soal dari koleksi
@app.route('/api/collections/<int:collection_id>/questions/<int:question_id>', methods=['DELETE'])
@login_required
def delete_question_from_collection(collection_id, question_id):
    if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
        return jsonify({"success": False, "message": "Unauthorized"}), 403
    
    try:
        # Verifikasi koleksi milik guru yang sedang login
        collection = QuestionCollection.query.filter_by(id=collection_id, guru_id=current_user.id).first()
        if not collection:
            return jsonify({"success": False, "message": "Koleksi tidak ditemukan"}), 404
        
        # Verifikasi soal milik guru yang sedang login
        question = Question.query.filter_by(id=question_id, guru_id=current_user.id).first()
        if not question:
            return jsonify({"success": False, "message": "Soal tidak ditemukan"}), 404
        
        # Cek apakah ada jawaban siswa untuk soal ini
        student_answers_count = SiswaAnswer.query.filter_by(question_id=question_id).count()
        
        if student_answers_count > 0:
            # Jika ada jawaban siswa, hanya hapus dari koleksi, jangan hapus soal
            collection_question = CollectionQuestion.query.filter_by(
                collection_id=collection_id, 
                question_id=question_id
            ).first()
            
            if not collection_question:
                return jsonify({"success": False, "message": "Soal tidak ada dalam koleksi ini"}), 404
            
            # Hapus relasi dari koleksi saja
            db.session.delete(collection_question)
            db.session.commit()
            
            return jsonify({
                "success": True,
                "message": f"Soal berhasil dihapus dari koleksi. Soal tidak dihapus dari database karena sudah ada {student_answers_count} jawaban siswa."
            })
        
        # Jika tidak ada jawaban siswa, lanjutkan dengan penghapusan normal
        collection_question = CollectionQuestion.query.filter_by(
            collection_id=collection_id, 
            question_id=question_id
        ).first()
        
        if not collection_question:
            return jsonify({"success": False, "message": "Soal tidak ada dalam koleksi ini"}), 404
        
        # Hapus relasi dari koleksi
        db.session.delete(collection_question)
        
        # Cek apakah soal ini masih digunakan di koleksi lain
        other_collections = CollectionQuestion.query.filter_by(question_id=question_id).count()
        
        if other_collections == 0:
            # Jika soal tidak digunakan di koleksi lain dan tidak ada jawaban siswa, hapus soal dari database
            # Hapus semua versi soal terlebih dahulu
            QuestionVersion.query.filter_by(question_id=question_id).delete()
            
            # Hapus soal utama
            db.session.delete(question)
            message = "Soal berhasil dihapus dari koleksi dan database (tidak digunakan di koleksi lain dan tidak ada jawaban siswa)"
        else:
            message = "Soal berhasil dihapus dari koleksi (masih digunakan di koleksi lain)"
        
        db.session.commit()
        
        return jsonify({
            "success": True,
            "message": message
        })
    
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error deleting question from collection: {str(e)}")
        return jsonify({"success": False, "message": str(e)}), 500

# Endpoint untuk memvalidasi semua soal dalam koleksi
@app.route('/api/collections/<int:collection_id>/validate-all', methods=['POST'])
@login_required
def validate_all_questions(collection_id):
    if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
        return jsonify({"success": False, "message": "Unauthorized"}), 403
    
    try:
        # Get collection and ensure it belongs to the current teacher
        collection = QuestionCollection.query.filter_by(id=collection_id, guru_id=current_user.id).first()
        
        if not collection:
            return jsonify({"success": False, "message": "Koleksi tidak ditemukan"}), 404
        
        # Get all questions in this collection
        collection_questions = CollectionQuestion.query.filter_by(collection_id=collection_id).all()
        question_ids = [cq.question_id for cq in collection_questions]
        
        # Update validation status of all questions
        questions = Question.query.filter(Question.id.in_(question_ids)).all()
        validated_ids = []
        
        for question in questions:
            question.is_validated = True
            validated_ids.append(question.id)
        
        db.session.commit()
        
        return jsonify({
            "success": True, 
            "message": f"{len(validated_ids)} soal berhasil divalidasi",
            "count": len(validated_ids),
            "validated": validated_ids
        })
    
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "message": str(e)}), 500

# Mendapatkan siswa yang memiliki akses ke koleksi
@app.route('/api/collections/<int:collection_id>/students', methods=['GET'])
@login_required
def get_collection_students(collection_id):
    try:
        # Verifikasi kepemilikan koleksi
        collection = QuestionCollection.query.filter_by(id=collection_id, guru_id=current_user.id).first()
        if not collection:
            return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
        
        # Dapatkan siswa yang memiliki akses ke koleksi ini
        students = db.session.query(Siswa).join(
            collection_students,
            collection_students.c.siswa_id == Siswa.id
        ).filter(
            collection_students.c.collection_id == collection_id
        ).all()
        
        # Format data siswa
        students_data = []
        for student in students:
            students_data.append({
                'id': student.id,
                'nama': student.nama,
                'username': student.username,
                'kelas': student.kelas
            })
        
        return jsonify({
            'success': True,
            'students': students_data
        })
        
    except Exception as e:
        print(f"Error getting collection students: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

# Menambahkan siswa ke koleksi
@app.route('/api/collections/<int:collection_id>/students', methods=['POST'])
@login_required
def add_student_to_collection(collection_id):
    try:
        # 1. Verifikasi koleksi dan kepemilikan
        collection = QuestionCollection.query.filter_by(
            id=collection_id, 
            guru_id=current_user.id
        ).first()
        
        if not collection:
            return jsonify({
                'success': False,
                'message': 'Koleksi tidak ditemukan atau akses ditolak'
            }), 404

        # 2. Validasi data request
        data = request.get_json()
        if not data or 'siswa_id' not in data:
            return jsonify({
                'success': False,
                'message': 'Data siswa tidak valid'
            }), 400

        siswa_id = data['siswa_id']

        # 2b. Cegah penambahan siswa jika belum terpenuhi matriks validasi minimal (level x kesulitan)
        ready, missing = check_validation_matrix(collection_id)
        if not ready:
            missing_str = ", ".join(missing) if isinstance(missing, list) else "-"
            return jsonify({
                'success': False,
                'message': 'Validasi minimal 1 soal untuk setiap kombinasi Level & Kesulitan diperlukan sebelum menambahkan siswa. Kombinasi yang belum tersedia: ' + missing_str,
                'missing': missing
            }), 200

        # 2b. Cegah penambahan siswa jika BELUM ada soal tervalidasi di koleksi ini
        validated_count = db.session.query(Question).join(
            CollectionQuestion, CollectionQuestion.question_id == Question.id
        ).filter(
            CollectionQuestion.collection_id == collection_id,
            Question.is_validated == True
        ).count()
        if validated_count == 0:
            return jsonify({
                'success': False,
                'message': 'Validasi soal terlebih dahulu sebelum menambahkan siswa ke koleksi ini.'
            }), 200

        # 3. Cek keberadaan siswa
        siswa = Siswa.query.get(siswa_id)
        if not siswa:
            return jsonify({
                'success': False,
                'message': 'Siswa tidak ditemukan'
            }), 404

        # 4. Cek duplikasi relasi
        existing = db.session.execute(
            collection_students.select().where(
                (collection_students.c.collection_id == collection_id) &
                (collection_students.c.siswa_id == siswa_id)
            )
        ).fetchone()

        if existing:
            return jsonify({
                'success': True,
                'message': 'Siswa sudah terdaftar dalam koleksi ini'
            })

        # 5. Tambahkan relasi
        db.session.execute(
            collection_students.insert().values(
                collection_id=collection_id,
                siswa_id=siswa_id
            )
        )

        # 6. Buat atau update hasil siswa
        siswa_result = SiswaResult.query.filter_by(
            siswa_id=siswa_id,
            collection_id=collection_id
        ).first()

        if not siswa_result:
            new_result = SiswaResult(
                siswa_id=siswa_id,
                collection_id=collection_id,
                current_level=1,
                correct=0,
                incorrect=0
            )
            db.session.add(new_result)

        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Siswa berhasil ditambahkan ke koleksi',
            'data': {
                'collection_id': collection_id,
                'siswa_id': siswa_id,
                'siswa_nama': siswa.nama,
                'collection_name': collection.name
            }
        })

    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Unexpected error: {str(e)}")
        return jsonify({
            'success': False,
            'message': 'Terjadi kesalahan server'
        }), 500

# Menghapus siswa dari koleksi
@app.route('/api/collections/<int:collection_id>/students/<int:student_id>', methods=['DELETE'])
@login_required
def remove_student_from_collection(collection_id, student_id):
    try:
        # Verifikasi kepemilikan koleksi
        collection = QuestionCollection.query.filter_by(id=collection_id, guru_id=current_user.id).first()
        if not collection:
            return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
        
        # Hapus akses
        stmt = collection_students.delete().where(
            collection_students.c.collection_id == collection_id,
            collection_students.c.siswa_id == student_id
        )
        result = db.session.execute(stmt)
        db.session.commit()
        
        if result.rowcount == 0:
            return jsonify({'success': False, 'message': 'Siswa tidak memiliki akses ke koleksi ini'}), 404
        
        return jsonify({
            'success': True,
            'message': 'Akses siswa berhasil dihapus'
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"Error removing student from collection: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/collections/<int:collection_id>/student-count', methods=['GET'])
def get_collection_student_count(collection_id):
    try:
        count = db.session.query(collection_students).filter_by(
            collection_id=collection_id
        ).count()
        return jsonify({'success': True, 'count': count})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

# Mendapatkan daftar semua siswa
@app.route('/api/siswa', methods=['GET'])
@login_required
def get_all_students():
    try:
        if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
            return jsonify({'success': False, 'message': 'Akses ditolak. Hanya guru yang dapat mengakses daftar siswa.'}), 403
        
        kelas_filter = request.args.get('kelas') # NEW: Ambil parameter kelas untuk filter
        
        students_query = Siswa.query

        if kelas_filter and kelas_filter != 'all':
            students_query = students_query.filter_by(kelas=kelas_filter)

        students = students_query.order_by(Siswa.kelas, Siswa.nama).all() # Urutkan berdasarkan kelas dan nama
        
        students_data = []
        for student in students:
            students_data.append({
                'id': student.id,
                'nama': student.nama,
                'username': student.username,
                'kelas': student.kelas
            })
        
        # NEW: Ambil daftar kelas unik yang ada di database
        available_classes = db.session.query(Siswa.kelas).distinct().order_by(Siswa.kelas).all()
        available_classes = [cls[0] for cls in available_classes if cls[0]] # Filter out None/empty classes
        
        return jsonify({
            'success': True,
            'students': students_data,
            'available_classes': available_classes # Kirim daftar kelas yang tersedia
        })
        
    except Exception as e:
        print(f"Error getting all students: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500
    
# Tambahkan endpoint baru untuk menambahkan siswa berdasarkan kelas
@app.route('/api/collections/<int:collection_id>/add_students_by_class', methods=['POST'])
@login_required
def add_students_by_class(collection_id):
    try:
        # 1. Verifikasi koleksi dan kepemilikan
        collection = QuestionCollection.query.filter_by(
            id=collection_id,
            guru_id=current_user.id
        ).first()

        if not collection:
            return jsonify({
                'success': False,
                'message': 'Koleksi tidak ditemukan atau akses ditolak'
            }), 404

        # 2. Validasi data request
        data = request.get_json()
        if not data or 'class_name' not in data:
            return jsonify({
                'success': False,
                'message': 'Nama kelas tidak valid'
            }), 400

        class_name = data['class_name']

        # 2b. Cegah penambahan siswa jika BELUM ada soal tervalidasi di koleksi ini
        ready, missing = check_validation_matrix(collection_id)
        if not ready:
            missing_str = ", ".join(missing) if isinstance(missing, list) else "-"
            return jsonify({
                'success': False,
                'message': 'Validasi minimal 1 soal untuk setiap kombinasi Level & Kesulitan diperlukan sebelum menambahkan siswa berdasarkan kelas. Kombinasi yang belum tersedia: ' + missing_str,
                'missing': missing
            }), 200

        # 3. Dapatkan semua siswa di kelas tersebut yang belum terdaftar di koleksi ini
        students_to_add = Siswa.query.filter(
            Siswa.kelas == class_name,
            # Filter siswa yang belum ada di koleksi ini
            ~Siswa.id.in_(
                db.session.query(collection_students.c.siswa_id).filter(
                    collection_students.c.collection_id == collection_id
                )
            )
        ).all()

        added_count = 0
        if students_to_add:
            # 4. Tambahkan siswa ke koleksi dan buat/update SiswaResult mereka
            for student in students_to_add:
                # Tambahkan relasi ke collection_students
                db.session.execute(
                    collection_students.insert().values(
                        collection_id=collection_id,
                        siswa_id=student.id
                    )
                )

                # Buat atau update SiswaResult
                siswa_result = SiswaResult.query.filter_by(
                    siswa_id=student.id,
                    collection_id=collection_id
                ).first()

                if not siswa_result:
                    new_result = SiswaResult(
                        siswa_id=student.id,
                        collection_id=collection_id,
                        current_level=1,
                        correct=0,
                        incorrect=0
                    )
                    db.session.add(new_result)
                added_count += 1
            db.session.commit()

        return jsonify({
            'success': True,
            'message': f'{added_count} siswa dari kelas {class_name} berhasil ditambahkan ke koleksi.',
            'added_count': added_count
        }), 200

    except Exception as e:
        db.session.rollback()
        print(f"Error adding students by class: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': 'Terjadi kesalahan server saat menambahkan siswa berdasarkan kelas'
        }), 500

@app.route('/api/siswa/collections', methods=['GET'])
@login_required
def get_student_collections():
    if not current_user.is_authenticated or not hasattr(current_user, 'user_type') or current_user.user_type != 'siswa':
        return jsonify({'success': False, 'message': 'Akses ditolak'}), 403
    
    try:
        # Ambil koleksi yang tersedia untuk siswa ini dari tabel relasi
        collections_query = db.session.query(
            QuestionCollection, Guru
        ).join(
            Guru, QuestionCollection.guru_id == Guru.id
        ).join(
            collection_students, collection_students.c.collection_id == QuestionCollection.id
        ).filter(
            collection_students.c.siswa_id == current_user.id
        ).all()
        
        # Format data untuk respon
        teachers_dict = {}
        for collection, guru in collections_query:
            if guru.id not in teachers_dict:
                teachers_dict[guru.id] = {
                    'id': guru.id,
                    'nama': guru.nama,
                    'email': guru.email,
                    'collections': []
                }
            
            # Cek apakah koleksi baru (dibuat dalam 1 hari terakhir)
            is_new = False
            if collection.created_at:
                delta = datetime.datetime.now() - collection.created_at
                is_new = delta.days < 1
            
            teachers_dict[guru.id]['collections'].append({
                'id': collection.id,
                'name': collection.name,
                'is_new': is_new,
                'description': collection.description
            })
        
        # Konversi ke list untuk response
        teachers_list = list(teachers_dict.values())
        
        return jsonify({
            'success': True, 
            'teachers': teachers_list
        })
        
    except Exception as e:
        print(f"Error getting student collections: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500
    
# =========================================
# MST ADAPTIVE ENDPOINTS
# =========================================

@app.route('/api/mst/start/<int:collection_id>', methods=['POST'])
@login_required
def start_mst_test(collection_id):
    """Start MST adaptive test for a student"""
    if not current_user.is_authenticated or not hasattr(current_user, 'user_type') or current_user.user_type != 'siswa':
        return jsonify({'success': False, 'message': 'Akses ditolak'}), 403
    
    try:
        # Check if collection exists and student has access
        collection = QuestionCollection.query.get(collection_id)
        if not collection:
            return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
        
        # Check if student is assigned to this collection
        access_check = db.session.query(collection_students).filter(
            collection_students.c.collection_id == collection_id,
            collection_students.c.siswa_id == current_user.id
        ).first()
        
        if not access_check:
            return jsonify({'success': False, 'message': 'Anda tidak memiliki akses ke koleksi ini'}), 403
        
        # Initialize MST state using new system
        mst_state = get_mst_state(current_user.id, collection_id)
        mst_status = {
            'stage': mst_state.current_stage,
            'level': mst_state.current_level,
            'is_complete': mst_state.test_completed
        }
        
        # Get first module questions (L1 Medium)
        questions = get_questions_for_module(collection_id, 1, "Medium")
        
        if not questions:
            return jsonify({'success': False, 'message': 'Tidak ada soal tersedia untuk modul ini'}), 404
        
        # Format questions for frontend
        question_data = []
        for q in questions:
            options = []
            if q.options:
                try:
                    options = json.loads(q.options)
                except:
                    options = []
            
            question_data.append({
                'id': q.id,
                'technology_level': q.technology_level,
                'difficulty': q.difficulty,
                'soal': q.soal,
                'options': options,
                'question_type': q.question_type,
                'explanation': q.explanation
            })
        
        return jsonify({
            'success': True,
            'message': 'MST test dimulai',
            'mst_status': mst_status,
            'current_module': {
                'technology_level': 1,
                'difficulty': 'Medium',
                'level_name': TECHNOLOGY_LEVELS[1]
            },
            'questions': question_data,
            'collection': {
                'id': collection.id,
                'name': collection.name,
                'description': collection.description
            }
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Terjadi kesalahan: {str(e)}'}), 500

@app.route('/api/mst/submit-answers/<int:collection_id>', methods=['POST'])
@login_required
def submit_mst_answers(collection_id):
    """Submit answers for current MST module and get next module"""
    if not current_user.is_authenticated or not hasattr(current_user, 'user_type') or current_user.user_type != 'siswa':
        return jsonify({'success': False, 'message': 'Akses ditolak'}), 403
    
    try:
        data = request.get_json()
        if not data or 'answers' not in data:
            return jsonify({'success': False, 'message': 'Data jawaban tidak valid'}), 400
        
        answers = data['answers']  # [{'question_id': 1, 'answer': 'A'}, ...]
        
        # Get current MST status using new system
        mst_state = get_mst_state(current_user.id, collection_id)
        current_stage = mst_state.current_stage
        current_path = "LEGACY"  # Legacy path for backward compatibility
        
        # Process answers and calculate score
        correct_count = 0
        total_count = len(answers)
        
        for answer_data in answers:
            question_id = answer_data['question_id']
            user_answer = answer_data['answer']
            
            # Get question
            question = Question.query.get(question_id)
            if not question:
                continue
            
            # Check if answer is correct
            is_correct = (user_answer.strip().upper() == question.jawaban_benar.strip().upper())
            if is_correct:
                correct_count += 1
            
            # Save answer
            existing_answer = SiswaAnswer.query.filter_by(
                siswa_id=current_user.id,
                question_id=question_id
            ).first()
            
            if existing_answer:
                existing_answer.siswa_answer = user_answer
                existing_answer.is_correct = is_correct
                existing_answer.technology_level = question.technology_level
                existing_answer.difficulty = question.difficulty
                existing_answer.stage = current_stage
                existing_answer.answered_at = db.func.now()
            else:
                new_answer = SiswaAnswer(
                    siswa_id=current_user.id,
                    question_id=question_id,
                    siswa_answer=user_answer,
                    is_correct=is_correct,
                    level=question.level,  # Keep for compatibility
                    technology_level=question.technology_level,
                    difficulty=question.difficulty,
                    stage=current_stage
                )
                db.session.add(new_answer)
        
        # Calculate pass/fail (>50%)
        score_percentage = (correct_count / total_count) * 100 if total_count > 0 else 0
        passed = score_percentage > 50
        
        # Simple completion check - this appears to be legacy code
        # The new MST system uses different endpoints and logic
        diagnosis_text = f"Legacy Test Completed - Score: {score_percentage:.1f}%"
        
        return jsonify({
            'success': True,
            'test_complete': True,
            'diagnosis': diagnosis_text,
            'score': score_percentage,
            'message': f'Legacy test completed with score: {score_percentage:.1f}%'
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'Terjadi kesalahan: {str(e)}'}), 500

@app.route('/api/mst/status/<int:collection_id>', methods=['GET'])
@login_required
def get_mst_status(collection_id):
    """Get current MST status for a student"""
    if not current_user.is_authenticated or not hasattr(current_user, 'user_type') or current_user.user_type != 'siswa':
        return jsonify({'success': False, 'message': 'Akses ditolak'}), 403
    
    try:
        # Get MST status using new system
        mst_state = get_mst_state(current_user.id, collection_id)
        mst_status = {
            'stage': mst_state.current_stage,
            'level': mst_state.current_level,
            'questions_answered': mst_state.questions_answered or 0,
            'questions_correct': mst_state.questions_correct or 0,
            'test_completed': mst_state.test_completed,
            'final_diagnosis': mst_state.final_diagnosis
        }
        
        return jsonify({
            'success': True,
            'mst_status': mst_status,
            'technology_levels': TECHNOLOGY_LEVELS
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Terjadi kesalahan: {str(e)}'}), 500

# =========================================
# ENDPOINT: GET_QUESTIONS (Dipanggil oleh result.html dan guru.html)
# Mengambil semua data soal dari database
# =========================================
@app.route("/api/get_questions", methods=["GET"])
@login_required  
def get_questions_api():  
    # Pastikan yang mengakses endpoint ini adalah guru
    if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
        return jsonify({'success': False, 'message': 'Akses ditolak. Hanya guru yang dapat melihat semua soal.'}), 403

    try:
        # Ambil semua soal yang dimiliki oleh guru yang sedang login
        # Tambahkan filter guru_id di sini
        questions_raw = (
            Question.query
            .filter_by(guru_id=current_user.id, is_current=True)
            .order_by(Question.level.asc(), Question.id.asc())
            .all()
        )
        
        print(f"Fetched {len(questions_raw)} questions for guru_id {current_user.id} via API.")
        
        questions_data = []
        for q in questions_raw:
            options = []
            if q.options:
                try:
                    options = json.loads(q.options)
                except json.JSONDecodeError:
                    pass
            
            questions_data.append({
                "id": q.id,
                "guru_id": q.guru_id,
                "collection_id": q.collection_id,
                "level": q.level,
                "soal": q.soal,
                "jawaban_benar": q.jawaban_benar,
                "options": options,
                "question_type": q.question_type,
                "p": float(q.p),
                "difficulty": q.difficulty if hasattr(q, 'difficulty') and q.difficulty else p_value_to_difficulty(float(q.p)),
                "bobot": q.bobot,
                "explanation": q.explanation,
                "created_at": q.created_at.isoformat() if q.created_at else None,
                "version": q.version,
                "is_current": q.is_current,
                "is_validated": q.is_validated  
            })
            
        return jsonify({"success": True, "data": questions_data}), 200
        
    except Exception as e:
        print(f"Error getting questions via API: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "message": f"Error: {str(e)}"}), 500

@app.route('/api/guru', methods=['GET'])
@login_required
def get_all_teachers():
    try:
        # Only siswa should access this endpoint
        if not hasattr(current_user, 'user_type') or current_user.user_type != 'siswa':
            return jsonify({'success': False, 'message': 'Akses ditolak'}), 403
        
        # Get all teachers
        teachers = Guru.query.all()
        
        # Format teacher data
        teachers_data = []
        for teacher in teachers:
            teachers_data.append({
                'id': teacher.id,
                'nama': teacher.nama,
                'email': teacher.email
            })
        
        return jsonify({
            'success': True,
            'teachers': teachers_data
        })
        
    except Exception as e:
        print(f"Error getting all teachers: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500
    
@app.route('/dashboard')
@login_required
def dashboard():
    # Jika pengguna adalah siswa, ambil koleksi yang tersedia
    collections = []
    if hasattr(current_user, 'user_type') and current_user.user_type == 'siswa':
        try:
            # Ambil koleksi yang tersedia untuk siswa ini
            collections_data = db.session.query(
                QuestionCollection, Guru
            ).join(
                Guru, QuestionCollection.guru_id == Guru.id
            ).join(
                collection_students, collection_students.c.collection_id == QuestionCollection.id
            ).filter(
                collection_students.c.siswa_id == current_user.id
            ).all()
            
            # Format data koleksi
            for collection, guru in collections_data:
                # Cek apakah koleksi baru (dibuat dalam 1 hari terakhir)
                is_new = False
                if collection.created_at:
                    delta = datetime.datetime.now() - collection.created_at
                    is_new = delta.days < 1
                
                collections.append({
                    'id': collection.id,
                    'name': collection.name,
                    'is_new': is_new,
                    'description': collection.description,
                    'guru_id': guru.id,
                    'guru_name': guru.nama,
                    'guru_email': guru.email
                })
        except Exception as e:
            print(f"Error fetching collections: {str(e)}")
    
    return render_template('DashboardSiswa.html', collections=collections)
# =========================================
# VERSIONING ENDPOINTS (Tempatkan setelah definisi model)
# =========================================
# Pastikan hanya ada satu implementasi, hapus yang lama
@app.route('/api/questions/<int:id>', methods=['PUT', 'POST', 'DELETE'])
@login_required
def update_question(id):
    try:
        # Handle DELETE request
        if request.method == 'DELETE':
            # Inline delete handler with safety checks
            if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
                return jsonify({"success": False, "message": "Unauthorized"}), 403

            try:
                # Verify question belongs to current guru
                question = Question.query.filter_by(id=id, guru_id=current_user.id).first()
                if not question:
                    return jsonify({"success": False, "message": "Soal tidak ditemukan"}), 404

                # Count student answers
                student_answers_count = SiswaAnswer.query.filter_by(question_id=id).count()

                # Always remove relation to collections for this question belonging to this guru's collections
                CollectionQuestion.query.filter_by(question_id=id).delete()

                if student_answers_count > 0:
                    # If there are student answers, don't delete the question; mark as not current (archived) and detach from collections
                    question.is_current = False
                    db.session.commit()
                    return jsonify({
                        "success": True,
                        "message": f"Soal dihapus dari semua koleksi dan diarsipkan (tidak ditampilkan lagi). Tidak dihapus dari database karena sudah ada {student_answers_count} jawaban siswa."
                    })

                # No student answers: safe to delete all versions and the question
                QuestionVersion.query.filter_by(question_id=id).delete()
                db.session.delete(question)
                db.session.commit()

                return jsonify({
                    "success": True,
                    "message": "Soal berhasil dihapus dari database"
                })
            except Exception as e:
                db.session.rollback()
                app.logger.error(f"Error deleting question: {str(e)}")
                return jsonify({"success": False, "message": str(e)}), 500
        
        # Log request method untuk debugging
        print(f"Request method: {request.method}")
        print(f"Request headers: {request.headers}")
        
        if request.method == 'POST' and request.headers.get('X-HTTP-Method-Override') == 'PUT':
            print("Using X-HTTP-Method-Override: PUT")
            # Lanjutkan sebagai PUT request
        
        data = request.get_json()
        print(f"Request data: {data}")
        
        # Dapatkan pertanyaan yang akan diupdate
        question = Question.query.get(id)
        if not question:
            return jsonify({'success': False, 'message': 'Pertanyaan tidak ditemukan'}), 404
        
        # Buat versi sebelum update
        prev_version = QuestionVersion(
            question_id=question.id,
            soal=question.soal,
            options=question.options,
            jawaban_benar=question.jawaban_benar,
            explanation=question.explanation,
            p=question.p,
            version=question.version  # Gunakan versi saat ini
        )
        db.session.add(prev_version)
        
        # Update pertanyaan utama
        question.soal = data.get('soal', question.soal)
        question.options = json.dumps(data.get('options')) if 'options' in data else question.options
        question.jawaban_benar = data.get('jawaban_benar', question.jawaban_benar)
        question.explanation = data.get('explanation', question.explanation)
        question.p = data.get('p', question.p)
        
        # Increment versi
        question.version += 1
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'question': {
                'id': question.id,
                'version': question.version,
                'soal': question.soal,
                'options': json.loads(question.options) if question.options else []
            }
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"Error updating question: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/questions/<int:id>/versions', methods=['GET'])
@login_required
def get_question_versions(id):
    try:
        print(f"Fetching versions for question {id}")
        versions = QuestionVersion.query.filter_by(question_id=id).order_by(QuestionVersion.version.desc()).all()
        print(f"Found {len(versions)} versions")
        
        result = []
        for version in versions:
            result.append({
                'id': version.id,
                'soal': version.soal,
                'options': json.loads(version.options) if version.options else [],
                'jawaban_benar': version.jawaban_benar,
                'explanation': version.explanation,
                'p': version.p,
                'version': version.version,
                'created_at': version.created_at.isoformat()
            })
            
        return jsonify({'success': True, 'versions': result})
    except Exception as e:
        print(f"Error getting versions: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/questions/in_collection', methods=['GET'])
@login_required
def get_questions_in_collection():
    try:
        # Periksa apakah collection_id disediakan sebagai parameter query
        collection_id = request.args.get('collection_id')
        
        if collection_id:
            # Filter berdasarkan collection_id tertentu
            collection_question_ids = db.session.query(CollectionQuestion.question_id)\
                .filter(CollectionQuestion.collection_id == collection_id).all()
        else:
            # Dapatkan semua ID soal yang ada dalam collection apapun
            collection_question_ids = db.session.query(CollectionQuestion.question_id).distinct().all()
            
        collection_question_ids = [id[0] for id in collection_question_ids]
        
        # Jika perlu detail soal, bukan hanya ID
        include_details = request.args.get('include_details') == 'true'
        if include_details:
            # Ambil detail soal berdasarkan ID yang sudah didapat
            questions = db.session.query(Question)\
                .filter(Question.id.in_(collection_question_ids))\
                .all()
                
            # Format data soal
            question_data = []
            for q in questions:
                options = []
                if hasattr(q, 'options') and q.options:
                    try:
                        options = json.loads(q.options)
                    except:
                        pass
                
                question_data.append({
                    'id': q.id,
                    'level': q.level,
                    'soal': q.soal,
                    'jawaban_benar': q.jawaban_benar,
                    'options': options,
                    'question_type': q.question_type if hasattr(q, 'question_type') else 'multiple_choice',
                    'p': float(q.p) if q.p else 0.5,
                    'bobot': q.bobot or 1,
                    'explanation': q.explanation or ''
                })
                
            return jsonify({
                'success': True,
                'questions': question_data,
                'question_ids': collection_question_ids
            })
        else:
            # Hanya mengembalikan ID soal
            return jsonify({
                'success': True,
                'question_ids': collection_question_ids
            })
            
    except Exception as e:
        print(f"Error getting questions in collection: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500
    
@app.route('/api/collections/<int:collection_id>/status', methods=['POST'])
@login_required
def set_collection_status(collection_id):
    try:
        # Verifikasi kepemilikan koleksi
        collection = QuestionCollection.query.filter_by(id=collection_id, guru_id=current_user.id).first()
        if not collection:
            return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
        
        data = request.json
        is_active = data.get('isActive', False)
        
        # Update status
        collection.is_active = bool(is_active)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'Koleksi berhasil {"diaktifkan" if is_active else "dinonaktifkan"}'
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500
    

@app.route('/guru/collections')
@login_required
def collection_page():
    # Verifikasi bahwa pengguna adalah guru
    if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
        flash('Halaman ini hanya untuk guru', 'error')
        return redirect(url_for('login'))
    
    return render_template('guru/collections.html')

@app.route('/api/export_result_excel', methods=['GET'])
@login_required
def export_result_excel():
    try:
        user_id = request.args.get('user_id')
        collection_id = request.args.get('collection_id')
        
        # Validasi parameter
        if not user_id or not collection_id:
            return jsonify({'success': False, 'message': 'Parameter tidak lengkap'}), 400
            
        # Ambil info siswa
        student = Siswa.query.get(user_id)
        if not student:
            return jsonify({'success': False, 'message': 'Siswa tidak ditemukan'}), 404
            
        # Ambil info koleksi
        collection = QuestionCollection.query.get(collection_id)
        if not collection:
            return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
            
        # Ambil hasil tes
        result = SiswaResult.query.filter_by(
            siswa_id=user_id,
            collection_id=collection_id
        ).first()
        
        if not result:
            return jsonify({'success': False, 'message': 'Hasil tes tidak ditemukan'}), 404
            
        # Buat file Excel
        output = BytesIO()
        
        # Format untuk membuat Excel
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            # Format untuk judul dan header
            workbook = writer.book
            title_format = workbook.add_format({
                'bold': True,
                'font_size': 16,
                'align': 'center',
                'valign': 'vcenter'
            })
            header_format = workbook.add_format({
                'bold': True,
                'text_wrap': True,
                'valign': 'top',
                'fg_color': '#D7E4BC',
                'border': 1
            })
            label_format = workbook.add_format({
                'bold': True,
                'align': 'right',
                'border': 1
            })
            data_format = workbook.add_format({
                'align': 'left',
                'border': 1
            })
            
            # Buat worksheet
            worksheet = workbook.add_worksheet('Ringkasan')
            
            # Judul utama
            worksheet.merge_range('A1:H1', f"HASIL TES DIAGNOSTIK - {collection.name}", title_format)
            
            # Detail siswa dalam format label-data
            worksheet.write('A3', 'Nama:', label_format)
            worksheet.write('B3', student.nama, data_format)
            
            worksheet.write('A4', 'Kelas:', label_format)
            worksheet.write('B4', student.kelas, data_format)
            
            worksheet.write('A5', 'Koleksi Soal:', label_format)
            worksheet.write('B5', collection.name, data_format)
            
            worksheet.write('D3', 'Level Terakhir:', label_format)
            norm_level = normalize_level(result.current_level)
            worksheet.write('E3', norm_level, data_format)
            
            worksheet.write('D4', 'Jawaban Benar:', label_format)
            worksheet.write('E4', result.correct, data_format)
            
            worksheet.write('D5', 'Jawaban Salah:', label_format)
            worksheet.write('E5', result.incorrect, data_format)
            
            total = result.correct + result.incorrect
            accuracy = f"{round((result.correct / total * 100), 2)}%" if total > 0 else "0%"
            
            # Tambahkan ringkas Nama Level (Taksonomi Teknologi) dan Basis Pengetahuan
            worksheet.write('G3', 'Nama Level:', label_format)
            try:
                worksheet.write('H3', TECHNOLOGY_LEVELS.get(int(norm_level), str(norm_level)), data_format)
            except Exception:
                worksheet.write('H3', str(norm_level), data_format)

            worksheet.write('G4', 'Basis Pengetahuan:', label_format)
            worksheet.write('H4', get_technology_taxonomy_short(norm_level), data_format)

            worksheet.write('G5', 'Total Soal:', label_format)
            worksheet.write('H5', total, data_format)
            
            worksheet.write('G6', 'Akurasi:', label_format)
            worksheet.write('H6', accuracy, data_format)
            
            worksheet.write('G7', 'Tanggal Pengerjaan:', label_format)
            worksheet.write('H7', result.updated_at.strftime('%d-%m-%Y %H:%M:%S') if result.updated_at else datetime.datetime.now().strftime('%d-%m-%Y %H:%M:%S'), data_format)

            # Status proses yang lebih spesifik (termasuk berhenti di Level X)
            progress_status = "Belum Mulai"
            if total > 0 and norm_level == 5:
                progress_status = "Selesai (Level 5)"
            elif total > 0 and 1 <= int(norm_level) < 5:
                progress_status = f"Berhenti di Level {int(norm_level)}"
            worksheet.write('G8', 'Status Proses:', label_format)
            worksheet.write('H8', progress_status, data_format)
            
            # --- PERUBAHAN: Tambah keterangan level singkat ---
            worksheet.write('A9', 'Keterangan Level:', label_format)
            level_desc = get_level_description_short(norm_level)
            data_format_wrap = workbook.add_format({
                'align': 'left',
                'border': 1,
                'text_wrap': True,
                'valign': 'vcenter'
            })
            worksheet.merge_range('B9:H9', level_desc, data_format_wrap)
            worksheet.set_row(8, 40)  # Atur tinggi baris
            # --- PERUBAHAN SELESAI ---
            
            # Sesuaikan lebar kolom dengan spacing yang lebih baik
            worksheet.set_column('A:A', 18)  # Label kolom
            worksheet.set_column('B:B', 28)  # Data siswa
            worksheet.set_column('C:C', 3)   # Spacer
            worksheet.set_column('D:D', 18)  # Label hasil
            worksheet.set_column('E:E', 15)  # Data hasil
            worksheet.set_column('F:F', 3)   # Spacer
            worksheet.set_column('G:G', 20)  # Label statistik
            worksheet.set_column('H:H', 25)  # Data statistik
            
            # Ambil riwayat jawaban
            answers = get_student_answer_history(user_id, collection_id)
            
            # Buat sheet untuk detail jawaban
            if answers:
                # Buat dataframe untuk jawaban
                answers_data = []
                for answer in answers:
                    ans_level = normalize_level(answer['level']) if 'level' in answer else None
                    answers_data.append({
                        'Soal': answer['question'],
                        'Jawaban Siswa': answer['user_answer'],
                        'Jawaban Benar': answer['correct_answer'],
                        'Status': 'Benar' if answer['is_correct'] else 'Salah',
                        'Level': ans_level,
                        'Nama Level': TECHNOLOGY_LEVELS.get(int(ans_level), str(ans_level)),
                        'Basis Pengetahuan': get_technology_taxonomy_short(ans_level),
                        'Waktu Menjawab': answer['answered_at']
                    })
                
                df_answers = pd.DataFrame(answers_data)
                # Tulis detail jawaban ke sheet kedua
                df_answers.to_excel(writer, sheet_name='Detail Jawaban', index=False)
                
                # Format sheet Detail Jawaban dengan spacing yang lebih baik
                worksheet2 = writer.sheets['Detail Jawaban']
                
                # Format headers
                header_format2 = workbook.add_format({
                    'bold': True,
                    'text_wrap': True,
                    'valign': 'top',
                    'fg_color': '#D7E4BC',
                    'border': 1
                })
                
                # Apply header format
                for col_num, value in enumerate(df_answers.columns.values):
                    worksheet2.write(0, col_num, value, header_format2)
                
                # Set column widths with proper spacing
                worksheet2.set_column('A:A', 50)  # Soal (wider untuk pertanyaan panjang)
                worksheet2.set_column('B:B', 25)  # Jawaban Siswa
                worksheet2.set_column('C:C', 25)  # Jawaban Benar
                worksheet2.set_column('D:D', 12)  # Status
                worksheet2.set_column('E:E', 8)   # Level
                worksheet2.set_column('F:F', 24)  # Nama Level
                worksheet2.set_column('G:G', 24)  # Basis Pengetahuan
                worksheet2.set_column('H:H', 20)  # Waktu Menjawab
                
                # Set row height for data rows
                for row in range(1, len(df_answers) + 1):
                    worksheet2.set_row(row, 25)
            else:
                # Tetap buat sheet Detail Jawaban dengan catatan jika belum ada jawaban tersimpan
                df_empty = pd.DataFrame([{
                    'Keterangan': 'Belum ada jawaban yang tersimpan untuk koleksi ini.'
                }])
                df_empty.to_excel(writer, sheet_name='Detail Jawaban', index=False)
                worksheet2 = writer.sheets['Detail Jawaban']
                for col_num, value in enumerate(df_empty.columns.values):
                    worksheet2.write(0, col_num, value, header_format)
                worksheet2.set_column('A:A', 60)
        
        # Kembali ke awal file
        output.seek(0)
        
        # Buat nama file
        filename = f"Hasil_Tes_{student.nama}_{collection.name}_{datetime.datetime.now().strftime('%Y%m%d')}.xlsx"
        
        # Kirim file
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        print(f"Error mengekspor ke Excel: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f"Error: {str(e)}"}), 500

@app.route('/api/export_result_excel_teacher', methods=['GET'])
@login_required
def export_result_excel_teacher():
    try:
        # Get parameters
        export_type = request.args.get('type', 'individual')  # 'individual' or 'class'
        user_id = request.args.get('user_id')
        class_id = request.args.get('class_id')
        collection_id = request.args.get('collection_id')
        
        # Validate parameters
        if not collection_id:
            return jsonify({'success': False, 'message': 'Parameter collection_id is required'}), 400
            
        if export_type == 'individual' and not user_id:
            return jsonify({'success': False, 'message': 'Parameter user_id is required for individual export'}), 400
            
        if export_type == 'class' and not class_id:
            return jsonify({'success': False, 'message': 'Parameter class_id is required for class export'}), 400
            
        # Create a BytesIO object to store the Excel file
        output = BytesIO()
        
        # Collection information
        collection = QuestionCollection.query.get(collection_id)
        if not collection:
            return jsonify({'success': False, 'message': 'Collection not found'}), 404
            
        # Create Excel with appropriate data based on type
        if export_type == 'individual':
            # Individual student export
            student = Siswa.query.get(user_id)
            if not student:
                return jsonify({'success': False, 'message': 'Student not found'}), 404
                
            # Get student results
            result = SiswaResult.query.filter_by(
                siswa_id=user_id,
                collection_id=collection_id
            ).first()
            
            if not result:
                return jsonify({'success': False, 'message': 'No results found for this student'}), 404
                
            # Create Excel file with student data
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                # Format for styling
                workbook = writer.book
                title_format = workbook.add_format({
                    'bold': True,
                    'font_size': 16,
                    'align': 'center',
                    'valign': 'vcenter'
                })
                header_format = workbook.add_format({
                    'bold': True,
                    'text_wrap': True,
                    'valign': 'top',
                    'fg_color': '#D7E4BC',
                    'border': 1
                })
                label_format = workbook.add_format({
                    'bold': True,
                    'align': 'right',
                    'border': 1
                })
                data_format = workbook.add_format({
                    'align': 'left',
                    'border': 1
                })
                
                # Create summary worksheet
                worksheet = workbook.add_worksheet('Summary')
                
                # Add title
                worksheet.merge_range('A1:H1', f"HASIL TES DIAGNOSTIK - {collection.name}", title_format)
                
                # Student details
                worksheet.write('A3', 'Nama:', label_format)
                worksheet.write('B3', student.nama, data_format)
                
                worksheet.write('A4', 'Kelas:', label_format)
                worksheet.write('B4', student.kelas, data_format)
                
                worksheet.write('A5', 'Koleksi Soal:', label_format)
                worksheet.write('B5', collection.name, data_format)
                
                # Result details
                worksheet.write('D3', 'Level Terakhir:', label_format)
                norm_level = normalize_level(result.current_level)
                worksheet.write('E3', norm_level, data_format)
                
                worksheet.write('D4', 'Jawaban Benar:', label_format)
                worksheet.write('E4', result.correct, data_format)
                
                worksheet.write('D5', 'Jawaban Salah:', label_format)
                worksheet.write('E5', result.incorrect, data_format)
                
                # Statistics
                total = result.correct + result.incorrect
                accuracy = f"{round((result.correct / total * 100), 2)}%" if total > 0 else "0%"
                
                # Tambah Nama Level dan Basis Pengetahuan
                worksheet.write('G3', 'Nama Level:', label_format)
                try:
                    worksheet.write('H3', TECHNOLOGY_LEVELS.get(int(norm_level), str(norm_level)), data_format)
                except Exception:
                    worksheet.write('H3', str(norm_level), data_format)
                worksheet.write('G4', 'Basis Pengetahuan:', label_format)
                worksheet.write('H4', get_technology_taxonomy_short(norm_level), data_format)

                worksheet.write('G5', 'Total Soal:', label_format)
                worksheet.write('H5', total, data_format)
                
                worksheet.write('G6', 'Akurasi:', label_format)
                worksheet.write('H6', accuracy, data_format)
                
                worksheet.write('G7', 'Tanggal Pengerjaan:', label_format)
                worksheet.write('H7', result.updated_at.strftime('%d-%m-%Y %H:%M:%S') if result.updated_at else datetime.datetime.now().strftime('%d-%m-%Y %H:%M:%S'), data_format)

                # Status proses yang lebih spesifik
                progress_status = "Belum Mulai"
                if total > 0 and norm_level == 5:
                    progress_status = "Selesai (Level 5)"
                elif total > 0 and 1 <= int(norm_level) < 5:
                    progress_status = f"Berhenti di Level {int(norm_level)}"
                worksheet.write('G8', 'Status Proses:', label_format)
                worksheet.write('H8', progress_status, data_format)
                
                # --- PERUBAHAN: Tambah keterangan level singkat ---
                worksheet.write('A9', 'Keterangan Level:', label_format)
                level_desc = get_level_description_short(norm_level)
                data_format_wrap = workbook.add_format({
                    'align': 'left',
                    'border': 1,
                    'text_wrap': True,
                    'valign': 'vcenter'
                })
                worksheet.merge_range('B9:H9', level_desc, data_format_wrap)
                worksheet.set_row(8, 40)  # Atur tinggi baris
                # --- PERUBAHAN SELESAI ---
                
                # Format columns
                worksheet.set_column('A:A', 15)
                worksheet.set_column('B:B', 25)
                worksheet.set_column('C:C', 5)
                worksheet.set_column('D:D', 15)
                worksheet.set_column('E:E', 15)
                worksheet.set_column('F:F', 5)
                worksheet.set_column('G:G', 18)
                worksheet.set_column('H:H', 25)
                
                # Get answer history
                answers = get_student_answer_history(user_id, collection_id)
                
                # Add answers to another sheet
                if answers:
                    # Create dataframe for answers
                    answers_data = []
                    for answer in answers:
                        ans_level = normalize_level(answer['level']) if 'level' in answer else None
                        answers_data.append({
                            'Soal': answer['question'],
                            'Jawaban Siswa': answer['user_answer'],
                            'Jawaban Benar': answer['correct_answer'],
                            'Status': 'Benar' if answer['is_correct'] else 'Salah',
                            'Level': ans_level,
                            'Nama Level': TECHNOLOGY_LEVELS.get(int(ans_level), str(ans_level)),
                            'Basis Pengetahuan': get_technology_taxonomy_short(ans_level),
                            'Waktu Menjawab': answer['answered_at']
                        })
                    
                    # Create DataFrame and write to Excel
                    if answers_data:
                        df_answers = pd.DataFrame(answers_data)
                        df_answers.to_excel(writer, sheet_name='Detail Jawaban', index=False)
                        
                        # Format Detail Jawaban sheet dengan spacing yang lebih baik
                        answer_sheet = writer.sheets['Detail Jawaban']
                        
                        # Apply header format
                        header_format_answer = workbook.add_format({
                            'bold': True,
                            'text_wrap': True,
                            'valign': 'top',
                            'fg_color': '#D7E4BC',
                            'border': 1
                        })
                        
                        for col_num, value in enumerate(df_answers.columns.values):
                            answer_sheet.write(0, col_num, value, header_format_answer)
                        
                        # Set column widths with proper spacing
                        answer_sheet.set_column('A:A', 50)  # Soal
                        answer_sheet.set_column('B:B', 25)  # Jawaban Siswa
                        answer_sheet.set_column('C:C', 25)  # Jawaban Benar
                        answer_sheet.set_column('D:D', 12)  # Status
                        answer_sheet.set_column('E:E', 8)   # Level
                        answer_sheet.set_column('F:F', 24)  # Nama Level
                        answer_sheet.set_column('G:G', 24)  # Basis Pengetahuan
                        answer_sheet.set_column('H:H', 20)  # Waktu Menjawab
                        
                        # Set row height for data rows
                        for row in range(1, len(df_answers) + 1):
                            answer_sheet.set_row(row, 25)
                else:
                    # Buat sheet kosong dengan keterangan bila tidak ada jawaban
                    df_empty = pd.DataFrame([{
                        'Keterangan': 'Belum ada jawaban yang tersimpan untuk koleksi ini.'
                    }])
                    df_empty.to_excel(writer, sheet_name='Detail Jawaban', index=False)
                    answer_sheet = writer.sheets['Detail Jawaban']
                    for col_num, value in enumerate(df_empty.columns.values):
                        answer_sheet.write(0, col_num, value, header_format)
                    answer_sheet.set_column('A:A', 60)
            
            # Generate filename
            filename = f"Hasil_Tes_{student.nama}_{collection.name}_{datetime.datetime.now().strftime('%Y%m%d')}.xlsx"
            
        else:
            # Class export
            # Get all students in the class
            students = Siswa.query.filter_by(kelas=class_id).all()
            if not students:
                return jsonify({'success': False, 'message': 'No students found in this class'}), 404
            
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                workbook = writer.book
                title_format = workbook.add_format({
                    'bold': True,
                    'font_size': 16,
                    'align': 'center',
                    'valign': 'vcenter'
                })
                header_format = workbook.add_format({
                    'bold': True,
                    'text_wrap': True,
                    'valign': 'top',
                    'fg_color': '#D7E4BC',
                    'border': 1
                })
                
                # Summary sheet
                worksheet = workbook.add_worksheet('Ringkasan Kelas')
                
                # Add title - disesuaikan untuk kolom baru
                worksheet.merge_range('A1:M1', f"RINGKASAN HASIL TES KELAS {class_id} - {collection.name}", title_format)
                
                # Set up headers for class summary - gunakan Nama Level dan Basis Pengetahuan
                headers = ['No', 'Nama', 'Kelas', 'Level', 'Nama Level', 'Keterangan Level', 'Basis Pengetahuan', 'Jawaban Benar', 'Jawaban Salah', 'Total', 'Akurasi (%)', 'Status', 'Tanggal Update']
                for col, header in enumerate(headers):
                    worksheet.write(2, col, header, header_format)
                
                # Gather student data
                class_data = []
                row = 3
                student_number = 1
                
                total_correct = 0
                total_incorrect = 0
                level_distribution = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
                completed_count = 0
                respondent_count = 0
                
                for student in students:
                    # Get student result
                    result = SiswaResult.query.filter_by(
                        siswa_id=student.id,
                        collection_id=collection_id
                    ).first()
                    
                    if result:
                        # Update statistics
                        correct = result.correct or 0
                        incorrect = result.incorrect or 0
                        total = correct + incorrect
                        accuracy = round((correct / total * 100), 2) if total > 0 else 0
                        current_level = normalize_level(result.current_level)
                        if total <= 0:
                            continue  # Skip students who haven't answered any question
                        
                        # Track level distribution
                        if current_level in level_distribution:
                            level_distribution[current_level] += 1
                        
                        # Track completion status
                        is_completed = (current_level == 5)
                        if is_completed:
                            completed_count += 1
                        respondent_count += 1
                            
                        # Update totals
                        total_correct += correct
                        total_incorrect += incorrect
                        
                        # --- PERUBAHAN: Tambah keterangan level singkat ---
                        level_desc = get_level_description_short(current_level)
                        # --- PERUBAHAN SELESAI ---
                        
                        # Write row data dengan format yang rapi
                        data_format_teacher = workbook.add_format({
                            'align': 'left',
                            'border': 1
                        })
                        description_format_teacher = workbook.add_format({
                            'align': 'left',
                            'border': 1,
                            'text_wrap': True,
                            'valign': 'top'
                        })
                        num_format_teacher = workbook.add_format({
                            'align': 'right',
                            'border': 1,
                            'num_format': '0'
                        })
                        
                        worksheet.write(row, 0, student_number, num_format_teacher)
                        worksheet.write(row, 1, student.nama, data_format_teacher)
                        worksheet.write(row, 2, student.kelas, data_format_teacher)
                        worksheet.write(row, 3, current_level, num_format_teacher)
                        worksheet.write(row, 4, TECHNOLOGY_LEVELS.get(int(current_level), str(current_level)), data_format_teacher)
                        worksheet.write(row, 5, level_desc, description_format_teacher)
                        worksheet.write(row, 6, get_technology_taxonomy_short(current_level), data_format_teacher)
                        worksheet.write(row, 7, correct, num_format_teacher)
                        worksheet.write(row, 8, incorrect, num_format_teacher)
                        worksheet.write(row, 9, total, num_format_teacher)
                        worksheet.write(row, 10, accuracy, num_format_teacher)
                        status_text = "Selesai" if is_completed else (f"Berhenti di Level {int(current_level)}" if total > 0 else "Belum Mulai")
                        worksheet.write(row, 11, status_text, data_format_teacher)
                        worksheet.write(row, 12, result.updated_at.strftime('%d-%m-%Y %H:%M:%S') if result.updated_at else '', data_format_teacher)
                        # --- PERUBAHAN SELESAI ---
                        
                        # Set row height for better readability
                        worksheet.set_row(row-1, 25)
                        
                        # Prepare data for DataFrame
                        class_data.append({
                            'Nama': student.nama,
                            'Kelas': student.kelas,
                            'Level': current_level,
                            'Nama Level': TECHNOLOGY_LEVELS.get(int(current_level), str(current_level)),
                            'Keterangan Level': level_desc,
                            'Jawaban Benar': correct,
                            'Jawaban Salah': incorrect,
                            'Total': total,
                            'Akurasi (%)': accuracy,
                            'Status': status_text,
                            'Tanggal Update': result.updated_at.strftime('%d-%m-%Y %H:%M:%S') if result.updated_at else ''
                        })
                        
                        row += 1
                        student_number += 1
                
                # Add totals row
                total_answers = total_correct + total_incorrect
                avg_accuracy = round((total_correct / total_answers * 100), 2) if total_answers > 0 else 0
                completion_rate = round((completed_count / respondent_count * 100), 2) if respondent_count > 0 else 0
                
                # Skip a row
                row += 1
                
                # Summary row - sesuaikan posisi kolom untuk Taksonomi Teknologi
                bold_format = workbook.add_format({'bold': True})
                worksheet.write(row, 0, "SUMMARY", bold_format)
                worksheet.write(row, 1, f"Total Responden: {respondent_count}")
                worksheet.write(row, 7, f"Total Benar: {total_correct}")
                worksheet.write(row, 8, f"Total Salah: {total_incorrect}")
                worksheet.write(row, 9, f"Total Jawaban: {total_answers}")
                worksheet.write(row, 10, f"Akurasi Rata-rata: {avg_accuracy}%")
                worksheet.write(row, 11, f"Penyelesaian: {completion_rate}%")
                
                # Format columns dengan spacing yang lebih baik - tambah Taksonomi Teknologi
                worksheet.set_column('A:A', 5)   # No
                worksheet.set_column('B:B', 25)  # Nama
                worksheet.set_column('C:C', 12)  # Kelas
                worksheet.set_column('D:D', 8)   # Level
                worksheet.set_column('E:E', 18)  # Nama Level
                worksheet.set_column('F:F', 22)  # Keterangan Level (singkat)
                worksheet.set_column('G:G', 20)  # Basis Pengetahuan (singkat)
                worksheet.set_column('H:H', 15)  # Jawaban Benar
                worksheet.set_column('I:I', 15)  # Jawaban Salah
                worksheet.set_column('J:J', 10)  # Total
                worksheet.set_column('K:K', 12)  # Akurasi (%)
                worksheet.set_column('L:L', 18)  # Status
                worksheet.set_column('M:M', 20)  # Tanggal Update
                
                # level Distribution Chart Sheet
                level_sheet = workbook.add_worksheet('Distribusi Level')
                
                # Create level distribution data
                level_data = []
                for level in [1,2,3,4,5]:
                    count = level_distribution.get(level, 0)
                    level_data.append({
                        'level': f"Level {level}",
                        'Jumlah Siswa': count,
                        'Persentase': round((count / respondent_count * 100), 2) if respondent_count > 0 else 0,
                        'Keterangan': get_level_description_short(level),
                        'Basis Pengetahuan': get_technology_taxonomy_short(level)
                    })
                
                # Write level distribution data
                level_sheet.merge_range('A1:E1', f"DISTRIBUSI LEVEL - {collection.name}", title_format)
                level_headers = ['Level', 'Jumlah Siswa', 'Persentase (%)', 'Keterangan', 'Basis Pengetahuan']
                for col, header in enumerate(level_headers):
                    level_sheet.write(2, col, header, header_format)
                
                # Create format for description column
                desc_format = workbook.add_format({
                    'align': 'left',
                    'border': 1,
                    'text_wrap': True,
                    'valign': 'top'
                })
                
                for i, data in enumerate(level_data):
                    level_sheet.write(i + 3, 0, data['level'], data_format_teacher)
                    level_sheet.write(i + 3, 1, data['Jumlah Siswa'], num_format_teacher)
                    level_sheet.write(i + 3, 2, data['Persentase'], num_format_teacher)
                    level_sheet.write(i + 3, 3, data['Keterangan'], desc_format)
                    level_sheet.write(i + 3, 4, data['Basis Pengetahuan'], data_format_teacher)
                    level_sheet.set_row(i + 3, 25)  # Set row height
                
                # Format columns dengan spacing yang lebih baik
                level_sheet.set_column('A:A', 12)  # Level
                level_sheet.set_column('B:B', 15)  # Jumlah Siswa
                level_sheet.set_column('C:C', 15)  # Persentase (%)
                level_sheet.set_column('D:D', 45)  # Keterangan
                level_sheet.set_column('E:E', 25)  # Taksonomi Teknologi
                
                # Create a chart for level distribution
                chart = workbook.add_chart({'type': 'column'})
                
                # Add data series - force include Level 1..5 rows (3..7)
                chart.add_series({
                    'name': 'Jumlah Siswa',
                    'categories': ['Distribusi Level', 3, 0, 7, 0],
                    'values': ['Distribusi Level', 3, 1, 7, 1],
                    'data_labels': {'value': True}
                })
                
                # Configure chart
                chart.set_title({'name': 'Distribusi Level'})
                chart.set_x_axis({'name': 'Level'})
                chart.set_y_axis({'name': 'Jumlah Siswa'})
                
                # Insert chart
                level_sheet.insert_chart('E3', chart, {'x_offset': 25, 'y_offset': 10, 'x_scale': 1.5, 'y_scale': 1.5})
                
                # If we have student data, create a full class data sheet
                if class_data:
                    df_class = pd.DataFrame(class_data)
                    df_class.to_excel(writer, sheet_name='Data Lengkap', index=False)
                    
                    # Format columns
                    class_sheet = writer.sheets['Data Lengkap']
                    class_sheet.set_column('A:A', 25)
                    class_sheet.set_column('B:B', 10)
                    class_sheet.set_column('C:C', 14)
                    class_sheet.set_column('D:D', 18)
                    class_sheet.set_column('E:E', 22)
                    class_sheet.set_column('F:F', 20)
                    class_sheet.set_column('G:G', 12)
                    class_sheet.set_column('H:H', 18)
            
            # Generate filename
            filename = f"Hasil_Kelas_{class_id}_{collection.name}_{datetime.datetime.now().strftime('%Y%m%d')}.xlsx"
        
        # Reset file pointer to beginning
        output.seek(0)
        
        # Send file as attachment
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        print(f"Error exporting to Excel: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f"Error: {str(e)}"}), 500

@app.route('/api/export_all_collection_data', methods=['GET'])
@login_required
def export_all_collection_data():
    try:
        collection_id = request.args.get('collection_id')
        
        if not collection_id:
            return jsonify({'success': False, 'message': 'Collection ID is required'}), 400
            
        # Get collection info
        collection = QuestionCollection.query.get(collection_id)
        if not collection:
            return jsonify({'success': False, 'message': 'Collection not found'}), 404
            
        # Verify ownership
        if collection.guru_id != current_user.id:
            return jsonify({'success': False, 'message': 'You do not have permission to export this collection'}), 403
            
        # Get all students assigned to this collection
        students_query = db.session.query(Siswa).join(
            collection_students,
            collection_students.c.siswa_id == Siswa.id
        ).filter(
            collection_students.c.collection_id == collection_id
        ).all()
        
        if not students_query:
            return jsonify({'success': False, 'message': 'No students assigned to this collection'}), 404
            
        # Create BytesIO object for the Excel file
        output = BytesIO()
        
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            workbook = writer.book
            
            # ===== FORMAT DEFINITIONS =====
            title_format = workbook.add_format({
                'bold': True,
                'font_size': 16,
                'align': 'center',
                'valign': 'vcenter'
            })
            header_format = workbook.add_format({
                'bold': True,
                'text_wrap': True,
                'valign': 'top',
                'fg_color': '#D7E4BC',
                'border': 1
            })
            label_format = workbook.add_format({
                'bold': True,
                'align': 'right',
                'border': 1
            })
            data_format = workbook.add_format({
                'align': 'left',
                'border': 1
            })
            num_format = workbook.add_format({
                'align': 'right',
                'border': 1,
                'num_format': '0'
            })
            percent_format = workbook.add_format({
                'align': 'right',
                'border': 1,
                'num_format': '0.00%'
            })
            date_format = workbook.add_format({
                'align': 'left',
                'border': 1,
                'num_format': 'dd/mm/yyyy hh:mm:ss'
            })
            heading_format = workbook.add_format({
                'bold': True,
                'font_size': 12,
                'align': 'left',
                'valign': 'vcenter'
            })
            
            # ===== COLLECTION SUMMARY SHEET =====
            summary_sheet = workbook.add_worksheet('Ringkasan Koleksi')
            
            # Title
            summary_sheet.merge_range('A1:H1', f"RINGKASAN KOLEKSI - {collection.name}", title_format)
            
            # Collection details
            summary_sheet.write('A3', 'Nama Koleksi:', label_format)
            summary_sheet.write('B3', collection.name, data_format)
            
            summary_sheet.write('A4', 'Deskripsi:', label_format)
            summary_sheet.write('B4', collection.description or 'Tidak ada deskripsi', data_format)
            
            summary_sheet.write('A5', 'Guru:', label_format)
            guru = Guru.query.get(collection.guru_id)
            summary_sheet.write('B5', guru.nama if guru else 'Unknown', data_format)
            
            summary_sheet.write('A6', 'Tanggal Dibuat:', label_format)
            summary_sheet.write('B6', collection.created_at.strftime('%d-%m-%Y %H:%M:%S') if collection.created_at else 'Unknown', data_format)
            
            # Calculate overall statistics for respondents only
            total_correct = 0
            total_incorrect = 0
            completed_count = 0
            level_distribution = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
            respondent_count = 0
            
            for student in students_query:
                # Get result
                result = SiswaResult.query.filter_by(
                    siswa_id=student.id,
                    collection_id=collection_id
                ).first()
                
                if result:
                    correct = result.correct or 0
                    incorrect = result.incorrect or 0
                    total = correct + incorrect
                    if total <= 0:
                        continue  # Skip non-respondents
                    total_correct += correct
                    total_incorrect += incorrect
                    
                    # Track level distribution
                    level = normalize_level(result.current_level)
                    if level in level_distribution:
                        level_distribution[level] += 1
                    
                    # Check if completed
                    if level == 5:
                        completed_count += 1
                    respondent_count += 1
            
            # Overall statistics (respondents only)
            total_answers = total_correct + total_incorrect
            avg_accuracy = (total_correct / total_answers * 100) if total_answers > 0 else 0
            completion_rate = (completed_count / respondent_count * 100) if respondent_count > 0 else 0
            
            summary_sheet.write('D3', 'Total Responden:', label_format)
            summary_sheet.write('E3', respondent_count, num_format)
            
            summary_sheet.write('D4', 'Jawaban Benar:', label_format)
            summary_sheet.write('E4', total_correct, num_format)
            
            summary_sheet.write('D5', 'Jawaban Salah:', label_format)
            summary_sheet.write('E5', total_incorrect, num_format)
            
            summary_sheet.write('D6', 'Total Jawaban:', label_format)
            summary_sheet.write('E6', total_answers, num_format)
            
            summary_sheet.write('G3', 'Akurasi Rata-rata:', label_format)
            summary_sheet.write('H3', avg_accuracy / 100, percent_format)
            
            summary_sheet.write('G4', 'Tingkat Penyelesaian:', label_format)
            summary_sheet.write('H4', completion_rate / 100, percent_format)
            
            summary_sheet.write('G5', 'Siswa Selesai:', label_format)
            summary_sheet.write('H5', f"{completed_count} / {respondent_count}", data_format)
            
            # level Distribution section
            summary_sheet.write('A9', 'Distribusi level:', heading_format)
            
            summary_sheet.write('A10', 'level', header_format)
            summary_sheet.write('B10', 'Jumlah Siswa', header_format)
            summary_sheet.write('C10', 'Persentase', header_format)
            
            row = 11
            # Tulis baris distribusi secara eksplisit untuk Level 1..5 agar urutan konsisten
            for level in [1, 2, 3, 4, 5]:
                count = level_distribution.get(level, 0)
                percentage = (count / respondent_count * 100) if respondent_count > 0 else 0

                summary_sheet.write(f'A{row}', f"level {level}", data_format)
                summary_sheet.write(f'B{row}', count, num_format)
                summary_sheet.write(f'C{row}', percentage / 100, percent_format)

                row += 1
            
            # ===== KETERANGAN LEVEL SECTION =====
            summary_sheet.write('A19', 'Keterangan Level, Basis Pengetahuan, dan Indikator Soal:', heading_format)
            
            # Level descriptions with Technology taxonomy
            level_descriptions = [
                (
                    'Level 1',
                    'Kesadaran Teknologi: mengenali istilah, alat, dan fungsi dasar.',
                    'Knowledge that',
                    'Mengidentifikasi istilah/ikon/alat dasar; menyebutkan fungsi sederhana.'
                ),
                (
                    'Level 2',
                    'Literasi Teknologi: memahami konsep, prinsip dasar, dan konteks penggunaan.',
                    'Knowledge that',
                    'Menjelaskan konsep/prinsip; memberi contoh penggunaan; membaca diagram sederhana.'
                ),
                (
                    'Level 3',
                    'Kemampuan Teknologi: menerapkan alat/metode untuk memecahkan masalah.',
                    'Knowledge that & how',
                    'Menerapkan prosedur/alat untuk menyelesaikan tugas; menyusun langkah-langkah.'
                ),
                (
                    'Level 4',
                    'Kreativitas Teknologi: memodifikasi/merancang solusi baru yang bermanfaat.',
                    'Knowledge that & how',
                    'Merancang/memodifikasi solusi; menggabungkan beberapa alat/konsep; membuat variasi.'
                ),
                (
                    'Level 5',
                    'Kritik Teknologi: mengevaluasi dampak, efektivitas, etika, dan trade-off.',
                    'Knowledge that, how & why',
                    'Mengevaluasi alternatif/efektivitas; memberi alasan dan pertimbangan etis; membandingkan trade-off.'
                )
            ]
            
            # Headers for level descriptions
            summary_sheet.write('A20', 'Level', header_format)
            summary_sheet.write('B20', 'Keterangan', header_format)
            summary_sheet.write('C20', 'Basis Pengetahuan', header_format)
            summary_sheet.write('D20', 'Indikator Soal', header_format)
            
            # Create format for level descriptions with text wrapping
            description_format = workbook.add_format({
                'align': 'left',
                'border': 1,
                'text_wrap': True,
                'valign': 'top'
            })
            
            row = 21
            for level_name, description, technology_taxonomy, indicator in level_descriptions:
                summary_sheet.write(f'A{row}', level_name, data_format)
                summary_sheet.write(f'B{row}', description, description_format)
                summary_sheet.write(f'C{row}', technology_taxonomy, description_format)
                summary_sheet.write(f'D{row}', indicator, description_format)
                summary_sheet.set_row(row-1, 28)  # Sedikit lebih tinggi untuk indikator
                row += 1

            # Lebarkan kolom indikator agar teks rapi
            summary_sheet.set_column('D:D', 45)
            
            # Format column widths with better spacing
            summary_sheet.set_column('A:A', 15)  # Level
            summary_sheet.set_column('B:B', 50)  # Keterangan
            summary_sheet.set_column('C:C', 30)  # Taksonomi Teknologi
            summary_sheet.set_column('D:D', 20)
            summary_sheet.set_column('E:E', 15)
            summary_sheet.set_column('F:F', 3)
            summary_sheet.set_column('G:G', 22)
            summary_sheet.set_column('H:H', 18)

            # ===== TOP SALAH QUESTIONS SECTION =====
            # Hitung soal dengan jumlah salah terbanyak (top 5)
            try:
                questions = db.session.query(Question).join(
                    CollectionQuestion,
                    CollectionQuestion.question_id == Question.id
                ).filter(
                    CollectionQuestion.collection_id == collection_id
                ).all()

                top_wrong = []
                for question in questions:
                    answers = db.session.query(SiswaAnswer).filter_by(
                        question_id=question.id
                    ).all()
                    incorrect_count = sum(1 for a in answers if not a.is_correct)
                    total_ans = len(answers)
                    accuracy_q = (1 - incorrect_count / total_ans) * 100 if total_ans > 0 else 0
                    q_level = normalize_level(question.level)
                    top_wrong.append({
                        'Soal': question.soal,
                        'Jumlah Salah': incorrect_count,
                        'Total Jawaban': total_ans,
                        'Akurasi (%)': round(accuracy_q, 2),
                        'Level': q_level,
                        'Nama Level': TECHNOLOGY_LEVELS.get(int(q_level), str(q_level))
                    })

                # Urutkan berdasarkan jumlah salah desc dan ambil 5 teratas
                top_wrong.sort(key=lambda x: x['Jumlah Salah'], reverse=True)
                top_wrong = top_wrong[:5]

                start_row = row + 1  # rapatkan tanpa spasi berlebih di bawah tabel sebelumnya
                summary_sheet.write(f'A{start_row}', 'Soal Paling Banyak Salah', heading_format)

                header_row = start_row + 1
                tw_headers = ['Soal', 'Jumlah Salah', 'Total Jawaban', 'Akurasi (%)', 'Level', 'Nama Level']
                for col, header in enumerate(tw_headers):
                    summary_sheet.write(header_row, col, header, header_format)

                # Format wrap untuk teks soal
                q_wrap_format = workbook.add_format({
                    'align': 'left',
                    'border': 1,
                    'text_wrap': True,
                    'valign': 'top'
                })

                data_row = header_row + 1
                for item in top_wrong:
                    summary_sheet.write(data_row, 0, item['Soal'], q_wrap_format)
                    summary_sheet.write(data_row, 1, item['Jumlah Salah'], num_format)
                    summary_sheet.write(data_row, 2, item['Total Jawaban'], num_format)
                    summary_sheet.write(data_row, 3, item['Akurasi (%)'], num_format)
                    summary_sheet.write(data_row, 4, item['Level'], num_format)
                    summary_sheet.write(data_row, 5, item['Nama Level'], data_format)
                    summary_sheet.set_row(data_row, 30)
                    data_row += 1
                # Lebarkan kolom agar rapi untuk tabel Top Wrong
                summary_sheet.set_column('A:A', 60)  # Soal
                summary_sheet.set_column('B:B', 16)  # Jumlah Salah
                summary_sheet.set_column('C:C', 16)  # Total Jawaban
                summary_sheet.set_column('D:D', 14)  # Akurasi (%)
                summary_sheet.set_column('E:E', 10)  # Level
                summary_sheet.set_column('F:F', 18)  # Nama Level
            except Exception as e:
                print(f"Error building top wrong questions section: {str(e)}")
            
            # ===== STUDENT RESULTS SHEET =====
            students_sheet = workbook.add_worksheet('Data Siswa')
            
            # Title - disesuaikan dengan jumlah kolom baru
            students_sheet.merge_range('A1:L1', f"DATA SISWA - {collection.name}", title_format)
            
            # Headers - tambah kolom Taksonomi Teknologi
            headers = ['No', 'Nama', 'Kelas', 'Level', 'Nama Level', 'Keterangan Level', 'Basis Pengetahuan', 'Jawaban Benar', 'Jawaban Salah', 'Total', 'Akurasi (%)', 'Status']
            for col, header in enumerate(headers):
                students_sheet.write(2, col, header, header_format)
            
            # Student data (respondents only); maintain continuous numbering
            student_data = []
            row = 3
            no_counter = 1
            for student in students_query:
                # Get result
                result = SiswaResult.query.filter_by(
                    siswa_id=student.id,
                    collection_id=collection_id
                ).first()
                
                if result:
                    correct = result.correct or 0
                    incorrect = result.incorrect or 0
                    total = correct + incorrect
                    if total <= 0:
                        continue  # Skip non-respondents
                    accuracy = (correct / total * 100) if total > 0 else 0
                    current_level = normalize_level(result.current_level) or 1
                    is_completed = (current_level == 5)
                    status = "Selesai" if is_completed else (f"Berhenti di Level {int(current_level)}" if total > 0 else "Belum Mulai")
                else:
                    continue  # Skip students without any result
                
                # Write to Excel - tambah kolom Taksonomi Teknologi
                students_sheet.write(row, 0, no_counter, num_format)
                students_sheet.write(row, 1, student.nama, data_format)
                students_sheet.write(row, 2, student.kelas or "-", data_format)
                students_sheet.write(row, 3, current_level, num_format)
                students_sheet.write(row, 4, TECHNOLOGY_LEVELS.get(int(current_level), str(current_level)), data_format)
                students_sheet.write(row, 5, get_level_description_short(current_level), data_format)
                students_sheet.write(row, 6, get_technology_taxonomy_short(current_level), data_format)
                students_sheet.write(row, 7, correct, num_format)
                students_sheet.write(row, 8, incorrect, num_format)
                students_sheet.write(row, 9, total, num_format)
                students_sheet.write(row, 10, accuracy, num_format)
                students_sheet.write(row, 11, status, data_format)
                
                # Set row height for better readability
                students_sheet.set_row(row-1, 25)
                
                # Add to data for pandas DataFrame
                student_data.append({
                    'Nama': student.nama,
                    'Kelas': student.kelas or "-",
                    'level': current_level,
                    'Jawaban Benar': correct,
                    'Jawaban Salah': incorrect,
                    'Total': total,
                    'Akurasi (%)': accuracy,
                    'Status': status
                })
                
                row += 1
                no_counter += 1
            
            # Format columns dengan spacing yang lebih baik - tambah kolom Taksonomi Bloom
            students_sheet.set_column('A:A', 5)   # No
            students_sheet.set_column('B:B', 25)  # Nama
            students_sheet.set_column('C:C', 12)  # Kelas
            students_sheet.set_column('D:D', 8)   # Level
            students_sheet.set_column('E:E', 18)  # Nama Level
            students_sheet.set_column('F:F', 22)  # Keterangan Level (singkat)
            students_sheet.set_column('G:G', 20)  # Basis Pengetahuan
            students_sheet.set_column('H:H', 15)  # Jawaban Benar
            students_sheet.set_column('I:I', 15)  # Jawaban Salah
            students_sheet.set_column('J:J', 10)  # Total
            students_sheet.set_column('K:K', 12)  # Akurasi (%)
            students_sheet.set_column('L:L', 18)  # Status
            
            # ===== CLASS SUMMARY SHEET =====
            # Group students by class
            classes = {}
            for student in students_query:
                class_name = student.kelas or "No Class"
                if class_name not in classes:
                    classes[class_name] = []
                classes[class_name].append(student.id)
            
            # Create class summary
            class_sheet = workbook.add_worksheet('Ringkasan Kelas')
            
            # Title
            class_sheet.merge_range('A1:H1', f"RINGKASAN PER KELAS - {collection.name}", title_format)
            
            # Headers
            class_headers = ['Kelas', 'Jumlah Responden', 'Jawaban Benar', 'Jawaban Salah', 'Total Jawaban', 'Akurasi (%)', 'Siswa Selesai', 'Persentase Selesai (%)']
            for col, header in enumerate(class_headers):
                class_sheet.write(2, col, header, header_format)
            
            # Class data
            row = 3
            for class_name, student_ids in classes.items():
                respondents = 0
                class_correct = 0
                class_incorrect = 0
                class_completed = 0
                
                # Calculate stats for this class
                for student_id in student_ids:
                    result = SiswaResult.query.filter_by(
                        siswa_id=student_id,
                        collection_id=collection_id
                    ).first()
                    
                    if result:
                        corr = result.correct or 0
                        inc = result.incorrect or 0
                        tot = corr + inc
                        if tot <= 0:
                            continue
                        class_correct += corr
                        class_incorrect += inc
                        lvl = normalize_level(result.current_level)
                        if lvl == 5:
                            class_completed += 1
                        respondents += 1
                
                class_total = class_correct + class_incorrect
                class_accuracy = (class_correct / class_total * 100) if class_total > 0 else 0
                completion_percentage = (class_completed / respondents * 100) if respondents > 0 else 0
                
                # Write to Excel
                class_sheet.write(row, 0, class_name, data_format)
                class_sheet.write(row, 1, respondents, num_format)
                class_sheet.write(row, 2, class_correct, num_format)
                class_sheet.write(row, 3, class_incorrect, num_format)
                class_sheet.write(row, 4, class_total, num_format)
                class_sheet.write(row, 5, class_accuracy, num_format)
                class_sheet.write(row, 6, f"{class_completed} / {respondents}", data_format)
                class_sheet.write(row, 7, completion_percentage, num_format)
                
                row += 1
            
            # Format columns with better spacing
            class_sheet.set_column('A:A', 20)  # Kelas
            class_sheet.set_column('B:B', 15)  # Jumlah Siswa
            class_sheet.set_column('C:C', 15)  # Jawaban Benar
            class_sheet.set_column('D:D', 15)  # Jawaban Salah
            class_sheet.set_column('E:E', 15)  # Total Jawaban
            class_sheet.set_column('F:F', 15)  # Akurasi (%)
            class_sheet.set_column('G:G', 18)  # Siswa Selesai
            class_sheet.set_column('H:H', 22)  # Persentase Selesai (%)
            
            # ===== QUESTIONS ANALYSIS SHEET =====
            # Get questions in this collection
            questions = db.session.query(Question).join(
                CollectionQuestion, 
                CollectionQuestion.question_id == Question.id
            ).filter(
                CollectionQuestion.collection_id == collection_id
            ).all()
            
            # Create question analysis sheet
            if questions:
                question_sheet = workbook.add_worksheet('Analisis Soal')
                
                # Title - disesuaikan dengan kolom baru
                question_sheet.merge_range('A1:H1', f"ANALISIS SOAL - {collection.name}", title_format)
                
                # Headers - tampilkan Nama Level dan Basis Pengetahuan
                question_headers = ['Level', 'Nama Level', 'Basis Pengetahuan', 'Soal', 'Jawaban Benar', 'Jawaban Salah', 'Total Jawaban', 'Akurasi (%)']
                for col, header in enumerate(question_headers):
                    question_sheet.write(2, col, header, header_format)
                
                # Question data
                row = 3
                for question in questions:
                    # Calculate question statistics
                    answers = db.session.query(SiswaAnswer).filter_by(
                        question_id=question.id
                    ).all()
                    
                    correct_count = sum(1 for a in answers if a.is_correct)
                    incorrect_count = len(answers) - correct_count
                    total_answers = len(answers)
                    accuracy = (correct_count / total_answers * 100) if total_answers > 0 else 0
                    
                    # Create format for question text with text wrapping
                    question_format = workbook.add_format({
                        'align': 'left',
                        'border': 1,
                        'text_wrap': True,
                        'valign': 'top'
                    })
                    
                    # Write to Excel - tambah kolom Taksonomi Teknologi
                    q_level = normalize_level(question.level)
                    question_sheet.write(row, 0, q_level, num_format)
                    question_sheet.write(row, 1, TECHNOLOGY_LEVELS.get(int(q_level), str(q_level)), data_format)
                    question_sheet.write(row, 2, get_technology_taxonomy_short(q_level), data_format)
                    question_sheet.write(row, 3, question.soal, question_format)
                    question_sheet.write(row, 4, correct_count, num_format)
                    question_sheet.write(row, 5, incorrect_count, num_format)
                    question_sheet.write(row, 6, total_answers, num_format)
                    question_sheet.write(row, 7, accuracy, num_format)
                    
                    # Set row height for better readability
                    question_sheet.set_row(row-1, 30)
                    
                    row += 1
                
                # Format columns dengan spacing yang lebih baik - tambah kolom Taksonomi Teknologi
                question_sheet.set_column('A:A', 8)   # Level
                question_sheet.set_column('B:B', 18)  # Nama Level
                question_sheet.set_column('C:C', 20)  # Basis Pengetahuan
                question_sheet.set_column('D:D', 45)  # Soal
                question_sheet.set_column('E:E', 15)  # Jawaban Benar
                question_sheet.set_column('F:F', 15)  # Jawaban Salah
                question_sheet.set_column('G:G', 15)  # Total Jawaban
                question_sheet.set_column('H:H', 15)  # Akurasi (%)
                
                # Create charts worksheet
                charts_sheet = workbook.add_worksheet('Grafik')
                
                # Title
                charts_sheet.merge_range('A1:H1', f"GRAFIK ANALISIS - {collection.name}", title_format)
                
                # Add level distribution chart
                level_chart = workbook.add_chart({'type': 'column'})
                
                # Configure the chart
                # Data tabel Distribusi level di 'Ringkasan Koleksi' dimulai di baris 11 (1-based) s.d. 15
                # XlsxWriter menggunakan indeks 0-based, jadi gunakan 10..14 agar Level 1 tidak terlewat
                level_chart.add_series({
                    'name': 'Jumlah Siswa',
                    'categories': ['Ringkasan Koleksi', 10, 0, 14, 0],
                    'values': ['Ringkasan Koleksi', 10, 1, 14, 1],
                    'data_labels': {'value': True}
                })
                
                level_chart.set_title({'name': 'Distribusi Level'})
                level_chart.set_x_axis({'name': 'Level'})
                level_chart.set_y_axis({'name': 'Jumlah Siswa'})
                
                # Insert chart
                charts_sheet.insert_chart('A3', level_chart, {'x_offset': 25, 'y_offset': 10, 'x_scale': 1.5, 'y_scale': 1.2})
                
                # Add class comparison chart
                class_chart = workbook.add_chart({'type': 'column'})
                
                # Configure the chart
                class_chart.add_series({
                    'name': 'Akurasi (%)',
                    'categories': ['Ringkasan Kelas', 3, 0, 3 + len(classes) - 1, 0],
                    'values': ['Ringkasan Kelas', 3, 5, 3 + len(classes) - 1, 5],
                    'data_labels': {'value': True}
                })
                
                class_chart.set_title({'name': 'Akurasi per Kelas'})
                class_chart.set_x_axis({'name': 'Kelas'})
                class_chart.set_y_axis({'name': 'Akurasi (%)'})
                
                # Insert chart
                charts_sheet.insert_chart('A20', class_chart, {'x_offset': 25, 'y_offset': 10, 'x_scale': 1.5, 'y_scale': 1.2})
        
        # Reset file pointer to beginning
        output.seek(0)
        
        # Create filename
        filename = f"Data_Koleksi_{collection.name}_{datetime.datetime.now().strftime('%Y%m%d')}.xlsx"
        
        # Send file as attachment
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        print(f"Error exporting collection data: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f"Error: {str(e)}"}), 500

@app.route('/api/export_result_pdf', methods=['GET'])
@login_required
def export_result_pdf():
    try:
        # Ambil parameter
        user_id = request.args.get('user_id')
        collection_id = request.args.get('collection_id')
        
        # Validasi dan ambil semua data
        if not user_id or not collection_id:
            return jsonify({'success': False, 'message': 'Parameter tidak lengkap'}), 400
        student = Siswa.query.get(user_id)
        if not student:
            return jsonify({'success': False, 'message': 'Siswa tidak ditemukan'}), 404
        collection = QuestionCollection.query.get(collection_id)
        if not collection:
            return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
        result = SiswaResult.query.filter_by(
            siswa_id=user_id,
            collection_id=collection_id
        ).first()
        if not result:
            return jsonify({'success': False, 'message': 'Hasil tes tidak ditemukan'}), 404
        
        mst_state = get_mst_state(user_id, collection_id)
        # Ambil jawaban siswa dan buat statistik
        answers = get_student_answer_history(user_id, collection_id)
        total = result.correct + result.incorrect
        accuracy = round((result.correct / total * 100), 2) if total > 0 else 0
        
        # Generate rekomendasi pembelajaran kontekstual berdasarkan jawaban siswa
        weakness_analysis = analyze_student_weaknesses(student.id, collection_id)
        contextual_recommendations = generate_contextual_recommendations(weakness_analysis)
        
        # Generate rekomendasi AI tambahan
        ai_recommendations = generate_ai_recommendations(student, result, collection_id)
        
        # Gabungkan kedua jenis rekomendasi
        combined_recommendations = {
            'contextual': contextual_recommendations,
            'ai_general': ai_recommendations,
            'weakness_summary': {
                'total_errors': sum(data["count"] for data in weakness_analysis.get("topic_statistics", {}).values()),
                'weak_topics': list(weakness_analysis.get("topic_statistics", {}).keys())[:3],
                'weak_levels': [level for level, data in weakness_analysis.get("level_statistics", {}).items() if data["count"] >= 2],
                'pattern_analysis': {
                    'concept_gaps': len(weakness_analysis.get("concept_gaps", [])),
                    'application_gaps': len(weakness_analysis.get("application_gaps", [])),
                    'analysis_gaps': len(weakness_analysis.get("analysis_gaps", []))
                }
            }
        }
        
        # Helper function untuk waktu sekarang
        def now():
            return datetime.datetime.now()
        
        # Render template HTML dengan data
        html = render_template(
            'siswa/export_pdf.html',
            student=student,
            collection=collection,
            result=result,
            answers=answers,
            accuracy=accuracy,
            recommendations=combined_recommendations,
            total_questions=total,
            now=now,
            mst_state=mst_state
        )
        
        # Fungsi untuk mengkonversi HTML ke PDF dengan xhtml2pdf yang ditingkatkan
        def html_to_pdf(source_html):
            try:
                # Log debugging info
                print("Starting PDF conversion...")
                
                # Buat PDF buffer untuk output
                result_buffer = BytesIO()
                
                # Atur konfigurasi PDF
                pdf_options = {
                    "page-size": "A4",
                    "margin-top": "2cm",
                    "margin-bottom": "2cm",
                    "margin-left": "2cm",
                    "margin-right": "2cm",
                    "encoding": "UTF-8",
                    # Tambahan opsi untuk stabilitas
                    "quiet": 0,
                }
                
                # Konversi HTML ke PDF menggunakan xhtml2pdf
                pisa_status = pisa.CreatePDF(
                    source_html,                # HTML sumber
                    dest=result_buffer,         # Buffer tujuan
                    encoding='UTF-8',           # Encoding untuk karakter internasional
                    options=pdf_options         # Opsi format
                )
                
                # Reset pointer buffer ke awal
                result_buffer.seek(0)
                
                # Cek status konversi
                if pisa_status.err:
                    print(f"Error pada konversi PDF: {pisa_status.err}")
                    return None
                    
                return result_buffer
                
            except Exception as pdf_error:
                print(f"Exception dalam konversi PDF: {str(pdf_error)}")
                import traceback
                traceback.print_exc()
                return None

        # Konversi HTML ke PDF
        pdf_buffer = html_to_pdf(html)
        
        if pdf_buffer is None:
            return jsonify({'success': False, 'message': 'Gagal membuat PDF'}), 500
        
        # Buat nama file yang lebih deskriptif untuk download
        current_date = datetime.datetime.now().strftime('%Y%m%d_%H%M')
        
        # Bersihkan nama file dari karakter yang tidak diperbolehkan
        safe_student_name = "".join(c for c in student.nama if c.isalnum() or c in ' _-').strip()
        safe_student_name = safe_student_name.replace(' ', '_')
        
        safe_collection_name = "".join(c for c in collection.name if c.isalnum() or c in ' _-').strip()
        safe_collection_name = safe_collection_name.replace(' ', '_')
        
        filename = f"Hasil_Tes_{safe_student_name}_{safe_collection_name}_{current_date}.pdf"
        
        # Kirim file PDF
        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/pdf'
        )
        
    except Exception as e:
        print(f"Error exporting to PDF: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f"Error: {str(e)}"}), 500

def generate_ai_recommendations(student, result, collection_id):
    try:
        # Mengambil nama koleksi dan konteks materi
        collection = QuestionCollection.query.get(collection_id)
        collection_name = collection.name if collection else "N/A"
        
        # Ambil konteks materi dan tujuan pembelajaran
        questions_context = get_questions_context_for_recommendations(collection_id, result.current_level)
        learning_objectives = extract_learning_objectives(collection_id)
        performance_analysis = analyze_student_performance(student.id, result, collection_id)
        
        # Hitung persentase keberhasilan untuk menentukan tingkat penguasaan
        total_answers = result.correct + result.incorrect
        mastery_percentage = (result.correct / total_answers * 100) if total_answers > 0 else 0
        
        # Tentukan tingkat penguasaan berdasarkan persentase
        if mastery_percentage >= 80:
            mastery_level = "tinggi"
        elif mastery_percentage >= 60:
            mastery_level = "sedang"
        else:
            mastery_level = "rendah"

        # Prompt AI untuk analisis jawaban dan rekomendasi pembelajaran yang humanize
        prompt = f"""
        Kamu adalah GURU yang baik dan peduli siswa. Berikan rekomendasi belajar yang SINGKAT, PRAKTIS, dan mudah dipahami.

        PROFIL SISWA:
        Nama: {student.nama}
        Materi: {collection_name}
        Level: {result.current_level} dari 7 level
        Benar: {result.correct} soal | Salah: {result.incorrect} soal
        Pemahaman: {mastery_level.upper()} ({mastery_percentage:.1f}%)

        TUGASMU:
        Berikan 3-4 saran belajar yang KONKRET dan SINGKAT. Maksimal 1-2 kalimat per saran.

        GAYA BAHASA:
        [OK] Gunakan "kamu" - akrab tapi sopan
        [OK] Bahasa sehari-hari, tidak formal
        [OK] Langsung to the point
        [OK] Positif dan mendukung
        [OK] Berikan alasan singkat kenapa penting

        HINDARI:
        [DITOLAK] Kalimat panjang dan bertele-tele
        [DITOLAK] Bahasa formal ("direkomendasikan", "hendaknya")
        [DITOLAK] Istilah teknis yang sulit
        [DITOLAK] Bullet points atau numbering
        [DITOLAK] Pengulangan yang tidak perlu

        CONTOH YANG BAGUS:
        "Kamu sudah bagus di bagian dasar. Sekarang fokus latihan soal aplikasi aja biar makin lancar."
        "Coba buat ringkasan sendiri pakai kata-katamu. Ini bikin kamu lebih paham konsep dasarnya."

        Buat rekomendasi yang sesuai dengan level pemahaman {student.nama} saat ini.
        """

        # Panggil AI untuk mendapatkan rekomendasi
        response = model.generate_content(prompt)
        recommendations_text = response.text
        
        # Parse hasil ke dalam format yang diinginkan
        recommendations = parse_ai_recommendations_for_teacher(recommendations_text)
        
        # Fallback jika terjadi kegagalan
        if not recommendations or len(recommendations) < 3:
            recommendations = get_teacher_style_recommendations(
                current_level=result.current_level,
                correct=result.correct,
                incorrect=result.incorrect,
                collection_name=collection_name,
                learning_objectives=learning_objectives,
                mastery_level=mastery_level
            )
        
        return recommendations
        
    except Exception as e:
        print(f"Error generating AI recommendations: {str(e)}")
        import traceback
        traceback.print_exc()
        
        # Fallback ke rekomendasi default dengan gaya guru
        return get_teacher_style_recommendations(
            current_level=result.current_level if 'result' in locals() else 1,
            correct=result.correct if 'result' in locals() else 0,
            incorrect=result.incorrect if 'result' in locals() else 0,
            collection_name=collection_name if 'collection_name' in locals() else "materi ini",
            learning_objectives=learning_objectives if 'learning_objectives' in locals() else [],
            mastery_level=mastery_level if 'mastery_level' in locals() else "sedang"
        )

# Fungsi helper untuk fokus rekomendasi berdasarkan tingkat penguasaan
def get_focus_by_mastery_level(mastery_level, domain_type):
    if domain_type == "konsep":
        if mastery_level == "rendah":
            return "memperkuat DASAR-DASAR materi dengan scaffolding dan visualisasi"
        elif mastery_level == "sedang":
            return "memperdalam HUBUNGAN ANTAR KONSEP dan aplikasi kontekstual"
        else:  # tinggi
            return "mengeksplorasi PERSPEKTIF LANJUTAN dan transfer pengetahuan"
    elif domain_type == "praktik":
        if mastery_level == "rendah":
            return "memberikan LATIHAN TERSTRUKTUR dengan panduan bertahap"
        elif mastery_level == "sedang":
            return "mencakup VARIASI KONTEKS dan pemecahan masalah mandiri"
        else:  # tinggi
            return "melibatkan KOMPLEKSITAS TINGGI dan potensi inovasi"
    elif domain_type == "evaluasi":
        if mastery_level == "rendah":
            return "berfokus pada IDENTIFIKASI KESENJANGAN pemahaman dan diagnosis kesulitan"
        elif mastery_level == "sedang":
            return "melibatkan EVALUASI STRATEGI belajar dan efektivitas pendekatan"
        else:  # tinggi
            return "mendorong SINTESIS PENGETAHUAN dan refleksi metakognitif"
    return ""

# Fungsi untuk mengekstrak tujuan pembelajaran dari modul
def extract_learning_objectives(collection_id):
    """
    Mengekstrak tujuan pembelajaran dari modul berdasarkan collection_id
    """
    try:
        # Ambil referensi modul dari koleksi soal
        questions = Question.query.filter_by(collection_id=collection_id).all()
        modul_references = []
        learning_objectives = {}
        
        # Pertama coba ekstrak dari question.modul_reference jika ada
        for question in questions:
            # Check if modul_reference attribute exists
            if hasattr(question, 'modul_reference') and question.modul_reference and question.modul_reference.strip():
                modul_ref = question.modul_reference.strip()
                modul_references.append(modul_ref)
        
        # Hapus duplikat
        unique_modules = list(set(modul_references))
        
        # Fallback: ekstrak dari konten soal
        for question in questions:
            if hasattr(question, 'explanation') and question.explanation:
                explanation_lower = question.explanation.lower()
                
                # Cari frasa yang mengandung kata tujuan pembelajaran
                if any(term in explanation_lower for term in ["tujuan pembelajaran", "learning objective", "capaian pembelajaran", "kompetensi"]):
                    # Coba ekstrak tujuan pembelajaran dari penjelasan soal
                    objective_text = extract_objective_from_text(question.explanation)
                    if objective_text:
                        ref = question.modul_reference if hasattr(question, 'modul_reference') else "unknown"
                        if ref not in learning_objectives:
                            learning_objectives[ref] = []
                        learning_objectives[ref].append(objective_text)
        
        # Format tujuan pembelajaran
        formatted_objectives = []
        for modul, objectives in learning_objectives.items():
            if isinstance(objectives, list):
                for obj in objectives:
                    formatted_objectives.append(f"Modul {modul}: {obj}")
            else:
                formatted_objectives.append(f"Modul {modul}: {objectives}")
        
        # Jika masih belum menemukan tujuan pembelajaran, coba ekstrak dari pertanyaan
        if not formatted_objectives:
            # Kategorikan pertanyaan berdasarkan level
            questions_by_level = {}
            for question in questions:
                if hasattr(question, 'level'):
                    level = question.level
                    if level not in questions_by_level:
                        questions_by_level[level] = []
                    questions_by_level[level].append(question)
            
            # Untuk setiap level, buat perkiraan tujuan pembelajaran
            for level, level_questions in sorted(questions_by_level.items()):
                # Ekstrak kata kunci dari soal-soal di level ini
                level_text = " ".join([q.soal for q in level_questions if hasattr(q, 'soal') and q.soal])
                keywords = extract_keywords_from_text(level_text, 3)
                
                if keywords:
                    keyword_phrase = ", ".join(keywords)
                    if level <= 2:
                        formatted_objectives.append(f"Level {level}: Memahami dan mengidentifikasi konsep {keyword_phrase}")
                    elif level <= 4:
                        formatted_objectives.append(f"Level {level}: Menerapkan dan menganalisis konsep {keyword_phrase}")
                    else:
                        formatted_objectives.append(f"Level {level}: Mengevaluasi dan mencipta berdasarkan konsep {keyword_phrase}")
        
        # Jika masih tidak menemukan tujuan pembelajaran, buat tujuan umum
        if not formatted_objectives:
            # Buat perkiraan tujuan pembelajaran dari konten soal
            all_text = " ".join([q.soal for q in questions if hasattr(q, 'soal') and q.soal])
            concepts = extract_keywords_from_text(all_text, 5)
            
            if concepts:
                for i, concept in enumerate(concepts[:3]):
                    if i == 0:
                        formatted_objectives.append(f"Memahami dan menerapkan konsep {concept}")
                    elif i == 1:
                        formatted_objectives.append(f"Menganalisis hubungan antara {concepts[0]} dan {concept}")
                    else:
                        formatted_objectives.append(f"Mengevaluasi penerapan {concept} dalam konteks nyata")
            else:
                # Jika tidak ada konsep spesifik, berikan tujuan generik
                formatted_objectives = [
                    "Memahami konsep-konsep dasar materi",
                    "Menerapkan pengetahuan dalam pemecahan masalah",
                    "Menganalisis dan mengevaluasi situasi berdasarkan konsep yang dipelajari"
                ]
        
        return formatted_objectives
    
    except Exception as e:
        print(f"Error extracting learning objectives: {str(e)}")
        import traceback
        traceback.print_exc()
        return ["Memahami dan menerapkan konsep dasar materi"]
    
def extract_keywords_from_text(text, count=5):
    """
    Ekstrak kata kunci dari teks berdasarkan frekuensi
    """
    if not text:
        return []
        
    # Bersihkan teks
    clean_text = re.sub(r'[^\w\s]', ' ', text.lower())
    
    # Hitung frekuensi kata
    words = clean_text.split()
    word_freq = {}
    
    stopwords = ["yang", "dan", "atau", "adalah", "untuk", "dalam", "pada", "ini", "itu", "dengan", 
                "akan", "dari", "jika", "maka", "oleh", "ada", "tidak", "kita", "saya", "kami", 
                "mereka", "dia", "suatu", "dapat", "bisa", "harus", "sebagai", "saat", "ketika"]
    
    for word in words:
        if len(word) > 3 and word not in stopwords:
            word_freq[word] = word_freq.get(word, 0) + 1
    
    # Ambil kata dengan frekuensi tertinggi
    sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
    return [word for word, freq in sorted_words[:count]]

def extract_objective_from_text(text):
    """
    Ekstrak tujuan pembelajaran dari teks penjelasan
    """
    # Cari kalimat yang memuat kata 'tujuan'
    sentences = re.split(r'[.!?]', text)
    for sentence in sentences:
        if re.search(r'\b(tujuan|kompetensi|capaian)\b', sentence.lower()):
            return sentence.strip()
    return None

def extract_key_concepts_from_questions(questions):
    """
    Ekstrak konsep kunci dari kumpulan soal
    """
    # Gabungkan semua soal dan penjelasan
    all_text = ""
    for question in questions:
        if hasattr(question, 'soal') and question.soal:
            all_text += question.soal + " "
        if hasattr(question, 'explanation') and question.explanation:
            all_text += question.explanation + " "
    
    # Bersihkan teks
    all_text = re.sub(r'[^\w\s]', '', all_text.lower())
    
    # Kata-kata yang sering muncul bisa jadi merupakan konsep kunci
    # Tambahkan logika NLP di sini untuk ekstrak konsep kunci
    # Sebagai contoh sederhana, kita menggunakan frekuensi kata dengan filter
    words = all_text.split()
    word_freq = {}
    
    stopwords = ["yang", "dan", "atau", "adalah", "untuk", "dalam", "pada", "ini", "itu", "dengan"]
    for word in words:
        if len(word) > 4 and word not in stopwords:
            word_freq[word] = word_freq.get(word, 0) + 1
    
    # Ambil kata dengan frekuensi tertinggi
    sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
    return [word for word, freq in sorted_words[:10]]  # Return top 10 konsep

def identify_learning_domains(collection_id):
    """
    Identifikasi domain pembelajaran yang diukur dalam koleksi soal
    """
    try:
        questions = Question.query.filter_by(collection_id=collection_id).all()
        domains = {
            "kognitif": 0,
            "praktis": 0,
            "reflektif": 0
        }
        
        for question in questions:
            # Analisis soal untuk menentukan domain
            if hasattr(question, 'soal') and question.soal:
                text = question.soal.lower()
                
                # Indikator domain kognitif
                if any(word in text for word in ["jelaskan", "definisikan", "sebutkan", "identifikasi"]):
                    domains["kognitif"] += 1
                
                # Indikator domain praktis
                if any(word in text for word in ["terapkan", "implementasikan", "lakukan", "bagaimana", "cara"]):
                    domains["praktis"] += 1
                
                # Indikator domain reflektif
                if any(word in text for word in ["evaluasi", "analisis", "bandingkan", "mengapa", "refleksikan"]):
                    domains["reflektif"] += 1
        
        # Format output
        domain_text = []
        total_questions = len(questions)
        if total_questions > 0:
            for domain, count in domains.items():
                if count > 0:
                    percentage = round((count / total_questions * 100), 1)
                    domain_text.append(f"{domain.capitalize()}: {percentage}%")
            
            return ", ".join(domain_text) if domain_text else "Domain tidak teridentifikasi"
        else:
            return "Domain Kognitif, Praktis, dan Reflektif"
    
    except Exception as e:
        print(f"Error identifying learning domains: {str(e)}")
        return "Domain Kognitif, Praktis, dan Reflektif"

def get_questions_context_for_recommendations(collection_id, current_level):
    """
    Ambil konteks dari soal-soal untuk rekomendasi yang lebih spesifik
    dengan fokus pada tujuan pembelajaran dari modul ajar
    """
    try:
        questions = Question.query.filter_by(collection_id=collection_id).all()
        
        if not questions:
            return "Tidak ada konteks materi yang tersedia."
        
        # Analisis konten soal
        context_parts = []
        
        # 1. Referensi modul ajar
        module_refs = []
        module_content = {}
        
        for question in questions:
            # Ekstrak referensi modul
            if hasattr(question, 'modul_reference') and question.modul_reference:
                module_ref = question.modul_reference.strip()
                module_refs.append(module_ref)
                
                # Kumpulkan konten berdasarkan modul
                if module_ref not in module_content:
                    module_content[module_ref] = []
                
                # Tambahkan soal dan penjelasan ke konten modul
                content = question.soal[:150] if hasattr(question, 'soal') and question.soal else ""
                if hasattr(question, 'explanation') and question.explanation:
                    content += " | " + question.explanation[:150]
                    
                if content:
                    module_content[module_ref].append(content)
        
        # 2. Ekstrak konsep-konsep kunci
        concepts = []
        for question in questions:
            # Ekstrak konsep dari explanation
            if hasattr(question, 'explanation') and question.explanation:
                # Identifikasi frasa kunci dalam penjelasan
                explanation = question.explanation.lower()
                
                # Cari frasa yang menandakan konsep penting
                for marker in ["konsep", "prinsip", "teori", "metode", "rumus"]:
                    pattern = fr'{marker}\s+(\w+[\s\w]+?)(?:\.|\,|\;|\:)'
                    matches = re.findall(pattern, explanation)
                    for match in matches:
                        if 3 < len(match) < 50:  # Filter panjang yang masuk akal
                            concepts.append(f"{marker.capitalize()}: {match.strip()}")
        
        # Format konteks
        # 1. Referensi modul
        if module_refs:
            unique_refs = list(set(module_refs))
            context_parts.append(f"REFERENSI MODUL AJAR: {', '.join(unique_refs)}")
        
        # 2. Konsep kunci
        if concepts:
            unique_concepts = list(set(concepts))[:5]  # Batasi 5 konsep paling penting
            context_parts.append(f"KONSEP KUNCI: {'; '.join(unique_concepts)}")
        
        # 3. Konten per modul
        for module, contents in module_content.items():
            if contents:
                # Ambil sample konten dari modul
                sample = contents[0]
                context_parts.append(f"ISI MODUL {module}: {sample}")
        
        # 4. Distribusi level
        level_dist = {}
        for question in questions:
            if hasattr(question, 'level'):
                level = question.level
                level_dist[level] = level_dist.get(level, 0) + 1
        
        level_info = [f"Level {level}: {count} soal" for level, count in sorted(level_dist.items())]
        context_parts.append(f"DISTRIBUSI LEVEL: {', '.join(level_info)}")
        
        return "\n".join(context_parts)
        
    except Exception as e:
        print(f"Error getting questions context: {str(e)}")
        import traceback
        traceback.print_exc()
        return "Tidak dapat menganalisis konteks materi."

def analyze_student_performance(student_id, result, collection_id):
    """
    Analisis performa siswa dengan gaya percakapan guru yang hangat dan mendukung
    """
    try:
        # Ambil semua jawaban siswa untuk collection ini
        answers = SiswaAnswer.query.filter_by(
            siswa_id=student_id
        ).join(
            Question, SiswaAnswer.question_id == Question.id
        ).join(
            CollectionQuestion, 
            (CollectionQuestion.question_id == Question.id) &
            (CollectionQuestion.collection_id == collection_id)
        ).all()
        
        if not answers:
            return "Belum bisa memberikan analisis detail karena data jawaban belum lengkap. Tapi gak apa-apa, nanti setelah kamu selesai mengerjakan beberapa soal lagi, aku bisa kasih feedback yang lebih spesifik."
        
        total_questions = len(answers)
        correct_answers = sum(1 for ans in answers if ans.is_correct)
        accuracy_percentage = (correct_answers / total_questions * 100) if total_questions > 0 else 0
        
        # Analisis per level dengan bahasa yang lebih manusiawi
        level_performance = {}
        for answer in answers:
            if hasattr(answer, 'question') and answer.question and hasattr(answer.question, 'level'):
                level = answer.question.level
                if level not in level_performance:
                    level_performance[level] = {'correct': 0, 'total': 0, 'questions': []}
                
                level_performance[level]['total'] += 1
                if answer.is_correct:
                    level_performance[level]['correct'] += 1
                level_performance[level]['questions'].append({
                    'question': answer.question.soal[:100] + "..." if len(answer.question.soal) > 100 else answer.question.soal,
                    'correct': answer.is_correct
                })
        
        # Buat narasi analisis yang humanize
        analysis_text = ""
        
        # Pembuka yang hangat berdasarkan performa overall
        if accuracy_percentage >= 80:
            analysis_text += f"Wah, kamu sudah menunjukkan pemahaman yang sangat baik! Dari {total_questions} soal, kamu berhasil menjawab {correct_answers} soal dengan benar ({accuracy_percentage:.1f}%). Ini pencapaian yang patut dibanggakan."
        elif accuracy_percentage >= 60:
            analysis_text += f"Kemajuan yang cukup bagus nih! Kamu sudah bisa menjawab {correct_answers} dari {total_questions} soal dengan benar ({accuracy_percentage:.1f}%). Masih ada ruang untuk berkembang, tapi ini sudah langkah yang baik."
        else:
            analysis_text += f"Aku lihat kamu udah berusaha mengerjakan {total_questions} soal dan berhasil {correct_answers} soal ({accuracy_percentage:.1f}%). Jangan berkecil hati ya, yang penting kamu udah berani mencoba. Sekarang kita fokus untuk memperbaiki area yang masih kurang."
        
        # Analisis per tingkat kesulitan
        analysis_text += "\n\nKalau dilihat dari tingkat kesulitannya:"
        
        strong_areas = []
        improvement_areas = []
        
        level_descriptions = {
            1: "soal-soal dasar (konsep fundamental)",
            2: "soal pemahaman (mengerti konsep)",  
            3: "soal aplikasi (menerapkan konsep)",
            4: "soal problem solving (memecahkan masalah)",
            5: "soal analisis (berpikir mendalam)",
            6: "soal evaluasi (menilai dan membandingkan)",
            7: "soal kreatif (menciptakan solusi baru)"
        }
        
        for level in sorted(level_performance.keys()):
            perf = level_performance[level]
            if perf['total'] > 0:
                accuracy = (perf['correct'] / perf['total'] * 100)
                level_desc = level_descriptions.get(level, f"soal level {level}")
                
                if accuracy >= 75:
                    strong_areas.append(f"{level_desc} ({perf['correct']}/{perf['total']})")
                elif accuracy < 50:
                    improvement_areas.append(f"{level_desc} ({perf['correct']}/{perf['total']})")
        
        if strong_areas:
            analysis_text += f"\n• Kamu sudah cukup kuat di: {', '.join(strong_areas)}"
        
        if improvement_areas:
            analysis_text += f"\n• Area yang perlu lebih diasah: {', '.join(improvement_areas)}"
        
        # Identifikasi pola belajar dan saran spesifik
        if len(level_performance) > 1:
            # Cek apakah ada tren performa menurun di level tinggi
            sorted_levels = sorted(level_performance.keys())
            if len(sorted_levels) >= 3:
                low_accuracy = sum(1 for level in sorted_levels[-2:] 
                                 if (level_performance[level]['correct'] / level_performance[level]['total']) < 0.5)
                if low_accuracy >= 1:
                    analysis_text += "\n\nAku perhatiin nih, kamu mulai kesulitan di soal-soal yang lebih kompleks. Ini normal kok! Artinya kamu perlu sedikit lebih banyak latihan untuk soal yang butuh analisis mendalam."
                    
        # Tips berdasarkan pola kesalahan
        if accuracy_percentage < 60:
            analysis_text += "\n\nSaran dari aku: mulai dengan memperkuat konsep dasarnya dulu, baru nanti kita naik ke tingkat yang lebih tinggi. Seperti membangun rumah, fondasinya harus kuat dulu."
        elif accuracy_percentage < 80:
            analysis_text += "\n\nKamu udah punya dasar yang bagus, sekarang tinggal diasah lagi kemampuan aplikasi dan problem solvingnya. Coba perbanyak latihan soal dengan variasi yang berbeda-beda."
        else:
            analysis_text += "\n\nKemampuan kamu sudah sangat baik! Sekarang waktunya untuk mengeksplorasi hal-hal yang lebih menantang dan kreatif."
        
        return analysis_text.strip()
        
    except Exception as e:
        print(f"Error analyzing student performance: {str(e)}")
        import traceback
        traceback.print_exc()
        return "Maaf ya, ada kendala teknis saat menganalisis jawaban kamu. Tapi jangan khawatir, aku tetap bisa kasih saran belajar yang berguna berdasarkan hasil tes kamu."

def parse_ai_recommendations_for_teacher(recommendations_text):
    """
    Mengubah teks rekomendasi dari AI menjadi format narasi yang bersih untuk PDF
    """
    try:
        # Bersihkan teks input
        cleaned_text = recommendations_text.strip()
        
        if not cleaned_text:
            return []
        
        # Hilangkan semua bullet points, numbering, dan format list
        clean_patterns = [
            r'^\s*[•◦▪▫–—\-\*\+]\s*',  # Bullet points
            r'^\s*\d+\.\s*',            # Numbering (1., 2., etc.)
            r'^\s*[a-zA-Z]\.\s*',       # Letter numbering (a., b., etc.)
            r'^\s*[ivxlcdm]+\.\s*',     # Roman numerals
            r'^[🔍💡📊[DITERIMA][DITOLAK]📈🎯📚💪🌟✨]\s*', # Emojis at start
            r'^\s*[-*]\s*',             # Dashes
            r'^\s*[[OK][DITOLAK]]\s*',           # Check marks
        ]
        
        # Split by lines and clean each line
        lines = cleaned_text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Skip section headers
            if any(header in line.lower() for header in [
                "bagian 1", "bagian 2", "analisis kondisi", "rekomendasi pembelajaran",
                "tugasmu sebagai", "prinsip komunikasi", "hindari penggunaan",
                "**", "###", "####"
            ]):
                continue
            
            # Clean the line from all patterns
            cleaned_line = line
            for pattern in clean_patterns:
                cleaned_line = re.sub(pattern, '', cleaned_line, flags=re.IGNORECASE)
            
            cleaned_line = cleaned_line.strip()
            
            # Only keep meaningful content
            if len(cleaned_line) > 15:
                cleaned_lines.append(cleaned_line)
        
        # Join lines into paragraphs with proper spacing
        if not cleaned_lines:
            return [cleaned_text]  # Fallback to original
        
        # Group lines into paragraphs (assuming double line breaks separate paragraphs)
        paragraphs = []
        current_paragraph = []
        
        for line in cleaned_lines:
            # Check if this line might be a new paragraph start
            if (line[0].isupper() and len(current_paragraph) > 0 and 
                any(end_marker in current_paragraph[-1] for end_marker in ['.', '!', '?', ':'])):
                # Start new paragraph
                if current_paragraph:
                    paragraphs.append(' '.join(current_paragraph))
                current_paragraph = [line]
            else:
                current_paragraph.append(line)
        
        # Add the last paragraph
        if current_paragraph:
            paragraphs.append(' '.join(current_paragraph))
        
        # Join paragraphs with double line breaks for PDF display
        result = '\n\n'.join(paragraphs)
        
        # Final cleanup - remove any remaining artifacts
        result = re.sub(r'\s+', ' ', result)  # Multiple spaces to single space
        result = re.sub(r'\n\s+\n', '\n\n', result)  # Clean up paragraph breaks
        
        return [result] if result.strip() else [cleaned_text]
        
    except Exception as e:
        print(f"Error parsing AI recommendations for PDF: {str(e)}")
        # Return original text as fallback, but clean basic patterns
        cleaned_basic = re.sub(r'^\s*[•\-\*\d+\.]\s*', '', recommendations_text, flags=re.MULTILINE)
        return [cleaned_basic] if cleaned_basic.strip() else [recommendations_text]

def get_teacher_style_recommendations(current_level, correct, incorrect, collection_name, learning_objectives=[], mastery_level="sedang"):
    """
    Membuat rekomendasi pembelajaran dengan gaya guru yang humanize dan personal
    """
    total_questions = correct + incorrect
    accuracy_percentage = (correct / total_questions * 100) if total_questions > 0 else 0
    
    # Ekstrak topik utama dari tujuan pembelajaran
    main_topics = []
    for objective in learning_objectives:
        # Ambil kata kunci penting dari tujuan pembelajaran
        words = re.findall(r'\b[a-zA-Z]{4,}\b', objective.lower())
        main_topics.extend([w for w in words if w not in ["siswa", "dapat", "mampu", "memahami", "menjelaskan", "menggunakan", "menerapkan"]])
    
    # Gunakan topik yang paling relevan
    from collections import Counter
    relevant_topics = [topic for topic, count in Counter(main_topics).most_common(3)] if main_topics else [collection_name]
    primary_topic = relevant_topics[0] if relevant_topics else collection_name
    
    # Analisis kondisi belajar siswa
    learning_condition = ""
    if mastery_level == "rendah":
        learning_condition = f"Kamu udah berusaha dengan baik! Dari {total_questions} soal, {correct} soal udah benar. Sekarang fokus perkuat dasarnya dulu ya."
    elif mastery_level == "sedang":
        learning_condition = f"Bagus! Kamu udah paham konsep dasar {collection_name}. Tinggal diasah lagi di beberapa area, pasti bisa naik level."
    else:  # tinggi
        learning_condition = f"Hebat! Kamu udah menguasai {collection_name} dengan baik. Sekarang saatnya tantangan yang lebih kreatif!"

    recommendations = []
    
    if mastery_level == "rendah":
        recommendations = [
            f"Buat catatan visual pakai diagram atau mind map. Ini bikin kamu lebih mudah ingat konsep {primary_topic}.",
            
            f"Belajar 15 menit setiap hari aja dulu. Konsisten lebih penting daripada lama tapi jarang-jarang.",
            
            f"Jangan malu tanya ke teman atau guru kalau bingung. Kadang penjelasan orang lain bikin kita langsung paham.",
            
            f"Mulai dari soal mudah dulu, baru naik ke yang sulit. Kayak main game, harus step by step."
        ]
        
    elif mastery_level == "sedang":
        recommendations = [
            f"Coba hubungkan konsep {primary_topic} dengan kehidupan sehari-hari. Ini bikin pemahaman lebih dalam.",
            
            f"Tantang dirimu untuk jelasin materi ini ke teman yang belum paham. Kalau bisa jelasin, berarti kamu udah ngerti.",
            
            f"Variasikan jenis soal yang dikerjakan. Jangan monoton, cari yang lebih menantang.",
            
            f"Bikin jadwal review rutin setiap akhir minggu. Belajar perlu diulang-ulang, bukan cuma sekali lewat."
        ]
        
    else:  # tinggi
        recommendations = [
            f"Saatnya eksplorasi {primary_topic} yang lebih mendalam. Cari artikel atau referensi tambahan.",
            
            f"Coba bikin project kecil yang nerapin beberapa konsep sekaligus. Misalnya presentasi atau infografis.",
            
            f"Jadilah mentor buat teman yang masih kesulitan. Ngajarin orang lain bikin pemahamanmu makin kuat.",
            
            f"Tantang dirimu dengan soal olimpiade atau kompetisi. Ini ngasah problem-solving skill yang berguna banget."
        ]
    
    # Tambahkan kata-kata motivasi dan dukungan yang personal
    motivational_messages = [
        f"Setiap orang punya kecepatan belajar yang beda, dan itu normal kok. Yang penting konsisten dan jangan malu tanya.",
        
        f"Aku percaya sama kemampuan kamu. Potensinya ada, tinggal diasah dengan cara yang tepat.",
        
        f"Jangan terlalu keras sama diri sendiri. Belajar itu proses, bukan hasil instan.",
        
        f"Kamu udah di jalur yang benar, cuma perlu penyesuaian strategi dikit aja. Tetap semangat!"
    ]
    
    import random
    selected_motivation = random.choice(motivational_messages)
    
    # Format hasil akhir dengan gaya percakapan guru yang peduli (tanpa numbering untuk PDF)
    result_parts = []
    
    # Tambahkan analisis kondisi
    result_parts.append(learning_condition.strip())
    
    # Tambahkan intro rekomendasi
    result_parts.append("Berdasarkan analisis hasil belajar kamu, berikut beberapa saran yang bisa membantu:")
    
    # Gabungkan semua rekomendasi menjadi paragraf narasi
    for rec in recommendations:
        result_parts.append(rec.strip())
    
    # Tambahkan motivasi
    result_parts.append(selected_motivation.strip())
    
    # Tambahkan penutup
    result_parts.append("Jika ada yang ingin didiskusikan atau butuh penjelasan lebih lanjut, jangan ragu untuk bertanya. Selalu ada dukungan untuk kemajuan belajarmu!")
    
    # Gabungkan semua bagian dengan paragraph breaks
    final_result = '\n\n'.join(result_parts)
    
    return [final_result]  # Return sebagai list dengan satu item untuk kompatibilitas

def get_contextual_default_recommendations(current_level, correct, incorrect, collection_name, 
                                           questions_context, learning_objectives, mastery_level="sedang"):
    """
    Rekomendasi default yang kontekstual dan spesifik berdasarkan tujuan pembelajaran
    """
    total = correct + incorrect
    accuracy = round((correct / total * 100), 2) if total > 0 else 0
    
    # Ekstrak kata kunci dari tujuan pembelajaran
    keywords = []
    for objective in learning_objectives:
        words = re.findall(r'\b\w{3,}\b', objective.lower())
        keywords.extend([w for w in words if w not in ["yang", "dan", "atau", "dengan", "untuk", "pada", "dari"]])
    
    # Gunakan keywords yang paling sering muncul
    from collections import Counter
    common_keywords = [word for word, count in Counter(keywords).most_common(5)] if keywords else []
    
    # Struktur rekomendasi berdasarkan domain
    recommendations = {
        "PENGUATAN KONSEP": [],
        "PRAKTIK TERAPAN": [],
        "EVALUASI & SINTESIS": []
    }
    
    # Rekomendasi berdasarkan tingkat penguasaan
    if mastery_level == "rendah":
        recommendations["PENGUATAN KONSEP"] = [
            f"Buat kartu belajar (flashcards) untuk konsep dasar {common_keywords[0] if common_keywords else collection_name}. Tulis definisi pada satu sisi kartu dan contoh sederhana pada sisi lainnya. Tinjau kartu-kartu ini selama 15 menit setiap hari untuk membangun pemahaman dasar yang kuat.",
            f"Gunakan teknik 'chunking' dengan membagi materi {collection_name} menjadi bagian-bagian kecil yang lebih mudah dipahami. Fokus pada satu konsep dalam satu waktu, dan buat diagram atau peta pikiran sederhana untuk visualisasi konsep."
        ]
        recommendations["PRAKTIK TERAPAN"] = [
            f"Selesaikan 5 soal latihan tingkat dasar tentang {common_keywords[0] if common_keywords else collection_name} dengan pendampingan. Tuliskan setiap langkah penyelesaian dan mintalah umpan balik langsung untuk mengidentifikasi kesalahan sedini mungkin.",
            "Ikuti tutorial terstruktur step-by-step dan praktikkan dengan contoh-contoh sederhana. Ulangi proses yang sama dengan variasi kecil untuk memastikan pemahaman dasar."
        ]
        recommendations["EVALUASI & SINTESIS"] = [
            f"Buat jurnal belajar harian yang mencatat: (1) Konsep {common_keywords[0] if common_keywords else 'utama'} yang dipelajari hari ini, (2) Pertanyaan spesifik yang masih membingungkan, dan (3) Contoh penerapan sederhana. Review jurnal ini setiap minggu untuk melihat kemajuan pemahaman."
        ]
    
    elif mastery_level == "sedang":
        recommendations["PENGUATAN KONSEP"] = [
            f"Buat peta hubungan (concept map) yang menghubungkan konsep {common_keywords[0] if common_keywords else 'utama'} dengan {common_keywords[1] if len(common_keywords) > 1 else 'konsep terkait'}. Identifikasi minimal 3 hubungan antar konsep dan jelaskan bagaimana konsep-konsep tersebut saling mempengaruhi.",
            f"Analisis 3-5 contoh kesalahan umum dalam penerapan {common_keywords[0] if common_keywords else collection_name}. Untuk setiap kesalahan, identifikasi miskonsepsi yang mendasarinya dan tuliskan penjelasan yang benar."
        ]
        recommendations["PRAKTIK TERAPAN"] = [
            f"Kerjakan 7 soal latihan dengan variasi konteks berbeda untuk konsep {common_keywords[0] if common_keywords else 'utama'}. Pilih soal dari tingkat kesulitan sedang dan catat pola/strategi penyelesaian untuk setiap jenis soal.",
            "Bentuk kelompok diskusi 3-4 orang untuk membahas dan memecahkan kasus penerapan dari materi yang dipelajari. Bergiliran menjelaskan pendekatan penyelesaian dan memberikan umpan balik konstruktif."
        ]
        recommendations["EVALUASI & SINTESIS"] = [
            f"Kembangkan mini proyek 3-hari yang mengintegrasikan konsep {common_keywords[0] if common_keywords else collection_name} dalam situasi praktis. Dokumentasikan: (1) Tujuan proyek, (2) Konsep yang diterapkan, (3) Tantangan yang dihadapi, dan (4) Solusi yang dikembangkan."
        ]
    
    else:  # tinggi
        recommendations["PENGUATAN KONSEP"] = [
            f"Telusuri minimal 2 sumber referensi lanjutan tentang {common_keywords[0] if common_keywords else collection_name} di luar materi kelas. Buat ringkasan perbandingan yang mengidentifikasi perspektif berbeda atau pengembangan terbaru dalam topik ini.",
            f"Kembangkan model analisis untuk mengklasifikasikan dan mengevaluasi berbagai pendekatan dalam {collection_name}. Sertakan kriteria evaluasi dan justifikasi untuk setiap kategori dalam model tersebut."
        ]
        recommendations["PRAKTIK TERAPAN"] = [
            f"Desain dan implementasikan proyek kompleks yang mengintegrasikan minimal 3 konsep utama: {', '.join(common_keywords[:3]) if len(common_keywords) > 2 else collection_name}. Proyek harus mencakup elemen kreatif dan solusi untuk masalah otentik.",
            "Ciptakan materi pembelajaran (tutorial video 5-10 menit atau panduan tertulis) untuk menjelaskan konsep kompleks kepada teman sekelas. Presentasikan materi ini dan kumpulkan umpan balik untuk perbaikan."
        ]
        recommendations["EVALUASI & SINTESIS"] = [
            f"Kembangkan portofolio digital yang mendokumentasikan perjalanan belajar Anda dalam {collection_name}. Sertakan refleksi mendalam tentang: (1) Perkembangan pemahaman konseptual, (2) Penerapan praktis yang telah dilakukan, (3) Kendala yang dihadapi dan cara mengatasinya, dan (4) Rencana pengembangan kompetensi selanjutnya."
        ]
    
    # Tambahkan konteks spesifik dari tujuan pembelajaran jika ada
    if learning_objectives and len(learning_objectives) > 0:
        # Ambil tujuan pembelajaran pertama untuk rekomendasi tambahan
        objective = learning_objectives[0]
        if ":" in objective:
            objective = objective.split(":", 1)[1].strip()
            
        recommendations["EVALUASI & SINTESIS"].append(
            f"Lakukan refleksi diri terkait tujuan pembelajaran '{objective}'. Buat rubrik evaluasi dengan skala 1-5 yang menilai tingkat pemahaman dan keterampilan saat ini. Identifikasi bukti konkret untuk setiap skor dan buat rencana spesifik untuk meningkatkan area yang masih lemah."
        )
    
    return recommendations
# =========================================
# MAIN ENTRY POINT
def generate_questions_with_gemini(prompt):
    """
    Generate soal menggunakan Gemini AI dan return hasil parsing
    """
    try:
        print(f"[{datetime.datetime.now()}] Mengirim prompt ke Gemini AI ({len(prompt)} chars)...")
        response = model.generate_content(prompt)
        raw_text = response.text.strip()
        print(f"[{datetime.datetime.now()}] Respons dari Gemini AI diterima.")
        
        # Parse questions menggunakan emergency parser
        questions = emergency_json_parser(raw_text)
        
        # Apply MST distribution enforcement
        final_questions = enforce_mst_distribution(questions)
        
        return final_questions
        
    except Exception as e:
        print(f"Error generating questions with Gemini: {e}")
        return []


@app.route('/test-create-questions', methods=['GET'])
def test_create_questions_api():
    """Test endpoint untuk generate MST soal sesuai distribusi 94 soal (no login)"""
    collection_name = "Test MST Collection"
    description = "Pengenalan Aplikasi Microsoft Word - Materi mengenai pengenalan interface Microsoft Word dan fungsi dasar seperti membuat dokumen baru, menyimpan, dan format teks dasar."
    
    try:
        # Generate mock module components for test
        mock_components = {
            'collection_name': collection_name,
            'mata_pelajaran': 'Informatika',
            'kelas': 'X',
            'tujuan_pembelajaran': [description],
            'kompetensi_awal': 'Siswa memahami penggunaan komputer dasar',
            'pemahaman_bermakna': ['Pengenalan interface aplikasi']
        }
        
        # Generate prompt
        prompt = create_optimized_prompt_with_good_structure(mock_components)
        
        # Generate dengan Gemini
        generated_questions = generate_questions_with_gemini(prompt)
        
        if not generated_questions:
            return jsonify({
                'success': False, 
                'message': 'Gagal generate soal',
                'prompt_used': prompt[:500] + "..." if len(prompt) > 500 else prompt
            }), 500
        
        # Format response dengan distribusi detail per level
        response_data = {
            'success': True,
            'message': f'Berhasil generate {len(generated_questions)} soal MST (target: 94 soal)',
            'total_questions': len(generated_questions),
            'target_questions': 94,
            'questions': generated_questions,
            'difficulty_distribution': {
                'Easy': sum(1 for q in generated_questions if q.get('difficulty') == 'Easy'),
                'Medium': sum(1 for q in generated_questions if q.get('difficulty') == 'Medium'), 
                'Hard': sum(1 for q in generated_questions if q.get('difficulty') == 'Hard')
            },
            'level_distribution': {
                f'L{i}': sum(1 for q in generated_questions if q.get('level') == i) for i in range(1, 6)
            },
            'expected_distribution': {
                'L1': 'E:5, M:5',
                'L2': 'E:12, H:8', 
                'L3': 'E:5, M:8, H:7',
                'L4': 'E:4, M:11, H:7',
                'L5': 'E:11, M:7, H:4',
                'Total': '94 soal (10+20+20+22+22)'
            },
            'prompt_used': prompt[:500] + "..." if len(prompt) > 500 else prompt
        }
        
        return jsonify(response_data), 200
        
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error: {str(e)}',
            'traceback': traceback.format_exc()
        }), 500


# =========================================
if __name__ == '__main__':  
    with app.app_context():
        try:
            # Inisialisasi database
            init_db()
            # Tes koneksi database
            test_db_connection()
        except Exception as e:
            print(f"Error inisialisasi database: {str(e)}")

    app.run(debug=True)