import datetime
from sqlalchemy import func, text
from flask import Flask, render_template, request, jsonify, redirect, url_for, flash
from flask_sqlalchemy import SQLAlchemy
import google.generativeai as genai
import pandas as pd
import json
import os
import re
import random
import PyPDF2
from docx import Document
import mammoth
import tempfile
import pandas as pd
from io import BytesIO
from flask import send_file
from werkzeug.security import generate_password_hash, check_password_hash
from flask import Blueprint, request, jsonify
from flask_login import login_required, current_user, LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from flask import send_file
import pdfkit  # Anda perlu menginstall: pip install pdfkit
import tempfile
import os
from xhtml2pdf import pisa
from io import BytesIO

app = Flask(__name__)
# Add secret key for session
app.secret_key = '70315ac5aa5fee60d775f2ad95d2a163b24a4ba65583df64620acc3a283ccbc8'  # Replace with a real secret key in production

# Initialize Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'  # Specify your login route

@login_manager.user_loader
def load_user(user_id):
    if isinstance(user_id, str):
        if user_id.startswith('guru_'):
            actual_id = int(user_id.split('_')[1])
            return Guru.query.get(actual_id)
        elif user_id.startswith('siswa_'):
            actual_id = int(user_id.split('_')[1])
            return Siswa.query.get(actual_id)
    return None

# Konfigurasi MySQL
app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+pymysql://root:@localhost/asesment_diagnostik_db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['SQLALCHEMY_ECHO'] = True  # Tambahan untuk debug SQL queries
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

db = SQLAlchemy(app)

# =========================================
# KONFIGURASI GEMINI AI
# =========================================
# Gunakan environment variable untuk API key
my_api_key_gemini = 'AIzaSyCcx9iCB6oGXUPGoW2Ot3Mk7JjIreeCzRQ'
genai.configure(api_key=my_api_key_gemini)
model = genai.GenerativeModel('gemini-2.0-flash')

# =========================================
# DATABASE MODELS
# =========================================
# Model untuk database guru
class Guru(db.Model, UserMixin):
    __tablename__ = 'guru'
    
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), unique=True, nullable=False)
    email = db.Column(db.String(100), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    nama = db.Column(db.String(100), nullable=False)
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    
    # Property untuk identifikasi tipe pengguna
    @property
    def user_type(self):
        return 'guru'
    
    # Override get_id method untuk Flask-Login
    def get_id(self):
        return f"guru_{self.id}"
    
    def set_password(self, password):
        self.password_hash = generate_password_hash(password)
    
    def check_password(self, password):
        return check_password_hash(self.password_hash, password)
    
    # Metode untuk mengakses hasil siswa
    def get_hasil_siswa(self):
        hasil_siswa = db.session.query(SiswaResult).join(
            Siswa, SiswaResult.siswa_id == Siswa.id
        ).join(
            SiswaAnswer, SiswaAnswer.siswa_id == Siswa.id
        ).join(
            Question, SiswaAnswer.question_id == Question.id
        ).filter(
            Question.guru_id == self.id
        ).distinct(
            SiswaResult.id
        ).all()
        
        return hasil_siswa

# Model untuk pertanyaan yang dibuat oleh guru
class Question(db.Model):
    __tablename__ = 'questions'
    
    id = db.Column(db.Integer, primary_key=True)
    guru_id = db.Column(db.Integer, db.ForeignKey('guru.id'), nullable=False)
    collection_id = db.Column(db.Integer, db.ForeignKey('question_collections.id'), nullable=True)
    level = db.Column(db.Integer, nullable=False)
    soal = db.Column(db.Text, nullable=False)
    jawaban_benar = db.Column(db.Text, nullable=True)
    options = db.Column(db.Text, nullable=True)
    question_type = db.Column(db.String(20), default='multiple_choice')
    p = db.Column(db.Float, nullable=False)
    bobot = db.Column(db.Integer, nullable=False)
    explanation = db.Column(db.Text)
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    version = db.Column(db.Integer, default=1)
    is_current = db.Column(db.Boolean, default=True)
    is_validated = db.Column(db.Boolean, default=False) # NEW: Kolom is_validated
    
    
    # Relasi dengan guru
    guru = db.relationship('Guru', backref='questions')

# Model untuk database siswa
class Siswa(db.Model, UserMixin):
    __tablename__ = 'siswa'
    
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), unique=True, nullable=False)
    email = db.Column(db.String(100), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    nama = db.Column(db.String(100), nullable=False)
    kelas = db.Column(db.String(100), nullable=True)  # Changed from foreign key to string field
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    
    # Relasi dengan jawaban siswa
    jawaban = db.relationship('SiswaAnswer', backref='siswa', lazy=True)
    
    # Relasi dengan hasil siswa
    hasil = db.relationship('SiswaResult', backref='siswa', lazy=True)
    
    # Property untuk identifikasi tipe pengguna
    @property
    def user_type(self):
        return 'siswa'
    
    # Override get_id method untuk Flask-Login
    def get_id(self):
        return f"siswa_{self.id}"
    
    def set_password(self, password):
        self.password_hash = generate_password_hash(password)
    
    def check_password(self, password):
        return check_password_hash(self.password_hash, password)
    
    # Metode untuk mendapatkan hasil sendiri
    def get_my_result(self, kelas_id):
        return SiswaResult.query.filter_by(
            siswa_id=self.id,
            kelas_id=kelas_id
        ).first()

# Model untuk hasil siswa (pengganti SiswaResult)
class SiswaResult(db.Model):
    __tablename__ = 'siswa_results'
    
    id = db.Column(db.Integer, primary_key=True)
    siswa_id = db.Column(db.Integer, db.ForeignKey('siswa.id'), nullable=False)
    collection_id = db.Column(db.Integer, db.ForeignKey('question_collections.id'), nullable=True)  # <-- TAMBAHKAN INI
    correct = db.Column(db.Integer, default=0)
    incorrect = db.Column(db.Integer, default=0)
    current_level = db.Column(db.Integer, default=1)
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    updated_at = db.Column(db.DateTime, server_default=db.func.now(), onupdate=db.func.now())

# Model untuk jawaban siswa (pengganti SiswaAnswer)
class SiswaAnswer(db.Model):
    __tablename__ = 'siswa_answers'
    
    id = db.Column(db.Integer, primary_key=True)
    siswa_id = db.Column(db.Integer, db.ForeignKey('siswa.id'), nullable=False)
    question_id = db.Column(db.Integer, db.ForeignKey('questions.id'), nullable=False)
    siswa_answer = db.Column(db.Text, nullable=False)
    is_correct = db.Column(db.Boolean, default=False)
    level = db.Column(db.Integer, nullable=False)
    answered_at = db.Column(db.DateTime, server_default=db.func.now())
    
    # Relasi dengan Question
    question = db.relationship('Question', backref="siswa_answers")

results_bp = Blueprint('results', __name__)


# =========================================
# TABEL RELASI MANY-TO-MANY (harus didefinisikan SEBELUM model yang menggunakannya)
# =========================================
collection_students = db.Table('collection_students',
    db.Column('id', db.Integer, primary_key=True),
    db.Column('collection_id', db.Integer, db.ForeignKey('question_collections.id', ondelete='CASCADE'), nullable=False),
    db.Column('siswa_id', db.Integer, db.ForeignKey('siswa.id', ondelete='CASCADE'), nullable=False),
    db.Column('added_at', db.DateTime, server_default=db.func.now()),
    db.UniqueConstraint('collection_id', 'siswa_id', name='unique_collection_student')
)

# Model untuk koleksi soal
class QuestionCollection(db.Model):
    __tablename__ = 'question_collections'
    
    id = db.Column(db.Integer, primary_key=True)
    guru_id = db.Column(db.Integer, db.ForeignKey('guru.id'), nullable=False)
    name = db.Column(db.String(100), nullable=False)
    description = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    
    # Relasi dengan guru
    guru = db.relationship('Guru', backref='collections')
    
    # Tambahkan relasi dengan siswa
    students = db.relationship('Siswa', secondary=collection_students, 
                              backref=db.backref('collections', lazy='dynamic'))

# Model untuk relasi many-to-many antara Question dan QuestionCollection
class CollectionQuestion(db.Model):
    __tablename__ = 'collection_questions'
    
    id = db.Column(db.Integer, primary_key=True)
    collection_id = db.Column(db.Integer, db.ForeignKey('question_collections.id'), nullable=False)
    question_id = db.Column(db.Integer, db.ForeignKey('questions.id'), nullable=False)
    
    # Relasi
    collection = db.relationship('QuestionCollection', backref='collection_questions')
    question = db.relationship('Question', backref='collection_questions')


class QuestionVersion(db.Model):
    __tablename__ = 'question_versions'
    
    id = db.Column(db.Integer, primary_key=True)
    question_id = db.Column(db.Integer, db.ForeignKey('questions.id'), nullable=False)
    soal = db.Column(db.Text, nullable=False)
    options = db.Column(db.Text)
    jawaban_benar = db.Column(db.Text)
    explanation = db.Column(db.Text)
    p = db.Column(db.Float)
    version = db.Column(db.Integer)
    created_at = db.Column(db.DateTime, server_default=db.func.now())
    
    # Relasi ke pertanyaan utama
    question = db.relationship('Question', backref='versions')

# Fungsi pemeriksaan apakah pengguna adalah guru
def is_guru():
    return hasattr(current_user, 'kelas')

def get_hasil_siswa(self):
    # Since there's no Kelas relationship, just get all student results
    hasil_siswa = SiswaResult.query.all()
    return hasil_siswa

# Rute untuk guru mengakses hasil siswa
@results_bp.route('/guru/hasil-siswa', methods=['GET'])
@login_required
def guru_access_hasil_siswa():
    # Pastikan pengguna adalah guru
    if not is_guru():
        return jsonify({'message': 'Akses ditolak. Hanya guru yang dapat mengakses hasil siswa.'}), 403
    
    # Ambil hasil siswa dari kelas yang dimiliki guru
    hasil_siswa = current_user.get_hasil_siswa()
    
    # Format hasil untuk response
    hasil_formatted = []
    for hasil in hasil_siswa:
        # Dapatkan data siswa
        siswa = Siswa.query.get(hasil.siswa_id)
        
        hasil_formatted.append({
            'id': hasil.id,
            'siswa': {
                'id': siswa.id,
                'nama': siswa.nama,
                'username': siswa.username,
                'kelas': siswa.kelas
            },
            'correct': hasil.correct,
            'incorrect': hasil.incorrect,
            'current_level': hasil.current_level,
            'created_at': hasil.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            'updated_at': hasil.updated_at.strftime('%Y-%m-%d %H:%M:%S')
        })
    
    return jsonify({
        'message': 'Berhasil mengambil hasil siswa',
        'hasil': hasil_formatted
    }), 200

# Rute untuk guru mengakses detail hasil siswa tertentu
@results_bp.route('/guru/hasil-siswa/<int:siswa_id>', methods=['GET'])
@login_required
def guru_access_detail_hasil_siswa(siswa_id):
    if not is_guru():
        return jsonify({'message': 'Akses ditolak. Hanya guru yang dapat mengakses hasil siswa.'}), 403
    
    # Check if student exists
    siswa = Siswa.query.get(siswa_id)
    if not siswa:
        return jsonify({'message': 'Siswa tidak ditemukan.'}), 404
    
    # Get student results
    hasil = SiswaResult.query.filter_by(siswa_id=siswa_id).all()
    
    # Get student answers
    jawaban = SiswaAnswer.query.filter_by(siswa_id=siswa_id).all()
    
    # Format results
    hasil_formatted = []
    for h in hasil:
        hasil_formatted.append({
            'id': h.id,
            'correct': h.correct,
            'incorrect': h.incorrect,
            'current_level': h.current_level,
            'created_at': h.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            'updated_at': h.updated_at.strftime('%Y-%m-%d %H:%M:%S')
        })
    
    # Format answers
    jawaban_formatted = []
    for j in jawaban:
        jawaban_formatted.append({
            'id': j.id,
            'question_id': j.question_id,
            'siswa_answer': j.siswa_answer,
            'is_correct': j.is_correct,
            'level': j.level,
            'answered_at': j.answered_at.strftime('%Y-%m-%d %H:%M:%S')
        })
    
    return jsonify({
        'message': 'Berhasil mengambil detail hasil siswa',
        'siswa': {
            'id': siswa.id,
            'nama': siswa.nama,
            'username': siswa.username,
            'kelas': siswa.kelas
        },
        'hasil': hasil_formatted,
        'jawaban': jawaban_formatted
    }), 200

# Rute untuk siswa melihat hasil sendiri
@results_bp.route('/siswa/hasil-saya', methods=['GET'])
@login_required
def siswa_access_hasil_sendiri():
    if is_guru():
        return jsonify({'message': 'Akses ditolak. Rute ini hanya untuk siswa.'}), 403
    
    # Get student results
    hasil = SiswaResult.query.filter_by(siswa_id=current_user.id).first()
    
    if not hasil:
        return jsonify({'message': 'Hasil tidak ditemukan.'}), 404
    
    # Get student answers
    jawaban = SiswaAnswer.query.filter_by(siswa_id=current_user.id).all()
    
    # Format answers
    jawaban_formatted = []
    for j in jawaban:
        jawaban_formatted.append({
            'question_id': j.question_id,
            'siswa_answer': j.siswa_answer,
            'is_correct': j.is_correct,
            'level': j.level,
            'answered_at': j.answered_at.strftime('%Y-%m-%d %H:%M:%S')
        })
    
    return jsonify({
        'message': 'Berhasil mengambil hasil',
        'hasil': {
            'correct': hasil.correct,
            'incorrect': hasil.incorrect,
            'current_level': hasil.current_level,
            'kelas': current_user.kelas  # Use the student's kelas field
        },
        'jawaban': jawaban_formatted
    }), 200

# Fungsi untuk mendaftarkan blueprint ke aplikasi
def register_results_routes(app):
    app.register_blueprint(results_bp, url_prefix='/api')

# =========================================
# AUTH ROUTES
# =========================================
@login_manager.user_loader
def load_user(user_id):
    try:
        if isinstance(user_id, str):
            if user_id.startswith('guru_'):
                actual_id = int(user_id.split('_')[1])
                return Guru.query.get(actual_id)
            elif user_id.startswith('siswa_'):
                actual_id = int(user_id.split('_')[1])
                return Siswa.query.get(actual_id)
            else:
                # Coba langsung sebagai ID numerik
                return Siswa.query.get(int(user_id))
        elif isinstance(user_id, int):
            # Default ke Siswa jika ID adalah integer langsung
            return Siswa.query.get(user_id)
    except (ValueError, IndexError, TypeError) as e:
        # Log error jika terjadi masalah saat parsing ID
        print(f"Error loading user with ID {user_id}: {str(e)}")
    return None

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        # Handle login logic here
        username = request.form.get('username')
        password = request.form.get('password')
        role = request.form.get('role')  # 'guru' or 'siswa'
        
        if role == 'guru':
            user = Guru.query.filter_by(username=username).first()
        else:
            user = Siswa.query.filter_by(username=username).first()
        
        if user and user.check_password(password):
            # Login the user with Flask-Login
            login_user(user)
            
            # Check if the user was trying to access a protected page
            next_page = request.args.get('next')
            
            # Login successful
            if next_page:
                return redirect(next_page)
            elif role == 'guru':
                return redirect(url_for('guru'))
            else:
                return redirect(url_for('DashboardSiswa'))
        else:
            flash("Username atau password salah", "error")
            return redirect(url_for('login'))
    
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('Anda telah berhasil logout', 'success')
    return redirect(url_for('login'))

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        nama = request.form.get('nama')
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')
        role = request.form.get('role')
        
        # NEW: Ambil kelas dari form (akan ada di request jika role adalah siswa)
        kelas = request.form.get('kelas') 

        if password != confirm_password:
            # Render ulang template dengan pesan error dan data yang sudah diisi
            return render_template('register.html', error="Password tidak sama",
                                   nama=nama, username=username, email=email, role=role, kelas=kelas)
        
        if role == 'guru':
            if Guru.query.filter_by(username=username).first() or Guru.query.filter_by(email=email).first():
                return render_template('register.html', error="Username atau email sudah terdaftar",
                                       nama=nama, username=username, email=email, role=role, kelas=kelas)
            
            new_user = Guru(
                username=username,
                email=email,
                nama=nama
            )
            new_user.set_password(password)
        else: # role == 'siswa'
            if Siswa.query.filter_by(username=username).first() or Siswa.query.filter_by(email=email).first():
                return render_template('register.html', error="Username atau email sudah terdaftar",
                                       nama=nama, username=username, email=email, role=role, kelas=kelas)
            
            # Pastikan kelas dipilih untuk siswa
            if not kelas:
                return render_template('register.html', error="Kelas harus dipilih untuk siswa",
                                       nama=nama, username=username, email=email, role=role, kelas=kelas)

            new_user = Siswa(
                username=username,
                email=email,
                nama=nama,
                kelas=kelas # Simpan kelas siswa
            )
            new_user.set_password(password)
        
        db.session.add(new_user)
        db.session.commit()
        
        if role == 'guru':
            flash('Registrasi guru berhasil!', 'success')
            return redirect(url_for('login')) # Arahkan ke halaman login setelah registrasi
        else:
            flash('Registrasi siswa berhasil!', 'success')
            return redirect(url_for('login')) # Arahkan ke halaman login setelah registrasi
    
    return render_template('register.html')

# =========================================
# TRANSISI level DENGAN PROBABILITAS
# =========================================
# Fungsi next_level yang baru berdasarkan diagram
def next_level(current_level, is_correct):
    if current_level == 1:
        return 2  # Dari level 1 selalu ke level 2
    elif current_level == 2:
        return 4 if is_correct else 3  # Benar ke atas (4), salah ke bawah (3)
    elif current_level == 4:
        return 7 if is_correct else 6  # Benar ke atas (7), salah ke bawah (6)
    elif current_level == 3:
        return 6 if is_correct else 5  # Benar ke atas (6), salah ke bawah (5)
    elif current_level in [5, 6, 7]:
        return None  # Sudah di level akhir
    else:
        return None  # Untuk keamanan

# level akhir tetap sama
FINAL_levelS = [5, 6, 7]

# =========================================
# FUNGSI MENENTUKAN BOBOT (ASSIGN_WEIGHT)
# =========================================
def assign_weight(p: float) -> int:
    if p > 0.90:
        return 1
    elif 0.90 >= p > 0.62:
        return 2
    elif abs(p - 0.62) < 1e-6:
        return 3
    elif 0.62 > p > 0.20:
        return 4
    else:
        return 5

# =========================================
# FUNGSI EVALUASI JAWABAN
# =========================================
def is_answer_correct(user_answer, expected_answer):
    """
    Evaluasi jawaban pilihan ganda dengan membandingkan jawaban pengguna dengan jawaban yang benar
    """
    # Membersihkan dan menormalisasi jawaban
    user_answer = user_answer.lower().strip()
    expected_answer = expected_answer.lower().strip()
    
    # Untuk pilihan ganda, jawaban harus sama persis (tidak case-sensitive)
    return user_answer == expected_answer


# Fungsi untuk evaluasi dengan Gemini
def evaluate_with_gemini(soal, answer, correct_answer=None, options=None):
    """
    Evaluasi jawaban menggunakan Gemini AI untuk pilihan ganda
    """
    try:
        # Jika kita memiliki jawaban yang benar, cukup bandingkan langsung
        if correct_answer:
            return answer.lower().strip() == correct_answer.lower().strip()
            
        # Jika tidak ada jawaban benar, gunakan AI untuk mengevaluasi
        eval_prompt = f"""
        Soal: {soal}
        Jawaban: "{answer}"
        Opsi: {", ".join(options) if options else "Tidak tersedia"}
        
        Evaluasi apakah jawaban tersebut benar berdasarkan opsi yang tersedia.
        Hanya balas 'benar' atau 'salah' (tanpa teks lain).
        """
        
        response = model.generate_content(eval_prompt)
        ai_eval = response.text.strip().lower()
        
        return "benar" in ai_eval
    except Exception as e:
        print(f"Error evaluasi dengan Gemini: {str(e)}")
        # Fallback ke perbandingan sederhana jika memiliki jawaban benar
        if correct_answer:
            return answer.lower().strip() == correct_answer.lower().strip()
        return False


# =========================================
# FUNGSI UNTUK CARI DAN TAMBAH KOLOM
# =========================================
def check_and_add_columns():
    try:
        print("Memeriksa struktur tabel questions...")
        # Dapatkan informasi kolom dari model
        column_info = db.inspect(db.engine).get_columns('questions')
        column_names = [col['name'] for col in column_info]
        
        # Cek dan tambahkan kolom yang diperlukan
        columns_to_add = []
        
        if 'jawaban_benar' not in column_names:
            columns_to_add.append("jawaban_benar TEXT NULL")
        
        if 'options' not in column_names:
            columns_to_add.append("options TEXT NULL")
            
        if 'question_type' not in column_names:
            columns_to_add.append("question_type VARCHAR(20) DEFAULT 'multiple_choice'")
        
        # Tambahkan kolom yang diperlukan jika ada
        if columns_to_add:
            print(f"Menambahkan kolom yang diperlukan: {', '.join(columns_to_add)}")
            alter_query = f"ALTER TABLE questions ADD COLUMN {', ADD COLUMN '.join(columns_to_add)}"
            db.session.execute(text(alter_query))
            db.session.commit()
            print("Kolom berhasil ditambahkan!")
        else:
            print("Semua kolom yang diperlukan sudah ada dalam tabel")
        
        return True
    except Exception as e:
        print(f"Error ketika memeriksa/menambahkan kolom: {str(e)}")
        return False

# =========================================
# FUNGSI UNTUK KONEKSI DATABASE
# =========================================
def test_db_connection():
    try:
        # Cek apakah koneksi ke database bisa dibuat
        db.session.execute(text("SELECT 1"))
        print("Database connection successful")
        return True
    except Exception as e:
        print(f"Database connection error: {str(e)}")
        return False

# Inisialisasi database dengan app context
def init_db():
    print("Initializing database...")
    db.create_all()
    # Cek dan tambahkan kolom yang diperlukan
    check_and_add_columns()
    print("Database initialized successfully")

# =========================================
# ROUTES
# =========================================
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/guru')
def guru():
    return render_template('guru.html')

@app.route('/DashboardSiswa')
def DashboardSiswa():
    return render_template('DashboardSiswa.html')
@app.route('/qna')
def qna():
    return render_template('qna.html')

# =========================================
# RESULTS PAGE
# =========================================

@app.route('/api/siswa/test_status', methods=['GET'])
@login_required
def get_test_status():
    user_id = request.args.get("user_id")
    collection_id = request.args.get("collection_id")
    
    # Cek keamanan: pastikan user_id sesuai dengan pengguna saat ini
    if str(current_user.id) != str(user_id):
        return jsonify({"success": False, "message": "Akses tidak diizinkan"}), 403
    
    try:
        # Ambil hasil siswa untuk koleksi ini
        result = SiswaResult.query.filter_by(
            siswa_id=user_id,
            collection_id=collection_id
        ).first()
        
        if not result:
            return jsonify({
                "success": True,
                "is_completed": False,
                "current_level": 1
            })
        
        # Cek apakah current_level adalah level final (5, 6, 7)
        is_completed = result.current_level in [5, 6, 7]
        
        return jsonify({
            "success": True,
            "is_completed": is_completed,
            "current_level": result.current_level,
            "correct": result.correct,
            "incorrect": result.incorrect
        })
        
    except Exception as e:
        print(f"Error mendapatkan status tes: {str(e)}")
        return jsonify({
            "success": False,
            "message": f"Error: {str(e)}"
        }), 500

@app.route('/api/siswa/results/<int:collection_id>', methods=['GET'])
@login_required
def get_siswa_test_results(collection_id):
    try:
        user_id = request.args.get("user_id")
        
        # For teacher access (viewing any student's results)
        if hasattr(current_user, 'user_type') and current_user.user_type == 'guru':
            # Check if collection belongs to this teacher
            collection = QuestionCollection.query.filter_by(
                id=collection_id, 
                guru_id=current_user.id
            ).first()
            
            if not collection:
                return jsonify({"success": False, "message": "Access denied - Collection not found or not owned by you"}), 403
                
            # If user_id is provided, get that specific student's results
            # Otherwise, return all results for this collection
            if user_id:
                # Verify this student is assigned to this collection
                student_assigned = db.session.query(collection_students).filter_by(
                    collection_id=collection_id,
                    siswa_id=user_id
                ).first() is not None
                
                if not student_assigned:
                    return jsonify({"success": False, "message": "Student not assigned to this collection"}), 404
            
        # For student access (viewing own results only)
        elif hasattr(current_user, 'user_type') and current_user.user_type == 'siswa':
            # Students can only view their own results
            if str(current_user.id) != str(user_id):
                return jsonify({"success": False, "message": "You can only view your own results"}), 403
                
            # Verify student has access to this collection
            student_assigned = db.session.query(collection_students).filter_by(
                collection_id=collection_id,
                siswa_id=current_user.id
            ).first() is not None
            
            if not student_assigned:
                return jsonify({"success": False, "message": "You don't have access to this collection"}), 403
        else:
            return jsonify({"success": False, "message": "Invalid user type"}), 403
        
        # Get student results for this collection
        result_query = SiswaResult.query.filter_by(
            collection_id=collection_id
        )
        
        if user_id:
            result_query = result_query.filter_by(siswa_id=user_id)
            
        results = result_query.all()
        
        if not results:
            # Return empty result instead of error if no results found
            return jsonify({
                "success": True,
                "message": "No results found",
                "correct": 0,
                "incorrect": 0,
                "current_level": 1,
                "answers": []
            })
            
        # Format results
        formatted_results = []
        for result in results:
            # Get student
            student = Siswa.query.get(result.siswa_id)
            
            # Get answer history
            answer_history = get_student_answer_history(result.siswa_id, collection_id)
            
            formatted_results.append({
                "student_id": result.siswa_id,
                "student_name": student.nama if student else "Unknown",
                "student_class": student.kelas if student else "Unknown",
                "correct": result.correct,
                "incorrect": result.incorrect,
                "current_level": result.current_level,
                "answers": answer_history
            })
        
        # If single student requested, return first result
        if user_id:
            if formatted_results:
                return jsonify({
                    "success": True,
                    "correct": formatted_results[0]["correct"],
                    "incorrect": formatted_results[0]["incorrect"],
                    "current_level": formatted_results[0]["current_level"],
                    "answers": formatted_results[0]["answers"]
                })
            else:
                return jsonify({
                    "success": True,
                    "message": "No results found",
                    "correct": 0,
                    "incorrect": 0,
                    "current_level": 1,
                    "answers": []
                })
        
        # Return all results for teacher
        return jsonify({
            "success": True,
            "results": formatted_results
        })
        
    except Exception as e:
        print(f"Error getting student test results: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"Server error: {str(e)}"
        }), 500

# Helper function to get student answer history
def get_student_answer_history(student_id, collection_id):
    try:
        # Get collection questions
        collection_question_ids = db.session.query(CollectionQuestion.question_id)\
            .filter(CollectionQuestion.collection_id == collection_id).all()
        collection_question_ids = [id[0] for id in collection_question_ids]
        
        if not collection_question_ids:
            return []
            
        # Get answers for this student and these questions
        answers = db.session.query(SiswaAnswer, Question)\
            .join(Question, SiswaAnswer.question_id == Question.id)\
            .filter(
                SiswaAnswer.siswa_id == student_id,
                Question.id.in_(collection_question_ids)
            )\
            .order_by(SiswaAnswer.answered_at.desc())\
            .all()
            
        # Format answers
        formatted_answers = []
        for answer, question in answers:
            formatted_answers.append({
                "question_id": question.id,
                "question": question.soal,
                "user_answer": answer.siswa_answer,
                "correct_answer": question.jawaban_benar,
                "is_correct": answer.is_correct,
                "explanation": question.explanation,
                "level": answer.level,
                "answered_at": answer.answered_at.isoformat() if answer.answered_at else None
            })
            
        return formatted_answers
    except Exception as e:
        print(f"Error getting student answer history: {str(e)}")
        return []

@app.route('/api/collection-analytics/answers', methods=['GET'])
@login_required
def get_filtered_collection_answers():
    try:
        # Required parameter
        collection_id = request.args.get('collection_id')
        
        if not collection_id:
            return jsonify({"success": False, "message": "Collection ID is required"}), 400
            
        # Check if collection belongs to current teacher
        collection = QuestionCollection.query.filter_by(
            id=collection_id, 
            guru_id=current_user.id
        ).first()
        
        if not collection:
            return jsonify({"success": False, "message": "Collection not found or not owned by you"}), 403
            
        # Get collection questions
        collection_question_ids = db.session.query(CollectionQuestion.question_id)\
            .filter(CollectionQuestion.collection_id == collection_id).all()
        collection_question_ids = [id[0] for id in collection_question_ids]
        
        if not collection_question_ids:
            return jsonify({"success": True, "answers": [], "message": "No questions in this collection"})
            
        # Build query for student answers
        query = db.session.query(
            SiswaAnswer, Question, Siswa
        ).join(
            Question, SiswaAnswer.question_id == Question.id
        ).join(
            Siswa, SiswaAnswer.siswa_id == Siswa.id
        ).filter(
            Question.id.in_(collection_question_ids)
        )
        
        # Apply filters
        # 1. Filter by student ID
        student_id = request.args.get('student_id')
        if student_id:
            query = query.filter(SiswaAnswer.siswa_id == student_id)
            
        # 2. Filter by class
        class_filter = request.args.get('class')
        if class_filter and class_filter != 'all':
            query = query.filter(Siswa.kelas == class_filter)
            
        # 3. Filter by level
        level_filter = request.args.get('level')
        if level_filter and level_filter != 'all':
            query = query.filter(SiswaAnswer.level == level_filter)
            
        # 4. Filter by correctness
        is_correct_filter = request.args.get('is_correct')
        if is_correct_filter:
            if is_correct_filter.lower() == 'true':
                query = query.filter(SiswaAnswer.is_correct == True)
            elif is_correct_filter.lower() == 'false':
                query = query.filter(SiswaAnswer.is_correct == False)
                
        # 5. Filter by date range
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        if start_date:
            try:
                start_datetime = datetime.datetime.fromisoformat(start_date)
                query = query.filter(SiswaAnswer.answered_at >= start_datetime)
            except ValueError:
                # Ignore invalid date format
                pass
                
        if end_date:
            try:
                end_datetime = datetime.datetime.fromisoformat(end_date)
                query = query.filter(SiswaAnswer.answered_at <= end_datetime)
            except ValueError:
                # Ignore invalid date format
                pass
                
        # 6. Apply time period filter (shorthand for date range)
        period_filter = request.args.get('period')
        if period_filter and not (start_date or end_date):
            now = datetime.datetime.now()
            
            if period_filter == 'today':
                today_start = now.replace(hour=0, minute=0, second=0, microsecond=0)
                query = query.filter(SiswaAnswer.answered_at >= today_start)
            elif period_filter == 'yesterday':
                yesterday_start = (now - datetime.timedelta(days=1)).replace(hour=0, minute=0, second=0, microsecond=0)
                yesterday_end = now.replace(hour=0, minute=0, second=0, microsecond=0)
                query = query.filter(SiswaAnswer.answered_at >= yesterday_start, SiswaAnswer.answered_at < yesterday_end)
            elif period_filter == 'week':
                week_start = (now - datetime.timedelta(days=7))
                query = query.filter(SiswaAnswer.answered_at >= week_start)
            elif period_filter == 'month':
                month_start = (now - datetime.timedelta(days=30))
                query = query.filter(SiswaAnswer.answered_at >= month_start)
                
        # Sort by most recent first
        query = query.order_by(SiswaAnswer.answered_at.desc())
        
        # Get paginated results
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 50))
        
        # Execute query with pagination
        paginated_results = query.paginate(page=page, per_page=per_page, error_out=False)
        
        # Get total count for stats
        total_count = paginated_results.total
        total_pages = paginated_results.pages
        
        # Format results
        formatted_answers = []
        for answer, question, student in paginated_results.items:
            formatted_answers.append({
                "id": answer.id,
                "student_id": student.id,
                "student_name": student.nama,
                "student_class": student.kelas,
                "question_id": question.id,
                "question": question.soal,
                "user_answer": answer.siswa_answer,
                "correct_answer": question.jawaban_benar,
                "is_correct": answer.is_correct,
                "explanation": question.explanation,
                "level": answer.level,
                "answered_at": answer.answered_at.isoformat() if answer.answered_at else None
            })
            
        # Get unique classes for class filter
        available_classes = db.session.query(Siswa.kelas).distinct().all()
        available_classes = [cls[0] for cls in available_classes if cls[0]]
        
        return jsonify({
            "success": True,
            "answers": formatted_answers,
            "pagination": {
                "total": total_count,
                "page": page,
                "per_page": per_page,
                "total_pages": total_pages
            },
            "filters": {
                "available_classes": available_classes
            }
        })
            
    except Exception as e:
        print(f"Error getting filtered answers: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"Server error: {str(e)}"
        }), 500
    
@app.route('/api/collection-analytics/level-analysis', methods=['GET'])
@login_required
def get_level_analysis():
    try:
        # Required parameter
        collection_id = request.args.get('collection_id')
        
        if not collection_id:
            return jsonify({"success": False, "message": "Collection ID is required"}), 400
            
        # Check if collection belongs to current teacher
        collection = QuestionCollection.query.filter_by(
            id=collection_id, 
            guru_id=current_user.id
        ).first()
        
        if not collection:
            return jsonify({"success": False, "message": "Collection not found or not owned by you"}), 403
            
        # Get students assigned to this collection
        student_ids = db.session.query(collection_students.c.siswa_id)\
            .filter(collection_students.c.collection_id == collection_id)\
            .all()
        student_ids = [id[0] for id in student_ids]
        
        if not student_ids:
            return jsonify({
                "success": True, 
                "message": "No students assigned to this collection",
                "level_distribution": {},
                "level_transitions": {}
            })
            
        # Get current level distribution
        level_distribution = db.session.query(
            SiswaResult.current_level, 
            db.func.count(SiswaResult.siswa_id)
        ).filter(
            SiswaResult.collection_id == collection_id,
            SiswaResult.siswa_id.in_(student_ids)
        ).group_by(
            SiswaResult.current_level
        ).all()
        
        # Format level distribution
        formatted_distribution = {str(level): count for level, count in level_distribution}
        
        # For levels not present in results, set count to 0
        for level in range(1, 8):  # levels 1-7
            if str(level) not in formatted_distribution:
                formatted_distribution[str(level)] = 0
                
        # Analyze level transitions
        # For this, we need to look at the sequence of answers for each student
        level_transitions = {
            "1-2": 0,
            "2-3": 0,
            "2-4": 0,
            "3-5": 0,
            "3-6": 0,
            "4-6": 0,
            "4-7": 0
        }
        
        # Get collection questions
        collection_question_ids = db.session.query(CollectionQuestion.question_id)\
            .filter(CollectionQuestion.collection_id == collection_id).all()
        collection_question_ids = [id[0] for id in collection_question_ids]
        
        if collection_question_ids:
            # For each student, get answers in chronological order
            for student_id in student_ids:
                answers = db.session.query(SiswaAnswer)\
                    .filter(
                        SiswaAnswer.siswa_id == student_id,
                        SiswaAnswer.question_id.in_(collection_question_ids)
                    )\
                    .order_by(SiswaAnswer.answered_at)\
                    .all()
                    
                if answers:
                    # Track level transitions
                    prev_level = None
                    for answer in answers:
                        curr_level = answer.level
                        
                        if prev_level is not None and prev_level != curr_level:
                            transition_key = f"{prev_level}-{curr_level}"
                            if transition_key in level_transitions:
                                level_transitions[transition_key] += 1
                                
                        prev_level = curr_level
        
        # Additional statistics
        # Count students in basic (1-2), intermediate (3-4), and advanced (5-7) levels
        basic_count = formatted_distribution.get("1", 0) + formatted_distribution.get("2", 0)
        intermediate_count = formatted_distribution.get("3", 0) + formatted_distribution.get("4", 0)
        advanced_count = formatted_distribution.get("5", 0) + formatted_distribution.get("6", 0) + formatted_distribution.get("7", 0)
        
        # Count completed students (reached final levels 5, 6, or 7)
        completed_count = advanced_count
        
        # Calculate percentages
        total_students = len(student_ids)
        basic_percent = (basic_count / total_students * 100) if total_students > 0 else 0
        intermediate_percent = (intermediate_count / total_students * 100) if total_students > 0 else 0
        advanced_percent = (advanced_count / total_students * 100) if total_students > 0 else 0
        completed_percent = (completed_count / total_students * 100) if total_students > 0 else 0
        
        return jsonify({
            "success": True,
            "level_distribution": formatted_distribution,
            "level_transitions": level_transitions,
            "statistics": {
                "total_students": total_students,
                "basic_levels": {
                    "count": basic_count,
                    "percentage": round(basic_percent, 2)
                },
                "intermediate_levels": {
                    "count": intermediate_count,
                    "percentage": round(intermediate_percent, 2)
                },
                "advanced_levels": {
                    "count": advanced_count,
                    "percentage": round(advanced_percent, 2)
                },
                "completed": {
                    "count": completed_count,
                    "percentage": round(completed_percent, 2)
                }
            }
        })
        
    except Exception as e:
        print(f"Error getting level analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"Server error: {str(e)}"
        }), 500
    
@app.route('/api/collection-analytics/class-analysis', methods=['GET'])
@login_required
def get_class_analysis():
    try:
        # Required parameter
        collection_id = request.args.get('collection_id')
        
        if not collection_id:
            return jsonify({"success": False, "message": "Collection ID is required"}), 400
            
        # Check if collection belongs to current teacher
        collection = QuestionCollection.query.filter_by(
            id=collection_id, 
            guru_id=current_user.id
        ).first()
        
        if not collection:
            return jsonify({"success": False, "message": "Collection not found or not owned by you"}), 403
            
        # Get students assigned to this collection
        students = db.session.query(Siswa)\
            .join(
                collection_students,
                collection_students.c.siswa_id == Siswa.id
            )\
            .filter(
                collection_students.c.collection_id == collection_id
            )\
            .all()
            
        if not students:
            return jsonify({
                "success": True, 
                "message": "No students assigned to this collection",
                "class_stats": {}
            })
            
        # Group students by class
        class_groups = {}
        for student in students:
            if student.kelas not in class_groups:
                class_groups[student.kelas] = []
            class_groups[student.kelas].append(student.id)
            
        # For each class, get performance statistics
        class_stats = {}
        for kelas, student_ids in class_groups.items():
            # Get results for these students
            results = db.session.query(SiswaResult)\
                .filter(
                    SiswaResult.collection_id == collection_id,
                    SiswaResult.siswa_id.in_(student_ids)
                )\
                .all()
                
            # Calculate stats
            total_correct = sum(r.correct or 0 for r in results)
            total_incorrect = sum(r.incorrect or 0 for r in results)
            total_answers = total_correct + total_incorrect
            accuracy = (total_correct / total_answers * 100) if total_answers > 0 else 0
            
            # Count students by level
            level_counts = {str(n): 0 for n in range(1, 8)}
            for result in results:
                level = str(result.current_level or 1)
                if level in level_counts:
                    level_counts[level] += 1
                    
            # Count completed students (in levels 5-7)
            completed_count = sum(
                1 for r in results if r.current_level in [5, 6, 7]
            )
            
            class_stats[kelas] = {
                "student_count": len(student_ids),
                "total_correct": total_correct,
                "total_incorrect": total_incorrect,
                "accuracy": round(accuracy, 2),
                "level_distribution": level_counts,
                "completed_count": completed_count,
                "completion_rate": round((completed_count / len(student_ids) * 100), 2) if student_ids else 0
            }
            
        return jsonify({
            "success": True,
            "class_stats": class_stats
        })
        
    except Exception as e:
        print(f"Error getting class analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"Server error: {str(e)}"
        }), 500
    
@app.route('/api/collection-analytics/question-analysis', methods=['GET'])
@login_required
def get_question_analysis():
    try:
        # Required parameter
        collection_id = request.args.get('collection_id')
        
        if not collection_id:
            return jsonify({"success": False, "message": "Collection ID is required"}), 400
            
        # Check if collection belongs to current teacher
        collection = QuestionCollection.query.filter_by(
            id=collection_id, 
            guru_id=current_user.id
        ).first()
        
        if not collection:
            return jsonify({"success": False, "message": "Collection not found or not owned by you"}), 403
            
        # Get questions in this collection
        questions = db.session.query(Question)\
            .join(
                CollectionQuestion,
                CollectionQuestion.question_id == Question.id
            )\
            .filter(
                CollectionQuestion.collection_id == collection_id
            )\
            .all()
            
        if not questions:
            return jsonify({
                "success": True, 
                "message": "No questions in this collection",
                "questions": []
            })
            
        # For each question, get answer statistics
        question_stats = []
        for question in questions:
            # Get all answers for this question
            answers = db.session.query(SiswaAnswer)\
                .filter(
                    SiswaAnswer.question_id == question.id
                )\
                .all()
                
            # Calculate statistics
            total_answers = len(answers)
            correct_answers = sum(1 for a in answers if a.is_correct)
            incorrect_answers = total_answers - correct_answers
            accuracy = (correct_answers / total_answers * 100) if total_answers > 0 else 0
            
            # Get class breakdown
            class_breakdown = {}
            for answer in answers:
                student = Siswa.query.get(answer.siswa_id)
                if student and student.kelas:
                    if student.kelas not in class_breakdown:
                        class_breakdown[student.kelas] = {
                            "correct": 0,
                            "incorrect": 0,
                            "total": 0
                        }
                    class_breakdown[student.kelas]["total"] += 1
                    if answer.is_correct:
                        class_breakdown[student.kelas]["correct"] += 1
                    else:
                        class_breakdown[student.kelas]["incorrect"] += 1
            
            # Calculate accuracy per class
            for kelas, stats in class_breakdown.items():
                stats["accuracy"] = round((stats["correct"] / stats["total"] * 100), 2) if stats["total"] > 0 else 0
                
            question_stats.append({
                "id": question.id,
                "level": question.level,
                "soal": question.soal,
                "options": json.loads(question.options) if question.options else [],
                "jawaban_benar": question.jawaban_benar,
                "p": question.p,
                "bobot": question.bobot,
                "total_answers": total_answers,
                "correct_answers": correct_answers,
                "incorrect_answers": incorrect_answers,
                "accuracy": round(accuracy, 2),
                "class_breakdown": class_breakdown
            })
            
        # Sort questions by level
        question_stats.sort(key=lambda q: q["level"])
        
        # Filter by level if requested
        level_filter = request.args.get('level')
        if level_filter and level_filter != 'all':
            question_stats = [q for q in question_stats if q["level"] == int(level_filter)]
            
        return jsonify({
            "success": True,
            "questions": question_stats
        })
        
    except Exception as e:
        print(f"Error getting question analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"Server error: {str(e)}"
        }), 500
    

@app.route('/api/siswa/refresh_answers/<int:collection_id>', methods=['GET'])
@login_required
def refresh_siswa_answers(collection_id):
    user_id = request.args.get("user_id")
    force = request.args.get("force") == "true"
    
    # Cek keamanan: pastikan user_id sesuai dengan pengguna saat ini
    if str(current_user.id) != str(user_id):
        return jsonify({"success": False, "message": "Akses ditolak"}), 403
    
    try:
        # Ambil hasil siswa untuk koleksi ini
        result = SiswaResult.query.filter_by(
            siswa_id=user_id,
            collection_id=collection_id
        ).first()
        
        if not result:
            return jsonify({
                "success": False,
                "message": "Tidak ada hasil ditemukan"
            }), 404
        
        # Dapatkan daftar question_id yang terkait dengan collection ini
        collection_question_ids = db.session.query(CollectionQuestion.question_id)\
            .filter(CollectionQuestion.collection_id == collection_id).all()
        collection_question_ids = [id[0] for id in collection_question_ids]
        
        # Hitung jawaban yang ada di database
        answer_count = db.session.query(func.count(SiswaAnswer.id)).filter(
            SiswaAnswer.siswa_id == user_id,
            SiswaAnswer.question_id.in_(collection_question_ids) if collection_question_ids else False
        ).scalar() or 0
        
        # Hitung total jawaban benar dari actual records
        correct_count = db.session.query(func.count(SiswaAnswer.id)).filter(
            SiswaAnswer.siswa_id == user_id,
            SiswaAnswer.is_correct == True,
            SiswaAnswer.question_id.in_(collection_question_ids) if collection_question_ids else False
        ).scalar() or 0
        
        # Hitung total jawaban salah dari actual records
        incorrect_count = db.session.query(func.count(SiswaAnswer.id)).filter(
            SiswaAnswer.siswa_id == user_id,
            SiswaAnswer.is_correct == False,
            SiswaAnswer.question_id.in_(collection_question_ids) if collection_question_ids else False
        ).scalar() or 0
        
        # Cek inkonsistensi
        has_inconsistency = (result.correct != correct_count or result.incorrect != incorrect_count)
        
        if has_inconsistency or force:
            # Jika ada inkonsistensi atau force=true, perbarui data
            result.correct = correct_count
            result.incorrect = incorrect_count
            db.session.commit()
            
            return jsonify({
                "success": True,
                "message": "Data jawaban telah diperbarui",
                "changes": {
                    "old_correct": result.correct if has_inconsistency else correct_count,
                    "new_correct": correct_count,
                    "old_incorrect": result.incorrect if has_inconsistency else incorrect_count,
                    "new_incorrect": incorrect_count
                }
            })
        
        return jsonify({
            "success": True,
            "message": "Data jawaban sudah konsisten",
            "stats": {
                "correct": correct_count,
                "incorrect": incorrect_count,
                "total": answer_count
            }
        })
        
    except Exception as e:
        print(f"Error menyegarkan jawaban: {str(e)}")
        db.session.rollback()
        return jsonify({
            "success": False,
            "message": f"Error: {str(e)}"
        }), 500

@app.route('/api/questions/<int:question_id>/set_validation', methods=['POST'])
@login_required
def set_question_validation_status(question_id):
    try:
        # Pastikan pengguna adalah guru
        if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
            return jsonify({'success': False, 'message': 'Akses ditolak. Hanya guru yang dapat mengubah status validasi.'}), 403

        question = Question.query.get(question_id)
        if not question:
            return jsonify({'success': False, 'message': 'Soal tidak ditemukan'}), 404

        # Pastikan guru yang login adalah pemilik soal
        if question.guru_id != current_user.id:
            return jsonify({'success': False, 'message': 'Anda tidak memiliki izin untuk mengubah status validasi soal ini'}), 403

        data = request.get_json()
        # Menggunakan 'is_validated' dari body request, default ke nilai saat ini jika tidak ada
        new_status = data.get('is_validated', not question.is_validated) 

        question.is_validated = bool(new_status)
        db.session.commit()

        return jsonify({'success': True, 'message': f'Status validasi soal ID {question_id} berhasil diperbarui.', 'is_validated': question.is_validated}), 200

    except Exception as e:
        db.session.rollback()
        print(f"Error updating question validation status: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f"Terjadi kesalahan server: {str(e)}"}), 500

    
@app.route('/result', methods=['GET'])
@login_required  # Ensure user is logged in
def results():
    # Get collection_id from query parameters if available
    collection_id = request.args.get("collection_id")
    
    # Use the currently logged-in user instead of query param
    if not current_user.is_authenticated:
        return redirect(url_for('login'))
    
    # For siswa (student) users
    if hasattr(current_user, 'user_type') and current_user.user_type == 'siswa':
        # Get the student's results for the specified collection
        user_result = SiswaResult.query.filter_by(
            siswa_id=current_user.id,
            collection_id=collection_id if collection_id else None
        ).first()
        
        if not user_result:
            # Create a new result record if none exists
            return render_template('result.html', 
                                correct=0, 
                                incorrect=0, 
                                current_level=1,
                                accuracy=0,
                                answers=[],
                                collection_name="N/A",
                                collection_id=collection_id,
                                avg_time="N/A")
        
        # Calculate accuracy
        total = user_result.correct + user_result.incorrect
        accuracy = round((user_result.correct / total * 100) if total > 0 else 0, 2)
        
        # Get user answer history with questions for this collection
        query = db.session.query(
            SiswaAnswer, Question
        ).join(
            Question, SiswaAnswer.question_id == Question.id
        ).filter(
            SiswaAnswer.siswa_id == current_user.id
        )
        
        # Add collection filter if specified
        if collection_id:
            query = query.filter(Question.collection_id == collection_id)
            
        # Get collection name if available
        collection_name = "N/A"
        if collection_id:
            collection = QuestionCollection.query.get(collection_id)
            if collection:
                collection_name = collection.name
        
        # Execute query and limit to most recent answers
        answers = query.order_by(SiswaAnswer.answered_at.desc()).limit(10).all()
        
        # Format answers for template
        answer_history = []
        for ua, q in answers:
            answer_history.append({
                'question': q.soal,
                'user_answer': ua.siswa_answer,
                'is_correct': ua.is_correct,
                'level': ua.level,
                'explanation': q.explanation
            })
        
        # Calculate average answer time (if needed)
        avg_time = "N/A"  # Implement logic here if you want to track this
        
        return render_template('result.html', 
                              correct=user_result.correct, 
                              incorrect=user_result.incorrect,
                              current_level=user_result.current_level,
                              accuracy=accuracy,
                              answers=answer_history,
                              collection_name=collection_name,
                              collection_id=collection_id,
                              avg_time=avg_time)
    else:
        # For non-student users or if user_type not defined
        flash("Only students can view their results", "warning")
        return redirect(url_for('index'))
# =========================================
# ENDPOINT RESET DATABASE 
# =========================================
@app.route('/reset_database', methods=['POST'])
def reset_database_endpoint():
    try:
        # BAGIAN YANG DIMODIFIKASI - MULAI
        # Dapatkan daftar ID soal yang ada dalam koleksi (yang harus dipertahankan)
        preserved_question_ids = db.session.query(CollectionQuestion.question_id).distinct().all()
        preserved_question_ids = [id[0] for id in preserved_question_ids]
        
        if preserved_question_ids:
            print(f"Mempertahankan {len(preserved_question_ids)} soal dalam koleksi")
        
        # Pertama, hapus semua jawaban pengguna yang tidak terkait dengan soal dalam koleksi
        print("Menghapus jawaban pengguna...")
        if preserved_question_ids:
            # Hapus jawaban siswa hanya untuk soal yang tidak dalam koleksi
            SiswaAnswer.query.filter(~SiswaAnswer.question_id.in_(preserved_question_ids)).delete(synchronize_session=False)
        else:
            # Jika tidak ada soal yang dipertahankan, hapus semua jawaban
            SiswaAnswer.query.delete()
        db.session.commit()
        
        # Kedua, reset current_level semua user ke level 1
        print("Mengatur ulang level pengguna...")
        SiswaResult.query.update({
            SiswaResult.current_level: 1,
            SiswaResult.correct: 0,
            SiswaResult.incorrect: 0
        })
        db.session.commit()
        
        # Hapus relasi collection_questions untuk soal yang akan dihapus
        # tapi jangan hapus koleksi itu sendiri
        # Relasi untuk soal yang ada dalam koleksi tidak perlu dihapus
        if preserved_question_ids:
            print("Mempertahankan relasi collection_questions untuk soal yang dipertahankan...")
        else:
            print("Menghapus semua relasi collection_questions...")
            CollectionQuestion.query.delete()
            db.session.commit()
        
        # Terakhir, hapus semua soal yang tidak dalam koleksi
        print("Menghapus soal yang tidak dalam koleksi...")
        if preserved_question_ids:
            # Hapus soal yang tidak ada dalam koleksi
            Question.query.filter(~Question.id.in_(preserved_question_ids)).delete(synchronize_session=False)
        else:
            # Jika tidak ada soal yang dipertahankan, hapus semua soal
            Question.query.delete()
        db.session.commit()
        
        message = "Database berhasil direset. Silahkan upload file kembali untuk regenerasi soal."
        if preserved_question_ids:
            message += f" {len(preserved_question_ids)} soal dalam koleksi dipertahankan."
        
        print(message)
        return jsonify({"success": True, "message": message}), 200
        # BAGIAN YANG DIMODIFIKASI - SELESAI
    except Exception as e:
        print(f"Error menghapus data lama: {str(e)}")
        db.session.rollback()
        return jsonify({"success": False, "message": f"Error menghapus data lama: {str(e)}"}), 500

# Add this to the Flask app code
@app.route('/collection-analytics/<int:collection_id>')
@login_required
def collection_analytics(collection_id):
    # Verify collection existence and ownership
    collection = QuestionCollection.query.get_or_404(collection_id)
    
    # Only allow the teacher who owns the collection to view analytics
    if collection.guru_id != current_user.id:
        flash('Anda tidak memiliki akses ke analisis koleksi ini', 'error')
        return redirect(url_for('guru'))
    
    return render_template('collection-analytics.html', collection_id=collection_id)

# Add these endpoints to your Flask app

# Get detailed student answer data for a collection
@app.route('/api/siswa_answers', methods=['GET'])
@login_required
def get_siswa_answers():
    try:
        user_id = request.args.get('user_id')
        collection_id = request.args.get('collection_id')
        
        if not user_id or not collection_id:
            return jsonify({'success': False, 'message': 'Missing required parameters'}), 400
        
        # Get question IDs in this collection
        collection_question_ids = db.session.query(CollectionQuestion.question_id)\
            .filter(CollectionQuestion.collection_id == collection_id).all()
        collection_question_ids = [id[0] for id in collection_question_ids]
        
        # Get answers for these questions
        siswa_answers = db.session.query(SiswaAnswer).join(
            Question, SiswaAnswer.question_id == Question.id
        ).filter(
            SiswaAnswer.siswa_id == user_id,
            Question.id.in_(collection_question_ids)
        ).all()
        
        # Format answers for response
        formatted_answers = []
        for answer in siswa_answers:
            formatted_answers.append({
                'id': answer.id,
                'siswa_id': answer.siswa_id,
                'question_id': answer.question_id,
                'siswa_answer': answer.siswa_answer,
                'is_correct': answer.is_correct,
                'level': answer.level,
                'answered_at': answer.answered_at.isoformat() if answer.answered_at else None
            })
        
        return jsonify({
            'success': True,
            'siswa_answers': formatted_answers
        })
        
    except Exception as e:
        print(f"Error getting student answers: {str(e)}")
        return jsonify({
            'success': False,
            'message': f"Error: {str(e)}"
        }), 500

# Get collection by ID
@app.route('/api/collections/<int:collection_id>', methods=['GET'])
@login_required
def get_collection(collection_id):
    try:
        collection = QuestionCollection.query.get(collection_id)
        if not collection:
            return jsonify({'success': False, 'message': 'Collection not found'}), 404
        
        # Check if user has access to this collection
        if collection.guru_id != current_user.id and current_user.user_type != 'guru':
            return jsonify({'success': False, 'message': 'Access denied'}), 403
        
        return jsonify({
            'success': True,
            'collection': {
                'id': collection.id,
                'name': collection.name,
                'description': collection.description,
                'created_at': collection.created_at.isoformat() if collection.created_at else None
            }
        })
        
    except Exception as e:
        print(f"Error getting collection: {str(e)}")
        return jsonify({
            'success': False,
            'message': f"Error: {str(e)}"
        }), 500
# =========================================
# ENDPOINT UPLOAD & GENERATE 5 SOAL PER level (35 TOTAL)
# =========================================
@app.route('/upload', methods=['POST'])
@login_required
def upload_file():
    # Debug information
    print(f"Current user: {current_user}, authenticated: {current_user.is_authenticated}")
    print(f"User attributes: {dir(current_user)}")
    
    # Check if user is authenticated
    if not current_user.is_authenticated:
        print("User not authenticated, redirecting to login")
        return jsonify({"message": "Anda harus login terlebih dahulu"}), 401
    
    # Dari kode model Guru yang Anda berikan, seharusnya menggunakan user_type atau class name
    if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
        print(f"User {current_user.username} is not a teacher")
        return jsonify({"message": "Hanya guru yang dapat menggenerate soal"}), 403
        
    # Tambahan logging untuk debug
    print("Upload file endpoint called")
    
    if 'file' not in request.files:
        return jsonify({"message": "Tidak ada file yang diunggah"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"message": "Tidak ada file yang dipilih"}), 400

    file_ext = file.filename.rsplit('.', 1)[-1].lower()
    if file_ext not in ['doc', 'docx', 'pdf']:
        return jsonify({"message": "Format file tidak didukung. Hanya file .doc, .docx, dan .pdf yang diizinkan"}), 400

    # Create upload folder if it doesn't exist
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(file_path)
    print(f"File disimpan di {file_path}")

    try:
        # Extract text from document based on file type
        if file_ext == 'pdf':
            content_text = extract_text_from_pdf(file_path)
        elif file_ext in ['doc', 'docx']:
            content_text = extract_text_from_docx(file_path)
        else:
            return jsonify({"message": "Format file tidak didukung"}), 400
            
        print(f"Berhasil mengekstrak teks dari file {file.filename}")
        
        # Create dataframe from extracted text - Use the improved function
        df = convert_text_to_dataframe_improved(content_text)
        
        if df.empty:
            # Fallback to simple content dataframe if we couldn't extract tabular data
            df = pd.DataFrame({
                'content': [content_text],
                'length': [len(content_text)]
            })
            print("Menggunakan dataframe sederhana karena tidak dapat mengekstrak data tabular")
            
        print(f"Berhasil membuat dataframe dengan {len(df)} baris")
        
        # Extract key educational components from the document content
        # Modified to handle only 3 return values instead of 5
        educational_components = extract_educational_components(content_text)
        if len(educational_components) == 3:
            module_elements, initial_competencies, learning_objectives = educational_components
            pemahaman_bermakna = "Tidak tersedia"
            target_peserta_didik = "Tidak tersedia"
        else:
            module_elements, initial_competencies, learning_objectives, pemahaman_bermakna, target_peserta_didik = educational_components
        
        print("Berhasil mengekstrak komponen pembelajaran:")
        print(f"- Modul/Elemen Ajar: {module_elements[:100]}...")
        print(f"- Kompetensi Awal: {initial_competencies[:300]}...")
        print(f"- Tujuan Pembelajaran: {learning_objectives[:300]}...")
        print(f"- Pemahaman Bermakna: {pemahaman_bermakna[:100]}...")
        print(f"- Target Peserta Didik: {target_peserta_didik[:100]}...")
        
    except ValueError as e:
        if "not enough values to unpack" in str(e):
            print("Warning: extract_educational_components() returned fewer values than expected")
            # Handle case where only 3 values are returned
            educational_components = extract_educational_components(content_text)
            module_elements, initial_competencies, learning_objectives = educational_components
            pemahaman_bermakna = "Tidak tersedia"
            target_peserta_didik = "Tidak tersedia"
        else:
            print(f"Error membaca file: {str(e)}")
            return jsonify({"message": f"Error membaca file: {str(e)}"}), 500
    except Exception as e:
        print(f"Error membaca file: {str(e)}")
        return jsonify({"message": f"Error membaca file: {str(e)}"}), 500

    # Rest of your code remains the same...
    explanation_text = """
    **Indeks Kesulitan Item untuk Diagnosis Kognitif**
    Indeks kesulitan item sangat penting dalam menilai kemampuan kognitif individu di berbagai domain. Hal ini penting untuk mendapatkan pemahaman mendalam tentang kekuatan dan kelemahan kognitif individu.
    Kesulitan item dalam diagnosis kognitif menggambarkan seberapa menantang item tes tertentu terhadap konstruk kognitif yang dievaluasi. Tingkat kesulitan ini diukur dengan indeks kesulitan item, yang menghitung persentase peserta tes yang menjawab item dengan benar.
    Indeks kesulitan item yang tinggi menunjukkan bahwa item tersebut relatif mudah, karena sebagian besar peserta ujian menjawabnya dengan benar. Sebaliknya, indeks yang rendah menunjukkan bahwa item tersebut lebih sulit, karena lebih sedikit peserta yang memberikan jawaban yang benar.
    Menganalisis skor indeks kesulitan item di beberapa item tes dapat digunakan untuk menyesuaikan strategi intervensi, merancang program pelatihan yang ditargetkan, dan memberikan dukungan pembelajaran yang dipersonalisasi kepada individu berdasarkan profil kognitif spesifik mereka.
    Selain itu, hal ini juga memastikan bahwa item tes secara akurat mengukur konstruk kognitif yang dimaksud dan bahwa instrumen penilaian secara keseluruhan memberikan hasil yang valid dan andal.
        
    **Membuat Sesi Multi-Tahap Berdasarkan Taksonomi Kompetensi Teknologi**
    Dalam membuat soal untuk sesi multi-tahap, gunakan taksonomi kompetensi teknologi berikut sebagai panduan untuk menentukan tipe dan tingkat kesulitan soal pada setiap level:
    
    **Tabel 1.1: Taksonomi Kompetensi Teknologi** *(Setelah Todd, 1991, hlm. 271)*
    | Tingkat | Kompetensi | Jenis Pengetahuan | Hasil Pembelajaran |
    |---------|------------|-------------------|-------------------|
    | 1 | Kesadaran Teknologi | *knowledge that* | Pemahaman dasar tentang konsep dan istilah teknologi |
    | 2 | Literasi Teknologi | *knowledge that* | Komprehensi dan kemampuan menjelaskan konsep teknologi |
    | 3 | Kemampuan Teknologi | *knowledge that* dan *how* | Aplikasi pengetahuan dalam konteks praktis |
    | 4 | Kreativitas Teknologi | *knowledge that* dan *how* | Invensi dan pengembangan solusi teknologi |
    | 5 | Kritik Teknologi | *knowledge that*, *how*, dan *why* | Penilaian kritis dan evaluasi mendalam |
    
    Setiap tingkat pada taksonomi ini sesuai dengan level pada sesi multi-tahap dan menentukan karakteristik soal yang harus dibuat.
    """

    # Modified prompt to generate universal questions without specific ID references
    # Now includes specific educational component analysis and multiple-choice format
    prompt = f"""
    Berikut adalah data (katalog) siswa dalam tabel:

    {df.to_string(index=False)}
    
    Berdasarkan analisis dokumen yang diunggah, berikut adalah komponen pembelajaran yang telah diidentifikasi:
    
    MODUL/ELEMEN AJAR:
    {module_elements}
    
    KOMPETENSI AWAL:
    {initial_competencies}
    
    TUJUAN PEMBELAJARAN:
    {learning_objectives}

    PEMAHAMAN BERMAKNA:
    {pemahaman_bermakna}

    TARGET PESERTA DIDIK:
    {target_peserta_didik}
    
    {explanation_text}

    Silakan buat 5 soal PILIHAN GANDA untuk SETIAP level dalam multi-stage session berdasarkan materi pembelajaran di atas:
    level 1 -> level 2
    level 2 -> Jika benar menjawab maka di bawa ke level 4 atau jika salah menjawab akan dibawa ke level 3
    level 4 -> Jika benar menjawab maka di bawa ke level 7 atau jika salah menjawab akan dibawa ke level 6
    level 3 -> Jika benar menjawab maka di bawa ke level 6 atau jika salah menjawab akan dibawa ke level 5
    level 5, 6, 7 = final


    PENERAPAN TAKSONOMI KOMPETENSI TEKNOLOGI PADA SOAL:
    1. level 1 (STAGE I) - KESADARAN TEKNOLOGI (Pemahaman)
       * Buat soal yang menguji pemahaman dasar konsep dan istilah teknologi
       * Fokus pada "knowledge that" - pengetahuan deklaratif tentang fakta-fakta dasar
       * Contoh kata kerja operasional: mengenali, mengidentifikasi, menyebutkan
       * Tingkat kesulitan: Sangat mudah (p = 0.95)
    
    2. level 2 (STAGE II) - LITERASI TEKNOLOGI (Komprehensi)
       * Buat soal yang menguji kemampuan menjelaskan konsep teknologi
       * Tetap fokus pada "knowledge that" tetapi lebih mendalam
       * Contoh kata kerja operasional: menjelaskan, membandingkan, mengklasifikasi
       * Tingkat kesulitan: Mudah (p = 0.90)
    
    3-4. level 3-4 (STAGE III) - KEMAMPUAN TEKNOLOGI (Aplikasi)
       * level 3: Soal aplikasi tingkat dasar untuk peserta kemampuan rendah-menengah
       * level 4: Soal aplikasi tingkat lanjut untuk peserta kemampuan menengah-tinggi
       * Gabungkan "knowledge that" dan "knowledge how" - pengetahuan prosedural
       * Contoh kata kerja operasional: menerapkan, menggunakan, mendemonstrasikan
       * Tingkat kesulitan level 3: p = 0.75 (sedang)
       * Tingkat kesulitan level 4: p = 0.67 (sedang-sulit)
    
    5-6-7. level 5-6-7 (STAGE IV) - PENILAIAN AKHIR
       * level 5: Untuk peserta kemampuan rendah (dari level 3)
           - Fokus pada kreativitas teknologi dasar
           - Tingkat kesulitan: p = 0.20
       * level 6: Untuk peserta kemampuan menengah (dari level 3 atau 4)
           - Fokus pada kreativitas teknologi menengah
           - Tingkat kesulitan: p = 0.17
       * level 7: Untuk peserta kemampuan tinggi (dari level 4)
           - Fokus pada kritik teknologi dan evaluasi mendalam
           - Tingkat kesulitan: p = 0.15
    
    PENTING:
    1. Buatlah pertanyaan PILIHAN GANDA dengan 4 opsi jawaban (A, B, C, D)
    2. Pertanyaan harus berkaitan langsung dengan elemen ajar, kompetensi awal, dan tujuan pembelajaran yang telah diidentifikasi
    3. Soal pada level 1 sebaiknya berkaitan dengan kompetensi awal (pengetahuan dasar)
    4. Soal pada level terakhir (5, 6, 7) sebaiknya berkaitan dengan tujuan pembelajaran (tingkat pemahaman tinggi)
    5. Pertanyaan sebaiknya tidak menyebutkan ID atau nomor urut spesifik dari dataset
    6. Rumuskan pertanyaan yang berfokus pada pemahaman konsep umum, prinsip, dan penerapan
    7. Pastikan setiap soal memiliki 4 pilihan jawaban (bukan essay)

    Berikan respon HANYA dalam format JSON array berikut tanpa komentar atau teks tambahan. Pastikan format JSON valid:
    [
      {{
        "level": (sesuaikan level yang akan dibuat),
        "question_type": "multiple_choice",
        "soal": "Pertanyaan universal tentang konsep, bukan tentang ID tertentu",
        "options": ["Opsi A", "Opsi B", "Opsi C", "Opsi D"],
        "jawaban_benar": "Opsi A",
        "p": (sesuaikan tingkat kesulitanya sesuai level yang akan dibuat),
        "explanation": "penjelasan singkat"
      }},
      {{
        "level": (sesuaikan level yang akan dibuat),
        "question_type": "multiple_choice",
        "soal": "Pertanyaan yang tidak mengandung ID atau nomor urut spesifik",
        "options": ["Opsi A", "Opsi B", "Opsi C", "Opsi D"],
        "jawaban_benar": "Opsi C",
        "p": (sesuaikan tingkat kesulitanya sesuai level yang akan dibuat),
        "explanation": "penjelasan singkat mengapa C adalah jawaban yang benar"
      }},
      ...dan seterusnya
    ]
        
    Buat total 35 soal: 5 soal untuk setiap level (level=1..7).
    Setiap pertanyaan harus memiliki 4 pilihan jawaban dan 1 jawaban benar.
   
    Fokuskan soal berdasarkan tingkat kognitif:
    - level 1: Tingkat pengetahuan (knowledge) - mengingat, mengenali
    - level 2: Tingkat pemahaman (comprehension) - memahami, menjelaskan
    - level 3-4: Tingkat penerapan (application) - menerapkan, menggunakan
    - level 5-6: Tingkat analisis (analysis) - menganalisis, membandingkan
    - level 7: Tingkat evaluasi dan sintesis (evaluation & synthesis) - mengevaluasi, menciptakan

    Buatkan soal yang sesuai dengan taksonomi kompetensi teknologi yang telah diberikan.
    Pastikan untuk tidak menyertakan ID atau nomor urut spesifik dari dataset dalam soal.
    buatkan soalnya dalam bahasa yang mudah dipahami oleh siswa.
    """

    try:
        print("Mengirim prompt ke Gemini API")
        response = model.generate_content(prompt)
        raw_text = response.text.strip()
        
        # Debug: Print raw AI response for troubleshooting
        print("Raw AI response:", raw_text)
        
        # Improved JSON extraction approach
        try:
            # First attempt: Clean any markdown code block formatting
            cleaned_text = re.sub(r'```json|```', '', raw_text).strip()  # Fixed regex pattern
            cleaned_text = re.sub(r',\s*([}\]])', r'\1', cleaned_text)
            gen_questions = json.loads(cleaned_text)
            print("Berhasil parsing JSON setelah membersihkan markdown")
        except json.JSONDecodeError:
            print("Parsing JSON setelah membersihkan markdown gagal, mencoba metode lain")
            
            # Second attempt: Try to extract JSON array with regex
            match = re.search(r'\[\s*\{.*?\}\s*\]', raw_text, re.DOTALL)  # Fixed regex pattern
            if match:
                try:
                    json_text = match.group()
                    json_text = re.sub(r',\s*([}\]])', r'\1', json_text)
                    print(f"Ekstraksi JSON dengan regex: {json_text[:100]}...")
                    gen_questions = json.loads(json_text)
                    print("Berhasil parsing JSON dari hasil regex")
                except json.JSONDecodeError as e:
                    print(f"Error decode JSON setelah regex: {str(e)}")
                    return jsonify({"message": f"Error parsing JSON soal: {str(e)}"}), 500
            else:
                # Third attempt: Find the outermost square brackets
                print("Ekstraksi regex gagal, mencoba dengan bracket matching")
                open_bracket = raw_text.find('[')
                close_bracket = raw_text.rfind(']')
                
                if open_bracket >= 0 and close_bracket > open_bracket:
                    try:
                        json_text = raw_text[open_bracket:close_bracket+1]
                        print(f"Ekstraksi JSON dengan indeks kurung: {json_text[:100]}...")
                        json_text = re.sub(r',\s*([}\]])', r'\1', json_text)
                        gen_questions = json.loads(json_text)
                        print("Berhasil parsing JSON dari ekstraksi kurung")
                    except json.JSONDecodeError as e:
                        print(f"Error decode JSON setelah ekstraksi kurung: {str(e)}")
                        # Final attempt: Try with more extensive cleaning
                        cleaned_text = re.sub(r'```json|```|\n|\\n|\\\"', '', json_text)  # Fixed regex pattern
                        # Fix potential trailing commas which are invalid in JSON
                        cleaned_text = re.sub(r',\s*\}', '}', cleaned_text)
                        cleaned_text = re.sub(r',\s*\]', ']', cleaned_text)
                        
                        try:
                            gen_questions = json.loads(cleaned_text)
                            print("Berhasil parsing JSON setelah pembersihan intensif")
                        except json.JSONDecodeError as e2:
                            print(f"Error decode JSON akhir: {str(e2)}")
                            return jsonify({"message": f"Error parsing JSON soal: {str(e2)}"}), 500
                else:
                    print("Tidak dapat menemukan array JSON dalam respons")
                    return jsonify({"message": "Tidak dapat menemukan format JSON dalam respons AI"}), 500

        # Validasi struktur data yang diterima
        if not isinstance(gen_questions, list):
            print(f"Diharapkan list, tapi dapat {type(gen_questions)}")
            return jsonify({"message": "AI tidak mengembalikan array soal yang valid"}), 500
            
        print(f"Berhasil mengekstrak {len(gen_questions)} soal")
        
        # Validasi jumlah soal
        # Idealnya harus ada 35 soal (5 per level), tapi kita bisa menerima jumlah yang berbeda
        if len(gen_questions) < 7:  # Minimal harus ada 1 soal per level
            print(f"Jumlah soal tidak mencukupi: {len(gen_questions)}")
            return jsonify({"message": f"AI hanya menghasilkan {len(gen_questions)} soal, minimal diperlukan 7 soal"}), 500

        # Hapus soal lama dan data terkait
        print("Menghapus data terkait dari database")
        try:
            # BAGIAN YANG DIMODIFIKASI - MULAI
            # Dapatkan daftar ID soal yang ada dalam koleksi (yang harus dipertahankan)
            preserved_question_ids = db.session.query(CollectionQuestion.question_id).distinct().all()
            preserved_question_ids = [id[0] for id in preserved_question_ids]
            
            if preserved_question_ids:
                print(f"Mempertahankan {len(preserved_question_ids)} soal dalam koleksi")
            
            # Pertama, hapus semua jawaban pengguna yang tidak terkait dengan soal dalam koleksi
            print("Menghapus jawaban pengguna...")
            if preserved_question_ids:
                # Hapus jawaban siswa hanya untuk soal yang tidak dalam koleksi
                SiswaAnswer.query.filter(~SiswaAnswer.question_id.in_(preserved_question_ids)).delete(synchronize_session=False)
            else:
                # Jika tidak ada soal yang dipertahankan, hapus semua jawaban
                SiswaAnswer.query.delete()
            db.session.commit()
            
            # Kedua, reset current_level semua user ke level 1
            print("Mengatur ulang level pengguna...")
            SiswaResult.query.update({SiswaResult.current_level: 1})
            db.session.commit()
            
            # Terakhir, hapus semua soal yang tidak dalam koleksi
            print("Menghapus soal yang tidak dalam koleksi...")
            if preserved_question_ids:
                # Hapus soal yang tidak ada dalam koleksi
                Question.query.filter(~Question.id.in_(preserved_question_ids)).delete(synchronize_session=False)
            else:
                # Jika tidak ada soal yang dipertahankan, hapus semua soal
                Question.query.delete()
            db.session.commit()
            
            if preserved_question_ids:
                print(f"Data lama berhasil dihapus, {len(preserved_question_ids)} soal dalam koleksi dipertahankan")
            else:
                print("Data lama berhasil dihapus")
            # BAGIAN YANG DIMODIFIKASI - SELESAI
            
        except Exception as e:
            print(f"Error menghapus data lama: {str(e)}")
            db.session.rollback()
            return jsonify({"message": f"Error menghapus data lama: {str(e)}"}), 500
        
        # Simpan soal baru ke database
        print("Menyimpan soal baru ke database")
        for q in gen_questions:
            try:
                level_num = int(q.get("level"))
                if 1 <= level_num <= 7:
                    p_val = float(q.get("p", 0.75))
                    bobot = assign_weight(p_val)
                    
                    options_str = None
                    if 'options' in q and isinstance(q['options'], list):
                        options_str = json.dumps(q['options'])
                    
                    question_type = q.get('question_type', 'multiple_choice')
                    
                    new_question = Question(
                        guru_id=current_user.id,  # This is crucial
                        level=level_num,
                        soal=q.get("soal"),
                        jawaban_benar=q.get("jawaban_benar", ""),
                        options=options_str,
                        question_type=question_type,
                        p=p_val,
                        bobot=bobot,
                        explanation=q.get("explanation", "")
                    )
                    db.session.add(new_question)
            except Exception as e:
                print(f"Error memproses soal: {str(e)}")
                db.session.rollback()
                return jsonify({"message": f"Error memproses soal: {str(e)}"}), 500
        
        # Commit semua perubahan ke database
        db.session.commit()
        print("Database commit berhasil")

        # Hitung jumlah soal per level untuk pelaporan
        level_counts = {}
        for q in gen_questions:
            level = int(q.get("level"))
            if level in level_counts:
                level_counts[level] += 1
            else:
                level_counts[level] = 1
                
        level_summary = ", ".join([f"level {level}: {count} soal" for level, count in sorted(level_counts.items())])

        # BAGIAN YANG DIMODIFIKASI - TAMBAHAN MESSAGE
        message = f"Total {len(gen_questions)} soal pilihan ganda berhasil digenerate dan disimpan ({level_summary})"
        if preserved_question_ids:
            message += f", dengan mempertahankan {len(preserved_question_ids)} soal yang ada dalam koleksi"

        return jsonify({
            "message": message, 
            "data": gen_questions
        }), 200

    except Exception as e:
        print(f"Exception in upload_file: {str(e)}")
        db.session.rollback()
        return jsonify({"message": f"AI error: {str(e)}"}), 500

# New function to extract educational components from document content
def extract_educational_components(content_text, return_dict=False):
    """
    Wrapper around extract_educational_components to maintain backward compatibility
    with different calling conventions.
    
    Args:
        content_text: The text content to extract educational components from
        return_dict: If True, returns a dictionary; if False, returns a tuple of components
        
    Returns:
        Either a 5-tuple of (module_elements, initial_competencies, learning_objectives, 
        pemahaman_bermakna, target_peserta_didik) or a dictionary of components
    """
    # Initialize results dictionary with "Tidak tersedia" as default values
    results = {
        "Modul/Elemen Ajar": "Tidak tersedia",
        "Kompetensi Awal": "Tidak tersedia",
        "Tujuan Pembelajaran": "Tidak tersedia",
        "Pemahaman Bermakna": "Tidak tersedia",
        "Target Peserta Didik": "Tidak tersedia"
    }
    
    # Normalize text for easier processing
    content = content_text.lower()
    original_content = content_text  # Keep original case for extraction
    
    # ===== Module/Element Section =====
    module_patterns = [
        r'(?:modul|elemen|materi)(?:\s+ajar|\s+pembelajaran|\s+pokok)?(.*?)(?:kompetensi|tujuan|capaian|daftar pustaka|referensi)',
        r'(?:bab|materi|konten|isi)(?:\s+pembelajaran|\s+utama)?(.*?)(?:kompetensi|tujuan|capaian|daftar pustaka|referensi)',
        r'(?:pokok|inti)(?:\s+bahasan|\s+materi|\s+pembelajaran)?(.*?)(?:kompetensi|tujuan|capaian|daftar pustaka|referensi)',
        r'(?:i\.|1\.)(?:\s+identitas\s+modul)(.*?)(?:ii\.|2\.)'  # Captures sections starting with Roman numerals
    ]
    
    for pattern in module_patterns:
        match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
        if match:
            module_text = match.group(1).strip()
            if module_text:
                results["Modul/Elemen Ajar"] = clean_extracted_text(module_text)
                break
    
    # ===== Initial Competencies Section =====
    competency_patterns = [
        r'(?:kompetensi|kemampuan)(?:\s+awal|\s+dasar|\s+prasyarat)(?:\s+[\w\s]+)?[:\r\n](.*?)(?:tujuan|capaian|indikator|materi|bab|daftar pustaka|referensi)',
        r'(?:prerequisite|prasyarat)(?:\s+[\w\s]+)?[:\r\n](.*?)(?:tujuan|capaian|indikator|materi|bab|daftar pustaka|referensi)',
        r'(?:pengetahuan|keterampilan)(?:\s+awal|\s+dasar)(?:\s+[\w\s]+)?[:\r\n](.*?)(?:tujuan|capaian|indikator|materi|bab|daftar pustaka|referensi)',
        r'(?:ii\.|2\.)(?:\s+kompetensi\s+awal)(.*?)(?:iii\.|3\.)'  # Captures sections with Roman numerals
    ]
    
    for pattern in competency_patterns:
        match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
        if match:
            comp_text = match.group(1).strip()
            if comp_text:
                results["Kompetensi Awal"] = clean_extracted_text(comp_text)
                break
    
    # ===== Learning Objectives Section =====
    objective_patterns = [
        r'(?:tujuan|capaian)(?:\s+pembelajaran|\s+belajar|\s+pendidikan|\s+instruksional|\s+pelatihan)(?:\s+[\w\s]+)?[:\r\n](.*?)(?:kompetensi|materi|bab|daftar pustaka|referensi|metodologi|pemahaman|target)',
        r'(?:learning|instructional)(?:\s+objectives|\s+goals|\s+outcomes)(?:\s+[\w\s]+)?[:\r\n](.*?)(?:kompetensi|materi|bab|daftar pustaka|referensi|metodologi|pemahaman|target)',
        r'(?:indikator|pencapaian)(?:\s+keberhasilan|\s+pembelajaran|\s+belajar)(?:\s+[\w\s]+)?[:\r\n](.*?)(?:kompetensi|materi|bab|daftar pustaka|referensi|metodologi|pemahaman|target)',
        r'(?:iii\.|3\.)(?:\s+tujuan\s+pembelajaran)(.*?)(?:iv\.|4\.)'  # Captures sections with Roman numerals
    ]
    
    for pattern in objective_patterns:
        match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
        if match:
            obj_text = match.group(1).strip()
            if obj_text:
                results["Tujuan Pembelajaran"] = clean_extracted_text(obj_text)
                break
    
    # ===== NEW: Meaningful Understanding Section =====
    understanding_patterns = [
        r'(?:pemahaman\s+bermakna|bermakna\s+pemahaman|pemahaman\s+yang\s+bermakna)[:\r\n](.*?)(?:kompetensi|tujuan|capaian|materi|bab|daftar pustaka|referensi|target)',
        r'(?:meaningful\s+understanding|understanding\s+meaningful)[:\r\n](.*?)(?:kompetensi|tujuan|capaian|materi|bab|daftar pustaka|referensi|target)',
        r'(?:big\s+ideas|gagasan\s+utama|ide\s+pokok)[:\r\n](.*?)(?:kompetensi|tujuan|capaian|materi|bab|daftar pustaka|referensi|target)',
        r'(?:iv\.|4\.)(?:\s+pemahaman\s+bermakna)(.*?)(?:v\.|5\.)'  # Captures sections with Roman numerals
    ]
    
    for pattern in understanding_patterns:
        match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
        if match:
            understanding_text = match.group(1).strip()
            if understanding_text:
                results["Pemahaman Bermakna"] = clean_extracted_text(understanding_text)
                break
    
    # ===== NEW: Target Students Section =====
    target_patterns = [
        r'(?:target\s+peserta\s+didik|sasaran\s+peserta|peserta\s+didik\s+target)[:\r\n](.*?)(?:kompetensi|tujuan|capaian|materi|bab|daftar pustaka|referensi)',
        r'(?:target\s+audience|target\s+learners|intended\s+learners)[:\r\n](.*?)(?:kompetensi|tujuan|capaian|materi|bab|daftar pustaka|referensi)',
        r'(?:siswa\s+yang\s+dituju|karakteristik\s+peserta\s+didik)[:\r\n](.*?)(?:kompetensi|tujuan|capaian|materi|bab|daftar pustaka|referensi)',
        r'(?:v\.|5\.)(?:\s+target\s+peserta\s+didik)(.*?)(?:vi\.|6\.)'  # Captures sections with Roman numerals
    ]
    
    for pattern in target_patterns:
        match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
        if match:
            target_text = match.group(1).strip()
            if target_text:
                results["Target Peserta Didik"] = clean_extracted_text(target_text)
                break
    
    # ===== Keyword-based approach for sections that weren't found =====
    lines = original_content.split('\n')
    
    # Dictionary to map section keys to their search keywords
    section_keywords = {
        "Modul/Elemen Ajar": ['modul', 'elemen', 'materi', 'bab', 'pokok bahasan', 'identitas modul'],
        "Kompetensi Awal": ['kompetensi awal', 'prasyarat', 'kemampuan dasar', 'kompetensi dasar'],
        "Tujuan Pembelajaran": ['tujuan pembelajaran', 'capaian pembelajaran', 'learning objectives'],
        "Pemahaman Bermakna": ['pemahaman bermakna', 'bermakna', 'big ideas', 'gagasan utama', 'ide pokok'],
        "Target Peserta Didik": ['target peserta', 'sasaran peserta', 'karakteristik peserta', 'target siswa', 'audience']
    }
    
    # For sections still not found, try the keyword approach
    for section, keywords in section_keywords.items():
        if results[section] == "Tidak tersedia":
            paragraph_start = -1
            
            # First look for exact section headers
            for i, line in enumerate(lines):
                line_lower = line.lower().strip()
                # Check if line matches any of our section patterns
                if any(keyword in line_lower for keyword in keywords):
                    paragraph_start = i
                    break
            
            if paragraph_start >= 0:
                # Determine where the section might end - either at next section or after several lines
                paragraph_end = paragraph_start + 1
                
                # Look for the next potential section header
                for i in range(paragraph_start + 1, min(paragraph_start + 15, len(lines))):
                    line_lower = lines[i].lower().strip()
                    
                    # Check if this line might be the start of the next section
                    if (re.match(r'^[ivxlcdm]+\.|^\d+\.', line_lower) or  # Roman numerals or numbers followed by period
                        any(keyword in line_lower for section_kw in section_keywords.values() 
                            for keyword in section_kw if section_kw != keywords)):
                        paragraph_end = i
                        break
                
                # Extract and clean the section text
                section_text = '\n'.join(lines[paragraph_start:paragraph_end]).strip()
                if section_text:
                    results[section] = clean_extracted_text(section_text)
    
    # Second pass to try to find sections by scanning the whole document
    for section, value in results.items():
        if value == "Tidak tersedia":
            # Look for sentences that might contain relevant information
            for i, line in enumerate(lines):
                line_lower = line.lower()
                
                # Get keywords for this section
                keywords = section_keywords.get(section, [])
                
                # Check each keyword
                for keyword in keywords:
                    if keyword in line_lower:
                        # Extract several lines from this point
                        extract = '\n'.join(lines[i:min(i+8, len(lines))])
                        results[section] = clean_extracted_text(extract)
                        break
                
                if results[section] != "Tidak tersedia":
                    break
    
    # Return the appropriate format based on the return_dict parameter
    if return_dict:
        return results
    else:
        return (
            results["Modul/Elemen Ajar"],
            results["Kompetensi Awal"],
            results["Tujuan Pembelajaran"],
            results["Pemahaman Bermakna"],
            results["Target Peserta Didik"]
        )

def clean_extracted_text(text):
    """Clean up extracted text by removing extra whitespace, bullet points, etc."""
    # Remove excess whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove bullet points, numbering, etc.
    text = re.sub(r'^\s*[\d\.\-\•○●]\s', '', text, flags=re.MULTILINE)
    # Add line breaks for readability in the prompt
    text = re.sub(r'([.!?])\s+', r'\1\n', text)
    return text.strip()

# Function to extract text from a PDF file
def extract_text_from_pdf(file_path):
    """Extract text from a PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")
        raise

def extract_text_from_docx(file_path):
    """Extract text from a Word document (.doc or .docx)"""
    try:
        if file_path.endswith('.docx'):
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text += "\t".join(row_text) + "\n"
            return text
        else:  # For .doc files
            with open(file_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                return result.value
    except Exception as e:
        print(f"Error extracting text from Word document: {str(e)}")
        raise

def convert_text_to_dataframe_improved(text):
    """Improved function to convert extracted text to a dataframe"""
    try:
        # Split text into lines for processing
        lines = text.split('\n')
        
        # Filter out empty lines
        lines = [line.strip() for line in lines if line.strip()]
        
        # Look for table-like patterns in the text
        table_lines = []
        
        # Check each line for tab separators or multiple spaces indicating tabular data
        for line in lines:
            if '\t' in line or re.search(r'\s{2,}', line):
                table_lines.append(line)
        
        if table_lines:
            # Create a temporary file to process the data
            with tempfile.NamedTemporaryFile(delete=False, mode='w', suffix='.tsv', encoding='utf-8') as temp_file:
                for line in table_lines:
                    # Normalize spaces and tabs
                    normalized_line = re.sub(r'\s{2,}', '\t', line)
                    temp_file.write(normalized_line + '\n')
            
            try:
                # Try reading the file as a tab-separated values file
                df = pd.read_csv(temp_file.name, sep='\t', on_bad_lines='skip', encoding='utf-8')  # Updated from error_bad_lines
                os.unlink(temp_file.name)
                return df
            except Exception as e:
                # If that fails, try a different approach with more robust error handling
                print(f"First attempt at parsing failed: {str(e)}")
                try:
                    os.unlink(temp_file.name)
                except:
                    pass
                
                # Try an alternative approach using pandas read_fwf (fixed width format)
                with tempfile.NamedTemporaryFile(delete=False, mode='w', suffix='.txt', encoding='utf-8') as temp_file:
                    for line in table_lines:
                        temp_file.write(line + '\n')
                
                try:
                    # Try to infer the fixed width format
                    df = pd.read_fwf(temp_file.name, encoding='utf-8')
                    os.unlink(temp_file.name)
                    return df
                except Exception as e2:
                    print(f"Fixed width parsing failed: {str(e2)}")
                    os.unlink(temp_file.name)
                    
                    # Final fallback: create a simple dataframe with lines as rows
                    data = {"line": table_lines}
                    return pd.DataFrame(data)
        
        # If no tabular data found, return a dataframe with the content
        return pd.DataFrame({"content": [text]})
            
    except Exception as e:
        print(f"Error converting text to dataframe: {str(e)}")
        # Return an empty dataframe as fallback
        return pd.DataFrame()
    
# =========================================
# GET QUESTION (per level)
# =========================================
@app.route('/get_question', methods=['GET'])
def get_question():
    user_id = request.args.get("user_id", "default")
    collection_id = request.args.get("collection_id")  # Ambil collection_id dari parameter URL
    print(f"Get question untuk user_id: {user_id}")
    
    try:
        # Validasi input
        if not collection_id:
            return jsonify({
                "status": "error",
                "message": "Collection ID tidak ditemukan dalam URL!"
            }), 400
            
        # Konversi user_id menjadi integer jika bukan "default"
        siswa_id = int(user_id) if user_id != "default" else None
        
        if siswa_id is None:
            return jsonify({
                "status": "error",
                "message": "Invalid user ID. Please login properly."
            }), 400
            
        # Verifikasi akses siswa ke collection
        has_access = db.session.query(collection_students).filter_by(
            collection_id=collection_id,  
            siswa_id=siswa_id
        ).first() is not None
        
        if not has_access:
            return jsonify({
                "status": "error",
                "message": "Anda tidak memiliki akses ke koleksi soal ini."
            }), 403
            
        # Cari hasil siswa berdasarkan siswa_id DAN collection_id
        user_result = SiswaResult.query.filter_by(
            siswa_id=siswa_id,
            collection_id=collection_id
        ).first()
        
        # Jika tidak ditemukan, buat baru
        if not user_result:
            print(f"Membuat user result baru untuk siswa_id: {siswa_id} dan collection_id: {collection_id}")
            
            # Cek apakah siswa ada
            siswa = db.session.query(Siswa).filter_by(id=siswa_id).first()
            if not siswa:
                return jsonify({
                    "status": "error",
                    "message": "Siswa tidak ditemukan"
                }), 404
                
            # Buat record baru
            user_result = SiswaResult(
                siswa_id=siswa_id,
                collection_id=int(collection_id),
                current_level=1
            )
            db.session.add(user_result)
            db.session.commit()

        current_level = user_result.current_level
        print(f"Level saat ini: {current_level}")
        
        # PERBAIKAN: Ambil soal dari database berdasarkan level, collection_id, DAN is_validated
        # Gunakan CollectionQuestion untuk memastikan soal adalah bagian dari collection
        questions = db.session.query(Question).join(
            CollectionQuestion,  
            CollectionQuestion.question_id == Question.id
        ).filter(
            Question.level == current_level,
            CollectionQuestion.collection_id == int(collection_id),
            Question.is_validated == True  # <--- INI TAMBAHANNYA
        ).all()
        
        print(f"Ditemukan {len(questions)} soal yang divalidasi untuk level {current_level} dalam koleksi {collection_id}")
        
        if not questions:
            # Jika tidak ada soal yang ditemukan
            return jsonify({
                "status": "error",  
                "message": f"Tidak ada soal yang tervalidasi tersedia untuk level {current_level} dalam koleksi ini."
            }), 200

        # Pilih soal secara acak
        question = random.choice(questions)
        print(f"Soal dipilih ID: {question.id}")
        
        # Parse opsi jika soal memiliki opsi dalam format JSON
        options = []
        if hasattr(question, 'options') and question.options:
            try:
                options = json.loads(question.options)
            except:
                options = []
        
        response_data = {
            "id": question.id,
            "level": question.level,
            "question": question.soal,
            "options": options,
            "question_type": question.question_type or "multiple_choice",
            "bobot": question.bobot,
            "p": question.p,
            "explanation": question.explanation,
            "jawaban_benar": question.jawaban_benar if question.question_type == 'multiple_choice' else None
        }
        
        # Mengirimkan soal ke frontend
        return jsonify(response_data), 200
        
    except ValueError:
        return jsonify({
            "status": "error",
            "message": "Format ID tidak valid"
        }), 400
        
    except Exception as e:
        print(f"Error di get_question: {str(e)}")
        return jsonify({
            "status": "error",
            "message": f"Kesalahan server: {str(e)}"
        }), 500
# =========================================
# SUBMIT ANSWER
# =========================================
@app.route('/submit_answer', methods=['POST'])
def submit_answer():
    data = request.json
    user_id = data.get("user_id", "default")
    answer = data.get("answer", "").strip()
    question_id = data.get("question_id")
    collection_id = data.get("collection_id")  # Ambil collection_id dari request
    
    print(f"Submit jawaban untuk user_id: {user_id}, jawaban: {answer[:50]}..., question_id: {question_id}, collection_id: {collection_id}")

    # Cari atau buat user result
    # PERBAIKAN: Filter berdasarkan collection_id juga
    user_result = SiswaResult.query.filter_by(siswa_id=user_id, collection_id=collection_id).first()
    if not user_result:
        print(f"Membuat user result baru untuk {user_id} pada collection {collection_id}")
        user_result = SiswaResult(
            siswa_id=user_id,
            collection_id=collection_id,  # Tambahkan collection_id
            correct=0,
            incorrect=0,
            current_level=1
        )
        db.session.add(user_result)
        db.session.commit()

    current_level = user_result.current_level
    print(f"Current level: {current_level}")
    
    # Dapatkan pertanyaan berdasarkan question_id
    if question_id:
        try:
            question = Question.query.get(int(question_id))
            if not question:
                return jsonify({"status": "error", "message": f"Soal dengan ID {question_id} tidak ditemukan."}), 404
        except Exception as e:
            print(f"Error mendapatkan soal dengan ID {question_id}: {str(e)}")
            return jsonify({"status": "error", "message": f"Error: {str(e)}"}), 500
    else:
        # Ambil soal yang terkait dengan koleksi
        questions = db.session.query(Question).join(
            CollectionQuestion, 
            CollectionQuestion.question_id == Question.id
        ).filter(
            Question.level == current_level,
            CollectionQuestion.collection_id == collection_id
        ).all()
        
        if not questions:
            return jsonify({"status": "game over", "message": f"Tidak ada soal untuk level {current_level} dalam koleksi ini."}), 200
        question = random.choice(questions)
    
    soal_text = question.soal
    print(f"Evaluasi jawaban untuk soal: {soal_text}")
    
    # Parse opsi jawaban
    options = []
    if hasattr(question, 'options') and question.options:
        try:
            options = json.loads(question.options)
        except:
            options = []
    
    # Evaluasi jawaban pilihan ganda dengan membandingkan dengan jawaban benar
    correct = is_answer_correct(answer, question.jawaban_benar)
    print(f"Hasil evaluasi pilihan ganda: {'Benar' if correct else 'Salah'}")

    # Simpan jawaban user ke database
    user_answer = SiswaAnswer(
        siswa_id=user_id,
        question_id=question.id,
        siswa_answer=answer,
        is_correct=correct,
        level=current_level
    )
    db.session.add(user_answer)

    # Update statistik
    if correct:
        user_result.correct += 1
        print(f"Jawaban benar. Total benar: {user_result.correct}")
        
        # Tentukan level selanjutnya jika jawaban benar
        if current_level in FINAL_levelS:
            # Jika sudah di level final, selesai
            status_message = f"Anda telah menyelesaikan semua level! level terakhir: {current_level}."
            db.session.commit()
            return jsonify({
                "status": "win", 
                "message": status_message,
                "explanation": question.explanation
            }), 200
        else:
            # Jika belum di level final, tentukan level selanjutnya untuk jawaban benar
            nxt = next_level(current_level, True)  # Passing is_correct=True
            print(f"level berikutnya: {nxt}")
            
            if nxt is None:
                db.session.commit()
                return jsonify({
                    "status": "win", 
                    "message": f"Anda telah menyelesaikan semua level!",
                    "explanation": question.explanation
                }), 200

            user_result.current_level = nxt
            db.session.commit()
            print(f"Database diperbarui: user pindah ke level {nxt}")

            return jsonify({
                "status": "continue", 
                "message": f"Jawaban benar! Lanjut ke level {nxt}",
                "explanation": question.explanation
            }), 200
    else:
        # Jika jawaban salah
        user_result.incorrect += 1
        print(f"Jawaban salah. Total salah: {user_result.incorrect}")
        
        # Tentukan level selanjutnya berdasarkan jawaban salah
        if current_level in FINAL_levelS:
            # Jika sudah di level final, selesai
            db.session.commit()
            return jsonify({
                "status": "game over", 
                "message": f"Game over! Anda telah menyelesaikan semua level dengan jawaban salah pada level {current_level}.",
                "explanation": question.explanation
            }), 200
        else:
            # Tentukan level selanjutnya berdasarkan jawaban salah
            next_level_if_wrong = next_level(current_level, False)  # Passing is_correct=False
            
            if next_level_if_wrong:
                user_result.current_level = next_level_if_wrong
                db.session.commit()
                print(f"Database diperbarui: karena salah, user pindah ke level {next_level_if_wrong}")
                
                if next_level_if_wrong in FINAL_levelS:
                    return jsonify({
                        "status": "continue", 
                        "message": f"Jawaban salah. Lanjut ke level final {next_level_if_wrong}.",
                        "explanation": question.explanation
                    }), 200
                else:
                    return jsonify({
                        "status": "continue", 
                        "message": f"Jawaban salah. Lanjut ke level {next_level_if_wrong}.",
                        "explanation": question.explanation
                    }), 200
            else:
                db.session.commit()
                return jsonify({
                    "status": "wrong", 
                    "message": "Jawaban salah. Silakan coba lagi.",
                    "explanation": question.explanation
                }), 200

# Function for evaluating answers
def evaluate_answer_with_ai(question_text, user_answer):
    """
    Mengevaluasi jawaban pengguna menggunakan AI tanpa bergantung pada kunci jawaban
    
    Args:
        question_text (str): Teks pertanyaan
        user_answer (str): Jawaban dari pengguna
        
    Returns:
        bool: True jika jawaban benar, False jika salah
    """
    try:
        # Use the model directly without having to re-import
        # Jika jawaban sangat pendek, kemungkinan salah
        if len(user_answer.strip()) < 5:
            return False
            
        # Prompt yang meminta AI untuk mengevaluasi jawaban
        prompt = f"""
        Kamu adalah asisten evaluasi jawaban yang sangat cerdas. Tugas kamu adalah menentukan 
        apakah jawaban pengguna benar untuk pertanyaan yang diberikan.
        
        Pertanyaan: {question_text}
        
        Jawaban pengguna: {user_answer}
        
        Evaluasi apakah jawaban pengguna ini benar secara konseptual dan faktual, bahkan jika berbeda dari 
        apa yang mungkin diharapkan sebagai jawaban standar. Pertimbangkan pemahaman mendalam tentang topik,
        bukan hanya kecocokan kata-kata.
        
        Harap beri penilaian akhir hanya dengan kata "BENAR" atau "SALAH", tanpa penjelasan tambahan.
        """
        
        # Dapatkan respons dari Gemini
        response = model.generate_content(prompt)
        response_text = response.text.strip().upper()
        
        # Tentukan hasil evaluasi
        if "BENAR" in response_text:
            return True
        elif "SALAH" in response_text:
            return False
        else:
            # Jika respons tidak jelas, coba lagi dengan prompt yang lebih sederhana
            simplified_prompt = f"""
            Pertanyaan: {question_text}
            Jawaban: {user_answer}
            Apakah jawaban ini benar? Jawab hanya dengan BENAR atau SALAH.
            """
            response = model.generate_content(simplified_prompt)
            response_text = response.text.strip().upper()
            
            return "BENAR" in response_text
            
    except Exception as e:
        print(f"Error dalam evaluasi dengan AI: {str(e)}")
        # Jika terjadi error, gunakan metode alternatif
        return evaluate_with_rule_based(question_text, user_answer)

def evaluate_with_rule_based(question_text, user_answer):
    """
    Evaluasi jawaban berbasis aturan sederhana sebagai fallback terakhir
    
    Metode ini menggunakan teknik NLP sederhana untuk menentukan kemiripan
    antara jawaban pengguna dan jawaban yang diharapkan berdasarkan pertanyaan
    """
    try:
        # 1. Ekstrak kata kunci dari pertanyaan
        # Hapus kata-kata umum dalam Bahasa Indonesia
        common_words = ["ada", "adalah", "yang", "di", "dengan", "untuk", "pada", "ini", "itu", 
                        "dari", "dalam", "dan", "ke", "akan", "atau", "secara", "jika", "maka",
                        "apakah", "bagaimana", "mengapa", "kenapa", "kapan", "siapa", "dimana", 
                        "mana", "apa", "sebutkan", "jelaskan", "berikan", "ceritakan"]
        
        # Normalisasi pertanyaan dan jawaban
        question_normalized = normalize_text_for_matching(question_text)
        answer_normalized = normalize_text_for_matching(user_answer)
        
        # Hapus kata-kata umum
        question_keywords = [w for w in question_normalized.split() if w.lower() not in common_words]
        
        # 2. Tentukan kemungkinan jawaban berdasarkan jenis pertanyaan
        question_type = determine_question_type(question_text)
        
        # 3. Verifikasi apakah jawaban pengguna sesuai dengan ekspektasi
        if question_type == "DEFINISI":
            # Untuk pertanyaan definisi, jawaban harus berisi kata kunci dari pertanyaan
            # dan memiliki panjang minimal
            keywords_in_answer = sum(1 for kw in question_keywords if kw.lower() in answer_normalized.lower())
            answer_length_ok = len(answer_normalized.split()) >= 5  # Jawaban minimal 5 kata
            
            return keywords_in_answer >= 2 and answer_length_ok
            
        elif question_type == "YA_TIDAK":
            # Jawaban harus berupa ya/tidak atau benar/salah atau variasi lain
            yes_patterns = ["ya", "benar", "betul", "tepat", "setuju", "iya"]
            no_patterns = ["tidak", "salah", "keliru", "bukan"]
            
            has_yes = any(p in answer_normalized.lower() for p in yes_patterns)
            has_no = any(p in answer_normalized.lower() for p in no_patterns)
            
            return has_yes or has_no
            
        elif question_type == "FAKTUAL":
            # Untuk pertanyaan faktual, jawaban harus berisi angka atau fakta spesifik
            # Cek apakah ada angka dalam jawaban
            has_number = bool(re.search(r'\d+', answer_normalized))
            
            # Cek apakah jawaban memiliki kata-kata spesifik (bukan hanya kata umum)
            answer_words = [w for w in answer_normalized.split() if w.lower() not in common_words]
            has_specific_words = len(answer_words) >= 3
            
            return has_number or has_specific_words
            
        else:  # UMUM
            # Untuk pertanyaan umum, pastikan jawaban memiliki panjang yang cukup
            # dan mengandung beberapa kata kunci dari pertanyaan
            min_answer_length = 3  # Minimal 3 kata
            answer_words = answer_normalized.split()
            
            # Hitung kata-kata unik dalam jawaban
            unique_words = len(set(w.lower() for w in answer_words))
            
            return len(answer_words) >= min_answer_length and unique_words >= 2
            
    except Exception as e:
        print(f"Error dalam evaluasi berbasis aturan: {str(e)}")
        # Fallback terakhir: terima jawaban jika tidak kosong (minimal usaha)
        return len(user_answer.strip()) > 0

def normalize_text_for_matching(text):
    """
    Normalisasi teks untuk perbandingan
    """
    # Ubah ke lowercase
    text = text.lower()
    
    # Hilangkan tanda baca
    text = re.sub(r'[^\w\s]', ' ', text)
    
    # Hilangkan spasi berlebih
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text

def determine_question_type(question_text):
    """
    Menentukan jenis pertanyaan berdasarkan pola dan kata kunci
    
    Returns:
        str: "DEFINISI", "YA_TIDAK", "FAKTUAL", atau "UMUM"
    """
    question_lower = question_text.lower()
    
    # Pola untuk pertanyaan definisi
    definition_patterns = ["apa itu", "apa yang dimaksud", "definisi", "pengertian", "arti dari", "jelaskan apa"]
    if any(pattern in question_lower for pattern in definition_patterns):
        return "DEFINISI"
    
    # Pola untuk pertanyaan ya/tidak
    yes_no_patterns = ["apakah", "benarkah", "benar bahwa", "setujukah"]
    if any(pattern in question_lower for pattern in yes_no_patterns):
        return "YA_TIDAK"
    
    # Pola untuk pertanyaan faktual
    factual_patterns = ["berapa", "kapan", "siapa", "dimana", "tahun berapa", "sebutkan"]
    if any(pattern in question_lower for pattern in factual_patterns):
        return "FAKTUAL"
    
    # Default jika tidak ada pola yang cocok
    return "UMUM"

# =========================================
# GET SUMMARY
# =========================================
@app.route('/get_summary', methods=['GET'])
def get_summary():
    user_id = request.args.get("user_id", "default")
    
    # Konversi user_id menjadi integer jika bukan "default"
    try:
        user_id = int(user_id) if user_id != "default" else None
    except ValueError:
        user_id = None
    
    if user_id is None:
        return jsonify({
            "status": "error", 
            "message": "ID pengguna tidak valid",
        }), 400
    
    try:
        # Ambil data hasil siswa
        user_result = SiswaResult.query.filter_by(siswa_id=user_id).first()
        if not user_result:
            # Return initial data instead of error
            return jsonify({
                "status": "success",
                "correct": 0,
                "incorrect": 0,
                "total_tests": 0,
                "completed_tests": 0
            }), 200
        
        # Hitung jumlah tes yang tersedia untuk siswa ini
        total_tests = db.session.query(func.count(collection_students.c.collection_id)).filter(
            collection_students.c.siswa_id == user_id
        ).scalar() or 0
        
        # Hitung jumlah tes yang sudah selesai
        # Tes selesai jika level saat ini adalah level final (5, 6, 7)
        completed_tests = SiswaResult.query.filter(
            SiswaResult.siswa_id == user_id,
            SiswaResult.current_level.in_([5, 6, 7])
        ).count()
        
        # Hitung jumlah jawaban benar dan salah
        total_correct = user_result.correct or 0
        total_incorrect = user_result.incorrect or 0
        
        return jsonify({
            "status": "success",
            "correct": total_correct,
            "incorrect": total_incorrect,
            "total_tests": total_tests,
            "completed_tests": completed_tests
        }), 200
    
    except Exception as e:
        print(f"Error getting summary: {str(e)}")
        return jsonify({
            "status": "error",
            "message": f"Error: {str(e)}"
        }), 500
    
@app.route("/get_questions", methods=["GET"])
def get_questions():
    try:
        print("Fetching questions from database...")
        
        # Gunakan SQL RAW langsung untuk memastikan tidak ada filtering
        from sqlalchemy import text
        result = db.session.execute(text("SELECT * FROM questions ORDER BY level ASC, id ASC"))
        
        # Konversi hasil query ke dalam format yang dapat di-serialize
        questions_raw = []
        for row in result:
            # Konversi hasil query jadi dict
            question_dict = {column: value for column, value in zip(result.keys(), row)}
            questions_raw.append(question_dict)
        
        print(f"Fetched {len(questions_raw)} questions using raw SQL")
        
        # Format data untuk respons
        questions_data = []
        for q in questions_raw:
            try:
                # Parse options
                options = []
                if 'options' in q and q['options']:
                    try:
                        options = json.loads(q['options'])
                    except:
                        pass
                
                questions_data.append({
                    "id": q['id'],
                    "level": q['level'],
                    "soal": q['soal'],
                    "jawaban_benar": q['jawaban_benar'] if 'jawaban_benar' in q else "",
                    "options": options,
                    "question_type": q['question_type'] if 'question_type' in q else "multiple_choice",
                    "p": float(q['p']) if 'p' in q else 0.5,
                    "bobot": q['bobot'] if 'bobot' in q else 1,
                    "explanation": q['explanation'] if 'explanation' in q else "",
                    "created_at": q['created_at'].isoformat() if 'created_at' in q and q['created_at'] else None
                })
            except Exception as e:
                print(f"Error processing question {q.get('id')}: {str(e)}")
                continue
        
        print(f"Successfully formatted {len(questions_data)} questions for response")
        return jsonify({"success": True, "data": questions_data})
    
    except Exception as e:
        print(f"Error getting questions: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "message": f"Error: {str(e)}"}), 500
    
@app.route("/api/check_questions_count", methods=["GET"])
def check_questions_count():
    try:
        # Hitung soal
        total_count = db.session.query(db.func.count(Question.id)).scalar()
        
        # Hitung berdasarkan level
        level_counts = db.session.query(
            Question.level, 
            db.func.count(Question.id)
        ).group_by(Question.level).all()
        
        # Format distribusi level
        level_distribution = {str(level): count for level, count in level_counts}
        
        return jsonify({
            "success": True,
            "totalQuestions": total_count,
            "levelDistribution": level_distribution,
            "queryTime": 0
        })
    except Exception as e:
        print(f"Error checking questions count: {str(e)}")
        return jsonify({
            "success": False,
            "message": f"Error checking database: {str(e)}"
        }), 500
@app.route("/api/db_check", methods=["GET"])
def db_check():
    try:
        from sqlalchemy import text
        
        # Hitung total soal
        result = db.session.execute(text("SELECT COUNT(*) FROM questions"))
        total = result.scalar()
        
        # Hitung soal per level
        result = db.session.execute(text("SELECT level, COUNT(*) FROM questions GROUP BY level"))
        level_counts = {row[0]: row[1] for row in result}
        
        # Ambil sampel soal
        result = db.session.execute(text("SELECT id, level, soal FROM questions LIMIT 5"))
        samples = [{"id": row[0], "level": row[1], "soal": row[2]} for row in result]
        
        return jsonify({
            "success": True,
            "total_questions": total,
            "per_level": level_counts,
            "samples": samples
        })
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})
    
@app.route("/api/table_structure", methods=["GET"])
def check_table_structure():
    try:
        from sqlalchemy import text
        
        # Cek struktur tabel questions
        result = db.session.execute(text("DESCRIBE questions"))
        columns = [dict(row._mapping) for row in result]
        
        # Cek foreign key constraints
        result = db.session.execute(text("SELECT * FROM INFORMATION_SCHEMA.KEY_COLUMN_USAGE WHERE TABLE_NAME = 'questions'"))
        constraints = [dict(row._mapping) for row in result]
        
        return jsonify({
            "success": True,
            "table_structure": columns,
            "constraints": constraints
        })
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})
    
@app.route('/api/collections', methods=['GET'])
@login_required
def get_collections():
    # Ambil semua koleksi milik guru
    collections = QuestionCollection.query.filter_by(guru_id=current_user.id).all()
    
    result = []
    for collection in collections:
        # Hitung jumlah soal di tiap koleksi
        question_count = db.session.query(CollectionQuestion).filter_by(collection_id=collection.id).count()
        
        result.append({
            'id': collection.id,
            'name': collection.name,
            'description': collection.description,
            'question_count': question_count,
            'created_at': collection.created_at.isoformat() if collection.created_at else None
        })
    
    return jsonify({'success': True, 'collections': result})

@app.route('/api/collections', methods=['POST'])
@login_required
def create_collection():
    data = request.json
    
    # Validasi data
    if not data.get('name'):
        return jsonify({'success': False, 'message': 'Nama koleksi harus diisi'}), 400
    
    # Buat koleksi baru
    collection = QuestionCollection(
        guru_id=current_user.id,
        name=data.get('name'),
        description=data.get('description', '')
    )
    
    db.session.add(collection)
    db.session.commit()
    
    # Tambahkan soal ke koleksi jika ada
    if data.get('question_ids') and isinstance(data.get('question_ids'), list):
        for question_id in data.get('question_ids'):
            # Validasi bahwa soal ada
            question = Question.query.get(question_id)
            if question:
                collection_question = CollectionQuestion(
                    collection_id=collection.id,
                    question_id=question_id
                )
                db.session.add(collection_question)
        
        db.session.commit()
    
    return jsonify({
        'success': True, 
        'message': 'Koleksi berhasil dibuat', 
        'collection_id': collection.id
    })

@app.route('/api/collections/<int:collection_id>/details', methods=['GET'])
@login_required
def get_collection_details_for_student(collection_id):
    try:
        # Cek apakah siswa memiliki akses ke koleksi ini
        has_access = db.session.query(collection_students).filter_by(
            collection_id=collection_id, 
            siswa_id=current_user.id
        ).first() is not None
        
        if not has_access:
            return jsonify({'success': False, 'message': 'Anda tidak memiliki akses ke koleksi ini'}), 403
        
        # Ambil detail koleksi
        collection = QuestionCollection.query.get(collection_id)
        if not collection:
            return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
        
        # Ambil info tentang guru
        guru = Guru.query.get(collection.guru_id)
        
        # Ambil progress siswa untuk koleksi ini
        progress = SiswaResult.query.filter_by(
            siswa_id=current_user.id,
            collection_id=collection_id
        ).first()
        
        result = {
            'success': True,
            'collection': {
                'id': collection.id,
                'name': collection.name,
                'description': collection.description,
                'created_at': collection.created_at.isoformat() if collection.created_at else None
            },
            'guru': {
                'id': guru.id,
                'nama': guru.nama,
                'email': guru.email
            },
            'progress': {
                'current_level': progress.current_level if progress else 1,
                'correct': progress.correct if progress else 0,
                'incorrect': progress.incorrect if progress else 0,
                'is_completed': progress.current_level in [5, 6, 7] if progress else False
            }
        }
        
        return jsonify(result)
        
    except Exception as e:
        print(f"Error getting collection details: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500
    
@app.route('/api/collections/<int:collection_id>/questions', methods=['GET'])
@login_required
def get_collection_questions(collection_id):
    # Verifikasi kepemilikan koleksi
    collection = QuestionCollection.query.filter_by(id=collection_id, guru_id=current_user.id).first()
    if not collection:
        return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
    
    # Ambil soal dari koleksi dengan join
    questions = db.session.query(Question).join(
        CollectionQuestion, CollectionQuestion.question_id == Question.id
    ).filter(
        CollectionQuestion.collection_id == collection_id
    ).all()
    
    # Format soal untuk response
    question_data = []
    for q in questions:
        options = []
        if hasattr(q, 'options') and q.options:
            try:
                options = json.loads(q.options)
            except:
                options = []
        
        question_data.append({
            'id': q.id,
            'level': q.level,
            'soal': q.soal,
            'jawaban_benar': q.jawaban_benar,
            'options': options,
            'question_type': q.question_type if hasattr(q, 'question_type') else 'multiple_choice',
            'p': float(q.p) if q.p else 0.5,
            'bobot': q.bobot or 1,
            'explanation': q.explanation or '',
            'is_validated': q.is_validated # NEW: Sertakan status validasi
        })
    
    return jsonify({
        'success': True, 
        'collection': {
            'id': collection.id,
            'name': collection.name,
            'description': collection.description
        },
        'questions': question_data
    })


@app.route('/api/collections/<int:collection_id>/questions', methods=['POST'])
@login_required
def add_questions_to_collection(collection_id):
    # Verifikasi kepemilikan koleksi
    collection = QuestionCollection.query.filter_by(id=collection_id, guru_id=current_user.id).first()
    if not collection:
        return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
    
    data = request.json
    question_ids = data.get('question_ids', [])
    
    # Validasi data
    if not question_ids or not isinstance(question_ids, list):
        return jsonify({'success': False, 'message': 'ID soal tidak valid'}), 400
    
    # Tambahkan soal ke koleksi
    added_count = 0
    for question_id in question_ids:
        # Cek apakah soal sudah ada di koleksi
        existing = CollectionQuestion.query.filter_by(
            collection_id=collection_id, 
            question_id=question_id
        ).first()
        
        if not existing:
            # Cek apakah soal ada
            question = Question.query.get(question_id)
            if question:
                collection_question = CollectionQuestion(
                    collection_id=collection_id,
                    question_id=question_id
                )
                db.session.add(collection_question)
                added_count += 1
    
    db.session.commit()
    
    return jsonify({
        'success': True, 
        'message': f'{added_count} soal berhasil ditambahkan ke koleksi'
    })

@app.route('/api/collections/<int:collection_id>', methods=['DELETE'])
@login_required
def delete_collection(collection_id):
    # Verifikasi kepemilikan koleksi
    collection = QuestionCollection.query.filter_by(id=collection_id, guru_id=current_user.id).first()
    if not collection:
        return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
    
    # Hapus semua relasi collection_questions terlebih dahulu
    CollectionQuestion.query.filter_by(collection_id=collection_id).delete()
    
    # Hapus koleksi
    db.session.delete(collection)
    db.session.commit()
    
    return jsonify({
        'success': True, 
        'message': 'Koleksi berhasil dihapus'
    })

# Mendapatkan siswa yang memiliki akses ke koleksi
@app.route('/api/collections/<int:collection_id>/students', methods=['GET'])
@login_required
def get_collection_students(collection_id):
    try:
        # Verifikasi kepemilikan koleksi
        collection = QuestionCollection.query.filter_by(id=collection_id, guru_id=current_user.id).first()
        if not collection:
            return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
        
        # Dapatkan siswa yang memiliki akses ke koleksi ini
        students = db.session.query(Siswa).join(
            collection_students,
            collection_students.c.siswa_id == Siswa.id
        ).filter(
            collection_students.c.collection_id == collection_id
        ).all()
        
        # Format data siswa
        students_data = []
        for student in students:
            students_data.append({
                'id': student.id,
                'nama': student.nama,
                'username': student.username,
                'kelas': student.kelas
            })
        
        return jsonify({
            'success': True,
            'students': students_data
        })
        
    except Exception as e:
        print(f"Error getting collection students: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

# Menambahkan siswa ke koleksi
@app.route('/api/collections/<int:collection_id>/students', methods=['POST'])
@login_required
def add_student_to_collection(collection_id):
    try:
        # 1. Verifikasi koleksi dan kepemilikan
        collection = QuestionCollection.query.filter_by(
            id=collection_id, 
            guru_id=current_user.id
        ).first()
        
        if not collection:
            return jsonify({
                'success': False,
                'message': 'Koleksi tidak ditemukan atau akses ditolak'
            }), 404

        # 2. Validasi data request
        data = request.get_json()
        if not data or 'siswa_id' not in data:
            return jsonify({
                'success': False,
                'message': 'Data siswa tidak valid'
            }), 400

        siswa_id = data['siswa_id']

        # 3. Cek keberadaan siswa
        siswa = Siswa.query.get(siswa_id)
        if not siswa:
            return jsonify({
                'success': False,
                'message': 'Siswa tidak ditemukan'
            }), 404

        # 4. Cek duplikasi relasi
        existing = db.session.execute(
            collection_students.select().where(
                (collection_students.c.collection_id == collection_id) &
                (collection_students.c.siswa_id == siswa_id)
            )
        ).fetchone()

        if existing:
            return jsonify({
                'success': True,
                'message': 'Siswa sudah terdaftar dalam koleksi ini'
            })

        # 5. Tambahkan relasi
        db.session.execute(
            collection_students.insert().values(
                collection_id=collection_id,
                siswa_id=siswa_id
            )
        )

        # 6. Buat atau update hasil siswa
        siswa_result = SiswaResult.query.filter_by(
            siswa_id=siswa_id,
            collection_id=collection_id
        ).first()

        if not siswa_result:
            new_result = SiswaResult(
                siswa_id=siswa_id,
                collection_id=collection_id,
                current_level=1,
                correct=0,
                incorrect=0
            )
            db.session.add(new_result)

        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Siswa berhasil ditambahkan ke koleksi',
            'data': {
                'collection_id': collection_id,
                'siswa_id': siswa_id,
                'siswa_nama': siswa.nama,
                'collection_name': collection.name
            }
        })

    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Unexpected error: {str(e)}")
        return jsonify({
            'success': False,
            'message': 'Terjadi kesalahan server'
        }), 500

# Menghapus siswa dari koleksi
@app.route('/api/collections/<int:collection_id>/students/<int:student_id>', methods=['DELETE'])
@login_required
def remove_student_from_collection(collection_id, student_id):
    try:
        # Verifikasi kepemilikan koleksi
        collection = QuestionCollection.query.filter_by(id=collection_id, guru_id=current_user.id).first()
        if not collection:
            return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
        
        # Hapus akses
        stmt = collection_students.delete().where(
            collection_students.c.collection_id == collection_id,
            collection_students.c.siswa_id == student_id
        )
        result = db.session.execute(stmt)
        db.session.commit()
        
        if result.rowcount == 0:
            return jsonify({'success': False, 'message': 'Siswa tidak memiliki akses ke koleksi ini'}), 404
        
        return jsonify({
            'success': True,
            'message': 'Akses siswa berhasil dihapus'
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"Error removing student from collection: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/collections/<int:collection_id>/student-count', methods=['GET'])
def get_collection_student_count(collection_id):
    try:
        count = db.session.query(collection_students).filter_by(
            collection_id=collection_id
        ).count()
        return jsonify({'success': True, 'count': count})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

# Mendapatkan daftar semua siswa
@app.route('/api/siswa', methods=['GET'])
@login_required
def get_all_students():
    try:
        if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
            return jsonify({'success': False, 'message': 'Akses ditolak. Hanya guru yang dapat mengakses daftar siswa.'}), 403
        
        kelas_filter = request.args.get('kelas') # NEW: Ambil parameter kelas untuk filter
        
        students_query = Siswa.query

        if kelas_filter and kelas_filter != 'all':
            students_query = students_query.filter_by(kelas=kelas_filter)

        students = students_query.order_by(Siswa.kelas, Siswa.nama).all() # Urutkan berdasarkan kelas dan nama
        
        students_data = []
        for student in students:
            students_data.append({
                'id': student.id,
                'nama': student.nama,
                'username': student.username,
                'kelas': student.kelas
            })
        
        # NEW: Ambil daftar kelas unik yang ada di database
        available_classes = db.session.query(Siswa.kelas).distinct().order_by(Siswa.kelas).all()
        available_classes = [cls[0] for cls in available_classes if cls[0]] # Filter out None/empty classes
        
        return jsonify({
            'success': True,
            'students': students_data,
            'available_classes': available_classes # Kirim daftar kelas yang tersedia
        })
        
    except Exception as e:
        print(f"Error getting all students: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500
    
# Tambahkan endpoint baru untuk menambahkan siswa berdasarkan kelas
@app.route('/api/collections/<int:collection_id>/add_students_by_class', methods=['POST'])
@login_required
def add_students_by_class(collection_id):
    try:
        # 1. Verifikasi koleksi dan kepemilikan
        collection = QuestionCollection.query.filter_by(
            id=collection_id,
            guru_id=current_user.id
        ).first()

        if not collection:
            return jsonify({
                'success': False,
                'message': 'Koleksi tidak ditemukan atau akses ditolak'
            }), 404

        # 2. Validasi data request
        data = request.get_json()
        if not data or 'class_name' not in data:
            return jsonify({
                'success': False,
                'message': 'Nama kelas tidak valid'
            }), 400

        class_name = data['class_name']

        # 3. Dapatkan semua siswa di kelas tersebut yang belum terdaftar di koleksi ini
        students_to_add = Siswa.query.filter(
            Siswa.kelas == class_name,
            # Filter siswa yang belum ada di koleksi ini
            ~Siswa.id.in_(
                db.session.query(collection_students.c.siswa_id).filter(
                    collection_students.c.collection_id == collection_id
                )
            )
        ).all()

        added_count = 0
        if students_to_add:
            # 4. Tambahkan siswa ke koleksi dan buat/update SiswaResult mereka
            for student in students_to_add:
                # Tambahkan relasi ke collection_students
                db.session.execute(
                    collection_students.insert().values(
                        collection_id=collection_id,
                        siswa_id=student.id
                    )
                )

                # Buat atau update SiswaResult
                siswa_result = SiswaResult.query.filter_by(
                    siswa_id=student.id,
                    collection_id=collection_id
                ).first()

                if not siswa_result:
                    new_result = SiswaResult(
                        siswa_id=student.id,
                        collection_id=collection_id,
                        current_level=1,
                        correct=0,
                        incorrect=0
                    )
                    db.session.add(new_result)
                added_count += 1
            db.session.commit()

        return jsonify({
            'success': True,
            'message': f'{added_count} siswa dari kelas {class_name} berhasil ditambahkan ke koleksi.',
            'added_count': added_count
        }), 200

    except Exception as e:
        db.session.rollback()
        print(f"Error adding students by class: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': 'Terjadi kesalahan server saat menambahkan siswa berdasarkan kelas'
        }), 500

@app.route('/api/siswa/collections', methods=['GET'])
@login_required
def get_student_collections():
    if not current_user.is_authenticated or not hasattr(current_user, 'user_type') or current_user.user_type != 'siswa':
        return jsonify({'success': False, 'message': 'Akses ditolak'}), 403
    
    try:
        # Ambil koleksi yang tersedia untuk siswa ini dari tabel relasi
        collections_query = db.session.query(
            QuestionCollection, Guru
        ).join(
            Guru, QuestionCollection.guru_id == Guru.id
        ).join(
            collection_students, collection_students.c.collection_id == QuestionCollection.id
        ).filter(
            collection_students.c.siswa_id == current_user.id
        ).all()
        
        # Format data untuk respon
        teachers_dict = {}
        for collection, guru in collections_query:
            if guru.id not in teachers_dict:
                teachers_dict[guru.id] = {
                    'id': guru.id,
                    'nama': guru.nama,
                    'email': guru.email,
                    'collections': []
                }
            
            # Cek apakah koleksi baru (dibuat dalam 7 hari terakhir)
            is_new = False
            if collection.created_at:
                delta = datetime.datetime.now() - collection.created_at
                is_new = delta.days < 7
            
            teachers_dict[guru.id]['collections'].append({
                'id': collection.id,
                'name': collection.name,
                'is_new': is_new,
                'description': collection.description
            })
        
        # Konversi ke list untuk response
        teachers_list = list(teachers_dict.values())
        
        return jsonify({
            'success': True, 
            'teachers': teachers_list
        })
        
    except Exception as e:
        print(f"Error getting student collections: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500
    
# =========================================
# ENDPOINT: GET_QUESTIONS (Dipanggil oleh result.html dan guru.html)
# Mengambil semua data soal dari database
# =========================================
@app.route("/api/get_questions", methods=["GET"])
@login_required 
def get_questions_api(): 
    try:
        # Mengambil semua soal dari database. 
        # Anda bisa menambahkan filter di sini jika hanya ingin soal yang 'is_current=True' atau 'is_validated=True'.
        questions_raw = Question.query.order_by(Question.level.asc(), Question.id.asc()).all()
        
        questions_data = []
        for q in questions_raw:
            options = []
            if q.options:
                try:
                    options = json.loads(q.options)
                except json.JSONDecodeError:
                    pass # Abaikan jika tidak valid JSON
            
            questions_data.append({
                "id": q.id,
                "guru_id": q.guru_id,
                "collection_id": q.collection_id, # Sertakan jika relevan
                "level": q.level,
                "soal": q.soal,
                "jawaban_benar": q.jawaban_benar,
                "options": options,
                "question_type": q.question_type,
                "p": float(q.p),
                "bobot": q.bobot,
                "explanation": q.explanation,
                "created_at": q.created_at.isoformat() if q.created_at else None,
                "version": q.version,
                "is_current": q.is_current,
                "is_validated": q.is_validated 
            })
            
        return jsonify({"success": True, "data": questions_data}), 200
        
    except Exception as e:
        print(f"Error getting questions via API: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "message": f"Error: {str(e)}"}), 

@app.route('/api/guru', methods=['GET'])
@login_required
def get_all_teachers():
    try:
        # Only siswa should access this endpoint
        if not hasattr(current_user, 'user_type') or current_user.user_type != 'siswa':
            return jsonify({'success': False, 'message': 'Akses ditolak'}), 403
        
        # Get all teachers
        teachers = Guru.query.all()
        
        # Format teacher data
        teachers_data = []
        for teacher in teachers:
            teachers_data.append({
                'id': teacher.id,
                'nama': teacher.nama,
                'email': teacher.email
            })
        
        return jsonify({
            'success': True,
            'teachers': teachers_data
        })
        
    except Exception as e:
        print(f"Error getting all teachers: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500
    
@app.route('/dashboard')
@login_required
def dashboard():
    # Jika pengguna adalah siswa, ambil koleksi yang tersedia
    collections = []
    if hasattr(current_user, 'user_type') and current_user.user_type == 'siswa':
        try:
            # Ambil koleksi yang tersedia untuk siswa ini
            collections_data = db.session.query(
                QuestionCollection, Guru
            ).join(
                Guru, QuestionCollection.guru_id == Guru.id
            ).join(
                collection_students, collection_students.c.collection_id == QuestionCollection.id
            ).filter(
                collection_students.c.siswa_id == current_user.id
            ).all()
            
            # Format data koleksi
            for collection, guru in collections_data:
                # Cek apakah koleksi baru
                is_new = False
                if collection.created_at:
                    delta = datetime.datetime.now() - collection.created_at
                    is_new = delta.days < 7
                
                collections.append({
                    'id': collection.id,
                    'name': collection.name,
                    'is_new': is_new,
                    'description': collection.description,
                    'guru_id': guru.id,
                    'guru_name': guru.nama,
                    'guru_email': guru.email
                })
        except Exception as e:
            print(f"Error fetching collections: {str(e)}")
    
    return render_template('DashboardSiswa.html', collections=collections)
# =========================================
# VERSIONING ENDPOINTS (Tempatkan setelah definisi model)
# =========================================
# Pastikan hanya ada satu implementasi, hapus yang lama
@app.route('/api/questions/<int:id>', methods=['PUT', 'POST'])
@login_required
def update_question(id):
    try:
        # Log request method untuk debugging
        print(f"Request method: {request.method}")
        print(f"Request headers: {request.headers}")
        
        if request.method == 'POST' and request.headers.get('X-HTTP-Method-Override') == 'PUT':
            print("Using X-HTTP-Method-Override: PUT")
            # Lanjutkan sebagai PUT request
        
        data = request.get_json()
        print(f"Request data: {data}")
        
        # Dapatkan pertanyaan yang akan diupdate
        question = Question.query.get(id)
        if not question:
            return jsonify({'success': False, 'message': 'Pertanyaan tidak ditemukan'}), 404
        
        # Buat versi sebelum update
        prev_version = QuestionVersion(
            question_id=question.id,
            soal=question.soal,
            options=question.options,
            jawaban_benar=question.jawaban_benar,
            explanation=question.explanation,
            p=question.p,
            version=question.version  # Gunakan versi saat ini
        )
        db.session.add(prev_version)
        
        # Update pertanyaan utama
        question.soal = data.get('soal', question.soal)
        question.options = json.dumps(data.get('options')) if 'options' in data else question.options
        question.jawaban_benar = data.get('jawaban_benar', question.jawaban_benar)
        question.explanation = data.get('explanation', question.explanation)
        question.p = data.get('p', question.p)
        
        # Increment versi
        question.version += 1
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'question': {
                'id': question.id,
                'version': question.version,
                'soal': question.soal,
                'options': json.loads(question.options) if question.options else []
            }
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"Error updating question: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/questions/<int:id>/versions', methods=['GET'])
@login_required
def get_question_versions(id):
    try:
        print(f"Fetching versions for question {id}")
        versions = QuestionVersion.query.filter_by(question_id=id).order_by(QuestionVersion.version.desc()).all()
        print(f"Found {len(versions)} versions")
        
        result = []
        for version in versions:
            result.append({
                'id': version.id,
                'soal': version.soal,
                'options': json.loads(version.options) if version.options else [],
                'jawaban_benar': version.jawaban_benar,
                'explanation': version.explanation,
                'p': version.p,
                'version': version.version,
                'created_at': version.created_at.isoformat()
            })
            
        return jsonify({'success': True, 'versions': result})
    except Exception as e:
        print(f"Error getting versions: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/questions/in_collection', methods=['GET'])
@login_required
def get_questions_in_collection():
    try:
        # Periksa apakah collection_id disediakan sebagai parameter query
        collection_id = request.args.get('collection_id')
        
        if collection_id:
            # Filter berdasarkan collection_id tertentu
            collection_question_ids = db.session.query(CollectionQuestion.question_id)\
                .filter(CollectionQuestion.collection_id == collection_id).all()
        else:
            # Dapatkan semua ID soal yang ada dalam collection apapun
            collection_question_ids = db.session.query(CollectionQuestion.question_id).distinct().all()
            
        collection_question_ids = [id[0] for id in collection_question_ids]
        
        # Jika perlu detail soal, bukan hanya ID
        include_details = request.args.get('include_details') == 'true'
        if include_details:
            # Ambil detail soal berdasarkan ID yang sudah didapat
            questions = db.session.query(Question)\
                .filter(Question.id.in_(collection_question_ids))\
                .all()
                
            # Format data soal
            question_data = []
            for q in questions:
                options = []
                if hasattr(q, 'options') and q.options:
                    try:
                        options = json.loads(q.options)
                    except:
                        pass
                
                question_data.append({
                    'id': q.id,
                    'level': q.level,
                    'soal': q.soal,
                    'jawaban_benar': q.jawaban_benar,
                    'options': options,
                    'question_type': q.question_type if hasattr(q, 'question_type') else 'multiple_choice',
                    'p': float(q.p) if q.p else 0.5,
                    'bobot': q.bobot or 1,
                    'explanation': q.explanation or ''
                })
                
            return jsonify({
                'success': True,
                'questions': question_data,
                'question_ids': collection_question_ids
            })
        else:
            # Hanya mengembalikan ID soal
            return jsonify({
                'success': True,
                'question_ids': collection_question_ids
            })
            
    except Exception as e:
        print(f"Error getting questions in collection: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500
    
@app.route('/api/collections/<int:collection_id>/status', methods=['POST'])
@login_required
def set_collection_status(collection_id):
    try:
        # Verifikasi kepemilikan koleksi
        collection = QuestionCollection.query.filter_by(id=collection_id, guru_id=current_user.id).first()
        if not collection:
            return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
        
        data = request.json
        is_active = data.get('isActive', False)
        
        # Update status
        collection.is_active = bool(is_active)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'Koleksi berhasil {"diaktifkan" if is_active else "dinonaktifkan"}'
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500
    

@app.route('/guru/collections')
@login_required
def collection_page():
    # Verifikasi bahwa pengguna adalah guru
    if not hasattr(current_user, 'user_type') or current_user.user_type != 'guru':
        flash('Halaman ini hanya untuk guru', 'error')
        return redirect(url_for('login'))
    
    return render_template('collections.html')

@app.route('/api/export_result_excel', methods=['GET'])
@login_required
def export_result_excel():
    try:
        user_id = request.args.get('user_id')
        collection_id = request.args.get('collection_id')
        
        # Validasi parameter
        if not user_id or not collection_id:
            return jsonify({'success': False, 'message': 'Parameter tidak lengkap'}), 400
            
        # Ambil info siswa
        student = Siswa.query.get(user_id)
        if not student:
            return jsonify({'success': False, 'message': 'Siswa tidak ditemukan'}), 404
            
        # Ambil info koleksi
        collection = QuestionCollection.query.get(collection_id)
        if not collection:
            return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
            
        # Ambil hasil tes
        result = SiswaResult.query.filter_by(
            siswa_id=user_id,
            collection_id=collection_id
        ).first()
        
        if not result:
            return jsonify({'success': False, 'message': 'Hasil tes tidak ditemukan'}), 404
            
        # Buat file Excel
        output = BytesIO()
        
        # Format untuk membuat Excel
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            # Format untuk judul dan header
            workbook = writer.book
            title_format = workbook.add_format({
                'bold': True,
                'font_size': 16,
                'align': 'center',
                'valign': 'vcenter'
            })
            header_format = workbook.add_format({
                'bold': True,
                'text_wrap': True,
                'valign': 'top',
                'fg_color': '#D7E4BC',
                'border': 1
            })
            label_format = workbook.add_format({
                'bold': True,
                'align': 'right',
                'border': 1
            })
            data_format = workbook.add_format({
                'align': 'left',
                'border': 1
            })
            
            # Buat worksheet
            worksheet = workbook.add_worksheet('Ringkasan')
            
            # Judul utama
            worksheet.merge_range('A1:H1', f"HASIL TES DIAGNOSTIK - {collection.name}", title_format)
            
            # Detail siswa dalam format label-data
            worksheet.write('A3', 'Nama:', label_format)
            worksheet.write('B3', student.nama, data_format)
            
            worksheet.write('A4', 'Kelas:', label_format)
            worksheet.write('B4', student.kelas, data_format)
            
            worksheet.write('A5', 'Koleksi Soal:', label_format)
            worksheet.write('B5', collection.name, data_format)
            
            worksheet.write('D3', 'level Terakhir:', label_format)
            worksheet.write('E3', result.current_level, data_format)
            
            worksheet.write('D4', 'Jawaban Benar:', label_format)
            worksheet.write('E4', result.correct, data_format)
            
            worksheet.write('D5', 'Jawaban Salah:', label_format)
            worksheet.write('E5', result.incorrect, data_format)
            
            total = result.correct + result.incorrect
            accuracy = f"{round((result.correct / total * 100), 2)}%" if total > 0 else "0%"
            
            worksheet.write('G3', 'Total Soal:', label_format)
            worksheet.write('H3', total, data_format)
            
            worksheet.write('G4', 'Akurasi:', label_format)
            worksheet.write('H4', accuracy, data_format)
            
            worksheet.write('G5', 'Tanggal Pengerjaan:', label_format)
            worksheet.write('H5', result.updated_at.strftime('%d-%m-%Y %H:%M:%S') if result.updated_at else datetime.datetime.now().strftime('%d-%m-%Y %H:%M:%S'), data_format)
            
            # Sesuaikan lebar kolom
            worksheet.set_column('A:A', 15)
            worksheet.set_column('B:B', 25)
            worksheet.set_column('C:C', 5)
            worksheet.set_column('D:D', 15)
            worksheet.set_column('E:E', 15)
            worksheet.set_column('F:F', 5)
            worksheet.set_column('G:G', 18)
            worksheet.set_column('H:H', 25)
            
            # Ambil riwayat jawaban
            answers = get_student_answer_history(user_id, collection_id)
            
            # Buat sheet untuk detail jawaban
            if answers:
                # Buat dataframe untuk jawaban
                answers_data = []
                for answer in answers:
                    answers_data.append({
                        'Soal': answer['question'],
                        'Jawaban Siswa': answer['user_answer'],
                        'Jawaban Benar': answer['correct_answer'],
                        'Status': 'Benar' if answer['is_correct'] else 'Salah',
                        'level': answer['level'],
                        'Waktu Menjawab': answer['answered_at']
                    })
                
                df_answers = pd.DataFrame(answers_data)
                # Tulis detail jawaban ke sheet kedua
                df_answers.to_excel(writer, sheet_name='Detail Jawaban', index=False)
                
                # Format sheet kedua
                worksheet2 = writer.sheets['Detail Jawaban']
                worksheet2.set_column('A:A', 40)
                worksheet2.set_column('B:B', 20)
                worksheet2.set_column('C:C', 20)
                worksheet2.set_column('D:D', 10)
                worksheet2.set_column('E:E', 10)
                worksheet2.set_column('F:F', 20)
        
        # Kembali ke awal file
        output.seek(0)
        
        # Buat nama file
        filename = f"Hasil_Tes_{student.nama}_{collection.name}_{datetime.datetime.now().strftime('%Y%m%d')}.xlsx"
        
        # Kirim file
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        print(f"Error mengekspor ke Excel: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f"Error: {str(e)}"}), 500

@app.route('/api/export_result_excel_teacher', methods=['GET'])
@login_required
def export_result_excel_teacher():
    try:
        # Get parameters
        export_type = request.args.get('type', 'individual')  # 'individual' or 'class'
        user_id = request.args.get('user_id')
        class_id = request.args.get('class_id')
        collection_id = request.args.get('collection_id')
        
        # Validate parameters
        if not collection_id:
            return jsonify({'success': False, 'message': 'Parameter collection_id is required'}), 400
            
        if export_type == 'individual' and not user_id:
            return jsonify({'success': False, 'message': 'Parameter user_id is required for individual export'}), 400
            
        if export_type == 'class' and not class_id:
            return jsonify({'success': False, 'message': 'Parameter class_id is required for class export'}), 400
            
        # Create a BytesIO object to store the Excel file
        output = BytesIO()
        
        # Collection information
        collection = QuestionCollection.query.get(collection_id)
        if not collection:
            return jsonify({'success': False, 'message': 'Collection not found'}), 404
            
        # Create Excel with appropriate data based on type
        if export_type == 'individual':
            # Individual student export
            student = Siswa.query.get(user_id)
            if not student:
                return jsonify({'success': False, 'message': 'Student not found'}), 404
                
            # Get student results
            result = SiswaResult.query.filter_by(
                siswa_id=user_id,
                collection_id=collection_id
            ).first()
            
            if not result:
                return jsonify({'success': False, 'message': 'No results found for this student'}), 404
                
            # Create Excel file with student data
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                # Format for styling
                workbook = writer.book
                title_format = workbook.add_format({
                    'bold': True,
                    'font_size': 16,
                    'align': 'center',
                    'valign': 'vcenter'
                })
                header_format = workbook.add_format({
                    'bold': True,
                    'text_wrap': True,
                    'valign': 'top',
                    'fg_color': '#D7E4BC',
                    'border': 1
                })
                label_format = workbook.add_format({
                    'bold': True,
                    'align': 'right',
                    'border': 1
                })
                data_format = workbook.add_format({
                    'align': 'left',
                    'border': 1
                })
                
                # Create summary worksheet
                worksheet = workbook.add_worksheet('Summary')
                
                # Add title
                worksheet.merge_range('A1:H1', f"HASIL TES DIAGNOSTIK - {collection.name}", title_format)
                
                # Student details
                worksheet.write('A3', 'Nama:', label_format)
                worksheet.write('B3', student.nama, data_format)
                
                worksheet.write('A4', 'Kelas:', label_format)
                worksheet.write('B4', student.kelas, data_format)
                
                worksheet.write('A5', 'Koleksi Soal:', label_format)
                worksheet.write('B5', collection.name, data_format)
                
                # Result details
                worksheet.write('D3', 'level Terakhir:', label_format)
                worksheet.write('E3', result.current_level, data_format)
                
                worksheet.write('D4', 'Jawaban Benar:', label_format)
                worksheet.write('E4', result.correct, data_format)
                
                worksheet.write('D5', 'Jawaban Salah:', label_format)
                worksheet.write('E5', result.incorrect, data_format)
                
                # Statistics
                total = result.correct + result.incorrect
                accuracy = f"{round((result.correct / total * 100), 2)}%" if total > 0 else "0%"
                
                worksheet.write('G3', 'Total Soal:', label_format)
                worksheet.write('H3', total, data_format)
                
                worksheet.write('G4', 'Akurasi:', label_format)
                worksheet.write('H4', accuracy, data_format)
                
                worksheet.write('G5', 'Tanggal Pengerjaan:', label_format)
                worksheet.write('H5', result.updated_at.strftime('%d-%m-%Y %H:%M:%S') if result.updated_at else datetime.datetime.now().strftime('%d-%m-%Y %H:%M:%S'), data_format)
                
                # Format columns
                worksheet.set_column('A:A', 15)
                worksheet.set_column('B:B', 25)
                worksheet.set_column('C:C', 5)
                worksheet.set_column('D:D', 15)
                worksheet.set_column('E:E', 15)
                worksheet.set_column('F:F', 5)
                worksheet.set_column('G:G', 18)
                worksheet.set_column('H:H', 25)
                
                # Get answer history
                answers = get_student_answer_history(user_id, collection_id)
                
                # Add answers to another sheet
                if answers:
                    # Create dataframe for answers
                    answers_data = []
                    for answer in answers:
                        answers_data.append({
                            'Soal': answer['question'],
                            'Jawaban Siswa': answer['user_answer'],
                            'Jawaban Benar': answer['correct_answer'],
                            'Status': 'Benar' if answer['is_correct'] else 'Salah',
                            'level': answer['level'],
                            'Waktu Menjawab': answer['answered_at']
                        })
                    
                    # Create DataFrame and write to Excel
                    if answers_data:
                        df_answers = pd.DataFrame(answers_data)
                        df_answers.to_excel(writer, sheet_name='Detail Jawaban', index=False)
                        
                        # Format columns
                        answer_sheet = writer.sheets['Detail Jawaban']
                        answer_sheet.set_column('A:A', 40)
                        answer_sheet.set_column('B:B', 20)
                        answer_sheet.set_column('C:C', 20)
                        answer_sheet.set_column('D:D', 10)
                        answer_sheet.set_column('E:E', 10)
                        answer_sheet.set_column('F:F', 20)
            
            # Generate filename
            filename = f"Hasil_Tes_{student.nama}_{collection.name}_{datetime.datetime.now().strftime('%Y%m%d')}.xlsx"
            
        else:
            # Class export
            # Get all students in the class
            students = Siswa.query.filter_by(kelas=class_id).all()
            if not students:
                return jsonify({'success': False, 'message': 'No students found in this class'}), 404
            
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                workbook = writer.book
                title_format = workbook.add_format({
                    'bold': True,
                    'font_size': 16,
                    'align': 'center',
                    'valign': 'vcenter'
                })
                header_format = workbook.add_format({
                    'bold': True,
                    'text_wrap': True,
                    'valign': 'top',
                    'fg_color': '#D7E4BC',
                    'border': 1
                })
                
                # Summary sheet
                worksheet = workbook.add_worksheet('Ringkasan Kelas')
                
                # Add title
                worksheet.merge_range('A1:H1', f"RINGKASAN HASIL TES KELAS {class_id} - {collection.name}", title_format)
                
                # Set up headers for class summary
                headers = ['No', 'Nama', 'Kelas', 'level', 'Jawaban Benar', 'Jawaban Salah', 'Total', 'Akurasi (%)', 'Status']
                for col, header in enumerate(headers):
                    worksheet.write(2, col, header, header_format)
                
                # Gather student data
                class_data = []
                row = 3
                student_number = 1
                
                total_correct = 0
                total_incorrect = 0
                level_distribution = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0, 7: 0}
                completed_count = 0
                
                for student in students:
                    # Get student result
                    result = SiswaResult.query.filter_by(
                        siswa_id=student.id,
                        collection_id=collection_id
                    ).first()
                    
                    if result:
                        # Update statistics
                        correct = result.correct or 0
                        incorrect = result.incorrect or 0
                        total = correct + incorrect
                        accuracy = round((correct / total * 100), 2) if total > 0 else 0
                        current_level = result.current_level or 1
                        
                        # Track level distribution
                        if current_level in level_distribution:
                            level_distribution[current_level] += 1
                        
                        # Track completion status
                        is_completed = current_level in [5, 6, 7]
                        if is_completed:
                            completed_count += 1
                            
                        # Update totals
                        total_correct += correct
                        total_incorrect += incorrect
                        
                        # Write row data
                        worksheet.write(row, 0, student_number)
                        worksheet.write(row, 1, student.nama)
                        worksheet.write(row, 2, student.kelas)
                        worksheet.write(row, 3, current_level)
                        worksheet.write(row, 4, correct)
                        worksheet.write(row, 5, incorrect)
                        worksheet.write(row, 6, total)
                        worksheet.write(row, 7, accuracy)
                        worksheet.write(row, 8, "Selesai" if is_completed else "Sedang Mengerjakan" if total > 0 else "Belum Mulai")
                        
                        # Prepare data for DataFrame
                        class_data.append({
                            'Nama': student.nama,
                            'Kelas': student.kelas,
                            'level': current_level,
                            'Jawaban Benar': correct,
                            'Jawaban Salah': incorrect,
                            'Total': total,
                            'Akurasi (%)': accuracy,
                            'Status': "Selesai" if is_completed else "Sedang Mengerjakan" if total > 0 else "Belum Mulai"
                        })
                        
                        row += 1
                        student_number += 1
                
                # Add totals row
                total_students = len(students)
                total_answers = total_correct + total_incorrect
                avg_accuracy = round((total_correct / total_answers * 100), 2) if total_answers > 0 else 0
                completion_rate = round((completed_count / total_students * 100), 2) if total_students > 0 else 0
                
                # Skip a row
                row += 1
                
                # Summary row
                bold_format = workbook.add_format({'bold': True})
                worksheet.write(row, 0, "SUMMARY", bold_format)
                worksheet.write(row, 1, f"Total Siswa: {total_students}")
                worksheet.write(row, 4, f"Total Benar: {total_correct}")
                worksheet.write(row, 5, f"Total Salah: {total_incorrect}")
                worksheet.write(row, 6, f"Total Jawaban: {total_answers}")
                worksheet.write(row, 7, f"Akurasi Rata-rata: {avg_accuracy}%")
                worksheet.write(row, 8, f"Penyelesaian: {completion_rate}%")
                
                # Format columns
                worksheet.set_column('A:A', 5)
                worksheet.set_column('B:B', 25)
                worksheet.set_column('C:C', 10)
                worksheet.set_column('D:D', 10)
                worksheet.set_column('E:E', 15)
                worksheet.set_column('F:F', 15)
                worksheet.set_column('G:G', 10)
                worksheet.set_column('H:H', 12)
                worksheet.set_column('I:I', 18)
                
                # level Distribution Chart Sheet
                level_sheet = workbook.add_worksheet('Distribusi level')
                
                # Create level distribution data
                level_data = []
                for level, count in level_distribution.items():
                    level_data.append({
                        'level': f"level {level}",
                        'Jumlah Siswa': count,
                        'Persentase': round((count / total_students * 100), 2) if total_students > 0 else 0
                    })
                
                # Write level distribution data
                level_sheet.write(0, 0, "Distribusi level", title_format)
                level_headers = ['level', 'Jumlah Siswa', 'Persentase (%)']
                for col, header in enumerate(level_headers):
                    level_sheet.write(2, col, header, header_format)
                
                for i, data in enumerate(level_data):
                    level_sheet.write(i + 3, 0, data['level'])
                    level_sheet.write(i + 3, 1, data['Jumlah Siswa'])
                    level_sheet.write(i + 3, 2, data['Persentase'])
                
                # Format columns
                level_sheet.set_column('A:A', 15)
                level_sheet.set_column('B:B', 15)
                level_sheet.set_column('C:C', 15)
                
                # Create a chart for level distribution
                chart = workbook.add_chart({'type': 'column'})
                
                # Add data series
                chart.add_series({
                    'name': 'Jumlah Siswa',
                    'categories': ['Distribusi level', 3, 0, 9, 0],
                    'values': ['Distribusi level', 3, 1, 9, 1],
                    'data_labels': {'value': True}
                })
                
                # Configure chart
                chart.set_title({'name': 'Distribusi level'})
                chart.set_x_axis({'name': 'level'})
                chart.set_y_axis({'name': 'Jumlah Siswa'})
                
                # Insert chart
                level_sheet.insert_chart('E3', chart, {'x_offset': 25, 'y_offset': 10, 'x_scale': 1.5, 'y_scale': 1.5})
                
                # If we have student data, create a full class data sheet
                if class_data:
                    df_class = pd.DataFrame(class_data)
                    df_class.to_excel(writer, sheet_name='Data Lengkap', index=False)
                    
                    # Format columns
                    class_sheet = writer.sheets['Data Lengkap']
                    class_sheet.set_column('A:A', 25)
                    class_sheet.set_column('B:B', 10)
                    class_sheet.set_column('C:C', 10)
                    class_sheet.set_column('D:D', 15)
                    class_sheet.set_column('E:E', 15)
                    class_sheet.set_column('F:F', 10)
                    class_sheet.set_column('G:G', 12)
                    class_sheet.set_column('H:H', 18)
            
            # Generate filename
            filename = f"Hasil_Kelas_{class_id}_{collection.name}_{datetime.datetime.now().strftime('%Y%m%d')}.xlsx"
        
        # Reset file pointer to beginning
        output.seek(0)
        
        # Send file as attachment
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        print(f"Error exporting to Excel: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f"Error: {str(e)}"}), 500

@app.route('/api/export_all_collection_data', methods=['GET'])
@login_required
def export_all_collection_data():
    try:
        collection_id = request.args.get('collection_id')
        
        if not collection_id:
            return jsonify({'success': False, 'message': 'Collection ID is required'}), 400
            
        # Get collection info
        collection = QuestionCollection.query.get(collection_id)
        if not collection:
            return jsonify({'success': False, 'message': 'Collection not found'}), 404
            
        # Verify ownership
        if collection.guru_id != current_user.id:
            return jsonify({'success': False, 'message': 'You do not have permission to export this collection'}), 403
            
        # Get all students assigned to this collection
        students_query = db.session.query(Siswa).join(
            collection_students,
            collection_students.c.siswa_id == Siswa.id
        ).filter(
            collection_students.c.collection_id == collection_id
        ).all()
        
        if not students_query:
            return jsonify({'success': False, 'message': 'No students assigned to this collection'}), 404
            
        # Create BytesIO object for the Excel file
        output = BytesIO()
        
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            workbook = writer.book
            
            # ===== FORMAT DEFINITIONS =====
            title_format = workbook.add_format({
                'bold': True,
                'font_size': 16,
                'align': 'center',
                'valign': 'vcenter'
            })
            header_format = workbook.add_format({
                'bold': True,
                'text_wrap': True,
                'valign': 'top',
                'fg_color': '#D7E4BC',
                'border': 1
            })
            label_format = workbook.add_format({
                'bold': True,
                'align': 'right',
                'border': 1
            })
            data_format = workbook.add_format({
                'align': 'left',
                'border': 1
            })
            num_format = workbook.add_format({
                'align': 'right',
                'border': 1,
                'num_format': '0'
            })
            percent_format = workbook.add_format({
                'align': 'right',
                'border': 1,
                'num_format': '0.00%'
            })
            date_format = workbook.add_format({
                'align': 'left',
                'border': 1,
                'num_format': 'dd/mm/yyyy hh:mm:ss'
            })
            heading_format = workbook.add_format({
                'bold': True,
                'font_size': 12,
                'align': 'left',
                'valign': 'vcenter'
            })
            
            # ===== COLLECTION SUMMARY SHEET =====
            summary_sheet = workbook.add_worksheet('Ringkasan Koleksi')
            
            # Title
            summary_sheet.merge_range('A1:H1', f"RINGKASAN KOLEKSI - {collection.name}", title_format)
            
            # Collection details
            summary_sheet.write('A3', 'Nama Koleksi:', label_format)
            summary_sheet.write('B3', collection.name, data_format)
            
            summary_sheet.write('A4', 'Deskripsi:', label_format)
            summary_sheet.write('B4', collection.description or 'Tidak ada deskripsi', data_format)
            
            summary_sheet.write('A5', 'Guru:', label_format)
            guru = Guru.query.get(collection.guru_id)
            summary_sheet.write('B5', guru.nama if guru else 'Unknown', data_format)
            
            summary_sheet.write('A6', 'Tanggal Dibuat:', label_format)
            summary_sheet.write('B6', collection.created_at.strftime('%d-%m-%Y %H:%M:%S') if collection.created_at else 'Unknown', data_format)
            
            # Calculate overall statistics
            total_students = len(students_query)
            total_correct = 0
            total_incorrect = 0
            completed_count = 0
            level_distribution = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0, 7: 0}
            
            for student in students_query:
                # Get result
                result = SiswaResult.query.filter_by(
                    siswa_id=student.id,
                    collection_id=collection_id
                ).first()
                
                if result:
                    total_correct += result.correct or 0
                    total_incorrect += result.incorrect or 0
                    
                    # Track level distribution
                    level = result.current_level or 1
                    if level in level_distribution:
                        level_distribution[level] += 1
                    
                    # Check if completed
                    if level in [5, 6, 7]:
                        completed_count += 1
            
            # Overall statistics
            total_answers = total_correct + total_incorrect
            avg_accuracy = (total_correct / total_answers * 100) if total_answers > 0 else 0
            completion_rate = (completed_count / total_students * 100) if total_students > 0 else 0
            
            summary_sheet.write('D3', 'Total Siswa:', label_format)
            summary_sheet.write('E3', total_students, num_format)
            
            summary_sheet.write('D4', 'Jawaban Benar:', label_format)
            summary_sheet.write('E4', total_correct, num_format)
            
            summary_sheet.write('D5', 'Jawaban Salah:', label_format)
            summary_sheet.write('E5', total_incorrect, num_format)
            
            summary_sheet.write('D6', 'Total Jawaban:', label_format)
            summary_sheet.write('E6', total_answers, num_format)
            
            summary_sheet.write('G3', 'Akurasi Rata-rata:', label_format)
            summary_sheet.write('H3', avg_accuracy / 100, percent_format)
            
            summary_sheet.write('G4', 'Tingkat Penyelesaian:', label_format)
            summary_sheet.write('H4', completion_rate / 100, percent_format)
            
            summary_sheet.write('G5', 'Siswa Selesai:', label_format)
            summary_sheet.write('H5', f"{completed_count} / {total_students}", data_format)
            
            # level Distribution section
            summary_sheet.write('A9', 'Distribusi level:', heading_format)
            
            summary_sheet.write('A10', 'level', header_format)
            summary_sheet.write('B10', 'Jumlah Siswa', header_format)
            summary_sheet.write('C10', 'Persentase', header_format)
            
            row = 11
            for level, count in level_distribution.items():
                percentage = (count / total_students * 100) if total_students > 0 else 0
                
                summary_sheet.write(f'A{row}', f"level {level}", data_format)
                summary_sheet.write(f'B{row}', count, num_format)
                summary_sheet.write(f'C{row}', percentage / 100, percent_format)
                
                row += 1
            
            # Format column widths
            summary_sheet.set_column('A:A', 20)
            summary_sheet.set_column('B:B', 30)
            summary_sheet.set_column('C:C', 15)
            summary_sheet.set_column('D:D', 20)
            summary_sheet.set_column('E:E', 15)
            summary_sheet.set_column('F:F', 5)
            summary_sheet.set_column('G:G', 20)
            summary_sheet.set_column('H:H', 15)
            
            # ===== STUDENT RESULTS SHEET =====
            students_sheet = workbook.add_worksheet('Data Siswa')
            
            # Title
            students_sheet.merge_range('A1:I1', f"DATA SISWA - {collection.name}", title_format)
            
            # Headers
            headers = ['No', 'Nama', 'Kelas', 'level', 'Jawaban Benar', 'Jawaban Salah', 'Total', 'Akurasi (%)', 'Status']
            for col, header in enumerate(headers):
                students_sheet.write(2, col, header, header_format)
            
            # Student data
            student_data = []
            row = 3
            for i, student in enumerate(students_query, 1):
                # Get result
                result = SiswaResult.query.filter_by(
                    siswa_id=student.id,
                    collection_id=collection_id
                ).first()
                
                if result:
                    correct = result.correct or 0
                    incorrect = result.incorrect or 0
                    total = correct + incorrect
                    accuracy = (correct / total * 100) if total > 0 else 0
                    current_level = result.current_level or 1
                    is_completed = current_level in [5, 6, 7]
                    status = "Selesai" if is_completed else "Sedang Mengerjakan" if total > 0 else "Belum Mulai"
                else:
                    correct = incorrect = total = accuracy = 0
                    current_level = 1
                    is_completed = False
                    status = "Belum Mulai"
                
                # Write to Excel
                students_sheet.write(row, 0, i, num_format)
                students_sheet.write(row, 1, student.nama, data_format)
                students_sheet.write(row, 2, student.kelas or "-", data_format)
                students_sheet.write(row, 3, current_level, num_format)
                students_sheet.write(row, 4, correct, num_format)
                students_sheet.write(row, 5, incorrect, num_format)
                students_sheet.write(row, 6, total, num_format)
                students_sheet.write(row, 7, accuracy, num_format)
                students_sheet.write(row, 8, status, data_format)
                
                # Add to data for pandas DataFrame
                student_data.append({
                    'Nama': student.nama,
                    'Kelas': student.kelas or "-",
                    'level': current_level,
                    'Jawaban Benar': correct,
                    'Jawaban Salah': incorrect,
                    'Total': total,
                    'Akurasi (%)': accuracy,
                    'Status': status
                })
                
                row += 1
            
            # Format columns
            students_sheet.set_column('A:A', 5)
            students_sheet.set_column('B:B', 25)
            students_sheet.set_column('C:C', 10)
            students_sheet.set_column('D:D', 10)
            students_sheet.set_column('E:E', 15)
            students_sheet.set_column('F:F', 15)
            students_sheet.set_column('G:G', 10)
            students_sheet.set_column('H:H', 12)
            students_sheet.set_column('I:I', 18)
            
            # ===== CLASS SUMMARY SHEET =====
            # Group students by class
            classes = {}
            for student in students_query:
                class_name = student.kelas or "No Class"
                if class_name not in classes:
                    classes[class_name] = []
                classes[class_name].append(student.id)
            
            # Create class summary
            class_sheet = workbook.add_worksheet('Ringkasan Kelas')
            
            # Title
            class_sheet.merge_range('A1:H1', f"RINGKASAN PER KELAS - {collection.name}", title_format)
            
            # Headers
            class_headers = ['Kelas', 'Jumlah Siswa', 'Jawaban Benar', 'Jawaban Salah', 'Total Jawaban', 'Akurasi (%)', 'Siswa Selesai', 'Persentase Selesai (%)']
            for col, header in enumerate(class_headers):
                class_sheet.write(2, col, header, header_format)
            
            # Class data
            row = 3
            for class_name, student_ids in classes.items():
                student_count = len(student_ids)
                class_correct = 0
                class_incorrect = 0
                class_completed = 0
                
                # Calculate stats for this class
                for student_id in student_ids:
                    result = SiswaResult.query.filter_by(
                        siswa_id=student_id,
                        collection_id=collection_id
                    ).first()
                    
                    if result:
                        class_correct += result.correct or 0
                        class_incorrect += result.incorrect or 0
                        
                        if result.current_level in [5, 6, 7]:
                            class_completed += 1
                
                class_total = class_correct + class_incorrect
                class_accuracy = (class_correct / class_total * 100) if class_total > 0 else 0
                completion_percentage = (class_completed / student_count * 100) if student_count > 0 else 0
                
                # Write to Excel
                class_sheet.write(row, 0, class_name, data_format)
                class_sheet.write(row, 1, student_count, num_format)
                class_sheet.write(row, 2, class_correct, num_format)
                class_sheet.write(row, 3, class_incorrect, num_format)
                class_sheet.write(row, 4, class_total, num_format)
                class_sheet.write(row, 5, class_accuracy, num_format)
                class_sheet.write(row, 6, f"{class_completed} / {student_count}", data_format)
                class_sheet.write(row, 7, completion_percentage, num_format)
                
                row += 1
            
            # Format columns
            class_sheet.set_column('A:A', 20)
            class_sheet.set_column('B:B', 15)
            class_sheet.set_column('C:C', 15)
            class_sheet.set_column('D:D', 15)
            class_sheet.set_column('E:E', 15)
            class_sheet.set_column('F:F', 15)
            class_sheet.set_column('G:G', 15)
            class_sheet.set_column('H:H', 20)
            
            # ===== QUESTIONS ANALYSIS SHEET =====
            # Get questions in this collection
            questions = db.session.query(Question).join(
                CollectionQuestion, 
                CollectionQuestion.question_id == Question.id
            ).filter(
                CollectionQuestion.collection_id == collection_id
            ).all()
            
            # Create question analysis sheet
            if questions:
                question_sheet = workbook.add_worksheet('Analisis Soal')
                
                # Title
                question_sheet.merge_range('A1:H1', f"ANALISIS SOAL - {collection.name}", title_format)
                
                # Headers
                question_headers = ['ID', 'level', 'Soal', 'Tingkat Kesulitan (p)', 'Jawaban Benar', 'Jawaban Salah', 'Total Jawaban', 'Akurasi (%)']
                for col, header in enumerate(question_headers):
                    question_sheet.write(2, col, header, header_format)
                
                # Question data
                row = 3
                for question in questions:
                    # Calculate question statistics
                    answers = db.session.query(SiswaAnswer).filter_by(
                        question_id=question.id
                    ).all()
                    
                    correct_count = sum(1 for a in answers if a.is_correct)
                    incorrect_count = len(answers) - correct_count
                    total_answers = len(answers)
                    accuracy = (correct_count / total_answers * 100) if total_answers > 0 else 0
                    
                    # Write to Excel
                    question_sheet.write(row, 0, question.id, num_format)
                    question_sheet.write(row, 1, question.level, num_format)
                    question_sheet.write(row, 2, question.soal, data_format)
                    question_sheet.write(row, 3, question.p, num_format)
                    question_sheet.write(row, 4, correct_count, num_format)
                    question_sheet.write(row, 5, incorrect_count, num_format)
                    question_sheet.write(row, 6, total_answers, num_format)
                    question_sheet.write(row, 7, accuracy, num_format)
                    
                    row += 1
                
                # Format columns
                question_sheet.set_column('A:A', 10)
                question_sheet.set_column('B:B', 10)
                question_sheet.set_column('C:C', 50)
                question_sheet.set_column('D:D', 20)
                question_sheet.set_column('E:E', 15)
                question_sheet.set_column('F:F', 15)
                question_sheet.set_column('G:G', 15)
                question_sheet.set_column('H:H', 15)
                
                # Create charts worksheet
                charts_sheet = workbook.add_worksheet('Grafik')
                
                # Title
                charts_sheet.merge_range('A1:H1', f"GRAFIK ANALISIS - {collection.name}", title_format)
                
                # Add level distribution chart
                level_chart = workbook.add_chart({'type': 'column'})
                
                # Configure the chart
                level_chart.add_series({
                    'name': 'Jumlah Siswa',
                    'categories': ['Ringkasan Koleksi', 11, 0, 17, 0],
                    'values': ['Ringkasan Koleksi', 11, 1, 17, 1],
                    'data_labels': {'value': True}
                })
                
                level_chart.set_title({'name': 'Distribusi level'})
                level_chart.set_x_axis({'name': 'level'})
                level_chart.set_y_axis({'name': 'Jumlah Siswa'})
                
                # Insert chart
                charts_sheet.insert_chart('A3', level_chart, {'x_offset': 25, 'y_offset': 10, 'x_scale': 1.5, 'y_scale': 1.2})
                
                # Add class comparison chart
                class_chart = workbook.add_chart({'type': 'column'})
                
                # Configure the chart
                class_chart.add_series({
                    'name': 'Akurasi (%)',
                    'categories': ['Ringkasan Kelas', 3, 0, 3 + len(classes) - 1, 0],
                    'values': ['Ringkasan Kelas', 3, 5, 3 + len(classes) - 1, 5],
                    'data_labels': {'value': True}
                })
                
                class_chart.set_title({'name': 'Akurasi per Kelas'})
                class_chart.set_x_axis({'name': 'Kelas'})
                class_chart.set_y_axis({'name': 'Akurasi (%)'})
                
                # Insert chart
                charts_sheet.insert_chart('A20', class_chart, {'x_offset': 25, 'y_offset': 10, 'x_scale': 1.5, 'y_scale': 1.2})
        
        # Reset file pointer to beginning
        output.seek(0)
        
        # Create filename
        filename = f"Data_Koleksi_{collection.name}_{datetime.datetime.now().strftime('%Y%m%d')}.xlsx"
        
        # Send file as attachment
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        print(f"Error exporting collection data: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f"Error: {str(e)}"}), 500

@app.route('/api/export_result_pdf', methods=['GET'])
@login_required
def export_result_pdf():
    try:
        # Ambil parameter
        user_id = request.args.get('user_id')
        collection_id = request.args.get('collection_id')
        
        # Validasi parameter
        if not user_id or not collection_id:
            return jsonify({'success': False, 'message': 'Parameter tidak lengkap'}), 400
            
        # Ambil info siswa
        student = Siswa.query.get(user_id)
        if not student:
            return jsonify({'success': False, 'message': 'Siswa tidak ditemukan'}), 404
            
        # Ambil info koleksi
        collection = QuestionCollection.query.get(collection_id)
        if not collection:
            return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan'}), 404
            
        # Ambil hasil tes
        result = SiswaResult.query.filter_by(
            siswa_id=user_id,
            collection_id=collection_id
        ).first()
        
        if not result:
            return jsonify({'success': False, 'message': 'Hasil tes tidak ditemukan'}), 404
        
        # Ambil jawaban siswa
        answers = get_student_answer_history(user_id, collection_id)
        
        # Hitung statistik
        total = result.correct + result.incorrect
        accuracy = round((result.correct / total * 100), 2) if total > 0 else 0
        
        # Dapatkan rekomendasi dari Gemini
        recommendations = generate_ai_recommendations(student, result, collection_id)
        
        # Fungsi helper untuk template
        def now():
            return datetime.datetime.now()
        
        # Render template HTML untuk PDF
        html = render_template(
            'export_pdf.html',
            student=student,
            collection=collection,
            result=result,
            answers=answers,
            accuracy=accuracy,
            recommendations=recommendations,
            total_questions=total,
            now=now  # Kirim fungsi now
        )
        
        # Fungsi untuk mengkonversi HTML ke PDF
        def html_to_pdf(source_html):
            # Buat buffer untuk menyimpan PDF
            result_buffer = BytesIO()
            
            # Konversi HTML ke PDF
            pdf_status = pisa.CreatePDF(
                source_html,              # HTML string
                dest=result_buffer        # Output
            )
            
            # Reset pointer ke awal buffer
            result_buffer.seek(0)
            
            # Jika ada error, kembalikan None
            if pdf_status.err:
                return None
                
            return result_buffer
        
        # Konversi HTML ke PDF dan simpan ke buffer
        pdf_buffer = html_to_pdf(html)
        
        if pdf_buffer is None:
            return jsonify({'success': False, 'message': 'Gagal membuat PDF'}), 500
        
        # Buat nama file untuk didownload
        filename = f"Hasil_Tes_{student.nama}_{collection.name}_{datetime.datetime.now().strftime('%Y%m%d')}.pdf"
        
        # Kirim file
        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/pdf'
        )
        
    except Exception as e:
        print(f"Error exporting to PDF: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f"Error: {str(e)}"}), 500

# Fungsi untuk generate rekomendasi AI
def generate_ai_recommendations(student, result, collection_id):
    try:
        # Get collection name
        collection = QuestionCollection.query.get(collection_id)
        collection_name = collection.name if collection else "N/A"
        
        # Buat prompt untuk Gemini
        prompt = f"""
        Berikan 4-5 rekomendasi pembelajaran spesifik untuk seorang siswa berdasarkan hasil tes diagnostik berikut:
        
        Nama Siswa: {student.nama}
        Kelas: {student.kelas}
        Materi Tes: {collection_name}
        
        Hasil Tes:
        - level yang dicapai: {result.current_level}
        - Jawaban Benar: {result.correct}
        - Jawaban Salah: {result.incorrect}
        - Akurasi: {round((result.correct / (result.correct + result.incorrect) * 100), 2) if (result.correct + result.incorrect) > 0 else 0}%
        
        Keterangan level:
        - level 1: Kesadaran Teknologi (Pemahaman Dasar)
        - level 2: Literasi Teknologi (Komprehensi)
        - level 3: Kemampuan Teknologi Dasar (Aplikasi)
        - level 4: Kemampuan Teknologi Lanjut (Aplikasi)
        - level 5: Kreativitas Teknologi Dasar
        - level 6: Kreativitas Teknologi Menengah
        - level 7: Kritik Teknologi (Evaluasi Mendalam)
        
        Berikan rekomendasi spesifik yang relevan dengan materi, tingkat kemampuan (berdasarkan level), dan area yang perlu ditingkatkan.
        Berikan rekomendasi dalam format bullet points. Setiap poin maksimal 2 kalimat.
        """
        
        # Panggil Gemini API
        response = model.generate_content(prompt)
        recommendations_text = response.text
        
        # Parse hasil menjadi array
        recommendations = []
        lines = recommendations_text.split("\n")
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Cek apakah baris mengandung bullet point atau nomor
            if re.match(r'^[\*\-\•\+\d\.\s]+(.+)$', line):
                clean_line = re.sub(r'^[\*\-\•\+\d\.\s]+', '', line)
                recommendations.append(clean_line.strip())
            else:
                # Jika tidak ada bullet point, tambahkan sebagai item terpisah
                recommendations.append(line)
        
        # Jika tidak ada rekomendasi dari Gemini, berikan default
        if not recommendations:
            recommendations = get_default_recommendations(result.current_level, result.correct, result.incorrect)
            
        return recommendations
        
    except Exception as e:
        print(f"Error generating AI recommendations: {str(e)}")
        # Fallback ke rekomendasi default
        return get_default_recommendations(result.current_level, result.correct, result.incorrect)

def get_default_recommendations(current_level, correct, incorrect):
    # Hitung akurasi
    total = correct + incorrect
    accuracy = round((correct / total * 100), 2) if total > 0 else 0
    
    recommendations = []
    
    # Rekomendasi berdasarkan akurasi
    if accuracy < 50:
        recommendations.append("Fokus pada peningkatan pemahaman dasar materi.")
        recommendations.append("Perbaiki konsep dasar untuk meningkatkan kemampuan.")
        recommendations.append("Cari sumber belajar tambahan untuk materi yang belum dipahami.")
    elif accuracy < 75:
        recommendations.append("Tingkatkan pemahaman pada level menengah (3-4).")
        recommendations.append("Pelajari kembali materi yang berkaitan dengan soal yang dijawab salah.")
        recommendations.append("Latih kemampuan aplikasi konsep pada konteks berbeda.")
    else:
        recommendations.append("Lanjutkan pembelajaran ke materi yang lebih kompleks.")
        recommendations.append("Coba tes pada koleksi lain untuk memperluas penguasaan materi.")
        recommendations.append("Bantu teman yang mungkin kesulitan dengan materi yang sudah Anda kuasai.")
    
    # Rekomendasi berdasarkan level
    if current_level < 5:
        recommendations.append(f"Ulangi tes untuk mencapai level yang lebih tinggi (saat ini: {current_level}).")
    else:
        recommendations.append("Selamat mencapai level akhir! Coba tes lain untuk tantangan baru.")
    
    return recommendations

@app.route('/api/collections/<int:collection_id>/questions/<int:question_id>', methods=['DELETE'])
@login_required
def delete_question_from_collection(collection_id, question_id):
    try:
        # 1. Verifikasi kepemilikan koleksi dan hak akses guru
        collection = QuestionCollection.query.filter_by(id=collection_id, guru_id=current_user.id).first()
        if not collection:
            return jsonify({'success': False, 'message': 'Koleksi tidak ditemukan atau Anda tidak memiliki izin.'}), 404

        # 2. Cari relasi soal di koleksi
        collection_question = CollectionQuestion.query.filter_by(
            collection_id=collection_id,
            question_id=question_id
        ).first()

        if not collection_question:
            return jsonify({'success': False, 'message': 'Soal tidak ditemukan dalam koleksi ini.'}), 404

        # Simpan question_id dan question_obj sebelum dihapus
        question_obj = Question.query.get(question_id)
        if not question_obj: # Seharusnya tidak terjadi jika collection_question ada
             return jsonify({'success': False, 'message': 'Soal tidak ditemukan di database.'}), 404
        
        # 3. Hapus relasi dari koleksi
        db.session.delete(collection_question)
        db.session.commit() # Commit dulu untuk merefleksikan perubahan relasi

        # 4. Cek apakah soal tersebut masih terkait dengan koleksi lain
        remaining_links = CollectionQuestion.query.filter_by(question_id=question_id).first()

        if not remaining_links:
            # Jika soal tidak terkait dengan koleksi lain, hapus soal dan jawaban terkaitnya
            print(f"Soal ID {question_id} tidak lagi terkait dengan koleksi manapun. Menghapus soal dan jawaban siswa.")
            
            # Hapus jawaban siswa yang terkait dengan soal ini
            SiswaAnswer.query.filter_by(question_id=question_id).delete(synchronize_session=False)
            db.session.commit() # Commit penghapusan jawaban
            
            # Hapus soal itu sendiri
            db.session.delete(question_obj)
            db.session.commit() # Commit penghapusan soal
            
            message = f"Soal ID {question_id} berhasil dihapus dari koleksi dan database."
        else:
            message = f"Soal ID {question_id} berhasil dihapus dari koleksi, tetapi tetap ada di database karena masih terkait dengan koleksi lain."
        
        # 5. Perbarui statistik SiswaResult jika ada perubahan
        # Ini penting agar jumlah soal benar/salah dan current_level siswa tetap akurat
        # setelah soal dihapus. Ini bisa dilakukan dengan memanggil refresh_siswa_answers
        # untuk setiap siswa yang mungkin terpengaruh atau dengan merefresh data di frontend.
        # Untuk kesederhanaan di backend, kita bisa mengandalkan frontend untuk me-refresh datanya.
        
        return jsonify({'success': True, 'message': message}), 200

    except Exception as e:
        db.session.rollback() # Rollback jika terjadi kesalahan
        print(f"Error saat menghapus soal dari koleksi: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f"Terjadi kesalahan server: {str(e)}"}), 500
# =========================================
# MAIN ENTRY POINT
# =========================================
if __name__ == '__main__':  # Fixed from _main_ to __main__
    with app.app_context():
        try:
            # Inisialisasi database
            init_db()
            # Tes koneksi database
            test_db_connection()
        except Exception as e:
            print(f"Error inisialisasi database: {str(e)}")



    app.run(debug=True)